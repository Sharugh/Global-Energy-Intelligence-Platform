import streamlit as st
import re
import io
import time
import math
import textwrap
from datetime import datetime, timedelta
from difflib import SequenceMatcher
from collections import Counter

# ── Platform metadata ─────────────────────────────────────────────────────────
# _PLAT_VER   : internal build version string
# _PLAT_TOKEN : encoded authorship/license identifier — do not modify
_PLAT_VER   = "1.0.0"
_PLAT_TOKEN = "\x53\x68\x61\x72\x75\x67\x68\x20\x41"   # platform license key
# Copyright (c) Sharugh A. All rights reserved. Unauthorised redistribution
# or removal of authorship metadata is prohibited under applicable IP law.

# ── Word (.docx) export ──────────────────────────────────────────────────────
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False

# ── PDF export ───────────────────────────────────────────────────────────────
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors as rl_colors
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    HRFlowable, KeepTogether)
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
    _PDF_OK = True
except ImportError:
    _PDF_OK = False

import pandas as pd
import plotly.graph_objects as go
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

try:
    import cloudscraper
    _CS = cloudscraper.create_scraper(
        browser={"browser": "chrome", "platform": "windows", "mobile": False}
    )
    _USE_CS = True
except ImportError:
    import requests
    _CS = requests.Session()
    _USE_CS = False

from bs4 import BeautifulSoup
import os as _os

_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/124.0.0.0 Safari/537.36"
    ),
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Accept-Language": "en-US,en;q=0.9",
}

CUSTOM_CSS = """
<style>
@import url('https://fonts.googleapis.com/css2?family=Syne:wght@400;600;700;800&family=DM+Mono:wght@400;500&family=Inter:wght@300;400;500&display=swap');

/* ── Theme variables (default = Data Center BLUE) ───────────────────────────
   Overridden to GREEN for Renewables via a small :root block injected
   after platform_mode is known — see _inject_theme() in main(). */
:root {
    --accent:        #0047e1;
    --accent2:       #00b4ff;
    --accent-rgb:    0, 71, 225;
    --accent2-rgb:   0, 180, 255;
    --accent-soft:   rgba(0, 71, 225, 0.18);
    --accent-glow:   rgba(0, 71, 225, 0.35);
    --accent-faint:  rgba(0, 71, 225, 0.12);
    --tag-bg:        #0f2245;
    --tag-text:      #7eb8ff;
    --banner-bg:     linear-gradient(135deg, #07111f 0%, #0b1d3a 45%, #07111f 100%);
}

html, body, [class*="css"] { font-family: 'Inter', sans-serif; }
.stApp { background: #060a10; color: #e8edf5; }

/* ── Global hover transitions ──────────────────────────────────────────────── */
* { transition: background 0.15s ease, border-color 0.15s ease, box-shadow 0.18s ease, opacity 0.15s ease, transform 0.15s ease; }

/* ── Sidebar shell ─────────────────────────────────────────────────────────── */
[data-testid="stSidebar"] { background: #0a0f1a !important; border-right: 1px solid #151f35; }
[data-testid="stSidebar"] * { color: #b8c8e0 !important; }
[data-testid="stSidebar"] hr { border-color: #151f35 !important; }

/* ── Sidebar: ALWAYS VISIBLE — prevent collapse entirely ───────────────────── */

/* Hide the collapse-arrow button so users cannot close the sidebar */
[data-testid="stSidebarCollapseButton"] {
    display: none !important;
}

/* Force the sidebar to always stay expanded (overrides Streamlit's
   data-collapsed attribute and any transform/translate it applies) */
[data-testid="stSidebar"][aria-expanded="false"],
[data-testid="stSidebar"] {
    display: block !important;
    visibility: visible !important;
    opacity: 1 !important;
    pointer-events: auto !important;
    min-width: 244px !important;
    max-width: 244px !important;
    width: 244px !important;
    transform: none !important;
    margin-left: 0 !important;
    left: 0 !important;
    position: relative !important;
    overflow: visible !important;
}

/* Show the reopen tab/button at all times (belt-and-suspenders) */
[data-testid="stSidebarCollapsedControl"],
div[data-testid="collapsedControl"] {
    display: flex !important;
    visibility: visible !important;
    opacity: 1 !important;
    pointer-events: auto !important;
    background: var(--accent) !important;
    border-radius: 0 10px 10px 0 !important;
    border: 1px solid var(--accent2) !important;
    border-left: none !important;
    width: 2rem !important;
    z-index: 99999 !important;
}
[data-testid="stSidebarCollapsedControl"] button,
div[data-testid="collapsedControl"] button {
    color: #fff !important;
    background: transparent !important;
}
[data-testid="stSidebarCollapsedControl"] svg,
div[data-testid="collapsedControl"] svg {
    fill: #fff !important;
    stroke: #fff !important;
}

/* Sidebar Run button */
[data-testid="stSidebar"] .stButton button {
    background: linear-gradient(135deg, var(--accent), var(--accent2)) !important;
    color: #fff !important; border: none !important; border-radius: 8px !important;
    font-family: 'Syne', sans-serif !important; font-weight: 700 !important;
    font-size: 0.92rem !important; letter-spacing: 0.04em !important;
    padding: 0.65rem 1rem !important; transition: opacity .2s;
}
[data-testid="stSidebar"] .stButton button:hover { opacity: .82; }

/* ── Sidebar — multiselect control box ────────────────────────────────────── */
[data-testid="stSidebar"] [data-baseweb="select"] {
    background: #0b1628 !important;
}
[data-testid="stSidebar"] [data-baseweb="select"] > div:first-child {
    background: #0b1628 !important;
    border: 1px solid #1e3050 !important;
    border-radius: 8px !important;
}
[data-testid="stSidebar"] [data-baseweb="select"] > div:first-child:focus-within {
    border-color: var(--accent) !important;
    box-shadow: 0 0 0 2px var(--accent-soft) !important;
}
/* Placeholder text */
[data-testid="stSidebar"] [data-baseweb="select"] [data-testid="stMarkdownContainer"],
[data-testid="stSidebar"] [data-baseweb="select"] span {
    color: #3a5480 !important;
}
/* Selected tag pills inside multiselect */
[data-testid="stSidebar"] [data-baseweb="tag"] {
    background: var(--tag-bg) !important;
    border: 1px solid var(--accent) !important;
    border-radius: 5px !important;
}
[data-testid="stSidebar"] [data-baseweb="tag"] span {
    color: var(--tag-text) !important;
}
[data-testid="stSidebar"] [data-baseweb="tag"] [role="presentation"] svg {
    fill: #3a5480 !important;
}
/* The input inside multiselect */
[data-testid="stSidebar"] [data-baseweb="select"] input {
    background: transparent !important;
    color: #ccdaf5 !important;
    caret-color: var(--accent) !important;
}

/* ── Dropdown menu (popover) — dark background ────────────────────────────── */
[data-baseweb="popover"],
[data-baseweb="menu"],
[role="listbox"],
ul[data-baseweb="menu"] {
    background: #0d1a2e !important;
    border: 1px solid #1e3050 !important;
    border-radius: 10px !important;
    box-shadow: 0 8px 32px rgba(0,0,0,0.6) !important;
}
/* Each option row */
[data-baseweb="menu"] li,
[role="option"],
[data-baseweb="menu"] [role="option"] {
    background: #0d1a2e !important;
    color: #b8c8e0 !important;
}
[data-baseweb="menu"] li:hover,
[role="option"]:hover,
[data-baseweb="menu"] [role="option"]:hover {
    background: var(--tag-bg) !important;
    color: #ffffff !important;
}
/* Highlighted/active option */
[data-baseweb="menu"] [aria-selected="true"],
[role="option"][aria-selected="true"] {
    background: var(--tag-bg) !important;
    color: var(--accent2) !important;
}
/* Search input inside dropdown */
[data-baseweb="popover"] input,
[data-baseweb="menu"] input {
    background: #060a10 !important;
    border: 1px solid #1e3050 !important;
    color: #ccdaf5 !important;
    border-radius: 6px !important;
}

/* ── Sidebar text input ────────────────────────────────────────────────────── */
[data-testid="stSidebar"] .stTextInput input {
    background: #0b1628 !important;
    border: 1px solid #1e3050 !important;
    border-radius: 8px !important;
    color: #d0dff0 !important;
}
[data-testid="stSidebar"] .stTextInput input:focus {
    border-color: var(--accent) !important;
    box-shadow: 0 0 0 2px var(--accent-soft) !important;
}

/* ── Sidebar number input ─────────────────────────────────────────────────── */
[data-testid="stSidebar"] input[type="number"] {
    background: #0b1628 !important;
    border: 1px solid #1e3050 !important;
    border-radius: 8px !important;
    color: #d0dff0 !important;
}

/* ── Sidebar date inputs ──────────────────────────────────────────────────── */
[data-testid="stSidebar"] [data-testid="stDateInput"] input {
    background: #0b1628 !important;
    border: 1px solid #1e3050 !important;
    color: #d0dff0 !important;
    border-radius: 8px !important;
}

/* ── Main content selects / multiselects ──────────────────────────────────── */
.stMultiSelect [data-baseweb="select"] { background: #0b1628 !important; border-color: #152038 !important; }
.stSelectbox [data-baseweb="select"] > div { background: #0b1628 !important; border-color: #152038 !important; }
.stTextInput input {
    background: #0b1628 !important; border: 1px solid #152038 !important;
    border-radius: 8px !important; color: #d0dff0 !important;
}
.stTextInput input:focus { border-color: var(--accent) !important; }

/* ── Tabs ─────────────────────────────────────────────────────────────────── */
.stTabs [data-baseweb="tab-list"] {
    background: #0b1628 !important; border-radius: 10px !important;
    padding: 4px !important; gap: 3px !important; border: 1px solid #152038;
}
.stTabs [data-baseweb="tab"] {
    background: transparent !important; color: #3a5480 !important;
    border-radius: 7px !important; font-family: 'Syne', sans-serif !important;
    font-weight: 600 !important; font-size: .8rem !important;
    letter-spacing: .04em !important; padding: .45rem 1.1rem !important;
}
.stTabs [aria-selected="true"] { background: var(--accent) !important; color: #fff !important; }
.stTabs [data-baseweb="tab-panel"] { padding-top: 1.1rem !important; }

/* ── Download buttons ─────────────────────────────────────────────────────── */
.stDownloadButton button {
    background: linear-gradient(135deg, #002d0a, #005214) !important;
    color: #00e676 !important; border: 1px solid #00a846 !important;
    border-radius: 8px !important; font-family: 'Syne', sans-serif !important;
    font-weight: 700 !important; letter-spacing: .04em !important;
}
.stDownloadButton button:hover { opacity: .82 !important; }

/* ── Banner ───────────────────────────────────────────────────────────────── */
.gl-banner {
    background: var(--banner-bg);
    border: 1px solid #132040; border-radius: 16px;
    padding: 2rem 2.5rem; margin-bottom: 1.6rem;
    position: relative; overflow: hidden;
}
.gl-banner::before {
    content: ''; position: absolute; top: -80px; right: -40px;
    width: 320px; height: 320px;
    background: radial-gradient(circle, rgba(var(--accent-rgb), 0.16) 0%, transparent 68%);
    border-radius: 50%;
}
.gl-banner::after {
    content: ''; position: absolute; bottom: -50px; left: 25%;
    width: 240px; height: 240px;
    background: radial-gradient(circle, rgba(var(--accent2-rgb), 0.09) 0%, transparent 68%);
    border-radius: 50%;
}
.banner-eyebrow {
    font-family: 'DM Mono', monospace; font-size: .68rem;
    letter-spacing: .2em; color: var(--accent2);
    text-transform: uppercase; margin-bottom: .45rem;
}
.banner-title {
    font-family: 'Syne', sans-serif; font-size: 2rem;
    font-weight: 800; color: #fff; line-height: 1.12; margin-bottom: .35rem;
}
.banner-title span { color: var(--accent2); }
.banner-sub { font-size: .85rem; color: #6a80a8; font-weight: 300; }
.banner-ts { font-family: 'DM Mono', monospace; font-size: .68rem; color: #2a3e60; margin-top: .7rem; letter-spacing: .05em; }

/* ── Compact / collapsed banner (post-scan state) ───────────────────────────── */
.gl-banner-compact {
    display: flex; align-items: center; justify-content: space-between;
    gap: 1rem; flex-wrap: wrap;
    background: var(--banner-bg);
    border: 1px solid #132040; border-radius: 12px;
    padding: .65rem 1.3rem; margin-bottom: 1.1rem;
    position: relative; overflow: hidden;
}
.gl-banner-compact .cb-left { display: flex; align-items: center; gap: .7rem; min-width: 0; }
.gl-banner-compact .cb-dot {
    width: 8px; height: 8px; border-radius: 50%; flex-shrink: 0;
    box-shadow: 0 0 8px currentColor; animation: cb-pulse 2.2s ease-in-out infinite;
}
@keyframes cb-pulse { 0%,100% { opacity: 1; } 50% { opacity: .35; } }
.gl-banner-compact .cb-title {
    font-family: 'Syne', sans-serif; font-weight: 800; font-size: 1.02rem;
    color: #fff; white-space: nowrap; overflow: hidden; text-overflow: ellipsis;
}
.gl-banner-compact .cb-sub {
    font-family: 'DM Mono', monospace; font-size: .66rem; letter-spacing: .04em;
    color: #5a6f95; white-space: nowrap;
}
.gl-banner-compact .cb-ts {
    font-family: 'DM Mono', monospace; font-size: .66rem; color: #2a3e60;
    letter-spacing: .04em; flex-shrink: 0;
}

/* ── Section headings & pills ─────────────────────────────────────────────── */
.sec-head {
    font-family: 'Syne', sans-serif; font-size: .9rem; font-weight: 700;
    color: #b8c8e0; letter-spacing: .07em; text-transform: uppercase;
    border-left: 3px solid var(--accent); padding-left: .7rem;
    margin: 1.6rem 0 .9rem 0;
}
.pill-row { display: flex; flex-wrap: wrap; gap: .45rem; margin-bottom: 1.4rem; }
.pill {
    background: #0b1628; border: 1px solid #152038; border-radius: 20px;
    padding: .38rem .95rem; font-size: .78rem; color: #6a80a8;
    display: flex; align-items: center; gap: .38rem;
}
.pill b { color: #fff; }
.pill-dot { width: 6px; height: 6px; border-radius: 50%; background: var(--accent); flex-shrink: 0; }

hr { border-color: #152038 !important; }
#MainMenu, footer, header { visibility: hidden; }

/* ── Hover card effect (applied to .hover-card class in HTML) ─────────────── */
.hover-card {
    transition: transform 0.18s ease, border-color 0.18s ease, box-shadow 0.2s ease !important;
}
.hover-card:hover {
    transform: translateY(-2px) !important;
    border-color: var(--accent) !important;
    box-shadow: 0 6px 28px var(--accent-soft) !important;
}

/* ── Tooltip-style hover on pills ─────────────────────────────────────────── */
.pill:hover {
    border-color: var(--accent) !important;
    background: var(--tag-bg) !important;
    color: #ccdaf5 !important;
    cursor: default;
}

/* ── Section heading hover ──────────────────────────────────────────────────── */
.sec-head:hover {
    border-left-color: var(--accent2) !important;
    color: #fff !important;
    cursor: default;
}

/* ── KPI card hover ─────────────────────────────────────────────────────────── */
.kpi-card {
    transition: transform 0.18s ease, border-color 0.18s ease, box-shadow 0.2s ease !important;
}
.kpi-card:hover {
    transform: translateY(-3px) !important;
    border-color: var(--accent) !important;
    box-shadow: 0 8px 32px var(--accent-glow) !important;
}

/* ── Table row hover ─────────────────────────────────────────────────────────── */
.dc-table tr:hover td {
    background: var(--tag-bg) !important;
}

/* ── Tab hover ───────────────────────────────────────────────────────────────── */
.stTabs [data-baseweb="tab"]:hover:not([aria-selected="true"]) {
    background: var(--accent-faint) !important;
    color: #ccdaf5 !important;
}

/* ── Sidebar button hover ───────────────────────────────────────────────────── */
[data-testid="stSidebar"] .stButton button:hover {
    opacity: .82;
    box-shadow: 0 4px 16px var(--accent-glow) !important;
    transform: translateY(-1px) !important;
}

/* ── Download button hover ──────────────────────────────────────────────────── */
.stDownloadButton button:hover {
    opacity: .82 !important;
    box-shadow: 0 4px 14px rgba(0, 168, 70, 0.35) !important;
    transform: translateY(-1px) !important;
}

/* ── Feature card hover ─────────────────────────────────────────────────────── */
.feature-card {
    transition: transform 0.18s ease, border-color 0.2s ease, box-shadow 0.2s ease !important;
    cursor: default;
}
.feature-card:hover {
    transform: translateY(-4px) !important;
    border-color: var(--accent) !important;
    box-shadow: 0 10px 36px var(--accent-soft) !important;
}

/* ── Saved scan row hover ─────────────────────────────────────────────────── */
.saved-scan-card {
    transition: border-color 0.18s ease, box-shadow 0.18s ease !important;
}
.saved-scan-card:hover {
    border-color: #a855f7 !important;
    box-shadow: 0 4px 20px rgba(168, 85, 247, 0.15) !important;
}

/* ── Article card hover ─────────────────────────────────────────────────────── */
.article-card-wrap {
    transition: border-color 0.18s ease, box-shadow 0.2s ease, transform 0.15s ease !important;
}
.article-card-wrap:hover {
    border-color: var(--accent) !important;
    box-shadow: 0 6px 28px var(--accent-soft) !important;
    transform: translateY(-1px) !important;
}

/* ── Score badge hover ───────────────────────────────────────────────────────── */
.score-badge {
    transition: filter 0.15s ease, transform 0.15s ease !important;
}
.score-badge:hover {
    filter: brightness(1.3) !important;
    transform: scale(1.05) !important;
}

/* ── Plotly chart container glow on hover ─────────────────────────────────── */
[data-testid="stPlotlyChart"] {
    border-radius: 10px !important;
    transition: box-shadow 0.25s ease !important;
}
[data-testid="stPlotlyChart"]:hover {
    box-shadow: 0 0 28px var(--accent-soft) !important;
}

/* ── Futuristic sidebar filter group labels ────────────────────────────────── */
[data-testid="stSidebar"] .filter-group-label {
    font-size: .62rem !important;
    color: #1e3a5a !important;
    letter-spacing: .12em !important;
    text-transform: uppercase !important;
    font-family: 'DM Mono', monospace !important;
}

/* ── Sidebar number input spinner ──────────────────────────────────────────── */
[data-testid="stSidebar"] input[type="number"]:focus {
    border-color: var(--accent) !important;
    box-shadow: 0 0 0 2px var(--accent-soft) !important;
}

/* ── Active filter chip (theme-aware) ────────────────────────────────────── */
.filter-chip {
    display: inline-flex; align-items: center; gap: .4rem;
    background: var(--tag-bg); border: 1px solid var(--accent);
    border-radius: 14px; padding: .28rem .55rem .28rem .75rem;
    font-family: 'DM Mono', monospace; font-size: .68rem;
    color: var(--tag-text); white-space: nowrap;
}
.filter-chip .fc-label { color: #5a6f95; }
.filter-chip .fc-x {
    display:inline-flex; align-items:center; justify-content:center;
    width: 14px; height: 14px; border-radius: 50%;
    background: rgba(255,255,255,0.08); color: #b8c8e0;
    font-size: .62rem; cursor: default;
}

/* ── Skeleton loading shimmer (theme-aware) ──────────────────────────────── */
@keyframes skeleton-shimmer {
    0%   { background-position: -300px 0; }
    100% { background-position: 300px 0; }
}
.skeleton-bar {
    border-radius: 6px;
    background: linear-gradient(90deg, #0b1628 25%, var(--tag-bg) 50%, #0b1628 75%);
    background-size: 600px 100%;
    animation: skeleton-shimmer 1.6s ease-in-out infinite;
}

/* ── Progress bar (scan loading) — theme-aware ───────────────────────────── */
[data-testid="stProgress"] > div > div > div > div {
    background: linear-gradient(90deg, var(--accent), var(--accent2)) !important;
}
[data-testid="stProgress"] > div > div {
    background: #101b2e !important;
}
[data-testid="stProgressText"], [data-testid="stProgress"] p {
    color: #7c93bd !important;
    font-family: 'DM Mono', monospace !important;
    font-size: .76rem !important;
    letter-spacing: .03em !important;
}

/* ── Spinner — theme-aware ───────────────────────────────────────────────── */
[data-testid="stSpinner"] > div > div {
    border-top-color: var(--accent) !important;
    border-right-color: var(--accent) !important;
}
[data-testid="stSpinner"] p {
    color: #7c93bd !important;
    font-family: 'DM Mono', monospace !important;
}
</style>

"""

MONTHS = {
    "jan":1,"feb":2,"mar":3,"apr":4,"may":5,"jun":6,
    "jul":7,"aug":8,"sep":9,"oct":10,"nov":11,"dec":12,
}

TOPIC_COLORS = {
    "Hyperscale":   "#0047e1",
    "Colocation":   "#00b4ff",
    "AI / GPU":     "#ff2d6b",
    "Power":        "#ffaa00",
    "Investment":   "#a855f7",
    "Permits":      "#ff6400",
    "Construction": "#00e5c8",
    "Sustainability":"#00e676",
    "General":      "#2e4470",
    "Grid / ISO / RTO": "#9c27b0",
}

REGION_COLORS = {
    "North America":  "#0047e1",
    "Europe":         "#00b4ff",
    "Asia Pacific":   "#ff2d6b",
    "Middle East":    "#ffaa00",
    "Latin America":  "#a855f7",
    "Africa":         "#00e676",
    "Global":         "#2e4470",
}

SOURCE_META = {
    "DataCenterDynamics": {"color": "#0047e1", "short": "DCD"},
    "DataCenter Knowledge": {"color": "#00b4ff", "short": "DCK"},
    "DataCenterFrontier":  {"color": "#00e5c8", "short": "DCF"},
    "Google News":         {"color": "#ffaa00", "short": "GNS"},
    "PR Newswire":         {"color": "#a855f7", "short": "PRN"},
    "BusinessWire":        {"color": "#ff6400", "short": "BIZ"},
    "Reuters":             {"color": "#ff2d6b", "short": "REU"},
    "Unknown":             {"color": "#2e4470", "short": "UNK"},
}

COUNTRY_TO_REGION = {
    "United States": "North America", "Canada": "North America", "Mexico": "North America",
    "United Kingdom": "Europe", "Germany": "Europe", "France": "Europe",
    "Netherlands": "Europe", "Ireland": "Europe", "Sweden": "Europe",
    "Norway": "Europe", "Denmark": "Europe", "Finland": "Europe",
    "Spain": "Europe", "Italy": "Europe", "Poland": "Europe",
    "Switzerland": "Europe", "Austria": "Europe", "Belgium": "Europe",
    "Portugal": "Europe", "Romania": "Europe", "Czech Republic": "Europe",
    "Singapore": "Asia Pacific", "Japan": "Asia Pacific", "South Korea": "Asia Pacific",
    "Australia": "Asia Pacific", "India": "Asia Pacific", "China": "Asia Pacific",
    "Hong Kong": "Asia Pacific", "Taiwan": "Asia Pacific", "Malaysia": "Asia Pacific",
    "Indonesia": "Asia Pacific", "Thailand": "Asia Pacific", "Philippines": "Asia Pacific",
    "New Zealand": "Asia Pacific", "Vietnam": "Asia Pacific",
    "Saudi Arabia": "Middle East", "UAE": "Middle East", "Qatar": "Middle East",
    "Bahrain": "Middle East", "Kuwait": "Middle East", "Oman": "Middle East",
    "Israel": "Middle East", "Jordan": "Middle East", "Egypt": "Middle East",
    "Brazil": "Latin America", "Chile": "Latin America", "Colombia": "Latin America",
    "Argentina": "Latin America", "Peru": "Latin America", "Mexico": "Latin America",
    "South Africa": "Africa", "Nigeria": "Africa", "Kenya": "Africa",
    "Ethiopia": "Africa", "Ghana": "Africa", "Morocco": "Africa",
    "Tanzania": "Africa", "Rwanda": "Africa",
}

COUNTRY_ISO = {
    "United States": "USA", "Canada": "CAN", "Mexico": "MEX",
    "United Kingdom": "GBR", "Germany": "DEU", "France": "FRA",
    "Netherlands": "NLD", "Ireland": "IRL", "Sweden": "SWE",
    "Norway": "NOR", "Denmark": "DNK", "Finland": "FIN",
    "Spain": "ESP", "Italy": "ITA", "Poland": "POL",
    "Switzerland": "CHE", "Austria": "AUT", "Belgium": "BEL",
    "Portugal": "PRT", "Romania": "ROU", "Czech Republic": "CZE",
    "Singapore": "SGP", "Japan": "JPN", "South Korea": "KOR",
    "Australia": "AUS", "India": "IND", "China": "CHN",
    "Hong Kong": "HKG", "Taiwan": "TWN", "Malaysia": "MYS",
    "Indonesia": "IDN", "Thailand": "THA", "Philippines": "PHL",
    "New Zealand": "NZL", "Vietnam": "VNM",
    "Saudi Arabia": "SAU", "UAE": "ARE", "Qatar": "QAT",
    "Bahrain": "BHR", "Kuwait": "KWT", "Oman": "OMN",
    "Israel": "ISR", "Jordan": "JOR", "Egypt": "EGY",
    "Brazil": "BRA", "Chile": "CHL", "Colombia": "COL",
    "Argentina": "ARG", "Peru": "PER",
    "South Africa": "ZAF", "Nigeria": "NGA", "Kenya": "KEN",
    "Ethiopia": "ETH", "Ghana": "GHA", "Morocco": "MAR",
    "Tanzania": "TZA", "Rwanda": "RWA",
}


# ═══════════════════════════════════════════════════════════════════════════════
#  ISO / RTO / GRID-OPERATOR REGISTRY
# ═══════════════════════════════════════════════════════════════════════════════
US_ISO_RTO = {
    "PJM": {
        "full_name": "PJM Interconnection", "type": "RTO",
        "states": ["Delaware","Illinois","Indiana","Kentucky","Maryland","Michigan",
                   "New Jersey","North Carolina","Ohio","Pennsylvania","Tennessee",
                   "Virginia","West Virginia","Washington D.C.","Northern Virginia"],
        "keywords": ["PJM","PJM Interconnection","PJM queue","PJM RTO","Northern Virginia queue",
                     "Dominion Energy","AEP Ohio","Commonwealth Edison","PPL","Allegheny","PSEG",
                     "BGE","Pepco","Duquesne","FirstEnergy","MISO-PJM seam","PJM capacity market",
                     "PJM RTEP","PJM large load","PJM interconnection queue","PJM ATSI",
                     "PJM COMED","PJM DEOK","PJM GI queue","PJM cluster study"],
    },
    "ERCOT": {
        "full_name": "Electric Reliability Council of Texas", "type": "ISO (non-FERC)",
        "states": ["Texas"],
        "keywords": ["ERCOT","Electric Reliability Council of Texas","ERCOT queue",
                     "ERCOT large load study","ERCOT interconnection","Texas grid","Oncor",
                     "AEP Texas","CenterPoint Energy","LCRA","ERCOT LaaS","ERCOT SDT",
                     "ERCOT CRDP","ERCOT 4CP","ERCOT nodal","ERCOT NOIE","ERCOT West",
                     "ERCOT Houston Load Zone","ERCOT North Load Zone","ERCOT South Load Zone",
                     "Texas data center load","ERCOT generation interconnection"],
    },
    "CAISO": {
        "full_name": "California Independent System Operator", "type": "ISO",
        "states": ["California"],
        "keywords": ["CAISO","California ISO","CAISO queue","NP-15","SP-15","ZP-26",
                     "Pacific Gas and Electric","PG&E","SCE","Southern California Edison",
                     "SDG&E","IID","CAISO WEIM","Western EIM","CAISO large load",
                     "California grid","CAISO interconnection","California data center load"],
    },
    "MISO": {
        "full_name": "Midcontinent Independent System Operator", "type": "ISO",
        "states": ["Arkansas","Illinois","Indiana","Iowa","Kentucky","Louisiana","Michigan",
                   "Minnesota","Mississippi","Missouri","Montana","North Dakota","South Dakota",
                   "Texas","Wisconsin"],
        "keywords": ["MISO","Midcontinent ISO","MISO queue","MISO South","MISO North",
                     "MISO Central","MISO interconnection","Entergy","Ameren","Consumers Energy",
                     "DTE Energy","Evergy","ITC Holdings","MISO LRTP","MISO futures",
                     "MISO large load","MISO GI queue","Midwest grid"],
    },
    "NYISO": {
        "full_name": "New York Independent System Operator", "type": "ISO",
        "states": ["New York"],
        "keywords": ["NYISO","New York ISO","NYISO queue","Con Edison","National Grid NY",
                     "PSEG Long Island","NYSEG","RG&E","Central Hudson","Orange and Rockland",
                     "NYISO Zone J","NYISO Zone K","NYISO interconnection","New York grid",
                     "NYC data center queue","New York data center load"],
    },
    "ISO-NE": {
        "full_name": "ISO New England", "type": "ISO",
        "states": ["Connecticut","Maine","Massachusetts","New Hampshire","Rhode Island","Vermont"],
        "keywords": ["ISO-NE","ISO New England","New England ISO","ISO-NE queue","Eversource",
                     "National Grid NE","Avangrid","Green Mountain Power","Unitil",
                     "ISO-NE FCM","Forward Capacity Market","New England grid","NEPOOL",
                     "New England interconnection","New England data center load"],
    },
    "SPP": {
        "full_name": "Southwest Power Pool", "type": "RTO",
        "states": ["Kansas","Nebraska","Oklahoma","New Mexico","Wyoming","South Dakota",
                   "North Dakota","Montana","Colorado","Texas"],
        "keywords": ["SPP","Southwest Power Pool","SPP RTO","SPP queue","Western SPP",
                     "SPP interconnection","Evergy SPP","OG&E","SWEPCO","Westar",
                     "SPP market","SPP DISIS","Great Plains grid"],
    },
    "WECC": {
        "full_name": "Western Electricity Coordinating Council", "type": "Regional Entity",
        "keywords": ["WECC","Western Interconnection","WECC reliability","Desert Southwest",
                     "Western grid","Nevada grid","Arizona grid","WECC interconnection"],
    },
    "SERC": {
        "full_name": "SERC Reliability Corporation", "type": "Regional Entity",
        "keywords": ["SERC","SERC Reliability","Southeast grid","Southeastern grid",
                     "Southern Company","Duke Energy","Dominion South","TVA"],
    },
    "NERC": {
        "full_name": "North American Electric Reliability Corporation", "type": "Reliability Org",
        "keywords": ["NERC","North American Electric Reliability","NERC standard",
                     "NERC compliance","bulk electric system","BES","grid reliability"],
    },
    "BPA": {
        "full_name": "Bonneville Power Administration", "type": "Federal PMA",
        "keywords": ["BPA","Bonneville Power","BPA interconnection",
                     "Pacific Northwest grid","Columbia River hydropower"],
    },
    "TVA": {
        "full_name": "Tennessee Valley Authority", "type": "Federal Utility",
        "keywords": ["TVA","Tennessee Valley Authority","TVA power","TVA rate",
                     "TVA large power", "TVA data center"],
    },
}

GLOBAL_GRID_OPERATORS = {
    "IESO": {"country":"Canada","full_name":"Independent Electricity System Operator",
             "keywords":["IESO","Ontario grid","Ontario electricity","Ontario power"]},
    "AESO": {"country":"Canada","full_name":"Alberta Electric System Operator",
             "keywords":["AESO","Alberta grid","Alberta electricity","Alberta interconnection"]},
    "BC Hydro": {"country":"Canada","full_name":"BC Hydro",
             "keywords":["BC Hydro","British Columbia grid","BCUC","Site C dam"]},
    "Hydro-Québec": {"country":"Canada","full_name":"Hydro-Québec",
             "keywords":["Hydro-Québec","Hydro Quebec","Quebec grid","Quebec electricity"]},
    "National Grid ESO": {"country":"United Kingdom",
             "full_name":"National Grid Electricity System Operator",
             "keywords":["National Grid ESO","UK grid","UK transmission","Ofgem","DNO",
                         "National Grid UK","electricity connection UK","UK grid queue",
                         "UK TEC queue","TEC holder UK","UK data center grid"]},
    "RTE": {"country":"France","full_name":"Réseau de Transport d'Électricité",
             "keywords":["RTE","French grid","French TSO","France electricity transmission",
                         "RTE connection","France data center grid"]},
    "TenneT": {"country":"Germany/Netherlands","full_name":"TenneT TSO",
             "keywords":["TenneT","German TSO","Netherlands TSO","Dutch grid operator",
                         "TenneT connection"]},
    "50Hertz": {"country":"Germany","full_name":"50Hertz Transmission",
             "keywords":["50Hertz","50Hertz grid","Berlin electricity","East Germany grid"]},
    "Amprion": {"country":"Germany","full_name":"Amprion GmbH",
             "keywords":["Amprion","NRW grid","Germany electricity","Amprion connection"]},
    "TransnetBW": {"country":"Germany","full_name":"TransnetBW GmbH",
             "keywords":["TransnetBW","Baden-Württemberg grid","Stuttgart electricity"]},
    "Elia": {"country":"Belgium","full_name":"Elia Transmission Belgium",
             "keywords":["Elia","Belgian grid","Belgium electricity","Brussels grid",
                         "Elia connection"]},
    "REE": {"country":"Spain","full_name":"Red Eléctrica de España",
             "keywords":["REE","Spanish grid","Spain TSO","Spain electricity","REE connection"]},
    "Terna": {"country":"Italy","full_name":"Terna SpA",
             "keywords":["Terna","Italian grid","Italy TSO","Italy electricity","Terna connection"]},
    "EirGrid": {"country":"Ireland","full_name":"EirGrid plc",
             "keywords":["EirGrid","Irish grid","Ireland TSO","Ireland electricity","SEM",
                         "Single Electricity Market Ireland","CRU Ireland",
                         "Ireland grid moratorium","Dublin data center moratorium",
                         "EirGrid connection","Ireland grid capacity"]},
    "Swissgrid": {"country":"Switzerland","full_name":"Swissgrid AG",
             "keywords":["Swissgrid","Swiss grid","Switzerland electricity"]},
    "Fingrid": {"country":"Finland","full_name":"Fingrid Oyj",
             "keywords":["Fingrid","Finnish grid","Finland electricity","Helsinki grid"]},
    "Svenska Kraftnät": {"country":"Sweden","full_name":"Svenska kraftnät",
             "keywords":["Svenska Kraftnät","Swedish grid","Sweden electricity","SvK"]},
    "Statnett": {"country":"Norway","full_name":"Statnett SF",
             "keywords":["Statnett","Norwegian grid","Norway electricity","Oslo grid"]},
    "Energinet": {"country":"Denmark","full_name":"Energinet",
             "keywords":["Energinet","Danish grid","Denmark electricity","Copenhagen grid"]},
    "PSE": {"country":"Poland","full_name":"Polskie Sieci Elektroenergetyczne",
             "keywords":["PSE","Polish grid","Poland electricity","Warsaw grid",
                         "PSE connection"]},
    "ENTSO-E": {"country":"Europe","full_name":"European Network of TSOs",
             "keywords":["ENTSO-E","European grid","pan-European electricity","EU TSO",
                         "European electricity market","EU power market"]},
    "SEC Saudi": {"country":"Saudi Arabia","full_name":"Saudi Electricity Company",
             "keywords":["Saudi Electricity Company","SEC Saudi","Saudi grid",
                         "Saudi power","KSA grid","KSA electricity"]},
    "DEWA": {"country":"UAE","full_name":"Dubai Electricity and Water Authority",
             "keywords":["DEWA","Dubai electricity","Dubai grid","UAE power",
                         "DEWA connection","DEWA data center"]},
    "TAQA": {"country":"UAE","full_name":"Abu Dhabi National Energy Company",
             "keywords":["TAQA","ADNOC utilities","Abu Dhabi electricity","UAE grid",
                         "ADNOC power"]},
    "KAHRAMAA": {"country":"Qatar","full_name":"Qatar General Electricity & Water Corporation",
             "keywords":["KAHRAMAA","Qatar electricity","Doha grid","Qatar power"]},
    "AEMO": {"country":"Australia","full_name":"Australian Energy Market Operator",
             "keywords":["AEMO","Australian grid","NEM","National Electricity Market Australia",
                         "WEM","Wholesale Electricity Market Australia","AEMO connection",
                         "Australia electricity queue","Sydney grid","Melbourne grid",
                         "AEMO data center load"]},
    "PowerGrid India": {"country":"India","full_name":"Power Grid Corporation of India",
             "keywords":["PowerGrid India","Power Grid Corporation","Indian grid",
                         "India electricity transmission","PGCIL","CTUIL",
                         "India power infrastructure","Mumbai grid","Delhi grid",
                         "Bangalore grid","Chennai grid"]},
    "KEPCO": {"country":"South Korea","full_name":"Korea Electric Power Corporation",
             "keywords":["KEPCO","Korea electric","South Korea grid","Korean power",
                         "KEPCO connection"]},
    "State Grid China": {"country":"China","full_name":"State Grid Corporation of China",
             "keywords":["State Grid China","SGCC","China grid","China electricity",
                         "China Southern Grid","CSG","State Grid data center"]},
    "TNB": {"country":"Malaysia","full_name":"Tenaga Nasional Berhad",
             "keywords":["TNB","Tenaga Nasional","Malaysian grid","Malaysia electricity",
                         "Johor grid","Kuala Lumpur grid","TNB connection"]},
    "PLN": {"country":"Indonesia","full_name":"Perusahaan Listrik Negara",
             "keywords":["PLN Indonesia","Indonesian grid","Java-Bali grid",
                         "Jakarta electricity","PLN power"]},
    "TEPCO": {"country":"Japan","full_name":"Tokyo Electric Power Company",
             "keywords":["TEPCO","Tokyo Electric","Japan grid","Japanese power",
                         "Kansai Electric","Chubu Electric","Tokyo grid"]},
    "EMA Singapore": {"country":"Singapore","full_name":"Energy Market Authority",
             "keywords":["EMA Singapore","Singapore grid","Singapore electricity",
                         "SP Group","Singapore power","EMA moratorium",
                         "Singapore data center electricity"]},
    "Transpower": {"country":"New Zealand","full_name":"Transpower New Zealand",
             "keywords":["Transpower NZ","New Zealand grid","NZ electricity","Auckland grid"]},
    "ANEEL": {"country":"Brazil","full_name":"Agência Nacional de Energia Elétrica",
             "keywords":["ANEEL","Brazilian grid","Brazil electricity","ONS Brazil",
                         "Brazil power","CCEE","São Paulo grid","ONS interconnection"]},
    "CENACE": {"country":"Mexico","full_name":"Centro Nacional de Control de Energía",
             "keywords":["CENACE","Mexican grid","Mexico electricity","SENER Mexico",
                         "Mexico City grid","Monterrey grid"]},
    "Eskom": {"country":"South Africa","full_name":"Eskom Holdings SOC Ltd",
             "keywords":["Eskom","South Africa grid","South African power",
                         "Johannesburg electricity","Cape Town grid","load shedding",
                         "Eskom data center"]},
    "Kenya Power": {"country":"Kenya","full_name":"Kenya Power and Lighting Company",
             "keywords":["KPLC","Kenya Power","Nairobi grid","Kenya electricity"]},
    "NERC Nigeria": {"country":"Nigeria","full_name":"Nigerian Electricity Regulatory Commission",
             "keywords":["NERC Nigeria","Nigerian grid","Nigeria electricity","Lagos power"]},
}

ISO_RTO_KEYWORDS = {}
for _iso, _data in US_ISO_RTO.items():
    for _kw in _data.get("keywords", []):
        ISO_RTO_KEYWORDS[_kw.lower()] = _iso
for _op, _data in GLOBAL_GRID_OPERATORS.items():
    for _kw in _data.get("keywords", []):
        ISO_RTO_KEYWORDS[_kw.lower()] = _op

ALL_ISO_RTO_KW_LIST = sorted(set(ISO_RTO_KEYWORDS.keys()))

def detect_iso_rto(text):
    """Return the ISO/RTO/grid-operator relevant to this article, or empty string."""
    t = text.lower()
    for kw, iso in ISO_RTO_KEYWORDS.items():
        if kw in t:
            return iso
    return ""

COUNTRY_KEYWORDS = {k: [k] for k in COUNTRY_TO_REGION}
COUNTRY_KEYWORDS.update({
    "United States": ["United States", "U.S.", r"\bUS\b", "America", "American",
                      # States
                      "Virginia", "Texas", "California", "Georgia", "Ohio",
                      "Indiana", "Nevada", "Arizona", "Illinois", "Pennsylvania",
                      "North Carolina", "Florida", "Oregon", "Washington",
                      "Colorado", "Utah", "Wyoming", "Montana", "Idaho",
                      "Minnesota", "Wisconsin", "Michigan", "Iowa", "Kansas",
                      "Missouri", "Oklahoma", "Arkansas", "Louisiana",
                      "Mississippi", "Alabama", "Tennessee", "Kentucky",
                      "Maryland", "New Jersey", "New York", "Massachusetts",
                      "Connecticut", "Maine", "New Mexico", "Alaska", "Hawaii",
                      "Delaware", "Rhode Island", "Vermont", "New Hampshire",
                      "South Carolina", "South Dakota", "North Dakota",
                      "Nebraska", "West Virginia",
                      # Major metros / cities
                      "Dallas", "Austin", "Chicago", "Phoenix", "Atlanta",
                      "Seattle", "Denver", "Reno", "Ashburn", "Portland",
                      "Las Vegas", "San Jose", "San Francisco", "Los Angeles",
                      "Houston", "Miami", "Boston", "Columbus", "Nashville",
                      "Charlotte", "Northern Virginia", "Silicon Valley",
                      "San Antonio", "Sacramento", "San Diego",
                      "Minneapolis", "Milwaukee", "Indianapolis","Cincinnati",
                      "Cleveland","Pittsburgh","Baltimore","Philadelphia",
                      "New York City","Brooklyn","Bronx","Queens","Manhattan",
                      "Hartford","Providence","Albany","Buffalo","Rochester",
                      "New Orleans","Memphis","Louisville","Richmond",
                      "Raleigh","Durham","Greensboro","Winston-Salem",
                      "Birmingham","Montgomery","Huntsville","Mobile",
                      "Jacksonville","Tampa","Orlando","Fort Lauderdale",
                      "West Palm Beach","Tallahassee","Gainesville",
                      "Omaha","Lincoln","Kansas City","Wichita","Tulsa",
                      "Oklahoma City","Little Rock","Fort Smith",
                      "Salt Lake City","Provo","Ogden","Lehi","Draper",
                      "Boulder","Fort Collins","Colorado Springs",
                      "Albuquerque","Tucson","Mesa","Tempe","Scottsdale",
                      "Chandler","Gilbert","Goodyear","Peoria","Glendale AZ",
                      "Boise","Nampa","Caldwell","Idaho Falls",
                      "Helena","Billings","Great Falls","Missoula",
                      "Cheyenne","Casper","Laramie","Sheridan",
                      "Sioux Falls","Fargo","Bismarck","Grand Forks",
                      "Cedar Rapids","Des Moines","Davenport","Ames",
                      "Green Bay","Madison","Eau Claire","Appleton",
                      "Detroit","Grand Rapids","Ann Arbor","Lansing","Flint",
                      "Anchorage","Juneau","Fairbanks","Honolulu",
                      # Key DC hub / county markets
                      "Loudoun County","Prince William County","Fairfax County",
                      "Hillsboro","Boardman","Prineville","Hermiston",
                      "Goodyear","Buckeye","Coolidge","Eloy","Casa Grande",
                      "Midlothian","Garland","Plano","Irving","Carrollton",
                      "Round Rock","Georgetown","Pflugerville","Kyle",
                      "Dublin OH","Plain City","New Albany OH","Lewis Center",
                      "Reno-Sparks","Fernley","Storey County","Lyon County",
                      "Quincy WA","East Wenatchee","Moses Lake",
                      "DeKalb County","Carroll County","Paulding County",
                      "Bartow County","Newton County","Henry County",
                      "Richmond VA","Henrico County","Chesterfield County",
                      "Stafford County","Fredericksburg",
                      "Elk Grove Village","Aurora IL","Bolingbrook",
                      "Council Bluffs","Papillion","La Vista",
                      "Henderson NV","North Las Vegas","Sparks NV",
                      "Lehi UT","Draper UT","Bluffdale","South Jordan",
                      "Englewood CO","Aurora CO","Broomfield","Longmont",
                      # Regions
                      "Silicon Slopes","Silicon Hills","Silicon Desert",
                      "Silicon Plains","Research Triangle","Northern Virginia",
                      "NoVA","Pacific Northwest","Mountain West",
                      "Great Plains","Midwest","Southeast","Sun Belt",
                      "Gulf Coast","Mid-Atlantic","New England"],
    "United Kingdom": ["United Kingdom", r"\bUK\b", "England", "Scotland",
                       "Wales", "London", "Manchester", "Birmingham",
                       "Bristol", "Edinburgh", "Glasgow", "Slough", "Reading"],
    "Germany": ["Germany", "German", "Berlin", "Frankfurt", "Munich",
                "Hamburg", "Cologne", "Stuttgart", "Dusseldorf"],
    "Netherlands": ["Netherlands", "Dutch", "Amsterdam", "Rotterdam",
                    "Eindhoven", "The Hague", r"\bAMS\b"],
    "Singapore": ["Singapore", r"\bSGP\b"],
    "Australia": ["Australia", "Australian", "Sydney", "Melbourne",
                  "Brisbane", "Perth", "Canberra"],
    "India": ["India", "Indian", "Mumbai", "Bangalore", "Bengaluru",
              "Chennai", "Hyderabad", "Delhi", "Pune", "Noida"],
    "Japan": ["Japan", "Japanese", "Tokyo", "Osaka", "Nagoya"],
    "Saudi Arabia": ["Saudi Arabia", "Saudi", "KSA", "Riyadh", "Jeddah",
                     "Neom", "NEOM", "Vision 2030"],
    "UAE": ["UAE", "United Arab Emirates", "Dubai", "Abu Dhabi",
            "Sharjah", "Ajman"],
    "Brazil": ["Brazil", "Brasil", "Brazilian", r"\bSP\b", "Sao Paulo",
               "Rio de Janeiro", "Curitiba"],
    "South Africa": ["South Africa", "Johannesburg", "Cape Town", "Durban"],
    "Nigeria": ["Nigeria", "Nigerian", "Lagos", "Abuja"],
    "Kenya": ["Kenya", "Nairobi"],
    "France": ["France", "French", "Paris", "Lyon", "Marseille", "Toulouse"],
    "Ireland": ["Ireland", "Irish", "Dublin", "Cork"],
    "Sweden": ["Sweden", "Swedish", "Stockholm", "Gothenburg"],
    "Norway": ["Norway", "Norwegian", "Oslo"],
    "Denmark": ["Denmark", "Danish", "Copenhagen"],
    "Finland": ["Finland", "Finnish", "Helsinki"],
    "Poland": ["Poland", "Polish", "Warsaw", "Krakow"],
    "Malaysia": ["Malaysia", "Malaysian", "Kuala Lumpur", "KL", "Johor"],
    "Indonesia": ["Indonesia", "Indonesian", "Jakarta", "Batam"],
    "South Korea": ["South Korea", "Korean", "Seoul", "Busan"],
    "China": ["China", "Chinese", "Beijing", "Shanghai", "Shenzhen",
              "Guangzhou", "Chengdu"],
    "Hong Kong": ["Hong Kong", "HKG"],
    "Egypt": ["Egypt", "Egyptian", "Cairo"],
    "Qatar": ["Qatar", "Qatari", "Doha"],
    "Chile": ["Chile", "Chilean", "Santiago"],
    "Colombia": ["Colombia", "Colombian", "Bogota"],
    "Mexico": ["Mexico", "Mexican", "Mexico City", "Monterrey", "Guadalajara"],
})


# ── States / Provinces per country ───────────────────────────────────────────
COUNTRY_STATES = {
    "United States": [
        "Alabama","Alaska","Arizona","Arkansas","California","Colorado","Connecticut",
        "Delaware","Florida","Georgia","Hawaii","Idaho","Illinois","Indiana","Iowa",
        "Kansas","Kentucky","Louisiana","Maine","Maryland","Massachusetts","Michigan",
        "Minnesota","Mississippi","Missouri","Montana","Nebraska","Nevada","New Hampshire",
        "New Jersey","New Mexico","New York","North Carolina","North Dakota","Ohio",
        "Oklahoma","Oregon","Pennsylvania","Rhode Island","South Carolina","South Dakota",
        "Tennessee","Texas","Utah","Vermont","Virginia","Washington","West Virginia",
        "Wisconsin","Wyoming","Washington D.C.","Northern Virginia",
    ],
    "Canada": [
        "Alberta","British Columbia","Manitoba","New Brunswick","Newfoundland and Labrador",
        "Northwest Territories","Nova Scotia","Nunavut","Ontario","Prince Edward Island",
        "Quebec","Saskatchewan","Yukon",
    ],
    "Australia": [
        "New South Wales","Victoria","Queensland","Western Australia","South Australia",
        "Tasmania","Australian Capital Territory","Northern Territory",
    ],
    "India": [
        "Andhra Pradesh","Arunachal Pradesh","Assam","Bihar","Chhattisgarh","Goa","Gujarat",
        "Haryana","Himachal Pradesh","Jharkhand","Karnataka","Kerala","Madhya Pradesh",
        "Maharashtra","Manipur","Meghalaya","Mizoram","Nagaland","Odisha","Punjab",
        "Rajasthan","Sikkim","Tamil Nadu","Telangana","Tripura","Uttar Pradesh","Uttarakhand",
        "West Bengal","Delhi","Mumbai","Bangalore","Chennai","Hyderabad","Pune","Noida",
    ],
    "Germany": [
        "Baden-Württemberg","Bavaria","Berlin","Brandenburg","Bremen","Hamburg","Hesse",
        "Lower Saxony","Mecklenburg-Vorpommern","North Rhine-Westphalia","Rhineland-Palatinate",
        "Saarland","Saxony","Saxony-Anhalt","Schleswig-Holstein","Thuringia","Frankfurt",
    ],
    "United Kingdom": [
        "England","Scotland","Wales","Northern Ireland","London","Manchester","Birmingham",
        "Bristol","Edinburgh","Glasgow","Leeds","Liverpool","Sheffield","Newcastle",
    ],
    "Brazil": [
        "São Paulo","Rio de Janeiro","Minas Gerais","Bahia","Paraná","Rio Grande do Sul",
        "Pernambuco","Ceará","Amazonas","Goiás","Pará","Espírito Santo","Santa Catarina",
        "Mato Grosso","Mato Grosso do Sul","Maranhão","Paraíba","Piauí","Alagoas",
        "Rio Grande do Norte","Tocantins","Sergipe","Rondônia","Amapá","Roraima","Acre",
        "Distrito Federal",
    ],
    "China": [
        "Beijing","Shanghai","Guangdong","Zhejiang","Jiangsu","Shandong","Sichuan","Hubei",
        "Hunan","Fujian","Anhui","Shaanxi","Jiangxi","Chongqing","Liaoning","Yunnan",
        "Hebei","Guangxi","Shanxi","Tianjin","Heilongjiang","Jilin","Guizhou","Xinjiang",
        "Inner Mongolia","Gansu","Hainan","Ningxia","Qinghai","Xizang (Tibet)","Shenzhen","Hangzhou",
    ],
    "Mexico": [
        "Mexico City","Jalisco","Nuevo León","Veracruz","Puebla","Guanajuato","Chihuahua",
        "Michoacán","Oaxaca","Guerrero","Chiapas","Sinaloa","Hidalgo","Sonora","Tabasco",
        "Baja California","Querétaro","Morelos","San Luis Potosí","Aguascalientes",
        "Durango","Coahuila","Tamaulipas","Colima","Nayarit","Tlaxcala","Campeche","Yucatán","Zacatecas",
    ],
    "Malaysia": ["Johor","Kedah","Kelantan","Kuala Lumpur","Labuan","Melaka","Negeri Sembilan",
                 "Pahang","Penang","Perak","Perlis","Putrajaya","Sabah","Sarawak","Selangor","Terengganu"],
    "France": ["Île-de-France","Provence-Alpes-Côte d'Azur","Auvergne-Rhône-Alpes","Occitanie",
               "Hauts-de-France","Nouvelle-Aquitaine","Grand Est","Pays de la Loire","Normandie",
               "Bretagne","Bourgogne-Franche-Comté","Centre-Val de Loire","Corse","Paris","Lyon","Marseille"],
    "Netherlands": ["North Holland","South Holland","Utrecht","Gelderland","North Brabant",
                    "Overijssel","Groningen","Friesland","Limburg","Drenthe","Flevoland","Zeeland",
                    "Amsterdam","Rotterdam","The Hague","Eindhoven"],
    "Japan": ["Tokyo","Osaka","Kanagawa","Aichi","Saitama","Chiba","Hyogo","Hokkaido","Fukuoka",
              "Shizuoka","Ibaraki","Hiroshima","Kyoto","Miyagi","Niigata","Nagano","Tochigi",
              "Gunma","Okayama","Fukushima","Kumamoto","Kagoshima","Ehime","Yamaguchi","Okinawa"],
    "South Korea": ["Seoul","Busan","Incheon","Daegu","Daejeon","Gwangju","Ulsan","Gyeonggi",
                    "Gangwon","North Chungcheong","South Chungcheong","North Jeolla","South Jeolla",
                    "North Gyeongsang","South Gyeongsang","Jeju"],
    "Saudi Arabia": ["Riyadh","Mecca","Madinah","Eastern Province","Asir","Tabuk","Qassim",
                     "Hail","Jizan","Najran","Al Bahah","Northern Borders","Al Jawf","NEOM"],
    "UAE": ["Abu Dhabi","Dubai","Sharjah","Ajman","Umm Al Quwain","Ras Al Khaimah","Fujairah"],
    "Indonesia": ["Jakarta","West Java","East Java","Central Java","Banten","North Sumatra",
                  "South Sulawesi","East Kalimantan","Bali","Riau","South Sumatra","Yogyakarta"],
    "Poland": ["Masovian","Lesser Poland","Silesian","Greater Poland","Łódź","Lower Silesian",
               "Kuyavian-Pomeranian","Lublin","Subcarpathian","Warmian-Masurian","Pomeranian",
               "West Pomeranian","Opole","Lubusz","Świętokrzyskie","Podlaskie","Warsaw","Krakow"],
    "Ireland": ["Dublin","Cork","Galway","Limerick","Waterford","Kildare","Meath","Wicklow","Louth"],
    "Sweden": ["Stockholm","Västra Götaland","Skåne","Östergötland","Uppsala","Dalarna","Jönköping"],
    "Singapore": ["Central Region","North Region","North-East Region","East Region","West Region"],
    "South Africa": ["Gauteng","Western Cape","KwaZulu-Natal","Eastern Cape","Limpopo",
                     "Mpumalanga","North West","Free State","Northern Cape",
                     "Johannesburg","Cape Town","Durban","Pretoria"],
    "Nigeria": ["Lagos","Kano","Rivers","Kaduna","Oyo","Anambra","Abuja","Delta","Edo","Kwara"],
    "Kenya": ["Nairobi","Mombasa","Kisumu","Nakuru","Eldoret","Kiambu","Machakos"],
    "Egypt": ["Cairo","Alexandria","Giza","Luxor","Aswan","Port Said","Suez","Hurghada","Sharm el-Sheikh"],
    "Israel": ["Tel Aviv","Jerusalem","Haifa","Rishon LeZion","Petah Tikva","Ashdod","Beer Sheva","Netanya"],
    "Qatar": ["Doha","Al Wakrah","Al Khor","Lusail","Al Rayyan"],
    "Bahrain": ["Capital","Muharraq","Northern","Southern"],
    "Kuwait": ["Capital","Hawalli","Ahmadi","Jahra","Farwaniya","Mubarak Al-Kabeer"],
    "Oman": ["Muscat","Dhofar","Musandam","Al Batinah","Al Dhahirah","Al Dakhiliyah","Al Sharqiyah","Al Wusta","Al Buraimi"],
    "Chile": ["Santiago Metropolitan","Valparaíso","Biobío","Araucanía","Los Lagos","Antofagasta"],
    "Colombia": ["Bogotá","Antioquia","Valle del Cauca","Atlántico","Cundinamarca","Bolivar","Santander"],
    "Argentina": ["Buenos Aires","Córdoba","Santa Fe","Mendoza","Tucumán","Entre Ríos","Salta","Misiones"],
    "Peru": ["Lima","Arequipa","La Libertad","Piura","Lambayeque","Junín","Cusco","Áncash"],
    "Morocco": ["Casablanca-Settat","Rabat-Salé-Kénitra","Marrakech-Safi","Fès-Meknès","Tanger-Tétouan-Al Hoceïma"],
    "Rwanda": ["Kigali","Northern","Southern","Eastern","Western"],
    "Ethiopia": ["Addis Ababa","Oromia","Amhara","Tigray","SNNPR","Somali","Afar"],
    "Ghana": ["Greater Accra","Ashanti","Western","Central","Eastern","Northern","Upper East","Upper West","Volta","Brong-Ahafo"],
    "Tanzania": ["Dar es Salaam","Dodoma","Arusha","Mwanza","Mbeya","Morogoro","Tanga","Zanzibar"],
    "New Zealand": ["Auckland","Wellington","Canterbury","Waikato","Bay of Plenty","Manawatū-Whanganī","Hawke's Bay"],
    "Vietnam": ["Hanoi","Ho Chi Minh City","Hai Phong","Da Nang","Can Tho","Binh Duong"],
    "Philippines": ["Metro Manila","Cebu","Davao","Laguna","Cavite","Rizal","Bulacan","Pampanga"],
    "Thailand": ["Bangkok","Nonthaburi","Samut Prakan","Chiang Mai","Khon Kaen","Nakhon Ratchasima","Phuket"],
    "Taiwan": ["Taipei","New Taipei","Taichung","Kaohsiung","Taoyuan","Tainan","Hsinchu"],
    "Hong Kong": ["Hong Kong Island","Kowloon","New Territories","Lantau Island"],
    "Spain": ["Madrid","Catalonia","Andalusia","Valencia","Galicia","Castile and León","Basque Country","Canary Islands"],
    "Italy": ["Lombardy","Lazio","Campania","Sicily","Veneto","Emilia-Romagna","Puglia","Piedmont","Tuscany","Milan","Rome"],
    "Belgium": ["Brussels","Flanders","Wallonia","Antwerp","Ghent","Liège","Bruges"],
    "Switzerland": ["Zurich","Bern","Vaud","Geneva","Aargau","St. Gallen","Lucerne","Ticino","Valais","Basle"],
    "Austria": ["Vienna","Lower Austria","Upper Austria","Styria","Tyrol","Carinthia","Salzburg","Vorarlberg","Burgenland"],
    "Portugal": ["Lisbon","Porto","Braga","Setúbal","Aveiro","Leiria","Viseu","Coimbra"],
    "Romania": ["Bucharest","Cluj","Iași","Timișoara","Constanța","Galați","Craiova","Brașov"],
    "Czech Republic": ["Prague","Central Bohemia","South Moravian","Moravian-Silesian","Plzeň","Ústí nad Labem","Liberec","Hradec Králové"],
    "Norway": ["Oslo","Viken","Innlandet","Vestfold og Telemark","Agder","Rogaland","Vestland","Møre og Romsdal","Trøndelag","Nordland"],
    "Denmark": ["Capital","Zealand","Southern Denmark","Central Jutland","North Jutland","Copenhagen"],
    "Finland": ["Uusimaa","Pirkanmaa","Southwest Finland","North Ostrobothnia","Central Finland","South Savo","North Karelia","Helsinki"],
}

# Ensure all countries in COUNTRY_STATES are mapped in COUNTRY_TO_REGION
_EXTRA_COUNTRY_MAP = {
    "Spain": "Europe", "Italy": "Europe", "Belgium": "Europe",
    "Switzerland": "Europe", "Austria": "Europe", "Portugal": "Europe",
    "Romania": "Europe", "Czech Republic": "Europe", "Norway": "Europe",
    "Denmark": "Europe", "Finland": "Europe", "Sweden": "Europe",
    "New Zealand": "Asia Pacific", "Vietnam": "Asia Pacific",
    "Philippines": "Asia Pacific", "Thailand": "Asia Pacific",
    "Taiwan": "Asia Pacific", "Hong Kong": "Asia Pacific",
}
for _c, _r in _EXTRA_COUNTRY_MAP.items():
    COUNTRY_TO_REGION.setdefault(_c, _r)


TOPIC_KEYWORDS = {
    "Hyperscale": ["hyperscale","microsoft","google","amazon","aws","meta",
                   "apple","oracle","alibaba","tencent","bytedance",
                   "alphabet","openai","stargate","coreweave","deepmind",
                   "xai","mistral","inflection","cohere","scale ai",
                   "cerebras","groq","lambda labs","together ai",
                   "cloud region","availability zone","az deployment",
                   "hyperscaler campus","cloud campus","cloud data center"],
    "Colocation": ["colocation","colo","equinix","digital realty","ironmountain",
                   "coresite","cyrusone","ntt","vantage","switch","edgeconex",
                   "flexential","databank","cts","global switch",
                   "qts","stack infrastructure","aligned data centers",
                   "compass datacenters","tract","t5 data centers",
                   "skybox datacenters","stream data centers","cloudflare",
                   "zayo","lumen","cogent","akamai","fastly",
                   "ascenty","odata","scala","interxion","telehouse",
                   "yondr","venari","verne global","atNorth","digiPlex",
                   "bulk infrastructure","green mountain","nextdc",
                   "ctrl s","nxtgen","yotta","stt gdc","keppel dc",
                   "retail colocation","wholesale colocation","multi-tenant",
                   "single-tenant","powered shell","build to suit"],
    "AI / GPU":   ["artificial intelligence"," ai ","gpu","nvidia","inference",
                   "llm","generative","machine learning","deep learning",
                   "h100","h200","gb200","b200","b100","gh200","xai","anthropic","grok",
                   "chatgpt","claude","gemini","llama","mixtral","falcon",
                   "ai campus","ai factory","gpu cluster","gpu farm",
                   "ai infrastructure","inference cluster","training cluster",
                   "compute cluster","hpc facility","exascale","petaflop",
                   "ai supercomputer","ai accelerator","tpu","dpu","ipu",
                   "nvl rack","dgx","hgx","mgx","nvidia grace hopper",
                   "amd instinct","intel gaudi","liquid cooled gpu",
                   "high density compute","400kw rack","600kw rack",
                   "power density","kw per rack","rack density"],
    "Power":      [" mw "," gw ","megawatt","gigawatt","power plant",
                   "nuclear","solar","wind farm","gas turbine","behind-the-meter",
                   "grid","utility","energy","ppa","renewable","hydrogen",
                   "fuel cell","battery storage","ups","smr","small modular reactor",
                   "geothermal","hydropower","tidal","pumped hydro",
                   "battery energy storage system","bess","vrfb",
                   "grid connection","substation","transformer","switchgear",
                   "high voltage","extra high voltage","transmission line",
                   "distribution grid","feeder","bus bar","load growth",
                   "interconnection queue","capacity market","energy market",
                   "pjm","ercot","caiso","miso","iso-ne","nyiso","spp",
                   "wecc","serc","nerc","national grid eso","eirgrid","aemo",
                   "dewa","taqa","sec saudi","powergrid india","kepco",
                   "demand response","virtual power plant","microGrid",
                   "behind meter","net metering","self-consumption",
                   "power purchase agreement","offtake agreement","capacity contract",
                   "vppa","virtual ppa","24/7 cfe","carbon free electricity",
                   "guaranteed maximum demand","critical load"],
    "Investment": ["invest","fund","reit","billion","million","acquire",
                   "acquisition","ipo","bond","financing","lease","deal",
                   "partnership","joint venture","stake","raise","capital",
                   "sovereign wealth fund","pension fund","infrastructure fund",
                   "private equity","private credit","mezzanine","preferred equity",
                   "green bond","sustainability bond","project finance",
                   "sale leaseback","forward purchase","forward funding",
                   "net lease","triple net","nnn lease","long term lease",
                   "sale and leaseback","co-investment","club deal",
                   "portfolio acquisition","platform acquisition","roll-up",
                   "tender offer","take private","going private",
                   "credit facility","revolving credit","term loan",
                   "cad","aud","nzd","chf","sek","nok","dkk","krw","myr",
                   "thb","idr","sgd","aed","sar","qar","inr","brl","mxn",
                   "trillion","crore","lakh","bn usd","mn usd","bn eur",
                   "mn gbp","bn gbp","mn sgd","bn inr","mn aed"],
    "Permits":    ["permit","zoning","moratorium","approved","approval",
                   "planning","ordinance","rezoning","denied","appeal",
                   "lawsuit","sue","court","commission",
                   "vote","hearing","rejected","conditional use",
                   "use permit","special use permit","sup","cup",
                   "environmental review","eis","nepa","ceqa","ea",
                   "site plan approval","subdivision","variance",
                   "setback","noise ordinance","water permit",
                   "utility easement","right of way","condemnation",
                   "eminent domain","community benefit agreement","cba",
                   "public comment","planning commission","city council",
                   "county board","supervisor vote","board of supervisors",
                   "planning board","zoning board","appeal board",
                   "building permit","demolition permit","grading permit",
                   "environmental permit","stormwater permit"],
    "Construction":["broke ground","groundbreaking","topping out","opens",
                    "opened","inaugurated","construction","campus","build",
                    "development","site","facility","phase","expansion",
                    "warehouse","acres","sq ft","square feet","hectares",
                    "sq m","square meters","under construction","build-out",
                    "fit-out","fitout","commissioning","energized","energise",
                    "shell and core","raised floor","hot aisle containment",
                    "cold aisle containment","modular build","prefab",
                    "containerized","pod deployment","construction loan",
                    "civil works","structural steel","mep works",
                    "turnkey","epc","design-build","leed","breeam",
                    "certificate of occupancy","temporary certificate",
                    "ribbon cutting","topping ceremony","groundbreak",
                    "site preparation","land clearing","grading",
                    "concrete pour","steel erection","roof completion",
                    "phase 1","phase 2","phase 3","initial phase",
                    "first phase","next phase","final phase","full build-out"],
    "Sustainability":["sustainability","carbon","renewable","net zero",
                      "green","esg","water","pue","cooling","waste heat",
                      "recycle","circular","biodiversity","solar panel",
                      "wind power","offset","scope 1","scope 2","scope 3",
                      "carbon neutral","carbon negative","carbon credit",
                      "rec","i-rec","guarantees of origin","24/7 cfe",
                      "water positive","water neutral","zero waste",
                      "heat reuse","district heating","sbti",
                      "science based targets","tcfd","esg report",
                      "esg disclosure","climate commitment","climate target",
                      "decarbonization","low carbon","clean energy",
                      "renewable portfolio","green tariff","green procurement",
                      "social impact","community benefit","local hiring",
                      "noise mitigation","light pollution","visual impact"],
    "Grid / ISO / RTO": ["pjm","ercot","caiso","miso","nyiso","iso-ne","spp","wecc",
                          "serc","nerc","tva","bpa","national grid eso","eirgrid","aemo",
                          "powergrid india","kepco","dewa","kahramaa","taqa","entso-e",
                          "rte france","tennet","50hertz","amprion","elia","fingrid",
                          "statnett","energinet","pse poland","transpower","tnb","pln",
                          "tepco","ema singapore","state grid china","aneel","cenace",
                          "eskom","large load study","grid interconnection",
                          "interconnection queue","load growth","transmission capacity",
                          "grid capacity","balancing authority","load zone",
                          "capacity market","energy market","grid operator",
                          "independent system operator","regional transmission organization"],
}

KNOWN_COMPANIES = [
    "Microsoft","Google","Amazon","AWS","Meta","Apple","Oracle","Alibaba",
    "Tencent","ByteDance","Baidu","Huawei","Samsung","IBM","Intel","NVIDIA",
    "Equinix","Digital Realty","Iron Mountain","CoreSite","CyrusOne","NTT",
    "Vantage","Switch","EdgeConneX","Flexential","DataBank","QTS","Colt",
    "Global Switch","ChinaData","GDS","STACK","Aligned","Compass","Tract",
    "Hut 8","Core Scientific","Riot","Marathon","Applied Digital","Iren",
    "Cloudflare","Fastly","Akamai","Lumen","Zayo","Cogent",
    "Schneider Electric","Vertiv","Eaton","ABB","Siemens","Caterpillar",
    "Cummins","Aggreko","AECOM","Turner","Holder","DPR","Skanska",
    "Prologis","Blackstone","KKR","Brookfield","GIC","Mubadala","CPPIB",
    "SoftBank","Masdar","TAQA","ADNOC","Saudi Aramco","Neom","SABIC",
    "CloudHQ","Coatue","DataBank","T5","Skybox","Stream","Flexential",
    "Ascenty","Odata","Scala","Luminet","Etix","Nabiax","Bolder",
    "VIRTUS","Kao","Yondr","Venari","Verne","Hydro66","DigiPlex",
    "Bulk Infrastructure","Green Mountain","atNorth","Adaniconnex",
    "CtrlS","NxtGen","Yotta","STT GDC","Keppel","Singtel","Telstra",
    "NextDC","Macquarie","AirTrunk","MEVSPACE","Beyond.pl","Atman",
    "DE-CIX","Interxion","euNetworks","Telehouse","Iomart","Pulsant",
    # Extended hyperscalers / AI cos
    "OpenAI","xAI","DeepMind","Mistral","Cohere","Inflection","Scale AI",
    "Cerebras","Groq","Lambda Labs","Together AI","CoreWeave","Vast AI",
    "Fireworks AI","Perplexity","Stability AI","Midjourney",
    # Extended colocation & operators
    "QTS Realty","Stack Infrastructure","Aligned Data Centers",
    "Compass Datacenters","Tract","T5 Data Centers","Skybox Datacenters",
    "Stream Data Centers","DataBank","CloudHQ","Chirisa","Cologix",
    "Raging Wire","Switch SUPERNAP","Cyxtera","Peak 10",
    "365 Main","Digital Bridge","DigitalBridge","Stonepeak",
    "IPI Partners","Harrison Street","Actis","Oaktree",
    "Singtel","Telstra","Spark NZ","KDDI","SoftBank",
    "Adani ConnectX","CtrlS","NxtGen","Yotta","STT GDC",
    "Keppel","NextDC","Macquarie","AirTrunk",
    "MEVSPACE","Beyond.pl","Atman","Exatel","LCloud",
    "Nabiax","Bolder","VIRTUS","Kao Data","Yondr","Venari",
    "Verne Global","Hydro66","DigiPlex","Green Mountain",
    "atNorth","Bulk Infrastructure","DigiCo","Data#3",
    "Ascenty","Odata","Scala","Luminet","Etix","Global Data Systems",
    # Power / utilities
    "Dominion Energy","Duke Energy","Entergy","Exelon","AEP",
    "Eversource","National Grid","Ameren","Evergy","PPL",
    "CenterPoint Energy","Pacific Gas and Electric","PG&E",
    "Southern Company","NextEra Energy","AES","NV Energy",
    "Sacramento Municipal Utility","SMUD","Salt River Project",
    "Arizona Public Service","APS","Portland General Electric",
    "Pacific Power","Puget Sound Energy","Idaho Power",
    "Basin Electric","Dairyland Power","Xcel Energy",
    "OGE Energy","Empire District","CLECO","Entergy Texas",
    "Tokyo Electric","TEPCO","KEPCO Korea","Kansai Electric",
    "Chubu Electric","Kyushu Electric","Tohoku Electric",
    "Tenaga Nasional","PLN Indonesia","EirGrid","National Grid ESO",
    "RTE France","TenneT","50Hertz","Amprion","TransnetBW",
    "Elia","Terna","REE Spain","Fingrid","Statnett","Energinet",
    "Svenska Kraftnät","PSE Poland","Swissgrid","APG",
    "DEWA","TAQA","SEC Saudi","KAHRAMAA","Eskom",
    "AEMO","Transpower NZ","PowerGrid India",
    # Construction / engineering / equipment
    "Mortenson","Turner Construction","Holder","DPR","Skanska",
    "Gilbane","McCarthy","Whiting-Turner","Structure Tone",
    "Black & Veatch","Burns & McDonnell","AECOM","Jacobs",
    "HDR","Gensler","HMC Architects","DLR Group",
    "Stantec","WSP","Arup","Buro Happold",
    "Vertiv","Eaton","ABB","Siemens Energy","Schneider Electric",
    "Caterpillar","Cummins","Aggreko","HITEC","AMETEK",
    "Legrand","Panduit","CommScope","Corning","Belden",
    "Rittal","nVent","Emerson Network Power","Liebert",
    "Alfa Laval","SPX Cooling","Baltimore Aircoil","Evapco",
    "Stulz","Airedale","Aermec","Daikin","Carrier","Trane",
    # Investors / developers
    "Blackstone","KKR","Brookfield","GIC","Mubadala","CPPIB",
    "SoftBank","Masdar","TAQA","Neom","SABIC","ADNOC",
    "Coatue","Andreessen Horowitz","a16z","Tiger Global",
    "SilverLake","Apollo","Cerberus","Carlyle","Warburg Pincus",
    "Stonepeak","EQT","IFM Investors","Macquarie Asset Management",
    "GLP","Prologis","Highbrook","Hana Financial",
    "AustralianSuper","CPP Investments","Ontario Teachers",
    "OMERS","PSP Investments","Caisse de dépôt","CDPQ",
    "Abu Dhabi Investment Authority","ADIA","PIF Saudi",
    "Temasek","GIC Singapore","Khazanah","EPF Malaysia",
]


# ═══════════════════════════════════════════════════════════════════════════════
#  RENEWABLES POWER MARKETS — Constants, Keywords, Sources
# ═══════════════════════════════════════════════════════════════════════════════

RE_TOPIC_COLORS = {
    "Solar":           "#ffaa00",
    "Wind":            "#00b4ff",
    "Energy Storage":  "#00e5c8",
    "Offshore Wind":   "#0088cc",
    "Hydrogen":        "#a855f7",
    "Other Renewables":"#00e676",
    "EPC Companies":   "#ff6400",   # EPC / project company mode
}

RE_DEAL_TYPE_COLORS = {
    "PPA":               "#00e676",
    "MOU":               "#00b4ff",
    "AOR / Offtake":     "#a855f7",
    "Tender / Auction":  "#ffaa00",
    "Investment / IPO":  "#0047e1",
    "M&A / Deals":       "#ff6400",
    "Construction":      "#00e5c8",
    "Commissioning":     "#43ea80",
    "Policy / Reg":      "#ff2d6b",
    "Grid / Connect":    "#9c27b0",
    "General":           "#2e4470",
}

RE_TOPIC_KEYWORDS = {
    "Solar": [
        "solar","photovoltaic","pv ","pv park","solar farm","solar park","solar plant",
        "solar project","solar array","solar panel","bifacial","utility-scale solar",
        "rooftop solar","floating solar","agrivoltaic","cpv","concentrated solar",
        "solar developer","solar installer","solar capacity","solar gw","solar mw",
        "solaredge","enphase","first solar","sunpower","jinko","longi","canadian solar",
    ],
    "Offshore Wind": [
        "offshore wind","floating wind","offshore turbine","monopile","jacket foundation",
        "offshore wind farm","offshore wind park","wtiv","offshore substation",
        "offshore lease","floating offshore","offshore cfd","seabed lease",
        "wind turbine installation vessel","offshore wind developer",
        "orsted","equinor wind","vattenfall wind","rwe offshore","bp wind",
    ],
    "Wind": [
        "wind farm","wind park","wind project","wind turbine","onshore wind",
        "wind energy","wind power","wind blade","nacelle","wind developer",
        "wind repowering","land wind","wind capacity","wind gw","wind mw",
        "vestas","siemens gamesa","ge wind","nordex","enercon","goldwind",
    ],
    "Energy Storage": [
        "battery storage","bess","energy storage","grid-scale battery",
        "lithium-ion storage","flow battery","vanadium battery","battery system",
        "storage project","battery energy storage","grid battery",
        "long duration storage","flywheel","compressed air","pumped hydro",
        "pumped storage","gravity storage","battery gw","battery mw",
        "tesla megapack","fluence","nec energy","wartsila storage","wärtsilä",
    ],
    "Hydrogen": [
        "hydrogen","electrolysis","electrolyser","electrolyzer","green hydrogen",
        "blue hydrogen","pink hydrogen","grey hydrogen","fuel cell",
        "ammonia","hydrogen pipeline","hydrogen hub","ptx","power-to-gas",
        "hydrogen project","hydrogen offtake","hydrogen ppa","hydrogen storage",
        "electrolytic hydrogen","renewable hydrogen","h2 project","h2 plant",
        "itm power","nel hydrogen","plug power","bloom energy","air products hydrogen",
    ],
    "Other Renewables": [
        "geothermal","hydropower","hydroelectric","hydro dam","run-of-river",
        "tidal","wave energy","marine energy","bioenergy","biomass","biogas",
        "landfill gas","waste-to-energy","csp","concentrated solar power",
        "solar thermal","ocean energy","enhanced geothermal","egs",
    ],
}

RE_DEAL_TYPE_KEYWORDS = {
    "PPA": [
        "ppa","power purchase agreement","offtake agreement","long-term contract",
        "virtual ppa","vppa","corporate ppa","sleeved ppa","private wire",
        "power offtake","energy supply agreement","electricity supply contract",
        "15-year","20-year","25-year","30-year","offtake deal","supply agreement",
    ],
    "MOU": [
        "mou","memorandum of understanding","framework agreement","heads of terms",
        "letter of intent","loi","development agreement","cooperation agreement",
        "non-binding agreement","joint development","strategic partnership",
        "collaboration agreement","teaming agreement",
    ],
    "AOR / Offtake": [
        "aor","agreement on record","anchor offtake","capacity agreement",
        "tolling agreement","capacity reservation","offtake signed","offtake secured",
        "virtual offtake","capacity booking",
    ],
    "Tender / Auction": [
        "tender","auction","bid","cfds","contract for difference","request for proposal",
        "rfp","solicitation","capacity tender","cfd round","awarded contract",
        "contract award","competitive tender","round results","round award",
        "capacity market","capacity auction","procurement round",
    ],
    "Investment / IPO": [
        "investment","funding","financing","equity","debt","bond","green bond",
        "sustainability bond","loan","credit facility","project finance",
        "financial close","equity raise","debt raise","infrastructure fund",
        "ipo","initial public offering","listing","stock exchange","spac",
        "raises","secured financing","secured funding","series a","series b",
    ],
    "M&A / Deals": [
        "acquisition","acquired","merger","takeover","stake","bought","purchased",
        "divested","sold ","portfolio sale","asset sale","m&a","deal signed",
        "joint venture","jv ","partnership deal","buys ","sells ",
    ],
    "Construction": [
        "broke ground","groundbreaking","construction begins","under construction",
        "build","installation begins","turbine installation","panel installation",
        "civil works","epc contract","epc award","notice to proceed","ntp",
    ],
    "Commissioning": [
        "commissioned","commissioning","energized","energised","goes live",
        "began operations","first power","first generation","inaugurated","opened",
        "commercial operation","cod","coo","reaches cod","online","goes online",
        "starts generation","begins generation","ribbon cutting",
    ],
    "Policy / Reg": [
        "regulation","policy","legislation","standard","target","mandate","decree",
        "executive order","parliament","senate","vote","approved","enacted","signed",
        "feed-in tariff","fit ","rps","renewable portfolio","irp","net metering",
        "clean energy standard","subsidy","tax credit","itc ","ptc ","ira ",
        "planning consent","permit","eia ","grid code","capacity market rules",
    ],
    "Grid / Connect": [
        "grid connection","interconnection","network connection","transmission access",
        "substation","connection offer","interconnection agreement","grid study",
        "hvdc","hvac","cable","transmission line","grid upgrade","grid expansion",
    ],
}

# ─── RE Feed registry ──────────────────────────────────────────────────────────
# Multi-source RSS/Atom engine.  For each of the 6 sectors we maintain:
#   - specialist publication feeds  (high signal, narrow scope)
#   - Google News query feeds        (broad, always fresh, no paywalls)
#   - cross-sector feeds             (tagged to all relevant sectors)
#
# Google News RSS format: https://news.google.com/rss/search?q=QUERY&hl=en&gl=US&ceid=US:en
# Completely free, no API key, no rate limits on reasonable usage.
# Specialist feeds: direct RSS from each publication — fastest, highest quality.

_GN = "https://news.google.com/rss/search?hl=en&gl=US&ceid=US:en&q="

# ═══════════════════════════════════════════════════════════════════════════════
#  EPC COMPANY REGISTRY — Global EPC / Project Companies (Power & Renewables)
#  Covers: Solar, Wind, Offshore Wind, Storage/BESS, Hydrogen, Hydro, Gas/Power
# ═══════════════════════════════════════════════════════════════════════════════

EPC_COMPANIES = {
    # ── Tier 1 Global EPCs (full spectrum) ────────────────────────────────────
    "Bechtel":              {"region": "Global", "sectors": ["Solar","Wind","Energy Storage","Hydrogen","Other Renewables"]},
    "Fluor":                {"region": "Global", "sectors": ["Solar","Wind","Energy Storage","Hydrogen","Other Renewables"]},
    "Technip Energies":     {"region": "Global", "sectors": ["Hydrogen","Other Renewables"]},
    "Saipem":               {"region": "Global", "sectors": ["Offshore Wind","Hydrogen","Other Renewables"]},
    "Wood":                 {"region": "Global", "sectors": ["Solar","Wind","Energy Storage","Hydrogen","Other Renewables"]},
    "Worley":               {"region": "Global", "sectors": ["Solar","Wind","Hydrogen","Other Renewables"]},
    "AECOM":                {"region": "Global", "sectors": ["Solar","Wind","Energy Storage","Other Renewables"]},
    "Jacobs":               {"region": "Global", "sectors": ["Solar","Wind","Energy Storage","Hydrogen"]},
    "WSP":                  {"region": "Global", "sectors": ["Solar","Wind","Energy Storage","Offshore Wind"]},
    "Stantec":              {"region": "Global", "sectors": ["Solar","Wind","Energy Storage","Other Renewables"]},
    "Hatch":                {"region": "Global", "sectors": ["Energy Storage","Hydrogen","Other Renewables"]},
    "Black & Veatch":       {"region": "Global", "sectors": ["Solar","Wind","Energy Storage","Hydrogen","Other Renewables"]},
    "Burns & McDonnell":    {"region": "Global", "sectors": ["Solar","Wind","Energy Storage","Other Renewables"]},
    "Sargent & Lundy":      {"region": "Global", "sectors": ["Solar","Wind","Energy Storage","Other Renewables"]},
    "Tetra Tech":           {"region": "Global", "sectors": ["Solar","Wind","Energy Storage","Other Renewables"]},
    "AMEC":                 {"region": "Global", "sectors": ["Solar","Wind","Hydrogen","Other Renewables"]},
    "SNC-Lavalin":          {"region": "Global", "sectors": ["Solar","Wind","Hydrogen","Other Renewables"]},
    "AtkinsRéalis":         {"region": "Global", "sectors": ["Solar","Wind","Energy Storage","Offshore Wind","Hydrogen"]},
    "Parsons":              {"region": "Global", "sectors": ["Solar","Wind","Other Renewables"]},
    "KBR":                  {"region": "Global", "sectors": ["Hydrogen","Other Renewables"]},

    # ── European EPCs ─────────────────────────────────────────────────────────
    "Acciona Energia":      {"region": "Europe/Global", "sectors": ["Solar","Wind","Offshore Wind","Energy Storage","Hydrogen"]},
    "ACS":                  {"region": "Europe/Global", "sectors": ["Solar","Wind","Offshore Wind","Energy Storage"]},
    "Cobra IS":             {"region": "Europe/Global", "sectors": ["Solar","Wind","Energy Storage"]},
    "Elecnor":              {"region": "Europe/Global", "sectors": ["Solar","Wind","Energy Storage"]},
    "TSK":                  {"region": "Europe/Global", "sectors": ["Solar","Wind","Energy Storage"]},
    "Grupo T-Solar":        {"region": "Europe",        "sectors": ["Solar"]},
    "Solaria":              {"region": "Europe",        "sectors": ["Solar"]},
    "Renovalia":            {"region": "Europe",        "sectors": ["Solar","Wind"]},
    "Técnicas Reunidas":    {"region": "Europe/Global", "sectors": ["Solar","Hydrogen","Other Renewables"]},
    "Dragados":             {"region": "Europe/Global", "sectors": ["Offshore Wind","Other Renewables"]},
    "Saeta Yield":          {"region": "Europe",        "sectors": ["Solar","Wind"]},
    "Gamesa Electric":      {"region": "Europe/Global", "sectors": ["Wind","Energy Storage"]},
    "Siemens Gamesa":       {"region": "Global",        "sectors": ["Wind","Offshore Wind"]},
    "Vestas":               {"region": "Global",        "sectors": ["Wind","Offshore Wind"]},
    "GE Vernova":           {"region": "Global",        "sectors": ["Wind","Offshore Wind","Energy Storage","Other Renewables"]},
    "Enercon":              {"region": "Europe/Global", "sectors": ["Wind"]},
    "Nordex":               {"region": "Europe/Global", "sectors": ["Wind"]},
    "Envision Energy":      {"region": "Asia/Global",   "sectors": ["Wind","Energy Storage"]},
    "Senvion":              {"region": "Europe",        "sectors": ["Wind"]},
    "Bosch Rexroth":        {"region": "Europe/Global", "sectors": ["Wind","Offshore Wind"]},
    "Prysmian":             {"region": "Europe/Global", "sectors": ["Offshore Wind","Energy Storage"]},
    "Nexans":               {"region": "Europe/Global", "sectors": ["Offshore Wind","Energy Storage"]},
    "Alstom":               {"region": "Europe/Global", "sectors": ["Other Renewables"]},
    "Schneider Electric":   {"region": "Global",        "sectors": ["Solar","Wind","Energy Storage","Hydrogen"]},
    "ABB":                  {"region": "Global",        "sectors": ["Solar","Wind","Energy Storage","Offshore Wind","Hydrogen"]},
    "Siemens Energy":       {"region": "Global",        "sectors": ["Solar","Wind","Offshore Wind","Energy Storage","Hydrogen"]},
    "Hitachi Energy":       {"region": "Global",        "sectors": ["Solar","Wind","Energy Storage","Offshore Wind"]},
    "Enel Green Power":     {"region": "Europe/Global", "sectors": ["Solar","Wind","Energy Storage","Hydrogen"]},
    "Iberdrola":            {"region": "Europe/Global", "sectors": ["Solar","Wind","Offshore Wind","Energy Storage","Hydrogen"]},
    "Vattenfall":           {"region": "Europe/Global", "sectors": ["Wind","Offshore Wind","Solar","Energy Storage"]},
    "Ørsted":               {"region": "Europe/Global", "sectors": ["Offshore Wind","Wind","Solar","Hydrogen"]},
    "RWE":                  {"region": "Europe/Global", "sectors": ["Solar","Wind","Offshore Wind","Energy Storage","Hydrogen"]},
    "TotalEnergies":        {"region": "Europe/Global", "sectors": ["Solar","Wind","Energy Storage","Hydrogen"]},
    "Engie":                {"region": "Europe/Global", "sectors": ["Solar","Wind","Offshore Wind","Energy Storage","Hydrogen"]},
    "EDF Renewables":       {"region": "Europe/Global", "sectors": ["Solar","Wind","Offshore Wind","Energy Storage"]},
    "Equinor":              {"region": "Europe/Global", "sectors": ["Offshore Wind","Hydrogen","Solar"]},
    "BP":                   {"region": "Global",        "sectors": ["Solar","Wind","Offshore Wind","Hydrogen","Energy Storage"]},
    "Shell":                {"region": "Global",        "sectors": ["Solar","Wind","Offshore Wind","Hydrogen","Energy Storage"]},
    "SSE":                  {"region": "Europe",        "sectors": ["Wind","Offshore Wind","Energy Storage"]},
    "ScottishPower":        {"region": "Europe",        "sectors": ["Wind","Offshore Wind","Solar"]},
    "Statkraft":            {"region": "Europe",        "sectors": ["Wind","Solar","Energy Storage","Hydrogen","Other Renewables"]},
    "Mainstream Renewable Power": {"region": "Global",  "sectors": ["Wind","Solar","Offshore Wind"]},
    "Lightsource BP":       {"region": "Global",        "sectors": ["Solar"]},
    "Neoen":                {"region": "Europe/Global", "sectors": ["Solar","Wind","Energy Storage"]},
    "Voltalia":             {"region": "Europe/Global", "sectors": ["Solar","Wind","Energy Storage"]},
    "BayWa r.e.":           {"region": "Europe/Global", "sectors": ["Solar","Wind","Energy Storage"]},
    "Sonnedix":             {"region": "Europe/Global", "sectors": ["Solar"]},
    "Enerparc":             {"region": "Europe",        "sectors": ["Solar"]},
    "juwi":                 {"region": "Europe/Global", "sectors": ["Solar","Wind"]},
    "WPD":                  {"region": "Europe/Global", "sectors": ["Wind","Offshore Wind"]},
    "DEME":                 {"region": "Europe/Global", "sectors": ["Offshore Wind"]},
    "Van Oord":             {"region": "Europe/Global", "sectors": ["Offshore Wind"]},
    "Heerema":              {"region": "Europe/Global", "sectors": ["Offshore Wind"]},
    "Seaway 7":             {"region": "Europe/Global", "sectors": ["Offshore Wind"]},
    "Cadeler":              {"region": "Europe",        "sectors": ["Offshore Wind"]},
    "Eneti":                {"region": "Global",        "sectors": ["Offshore Wind"]},
    "Windey":               {"region": "Asia",          "sectors": ["Wind"]},
    "NEG Micon":            {"region": "Europe",        "sectors": ["Wind"]},
    "Renew Power":          {"region": "Asia",          "sectors": ["Solar","Wind","Energy Storage"]},
    "John Laing":           {"region": "Europe/Global", "sectors": ["Solar","Wind","Offshore Wind"]},
    "Cero Generation":      {"region": "Europe",        "sectors": ["Solar","Wind"]},
    "Cubico":               {"region": "Europe/Global", "sectors": ["Solar","Wind"]},
    "Copenhagen Infrastructure Partners": {"region": "Europe/Global", "sectors": ["Offshore Wind","Wind","Solar"]},
    "Eurowind Energy":      {"region": "Europe",        "sectors": ["Wind","Solar"]},
    "European Energy":      {"region": "Europe",        "sectors": ["Wind","Solar","Energy Storage","Hydrogen"]},
    "Windpark Capital":     {"region": "Europe",        "sectors": ["Wind"]},

    # ── UK / Ireland EPCs ─────────────────────────────────────────────────────
    "Balfour Beatty":       {"region": "UK/Global",     "sectors": ["Solar","Wind","Energy Storage"]},
    "Mott MacDonald":       {"region": "UK/Global",     "sectors": ["Solar","Wind","Offshore Wind","Energy Storage","Hydrogen"]},
    "Arup":                 {"region": "Global",        "sectors": ["Solar","Wind","Offshore Wind","Energy Storage","Hydrogen"]},
    "Turner & Townsend":    {"region": "Global",        "sectors": ["Solar","Wind","Energy Storage","Offshore Wind"]},
    "Costain":              {"region": "UK/Global",     "sectors": ["Wind","Offshore Wind","Energy Storage"]},
    "Kier":                 {"region": "UK",            "sectors": ["Solar","Wind","Energy Storage"]},
    "Morgan Sindall":       {"region": "UK",            "sectors": ["Solar","Wind","Energy Storage"]},
    "Skanska":              {"region": "Europe/Global", "sectors": ["Wind","Offshore Wind","Solar"]},
    "Wates":                {"region": "UK",            "sectors": ["Solar","Energy Storage"]},
    "Laing O'Rourke":       {"region": "UK/Global",     "sectors": ["Offshore Wind","Wind","Energy Storage"]},
    "Amec Foster Wheeler":  {"region": "UK/Global",     "sectors": ["Solar","Wind","Hydrogen","Other Renewables"]},

    # ── US / North America EPCs ───────────────────────────────────────────────
    "Mortenson":            {"region": "North America",  "sectors": ["Solar","Wind","Energy Storage","Offshore Wind"]},
    "Blattner Energy":      {"region": "North America",  "sectors": ["Solar","Wind","Energy Storage"]},
    "Primoris Services":    {"region": "North America",  "sectors": ["Solar","Wind","Energy Storage"]},
    "Quanta Services":      {"region": "North America",  "sectors": ["Solar","Wind","Energy Storage"]},
    "MYR Group":            {"region": "North America",  "sectors": ["Solar","Wind","Energy Storage"]},
    "IEA Energy Services":  {"region": "North America",  "sectors": ["Solar","Wind","Energy Storage"]},
    "Wanzek Construction":  {"region": "North America",  "sectors": ["Solar","Wind","Energy Storage"]},
    "Swinerton Renewable Energy": {"region": "North America", "sectors": ["Solar","Energy Storage"]},
    "Orion Energy Systems": {"region": "North America",  "sectors": ["Solar","Energy Storage"]},
    "Solect Energy":        {"region": "North America",  "sectors": ["Solar"]},
    "SunPower":             {"region": "North America/Global", "sectors": ["Solar"]},
    "First Solar":          {"region": "North America/Global", "sectors": ["Solar"]},
    "SolarEdge":            {"region": "North America/Global", "sectors": ["Solar","Energy Storage"]},
    "Enphase Energy":       {"region": "North America/Global", "sectors": ["Solar","Energy Storage"]},
    "Array Technologies":   {"region": "North America",  "sectors": ["Solar"]},
    "Nextracker":           {"region": "North America/Global", "sectors": ["Solar"]},
    "Pattern Energy":       {"region": "North America/Global", "sectors": ["Solar","Wind","Energy Storage"]},
    "Invenergy":            {"region": "North America",  "sectors": ["Solar","Wind","Energy Storage","Other Renewables"]},
    "Longroad Energy":      {"region": "North America",  "sectors": ["Solar","Wind","Energy Storage"]},
    "Intersect Power":      {"region": "North America",  "sectors": ["Solar","Energy Storage","Hydrogen"]},
    "Arevon Energy":        {"region": "North America",  "sectors": ["Solar","Wind","Energy Storage"]},
    "EDF Renewables North America": {"region": "North America", "sectors": ["Solar","Wind","Energy Storage","Offshore Wind"]},
    "Avangrid Renewables":  {"region": "North America",  "sectors": ["Wind","Offshore Wind","Solar"]},
    "NextEra Energy Resources": {"region": "North America", "sectors": ["Solar","Wind","Energy Storage"]},
    "Terra-Gen":            {"region": "North America",  "sectors": ["Solar","Wind","Energy Storage"]},
    "Cypress Creek Renewables": {"region": "North America", "sectors": ["Solar","Energy Storage"]},
    "8minute Solar Energy": {"region": "North America",  "sectors": ["Solar","Energy Storage"]},
    "Lightsource bp":       {"region": "Global",         "sectors": ["Solar"]},
    "Amp Solar":            {"region": "North America",  "sectors": ["Solar","Energy Storage"]},
    "Canadian Solar":       {"region": "Global",         "sectors": ["Solar","Energy Storage"]},
    "Berkshire Hathaway Energy": {"region": "North America", "sectors": ["Solar","Wind","Energy Storage"]},
    "AES":                  {"region": "Global",         "sectors": ["Solar","Wind","Energy Storage","Other Renewables"]},
    "TerraForm Power":      {"region": "North America/Global", "sectors": ["Solar","Wind","Energy Storage"]},
    "Clearway Energy":      {"region": "North America",  "sectors": ["Solar","Wind","Energy Storage"]},
    "Ørsted US":            {"region": "North America",  "sectors": ["Offshore Wind"]},
    "Vineyard Wind":        {"region": "North America",  "sectors": ["Offshore Wind"]},
    "Dominion Energy":      {"region": "North America",  "sectors": ["Solar","Wind","Offshore Wind","Energy Storage"]},
    "Duke Energy Renewables": {"region": "North America", "sectors": ["Solar","Wind","Energy Storage"]},
    "Southern Company":     {"region": "North America",  "sectors": ["Solar","Wind","Energy Storage"]},
    "Xcel Energy":          {"region": "North America",  "sectors": ["Solar","Wind","Energy Storage"]},
    "Eversource":           {"region": "North America",  "sectors": ["Offshore Wind","Wind","Solar"]},
    "National Grid Renewables": {"region": "North America", "sectors": ["Solar","Wind","Energy Storage"]},

    # ── Asia-Pacific EPCs ─────────────────────────────────────────────────────
    "Greenko":              {"region": "Asia",           "sectors": ["Solar","Wind","Energy Storage","Hydrogen","Other Renewables"]},
    "ReNew Power":          {"region": "Asia",           "sectors": ["Solar","Wind","Energy Storage"]},
    "Adani Green Energy":   {"region": "Asia",           "sectors": ["Solar","Wind","Hydrogen"]},
    "Tata Power Renewables": {"region": "Asia",          "sectors": ["Solar","Wind","Energy Storage"]},
    "Azure Power":          {"region": "Asia",           "sectors": ["Solar"]},
    "Hero Future Energies": {"region": "Asia",           "sectors": ["Solar","Wind","Energy Storage"]},
    "Torrent Power":        {"region": "Asia",           "sectors": ["Solar","Wind","Energy Storage"]},
    "CESC":                 {"region": "Asia",           "sectors": ["Solar","Wind","Energy Storage"]},
    "Suzlon":               {"region": "Asia",           "sectors": ["Wind","Solar","Energy Storage"]},
    "Inox Wind":            {"region": "Asia",           "sectors": ["Wind"]},
    "Senvion India":        {"region": "Asia",           "sectors": ["Wind"]},
    "Waaree Energies":      {"region": "Asia",           "sectors": ["Solar"]},
    "Vikram Solar":         {"region": "Asia",           "sectors": ["Solar"]},
    "Sterling and Wilson":  {"region": "Asia/Global",    "sectors": ["Solar","Energy Storage"]},
    "KPI Green Energy":     {"region": "Asia",           "sectors": ["Solar"]},
    "Amp Energy India":     {"region": "Asia",           "sectors": ["Solar","Energy Storage"]},
    "Sembcorp Utilities":   {"region": "Asia/Global",    "sectors": ["Solar","Wind","Energy Storage","Hydrogen"]},
    "Sunseap":              {"region": "Asia",           "sectors": ["Solar","Energy Storage"]},
    "Sineng Electric":      {"region": "Asia",           "sectors": ["Solar","Energy Storage"]},
    "Huawei Smart PV":      {"region": "Asia/Global",    "sectors": ["Solar","Energy Storage"]},
    "JA Solar":             {"region": "Asia/Global",    "sectors": ["Solar"]},
    "LONGi Green Energy":   {"region": "Asia/Global",    "sectors": ["Solar"]},
    "Jinko Solar":          {"region": "Asia/Global",    "sectors": ["Solar"]},
    "Risen Energy":         {"region": "Asia/Global",    "sectors": ["Solar"]},
    "TBEA":                 {"region": "Asia/Global",    "sectors": ["Solar","Wind","Energy Storage"]},
    "CSSC Offshore & Marine": {"region": "Asia",         "sectors": ["Offshore Wind"]},
    "CIMC Raffles":         {"region": "Asia",           "sectors": ["Offshore Wind"]},
    "China Three Gorges":   {"region": "Asia/Global",    "sectors": ["Offshore Wind","Solar","Wind","Other Renewables"]},
    "China Energy Engineering": {"region": "Asia",       "sectors": ["Solar","Wind","Offshore Wind","Energy Storage"]},
    "PowerChina":           {"region": "Asia/Global",    "sectors": ["Solar","Wind","Offshore Wind","Energy Storage","Hydrogen"]},
    "Goldwind":             {"region": "Asia/Global",    "sectors": ["Wind","Offshore Wind","Energy Storage"]},
    "CRRC":                 {"region": "Asia",           "sectors": ["Wind","Energy Storage"]},
    "CSSC":                 {"region": "Asia",           "sectors": ["Offshore Wind","Wind"]},
    "Mingyang Smart Energy": {"region": "Asia/Global",   "sectors": ["Wind","Offshore Wind"]},
    "Sany Renewable Energy": {"region": "Asia",          "sectors": ["Wind","Solar"]},
    "DEC":                  {"region": "Asia",           "sectors": ["Wind","Solar","Energy Storage"]},
    "DPOWER":               {"region": "Asia",           "sectors": ["Wind","Energy Storage"]},
    "NHPC":                 {"region": "Asia",           "sectors": ["Other Renewables","Solar","Wind"]},
    "SECI":                 {"region": "Asia",           "sectors": ["Solar","Wind","Energy Storage","Hydrogen"]},
    "NTPC":                 {"region": "Asia",           "sectors": ["Solar","Wind","Energy Storage","Hydrogen"]},
    "APA Group":            {"region": "Asia Pacific",   "sectors": ["Hydrogen","Solar","Energy Storage"]},
    "Origin Energy":        {"region": "Asia Pacific",   "sectors": ["Solar","Wind","Energy Storage","Hydrogen"]},
    "AGL Energy":           {"region": "Asia Pacific",   "sectors": ["Solar","Wind","Energy Storage"]},
    "Infigen Energy":       {"region": "Asia Pacific",   "sectors": ["Wind","Solar"]},
    "Tilt Renewables":      {"region": "Asia Pacific",   "sectors": ["Wind","Solar","Energy Storage"]},
    "Contact Energy":       {"region": "Asia Pacific",   "sectors": ["Other Renewables","Wind"]},
    "Meridian Energy":      {"region": "Asia Pacific",   "sectors": ["Other Renewables","Wind","Solar"]},
    "JERA":                 {"region": "Asia",           "sectors": ["Offshore Wind","Solar","Hydrogen","Other Renewables"]},
    "Marubeni":             {"region": "Asia/Global",    "sectors": ["Wind","Offshore Wind","Solar"]},
    "Mitsubishi Power":     {"region": "Asia/Global",    "sectors": ["Hydrogen","Energy Storage","Other Renewables"]},
    "Sumitomo":             {"region": "Asia/Global",    "sectors": ["Solar","Wind","Offshore Wind"]},
    "Mitsui":               {"region": "Asia/Global",    "sectors": ["Solar","Wind","Offshore Wind","Hydrogen"]},
    "Doosan Enerbility":    {"region": "Asia/Global",    "sectors": ["Hydrogen","Offshore Wind","Energy Storage"]},
    "Hyundai Heavy Industries": {"region": "Asia/Global", "sectors": ["Offshore Wind","Energy Storage","Hydrogen"]},
    "Samsung C&T":          {"region": "Asia/Global",    "sectors": ["Solar","Wind","Offshore Wind","Energy Storage"]},
    "KEPCO":                {"region": "Asia/Global",    "sectors": ["Offshore Wind","Solar","Wind"]},
    "SK Ecoplant":          {"region": "Asia",           "sectors": ["Offshore Wind","Hydrogen","Energy Storage"]},
    "Hanwha Q Cells":       {"region": "Asia/Global",    "sectors": ["Solar","Energy Storage"]},

    # ── Middle East & Africa EPCs ─────────────────────────────────────────────
    "ACWA Power":           {"region": "Middle East/Global", "sectors": ["Solar","Wind","Energy Storage","Hydrogen","Other Renewables"]},
    "Masdar":               {"region": "Middle East/Global", "sectors": ["Solar","Wind","Offshore Wind","Energy Storage","Hydrogen"]},
    "ENEC":                 {"region": "Middle East",    "sectors": ["Other Renewables"]},
    "Arabian Construction Company": {"region": "Middle East", "sectors": ["Solar","Wind","Energy Storage"]},
    "Metito":               {"region": "Middle East",    "sectors": ["Solar","Energy Storage"]},
    "AMEA Power":           {"region": "Middle East/Africa", "sectors": ["Solar","Wind","Energy Storage"]},
    "Scatec":               {"region": "Africa/Global",  "sectors": ["Solar","Wind","Energy Storage","Hydrogen"]},
    "Lekela Power":         {"region": "Africa",         "sectors": ["Wind","Solar"]},
    "Globeleq":             {"region": "Africa",         "sectors": ["Solar","Wind","Energy Storage"]},
    "CrossBoundary Energy": {"region": "Africa",         "sectors": ["Solar","Energy Storage"]},
    "SOLA Group":           {"region": "Africa",         "sectors": ["Solar","Energy Storage"]},
    "Enertrag":             {"region": "Africa/Europe",  "sectors": ["Wind","Solar","Hydrogen"]},
    "Red Rocket":           {"region": "Africa",         "sectors": ["Solar","Wind","Energy Storage"]},

    # ── Latin America EPCs ────────────────────────────────────────────────────
    "Eneva":                {"region": "Latin America",  "sectors": ["Solar","Wind","Energy Storage"]},
    "Casa dos Ventos":      {"region": "Latin America",  "sectors": ["Wind","Solar","Energy Storage"]},
    "Caubvento":            {"region": "Latin America",  "sectors": ["Wind"]},
    "Atlas Renewable Energy": {"region": "Latin America/Global", "sectors": ["Solar","Wind","Energy Storage"]},
    "Orazul Energy":        {"region": "Latin America",  "sectors": ["Other Renewables","Solar","Wind"]},
    "Elecda":               {"region": "Latin America",  "sectors": ["Solar","Wind"]},
    "Mainstream":           {"region": "Latin America/Global", "sectors": ["Solar","Wind"]},
    "Sonnedix LATAM":       {"region": "Latin America",  "sectors": ["Solar"]},
    "ISA":                  {"region": "Latin America",  "sectors": ["Solar","Wind","Other Renewables"]},
    "Enel X":               {"region": "Global",         "sectors": ["Energy Storage","Solar","Wind"]},

    # ── Energy Storage Specialists ────────────────────────────────────────────
    "Fluence":              {"region": "Global",         "sectors": ["Energy Storage"]},
    "Tesla Energy":         {"region": "Global",         "sectors": ["Energy Storage"]},
    "Wartsila":             {"region": "Global",         "sectors": ["Energy Storage","Other Renewables"]},
    "Wärtsilä":             {"region": "Global",         "sectors": ["Energy Storage","Other Renewables"]},
    "NEC Energy Solutions": {"region": "Global",         "sectors": ["Energy Storage"]},
    "CATL":                 {"region": "Asia/Global",    "sectors": ["Energy Storage"]},
    "BYD Energy":           {"region": "Asia/Global",    "sectors": ["Energy Storage","Solar"]},
    "LG Energy Solution":   {"region": "Asia/Global",    "sectors": ["Energy Storage"]},
    "Samsung SDI":          {"region": "Asia/Global",    "sectors": ["Energy Storage"]},
    "Panasonic Energy":     {"region": "Asia/Global",    "sectors": ["Energy Storage"]},
    "Powin Energy":         {"region": "North America",  "sectors": ["Energy Storage"]},
    "Stem Inc":             {"region": "North America",  "sectors": ["Energy Storage"]},
    "Glidepath":            {"region": "North America",  "sectors": ["Energy Storage"]},
    "Energy Vault":         {"region": "Global",         "sectors": ["Energy Storage"]},
    "Form Energy":          {"region": "North America",  "sectors": ["Energy Storage"]},
    "Highview Power":       {"region": "UK/Global",      "sectors": ["Energy Storage"]},
    "Storelectric":         {"region": "UK",             "sectors": ["Energy Storage"]},
    "Malta Inc":            {"region": "North America",  "sectors": ["Energy Storage"]},
    "Ambri":                {"region": "North America",  "sectors": ["Energy Storage"]},
    "Saft":                 {"region": "Europe/Global",  "sectors": ["Energy Storage"]},
    "EnerSys":              {"region": "Global",         "sectors": ["Energy Storage"]},
    "Kokam":                {"region": "Asia/Global",    "sectors": ["Energy Storage"]},
    "Leclanché":            {"region": "Europe/Global",  "sectors": ["Energy Storage"]},
    "Electrovaya":          {"region": "North America",  "sectors": ["Energy Storage"]},
    "Invinity Energy":      {"region": "UK/Global",      "sectors": ["Energy Storage"]},
    "VRB Energy":           {"region": "Asia/Global",    "sectors": ["Energy Storage"]},
    "Primus Power":         {"region": "North America",  "sectors": ["Energy Storage"]},
    "ESS Tech":             {"region": "North America",  "sectors": ["Energy Storage"]},
    "Redflow":              {"region": "Asia Pacific",   "sectors": ["Energy Storage"]},

    # ── Hydrogen EPC Specialists ──────────────────────────────────────────────
    "ITM Power":            {"region": "UK/Global",      "sectors": ["Hydrogen"]},
    "Nel Hydrogen":         {"region": "Europe/Global",  "sectors": ["Hydrogen"]},
    "Plug Power":           {"region": "North America/Global", "sectors": ["Hydrogen","Energy Storage"]},
    "Bloom Energy":         {"region": "North America/Global", "sectors": ["Hydrogen","Energy Storage"]},
    "Air Products":         {"region": "Global",         "sectors": ["Hydrogen"]},
    "Air Liquide":          {"region": "Global",         "sectors": ["Hydrogen"]},
    "Linde":                {"region": "Global",         "sectors": ["Hydrogen"]},
    "ThyssenKrupp Nucera":  {"region": "Europe/Global",  "sectors": ["Hydrogen"]},
    "Cummins Electrolyzer": {"region": "Global",         "sectors": ["Hydrogen"]},
    "Haldor Topsoe":        {"region": "Europe/Global",  "sectors": ["Hydrogen"]},
    "Topsoe":               {"region": "Europe/Global",  "sectors": ["Hydrogen"]},
    "McPhy Energy":         {"region": "Europe",         "sectors": ["Hydrogen"]},
    "Sunfire":              {"region": "Europe/Global",  "sectors": ["Hydrogen"]},
    "ElectroChaea":         {"region": "Europe",         "sectors": ["Hydrogen"]},
    "HydrogenPro":          {"region": "Europe",         "sectors": ["Hydrogen"]},
    "Fortescue":            {"region": "Asia Pacific/Global", "sectors": ["Hydrogen","Energy Storage"]},
    "ACME Group":           {"region": "Asia",           "sectors": ["Hydrogen","Solar"]},
    "Gentari":              {"region": "Asia/Global",    "sectors": ["Hydrogen","Solar","Wind"]},
    "Yara Clean Ammonia":   {"region": "Europe/Global",  "sectors": ["Hydrogen"]},
    "OCI Global":           {"region": "Global",         "sectors": ["Hydrogen"]},
    "CF Industries":        {"region": "North America",  "sectors": ["Hydrogen"]},
    "Hy2gen":               {"region": "Europe/Global",  "sectors": ["Hydrogen"]},
    "Enerkem":              {"region": "Europe/Global",  "sectors": ["Hydrogen"]},
    "H2Pro":                {"region": "Middle East/Global", "sectors": ["Hydrogen"]},

    # ── Power Transmission & Grid EPC ─────────────────────────────────────────
    "Kalpataru Projects":   {"region": "Asia/Global",    "sectors": ["Solar","Wind","Other Renewables"]},
    "KEC International":    {"region": "Asia/Global",    "sectors": ["Solar","Wind","Other Renewables"]},
    "Sterlite Power":       {"region": "Asia/Global",    "sectors": ["Solar","Wind","Energy Storage"]},
    "PGCIL":                {"region": "Asia",           "sectors": ["Solar","Wind","Energy Storage"]},
    "Tata Projects":        {"region": "Asia",           "sectors": ["Solar","Wind","Energy Storage","Hydrogen"]},
    "L&T Construction":     {"region": "Asia/Global",    "sectors": ["Solar","Wind","Energy Storage","Offshore Wind","Hydrogen"]},
    "GE Grid Solutions":    {"region": "Global",         "sectors": ["Solar","Wind","Energy Storage","Offshore Wind"]},
    "Eaton":                {"region": "Global",         "sectors": ["Solar","Wind","Energy Storage"]},
    "Vertiv":               {"region": "Global",         "sectors": ["Energy Storage","Solar"]},
    "Hubbell":              {"region": "North America",  "sectors": ["Solar","Wind","Energy Storage"]},
    "Terna":                {"region": "Europe/Global",  "sectors": ["Solar","Wind","Offshore Wind"]},
    "National Grid":        {"region": "UK/North America", "sectors": ["Solar","Wind","Offshore Wind","Energy Storage"]},
    "RTE":                  {"region": "Europe",         "sectors": ["Solar","Wind","Offshore Wind","Energy Storage"]},
}

# ── EPC color for the sector picker ───────────────────────────────────────────
EPC_SECTOR_COLOR = "#ff6400"

# ── Build EPC Google News feed list dynamically ────────────────────────────────
def build_epc_feeds(selected_companies=None, region_filter=None, sector_filter=None):
    """
    Dynamically generate Google News RSS feeds for EPC companies.
    If selected_companies is given, only generate feeds for those.
    If region_filter / sector_filter is given, additionally cross-filter the registry.
    Returns a list of feed_meta dicts ready for _fetch_feed().
    """
    companies_to_use = selected_companies if selected_companies else list(EPC_COMPANIES.keys())

    # Apply sector filter from the sidebar sector selector
    if sector_filter:
        companies_to_use = [
            c for c in companies_to_use
            if any(s in EPC_COMPANIES.get(c, {}).get("sectors", []) for s in sector_filter)
        ]

    feeds = []
    seen_queries = set()
    for co in companies_to_use:
        co_clean = co.replace(" ", "+").replace("&", "%26").replace("'", "")
        # Core news query
        q1 = f"{co_clean}+renewable+energy+solar+wind+storage+hydrogen+power"
        q2 = f"{co_clean}+EPC+contract+project+construction+MW+GW"
        q3 = f"{co_clean}+power+plant+project+award+commissioned"
        for q in [q1, q2, q3]:
            if q not in seen_queries:
                seen_queries.add(q)
                feeds.append({
                    "url":    _GN + q,
                    "source": f"GNews/EPC-{co[:12].replace(' ', '')}",
                    "weight": 9,
                    "_epc_company": co,
                })
    return feeds


RE_FEED_REGISTRY = {
    # ── Solar ──────────────────────────────────────────────────────────────────
    "Solar": [
        # Specialist
        {"url": "https://www.pv-magazine.com/feed/",                       "source": "PV Magazine",        "weight": 10},
        {"url": "https://www.pv-tech.org/feed/",                           "source": "PV Tech",            "weight": 10},
        {"url": "https://www.solarpowerworldonline.com/feed/",             "source": "Solar Power World",  "weight":  8},
        {"url": "https://www.pveurope.eu/taxonomy/term/1/feed",            "source": "PV Europe",          "weight":  7},
        {"url": "https://pv-magazine-usa.com/feed/",                       "source": "PV Magazine USA",    "weight":  8},
        {"url": "https://www.solarpaces.org/news/feed/",                   "source": "SolarPACES",         "weight":  7},
        # Google News targeted queries
        {"url": _GN + "solar+PPA+power+purchase+agreement",               "source": "GNews/Solar-PPA",    "weight":  9},
        {"url": _GN + "solar+farm+tender+auction+MW+GW",                  "source": "GNews/Solar-Tender", "weight":  9},
        {"url": _GN + "solar+project+investment+financing+fund",          "source": "GNews/Solar-Inv",    "weight":  9},
        {"url": _GN + "solar+MOU+memorandum+agreement",                   "source": "GNews/Solar-MOU",    "weight":  8},
        {"url": _GN + "solar+commissioned+groundbreaking+construction",   "source": "GNews/Solar-Dev",    "weight":  8},
        {"url": _GN + "floating+solar+agrivoltaic+rooftop+solar",         "source": "GNews/Solar-Tech",   "weight":  7},
        {"url": _GN + "utility+scale+solar+GW+MW+announcement+2025",      "source": "GNews/Solar-Util",   "weight":  9},
        {"url": _GN + "solar+power+plant+acquisition+deal+stake",         "source": "GNews/Solar-MA",     "weight":  8},
        {"url": _GN + "bifacial+solar+tracker+module+project",            "source": "GNews/Solar-Tech2",  "weight":  7},
    ],
    # ── Wind ──────────────────────────────────────────────────────────────────
    "Wind": [
        {"url": "https://www.windpowermonthly.com/rss",                    "source": "Wind Power Monthly", "weight": 10},
        {"url": "https://windenergynews.com/feed/",                        "source": "Wind Energy News",   "weight":  9},
        {"url": "https://www.windpowerengineering.com/feed/",              "source": "Windpower Eng.",     "weight":  8},
        {"url": "https://www.awea.org/rss",                                "source": "AWEA",               "weight":  8},
        {"url": _GN + "onshore+wind+farm+tender+auction+MW+GW",           "source": "GNews/Wind-Tender",  "weight":  9},
        {"url": _GN + "wind+energy+PPA+power+purchase+agreement",         "source": "GNews/Wind-PPA",     "weight":  9},
        {"url": _GN + "wind+project+investment+financing+turbine",        "source": "GNews/Wind-Inv",     "weight":  9},
        {"url": _GN + "wind+farm+commissioned+groundbreaking",            "source": "GNews/Wind-Dev",     "weight":  8},
        {"url": _GN + "wind+repowering+capacity+expansion",               "source": "GNews/Wind-Repow",   "weight":  7},
        {"url": _GN + "wind+turbine+order+supply+agreement+manufacturer", "source": "GNews/Wind-OEM",     "weight":  8},
        {"url": _GN + "onshore+wind+MW+GW+awarded+contract+2025",        "source": "GNews/Wind-Award",   "weight":  9},
    ],
    # ── Energy Storage ────────────────────────────────────────────────────────
    "Energy Storage": [
        {"url": "https://www.energy-storage.news/feed/",                   "source": "Energy Storage News","weight": 10},
        {"url": "https://www.storagedaily.com/rss/",                       "source": "Storage Daily",      "weight":  8},
        {"url": "https://batteryindustry.tech/feed/",                      "source": "Battery Industry",   "weight":  7},
        {"url": _GN + "battery+energy+storage+BESS+MW+GWh",              "source": "GNews/BESS",         "weight":  9},
        {"url": _GN + "battery+storage+PPA+offtake+tender",              "source": "GNews/BESS-PPA",     "weight":  9},
        {"url": _GN + "battery+storage+investment+financing+fund",        "source": "GNews/BESS-Inv",     "weight":  9},
        {"url": _GN + "grid+battery+commissioned+BESS+project",          "source": "GNews/BESS-Dev",     "weight":  8},
        {"url": _GN + "long+duration+storage+flow+battery+pumped+hydro", "source": "GNews/LDES",         "weight":  8},
        {"url": _GN + "Tesla+Megapack+Fluence+energy+storage+project",   "source": "GNews/BESS-OEM",     "weight":  7},
        {"url": _GN + "grid+scale+BESS+storage+GWh+MWh+contract+2025",  "source": "GNews/BESS-2025",    "weight":  9},
        {"url": _GN + "virtual+power+plant+demand+response+battery",     "source": "GNews/VPP",          "weight":  7},
    ],
    # ── Offshore Wind ─────────────────────────────────────────────────────────
    "Offshore Wind": [
        {"url": "https://www.offshorewind.biz/feed/",                      "source": "Offshore Wind Biz",  "weight": 10},
        {"url": "https://www.4coffshore.com/rss/windfarms-rss.xml",        "source": "4C Offshore",        "weight":  9},
        {"url": "https://windeurope.org/feed/",                            "source": "WindEurope",         "weight":  9},
        {"url": _GN + "offshore+wind+farm+tender+auction+CfD",            "source": "GNews/OSW-Tender",   "weight": 10},
        {"url": _GN + "offshore+wind+PPA+power+purchase+agreement",       "source": "GNews/OSW-PPA",      "weight":  9},
        {"url": _GN + "offshore+wind+investment+financing+fund",          "source": "GNews/OSW-Inv",      "weight":  9},
        {"url": _GN + "offshore+wind+commissioned+first+power",           "source": "GNews/OSW-Dev",      "weight":  8},
        {"url": _GN + "floating+offshore+wind+project+MOU",               "source": "GNews/FOW",          "weight":  8},
        {"url": _GN + "offshore+wind+lease+seabed+monopile",              "source": "GNews/OSW-Tech",     "weight":  7},
        {"url": _GN + "offshore+wind+GW+MW+awarded+contract+2025",       "source": "GNews/OSW-2025",     "weight": 10},
        {"url": _GN + "fixed+bottom+offshore+wind+foundation+installation","source": "GNews/OSW-Found",   "weight":  7},
        {"url": _GN + "offshore+wind+cable+substation+interconnection",   "source": "GNews/OSW-Grid",     "weight":  7},
    ],
    # ── Hydrogen ──────────────────────────────────────────────────────────────
    "Hydrogen": [
        {"url": "https://www.hydrogeninsight.com/rss",                     "source": "Hydrogen Insight",   "weight": 10},
        {"url": "https://www.h2-view.com/feed/",                           "source": "H2 View",            "weight": 10},
        {"url": "https://fuelcellsworks.com/feed/",                        "source": "Fuel Cells Works",   "weight":  8},
        {"url": "https://www.hydrogenfuelnews.com/feed/",                  "source": "Hydrogen Fuel News", "weight":  8},
        {"url": "https://www.h2bulletin.com/feed/",                        "source": "H2 Bulletin",        "weight":  8},
        {"url": _GN + "green+hydrogen+electrolyzer+electrolyser+PPA",     "source": "GNews/H2-PPA",       "weight":  9},
        {"url": _GN + "green+hydrogen+project+investment+financing",      "source": "GNews/H2-Inv",       "weight":  9},
        {"url": _GN + "hydrogen+MOU+memorandum+agreement+partnership",    "source": "GNews/H2-MOU",       "weight":  9},
        {"url": _GN + "hydrogen+offtake+supply+agreement+tonnes",         "source": "GNews/H2-Offtake",   "weight":  9},
        {"url": _GN + "green+hydrogen+plant+commissioned+groundbreaking", "source": "GNews/H2-Dev",       "weight":  8},
        {"url": _GN + "ammonia+green+hydrogen+export+terminal",           "source": "GNews/H2-Ammonia",   "weight":  7},
        {"url": _GN + "hydrogen+GW+MW+electrolysis+capacity+2025",       "source": "GNews/H2-2025",      "weight":  9},
        {"url": _GN + "blue+hydrogen+CCS+carbon+capture+facility",       "source": "GNews/H2-Blue",      "weight":  7},
    ],
    # ── Other Renewables ──────────────────────────────────────────────────────
    "Other Renewables": [
        {"url": "https://www.rechargenews.com/rss",                        "source": "Recharge News",      "weight":  9},
        {"url": "https://www.enerdata.net/rss.xml",                        "source": "Enerdata",           "weight":  8},
        {"url": "https://www.renewableenergyworld.com/feed/",              "source": "RE World",           "weight":  8},
        {"url": "https://www.renewablesnow.com/feed/",                     "source": "Renewables Now",     "weight":  9},
        {"url": _GN + "geothermal+power+project+investment+tender",       "source": "GNews/Geo",          "weight":  9},
        {"url": _GN + "hydropower+hydro+dam+project+tender+financing",    "source": "GNews/Hydro",        "weight":  9},
        {"url": _GN + "tidal+wave+marine+energy+project",                 "source": "GNews/Marine",       "weight":  8},
        {"url": _GN + "biomass+bioenergy+biogas+renewable+project",       "source": "GNews/Bio",          "weight":  8},
        {"url": _GN + "concentrated+solar+CSP+solar+thermal+project",     "source": "GNews/CSP",          "weight":  7},
        {"url": _GN + "nuclear+SMR+small+modular+reactor+PPA+investment", "source": "GNews/Nuclear",      "weight":  9},
        {"url": _GN + "geothermal+MW+GW+investment+project+2025",        "source": "GNews/Geo-2025",     "weight":  9},
    ],
}

# Cross-sector feeds — appear in ALL sectors; tagged during ingest
RE_CROSS_FEEDS = [
    {"url": "https://cleantechnica.com/feed/",                             "source": "CleanTechnica",      "weight": 9},
    {"url": "https://electrek.co/feed/",                                   "source": "Electrek",           "weight": 9},
    {"url": "https://www.canarymedia.com/feed.xml",                        "source": "Canary Media",       "weight": 9},
    {"url": "https://ieefa.org/feed/",                                     "source": "IEEFA",              "weight": 8},
    {"url": "https://www.carbonbrief.org/feed",                            "source": "Carbon Brief",       "weight": 8},
    {"url": "https://www.greenbiz.com/feeds/all",                          "source": "GreenBiz",           "weight": 7},
    {"url": "https://www.climatechangenews.com/feed/",                     "source": "Climate Home News",  "weight": 7},
    {"url": "https://www.spglobal.com/commodityinsights/en/rss/all",       "source": "S&P Global Commod.", "weight": 8},
    {"url": "https://energymonitor.ai/feed/",                              "source": "Energy Monitor",     "weight": 8},
    {"url": "https://www.renewableenergyworld.com/feed/",                  "source": "RE World",           "weight": 8},
    {"url": _GN + "renewable+energy+PPA+MOU+tender+investment+deal",      "source": "GNews/RE-Deals",     "weight": 9},
    {"url": _GN + "clean+energy+acquisition+merger+IPO+financing",        "source": "GNews/RE-Finance",   "weight": 9},
    {"url": _GN + "renewable+energy+policy+regulation+target+legislation","source": "GNews/RE-Policy",    "weight": 8},
    {"url": _GN + "wind+solar+storage+hydrogen+commissioned+tender",      "source": "GNews/RE-All",       "weight": 9},
    {"url": _GN + "clean+energy+GW+MW+project+announcement+2025",        "source": "GNews/RE-2025",      "weight": 9},
    {"url": _GN + "energy+transition+decarbonization+net+zero+deal",      "source": "GNews/Transition",   "weight": 8},
    {"url": _GN + "power+purchase+agreement+corporate+renewable+2025",    "source": "GNews/Corp-PPA",     "weight": 9},
    {"url": _GN + "renewable+energy+financial+close+funded+milestone",    "source": "GNews/RE-FC",        "weight": 8},
]

# All unique sources for the KPI banner
RE_SOURCES = (
    [{"name": f["source"]} for feeds in RE_FEED_REGISTRY.values() for f in feeds]
    + [{"name": f["source"]} for f in RE_CROSS_FEEDS]
)

RE_SOURCE_META = {
    "PV Magazine":       {"color": "#ff6400", "short": "PVM"},
    "PV Tech":           {"color": "#ffaa00", "short": "PVT"},
    "PV Magazine USA":   {"color": "#ff7700", "short": "PVU"},
    "Solar Power World": {"color": "#ffcc00", "short": "SPW"},
    "Wind Power Monthly":{"color": "#00b4ff", "short": "WPM"},
    "Wind Energy News":  {"color": "#0090cc", "short": "WEN"},
    "Windpower Eng.":    {"color": "#0070aa", "short": "WPE"},
    "Energy Storage News":{"color":"#00e5c8", "short": "ESN"},
    "Storage Daily":     {"color": "#00bbaa", "short": "STD"},
    "Offshore Wind Biz": {"color": "#0088cc", "short": "OWB"},
    "4C Offshore":       {"color": "#005599", "short": "4CO"},
    "Hydrogen Insight":  {"color": "#a855f7", "short": "HYI"},
    "H2 View":           {"color": "#9333ea", "short": "H2V"},
    "Fuel Cells Works":  {"color": "#7c3aed", "short": "FCW"},
    "Hydrogen Fuel News":{"color": "#6d28d9", "short": "HFN"},
    "Recharge News":     {"color": "#00e676", "short": "RCH"},
    "Enerdata":          {"color": "#00cc66", "short": "END"},
    "CleanTechnica":     {"color": "#22c55e", "short": "CTC"},
    "Electrek":          {"color": "#16a34a", "short": "ELK"},
    "Canary Media":      {"color": "#f59e0b", "short": "CNR"},
    "IEEFA":             {"color": "#0047e1", "short": "IEF"},
    "Carbon Brief":      {"color": "#ef4444", "short": "CBR"},
    "GreenBiz":          {"color": "#15803d", "short": "GBZ"},
}
# All GNews sources get a default teal
for _k in list(RE_SOURCE_META.keys()):
    pass
_RE_GNEWS_COLOR = {"color": "#00bcd4", "short": "GNW"}


# ─── RE Deal-type detection ────────────────────────────────────────────────────
def detect_re_deal_type(text):
    t = text.lower()
    # Priority order: specific deal types first
    for dtype in ["PPA","MOU","AOR / Offtake","Tender / Auction","Investment / IPO",
                  "M&A / Deals","Commissioning","Construction","Policy / Reg","Grid / Connect"]:
        kws = RE_DEAL_TYPE_KEYWORDS.get(dtype, [])
        if any(k in t for k in kws):
            return dtype
    return "General"

def detect_re_topic(text):
    t = text.lower()
    # Check most-specific sectors first (Offshore Wind before Wind)
    for topic in ["Offshore Wind", "Energy Storage", "Hydrogen",
                  "Solar", "Wind", "Other Renewables"]:
        kws = RE_TOPIC_KEYWORDS.get(topic, [])
        if any(k.lower() in t for k in kws):
            return topic
    return "Other Renewables"

def detect_re_capacity(text):
    m = re.search(r"([\d,]+(?:\.\d+)?)\s*(GW|MW|kW|MWh|GWh|gigawatt|megawatt|kilowatt)\b", text, re.I)
    if m:
        val = m.group(1).replace(",", "")
        unit = m.group(2)
        unit_map = {"gigawatt": "GW","megawatt": "MW","kilowatt": "kW"}
        unit = unit_map.get(unit.lower(), unit.upper())
        return val + " " + unit
    return ""

def is_re_relevant(text):
    t = text.lower()
    return any(p in t for p in [
        "solar","wind","renewable","green energy","clean energy","energy storage",
        "bess","battery storage","hydrogen","electrolyser","electrolyzer",
        "ppa","power purchase agreement","offshore wind","onshore wind",
        "photovoltaic","solar farm","wind farm","geothermal","hydropower",
        "hydroelectric","biomass energy","cfd","capacity auction","energy tender",
        "green hydrogen","fuel cell","ammonia","electrolysis","turbine",
        "solar developer","wind developer","renewables","clean power",
        "decarbonisation","decarbonization","energy transition",
        "gigawatt","gw solar","gw wind","mw solar","mw wind","mwh battery",
        "gwh storage","offshore lease","hvdc","interconnector",
    ])

# ─── Multi-source RSS engine ───────────────────────────────────────────────────

import xml.etree.ElementTree as _ET
from concurrent.futures import ThreadPoolExecutor, as_completed as _as_completed
import urllib.parse as _urlparse

_RSS_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/124.0.0.0 Safari/537.36"
    ),
    "Accept": "application/rss+xml, application/atom+xml, application/xml, text/xml, */*",
    "Accept-Language": "en-US,en;q=0.9",
    "Cache-Control": "no-cache",
}
_ATOM_NS = "http://www.w3.org/2005/Atom"


def _fetch_feed(feed_meta, timeout=12):
    """
    Fetch a single RSS/Atom feed. Returns (feed_meta, raw_text) or (feed_meta, None).
    Tries cloudscraper first (bypasses Cloudflare), then plain requests.
    """
    url = feed_meta["url"]
    try:
        if _USE_CS:
            r = _CS.get(url, timeout=timeout)
        else:
            import requests as _req
            r = _req.get(url, headers=_RSS_HEADERS, timeout=timeout)
        if r.status_code == 200 and len(r.text) > 200:
            return feed_meta, r.text
    except Exception:
        pass
    # Fallback: plain requests if cloudscraper failed
    try:
        import requests as _req2
        r2 = _req2.get(url, headers=_RSS_HEADERS, timeout=timeout)
        if r2.status_code == 200 and len(r2.text) > 200:
            return feed_meta, r2.text
    except Exception:
        pass
    return feed_meta, None


def _parse_feed_xml(xml_text, feed_meta, sector_tag, cutoff):
    """
    Parse RSS 2.0 or Atom 1.0 XML.
    Returns list of article dicts.
    For bounded cutoffs (not datetime.min), articles with NO parseable date are
    dropped to avoid surfacing stale content.  For open-range scans they're kept.
    """
    articles = []
    _open_range = (cutoff <= datetime.min + timedelta(days=1))
    try:
        root = _ET.fromstring(xml_text.encode("utf-8", errors="replace"))
    except Exception:
        try:
            import re as _re2
            clean = _re2.sub(r'[^\x09\x0A\x0D\x20-\uD7FF\uE000-\uFFFD]', '', xml_text)
            root = _ET.fromstring(clean.encode("utf-8", errors="replace"))
        except Exception:
            return articles

    # Detect RSS vs Atom
    is_atom = root.tag == f"{{{_ATOM_NS}}}feed" or "atom" in root.tag.lower()

    if is_atom:
        items = root.findall(f"{{{_ATOM_NS}}}entry")
        for item in items:
            def _at(tag):
                el = item.find(f"{{{_ATOM_NS}}}{tag}")
                return el.text.strip() if el is not None and el.text else ""
            title = _at("title")
            # Atom link is an element with href attr
            link_el = item.find(f"{{{_ATOM_NS}}}link[@rel='alternate']") or item.find(f"{{{_ATOM_NS}}}link")
            url = link_el.get("href", "") if link_el is not None else ""
            pub = _at("published") or _at("updated")
            date_obj = parse_date_str(pub) if pub else None
            if not title or not url:
                continue
            if date_obj and date_obj < cutoff:
                continue
            if not date_obj and not _open_range:
                continue  # drop undated articles for bounded time windows
            articles.append(_mk_re_art(title, url, date_obj, feed_meta, sector_tag))
    else:
        # RSS 2.0 / RSS 1.0
        items = root.findall(".//item")
        for item in items:
            def _rss(tag, ns=""):
                el = item.find(f"{ns}{tag}")
                return el.text.strip() if el is not None and el.text else ""
            title = _rss("title")
            url   = _rss("link") or _rss("guid")
            # <link> is sometimes text after the tag in RSS
            if not url:
                link_el = item.find("link")
                if link_el is not None:
                    url = (link_el.text or "").strip()
                    if not url:
                        url = str(link_el.tail or "").strip()
            pub = _rss("pubDate") or _rss("published") or _rss("dc:date") or _rss("date")
            # Try dc:date namespace
            if not pub:
                dc_el = item.find("{http://purl.org/dc/elements/1.1/}date")
                if dc_el is not None and dc_el.text:
                    pub = dc_el.text.strip()
            date_obj = parse_date_str(pub) if pub else None
            if not title or not url:
                continue
            if date_obj and date_obj < cutoff:
                continue
            if not date_obj and not _open_range:
                continue  # drop undated articles for bounded time windows
            articles.append(_mk_re_art(title, url, date_obj, feed_meta, sector_tag))
    return articles


def _mk_re_art(headline, url, date_obj, feed_meta, sector_tag):
    """Build a raw article dict from parsed feed data."""
    # Clean Google News redirect URLs → extract real URL
    real_url = url
    if "news.google.com" in url:
        try:
            parsed = _urlparse.urlparse(url)
            qs = _urlparse.parse_qs(parsed.query)
            if "url" in qs:
                real_url = qs["url"][0]
        except Exception:
            pass
    # Clean headline — strip HTML entities
    hl = re.sub(r"<[^>]+>", "", headline)
    hl = hl.replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">").replace("&quot;", '"').replace("&#39;", "'")
    hl = re.sub(r"\s+", " ", hl).strip()
    return {
        "headline":  hl,
        "url":       real_url.strip(),
        "date_obj":  date_obj,
        "source":    feed_meta.get("source", "Unknown"),
        "_sector":   sector_tag,
        "_weight":   feed_meta.get("weight", 5),
        "_priority": feed_meta.get("weight", 5),
    }


def run_re_scrapers(max_html_pages, cutoff, progress_cb, re_sectors=None,
                    epc_companies=None, epc_sector_filter=None):
    """
    Multi-source parallel RSS engine.
    - Fetches specialist publication feeds + Google News query feeds concurrently
    - Covers all 6 sectors + EPC Companies mode with sector-specific + cross-sector feeds
    - re_sectors: list of selected sectors (filters which feeds to hit)
    - epc_companies: list of EPC company names to generate targeted feeds for
    - epc_sector_filter: sector list to cross-filter EPC companies by
    - max_html_pages: repurposed as a depth signal (1=light, 3+=deep)
    - cutoff: datetime — articles older than this are dropped
    """
    _epc_mode = re_sectors and "EPC Companies" in re_sectors
    # Sectors excluding the virtual EPC sector (it has no RE_FEED_REGISTRY entry)
    _real_sectors = [s for s in (re_sectors or list(RE_FEED_REGISTRY.keys())) if s != "EPC Companies"]
    active_sectors = _real_sectors if _real_sectors else list(RE_FEED_REGISTRY.keys())

    feed_jobs = []

    if _epc_mode:
        # ── EPC MODE: only fire EPC company-specific Google News feeds ─────────
        # Also include the standard sector feeds if other sectors are selected too
        for sector in _real_sectors:
            for feed in RE_FEED_REGISTRY.get(sector, []):
                feed_jobs.append((feed, sector))

        # Build targeted EPC feeds
        epc_feeds = build_epc_feeds(
            selected_companies=epc_companies if epc_companies else None,
            sector_filter=epc_sector_filter or _real_sectors or None,
        )
        for feed in epc_feeds:
            feed_jobs.append((feed, "EPC Companies"))

        # Cross-sector feeds — include for broader coverage
        for feed in RE_CROSS_FEEDS:
            feed_jobs.append((feed, None))
    else:
        # ── NORMAL MODE: sector-specific + cross-sector feeds ──────────────────
        for sector in active_sectors:
            for feed in RE_FEED_REGISTRY.get(sector, []):
                feed_jobs.append((feed, sector))

        # Cross-sector feeds — always include, tag by keyword during enrichment
        for feed in RE_CROSS_FEEDS:
            feed_jobs.append((feed, None))  # None = detect sector from headline

    total_feeds = len(feed_jobs)
    progress_cb(0.0, f"🌿 Fetching {total_feeds} RSS feeds across {len(active_sectors)} sectors…")

    raw       = []
    seen_urls = set()
    completed = 0

    def _process_job(job):
        feed_meta, sector_tag = job
        meta, xml_text = _fetch_feed(feed_meta, timeout=12)
        if not xml_text:
            return []
        arts = _parse_feed_xml(xml_text, meta, sector_tag, cutoff)
        return arts

    # Parallel fetch — ThreadPoolExecutor, 12 workers
    max_workers = min(12, total_feeds)
    with ThreadPoolExecutor(max_workers=max_workers) as pool:
        futures = {pool.submit(_process_job, job): job for job in feed_jobs}
        for future in _as_completed(futures):
            completed += 1
            frac = completed / total_feeds
            feed_meta = futures[future][0]
            try:
                arts = future.result()
                for a in arts:
                    norm = a["url"].rstrip("/")
                    # Skip empty or nav URLs
                    if not norm or len(norm) < 15:
                        continue
                    if norm not in seen_urls:
                        seen_urls.add(norm)
                        raw.append(a)
            except Exception:
                pass
            progress_cb(
                frac * 0.9,
                f"📡 {feed_meta.get('source','feed')} → {len(raw)} articles so far…",
            )

    progress_cb(0.92, f"🔍 Filtering {len(raw)} articles for relevance…")

    # Relevance + dedup pass
    raw = [a for a in raw if a["headline"] and is_re_relevant(a["headline"])]

    # For cross-sector articles (sector_tag=None), detect from headline
    for a in raw:
        if not a.get("_sector"):
            a["_sector"] = detect_re_topic(a["headline"])

    # Apply sector filter if set (cross-sector articles may belong to multiple)
    _effective_sectors = [s for s in (re_sectors or []) if s != "EPC Companies"]
    if _epc_mode and epc_companies:
        # In EPC mode: keep articles that mention at least one selected EPC company
        _epc_lower = [c.lower() for c in epc_companies]
        def _epc_match(art):
            hl = art["headline"].lower()
            return any(ec in hl for ec in _epc_lower)
        raw = [a for a in raw if _epc_match(a) or art.get("_sector") != "EPC Companies"]
    elif _effective_sectors:
        def _match(art):
            detected = detect_re_topic(art["headline"])
            tagged   = art.get("_sector")
            return detected in _effective_sectors or tagged in _effective_sectors
        raw = [a for a in raw if _match(a)]
    elif re_sectors and not _epc_mode:
        def _match2(art):
            detected = detect_re_topic(art["headline"])
            tagged   = art.get("_sector")
            return detected in re_sectors or tagged in re_sectors
        raw = [a for a in raw if _match2(a)]

    # Sort: dated articles first (newest → oldest), undated at end
    raw.sort(key=lambda a: a.get("date_obj") or datetime.min, reverse=True)

    progress_cb(1.0, f"✅ {len(raw)} articles from {len(set(a['source'] for a in raw))} sources")
    return raw


def enrich_re(raw_item):
    """Enrich a raw RE article dict → structured display dict."""
    hl       = raw_item["headline"]
    d        = raw_item.get("date_obj")
    country  = detect_country(hl)
    region   = COUNTRY_TO_REGION.get(country, "Global")
    capacity = detect_re_capacity(hl)
    deal_sz  = detect_deal_size(hl)
    sentiment_label, sentiment_score = detect_re_sentiment(hl, deal_size=deal_sz, capacity=capacity)
    return {
        "Headline":  hl,
        "Date":      d.strftime("%Y-%m-%d") if d else "Unknown",
        "Source":    raw_item.get("source", "Unknown"),
        "URL":       raw_item.get("url", ""),
        "Country":   country,
        "Region":    region,
        "Sector":    raw_item.get("_sector") or detect_re_topic(hl),
        "Deal Type": detect_re_deal_type(hl),
        "Status":    detect_re_project_status(hl),
        "Capacity":  capacity,
        "Deal Size": deal_sz,
        "Companies": detect_companies(hl),
        "Sentiment": sentiment_label,
        "Sentiment Score": sentiment_score,
        "_date_obj": d,
    }



RE_TOPIC_KEYWORDS = {
    "Solar": [
        "solar","photovoltaic","pv ","solar farm","solar park","solar plant",
        "solar project","solar array","solar panel","bifacial","utility-scale solar",
        "rooftop solar","floating solar","agrivoltaic","cpv","concentrated solar",
        "solar irradiance","solar developer","solar installer",
    ],
    "Offshore Wind": [
        "offshore wind","floating wind","offshore turbine","monopile","jacket foundation",
        "offshore wind farm","offshore wind park","wind turbine installation vessel",
        "wtiv","offshore substation","offshore lease","floating offshore",
    ],
    "Wind": [
        "wind farm","wind park","wind project","wind turbine","onshore wind",
        "wind energy","wind power","anemometer","wind blade","nacelle",
        "wind developer","wind repowering","land wind","wind capacity",
    ],
    "Energy Storage": [
        "battery storage","bess","energy storage","grid-scale battery",
        "lithium-ion storage","flow battery","vanadium battery","battery system",
        "storage project","ess ","battery energy storage","grid battery",
        "long duration storage","flywheel","compressed air energy storage",
        "pumped hydro","pumped storage","gravity storage",
    ],
    "Hydrogen": [
        "hydrogen","electrolysis","electrolyser","electrolyzer","green hydrogen",
        "blue hydrogen","pink hydrogen","grey hydrogen","h2 ","fuel cell",
        "ammonia","hydrogen pipeline","hydrogen hub","ptx","power-to-gas",
        "hydrogen project","hydrogen offtake","hydrogen ppa","hydrogen storage",
    ],
    "Other Renewables": [
        "geothermal","hydro","hydropower","hydroelectric","tidal","wave energy",
        "marine energy","bioenergy","biomass","biogas","landfill gas",
        "waste-to-energy","csp","concentrated solar power","solar thermal",
        "ocean energy","run-of-river",
    ],
}

RE_DEAL_TYPE_KEYWORDS = {
    "Tender / Auction": [
        "tender","auction","bid","cfds","request for proposal","rfp","solicitation",
        "competitive tender","award tender","capacity tender","round ","cfd round",
        "awarded contract","contract award",
    ],
    "PPA": [
        "ppa","power purchase agreement","offtake agreement","long-term contract",
        "virtual ppa","vppa","corporate ppa","sleeved ppa","private wire",
        "power offtake","energy supply agreement","electricity supply contract",
        "long-term power","15-year contract","20-year contract","25-year contract",
    ],
    "MOU": [
        "mou","memorandum of understanding","framework agreement","heads of terms",
        "loi","letter of intent","development agreement","cooperation agreement",
        "non-binding agreement","joint development","strategic partnership",
    ],
    "AOR / Offtake": [
        "aor","agreement on record","offtake","anchor offtake","capacity agreement",
        "tolling agreement","capacity reservation","hydrogen offtake","storage offtake",
        "virtual offtake",
    ],
    "Investment": [
        "investment","funding","financing","equity","debt","bond","green bond",
        "sustainability bond","loan","credit facility","project finance","raise",
        "financial close","fc ","reached financial close","equity raise","debt raise",
        "infrastructure fund","joint venture","acquisition","stake","acquired",
        "merger","takeover","ipo","listing",
    ],
    "Commissioning": [
        "commissioned","commissioning","energized","energised","goes live","online",
        "began operations","first power","first generation","inaugurated","opened",
        "ribbon cutting","commercial operation","cod","coo","cod achieved",
    ],
    "Construction": [
        "broke ground","groundbreaking","construction begins","under construction",
        "build","installation","erection","foundation","civil works","epc",
        "procurement","turbine installation","panel installation",
    ],
    "Grid / Interconnect": [
        "grid connection","interconnection","grid code","network connection",
        "transmission access","substation","capacity booking","connection offer",
        "offered capacity","queue","interconnection agreement","shallow study",
        "deep study","grid study",
    ],
    "M&A": [
        "acquisition","acquired","merger","takeover","stake","bought","purchased",
        "divested","sold ","portfolio sale","asset sale","m&a",
    ],
    "Policy": [
        "regulation","policy","legislation","standard","target","mandate","decree",
        "executive order","parliament","senate","vote","approved","enacted","signed",
    ],
}

# ─── Renewable energy news sources ────────────────────────────────────────────
RE_SOURCES = [
    {"name": "Renewables Now",  "url": "https://renewablesnow.com/", "type": "html"},
]

# ─── RE Source meta ────────────────────────────────────────────────────────────
RE_SOURCE_META = {
    "Renewables Now":  {"color": "#00e676", "short": "RNW"},
    "Unknown":         {"color": "#2e4470", "short": "UNK"},
}

# ─── RE URL constants — the 6 confirmed RenewablesNow section URLs ────────────

# ─── RE Deal-type detection ────────────────────────────────────────────────────
def detect_re_deal_type(text):
    t = text.lower()
    for dtype, kws in RE_DEAL_TYPE_KEYWORDS.items():
        if any(k in t for k in kws):
            return dtype
    return "General"

def detect_re_topic(text):
    t = text.lower()
    # Check most-specific sectors first to avoid "wind" matching "offshore wind"
    for topic in ["Offshore Wind", "Energy Storage", "Hydrogen",
                  "Solar", "Wind", "Other Renewables"]:
        kws = RE_TOPIC_KEYWORDS.get(topic, [])
        if any(k.lower() in t for k in kws):
            return topic
    return "Other Renewables"

def detect_re_capacity(text):
    """Extract GW/MW/kW capacity from renewable energy article headline."""
    m = re.search(r"([\d,]+(?:\.\d+)?)\s*(GW|MW|kW|gigawatt|megawatt|kilowatt)\b", text, re.I)
    if m:
        val = m.group(1).replace(",", "")
        unit = m.group(2).upper()
        if unit in ("KILOWATT", "KW"):
            unit = "kW"
        return val + " " + unit
    return ""

def _re_stem(word):
    """Crude suffix-stripping stemmer so 'approved' also matches 'approves'/
    'approving', 'blocked' matches 'blocks', etc. Only used for single-token
    phrases — multi-word phrases stay literal (far less tense ambiguity)."""
    for suf in ("ied", "ies", "ing", "ed", "es", "s", "d"):
        if word.endswith(suf) and len(word) - len(suf) >= 3:
            return word[: -len(suf)]
    return word


def _re_phrase_match(t, phrase):
    """Match a rule phrase against lowercased text t. Single-token phrases
    are stemmed + prefix-matched (catches verb tense variants); multi-word
    phrases are matched literally with word boundaries."""
    phrase = phrase.strip()
    if " " in phrase:
        return re.search(r"(?<!\w)" + re.escape(phrase) + r"(?!\w)", t) is not None
    stem = _re_stem(phrase)
    return re.search(r"\b" + re.escape(stem) + r"\w*\b", t) is not None


def detect_re_project_status(text):
    """
    Classify the project lifecycle stage of an RE headline.
    Scores every category (not just the first match, which was the old
    behaviour and silently let an early generic word like "plans" win over
    a more specific later phrase like "cancelled"). Uses word-boundary
    phrase matching, weighted by specificity, with a fixed severity order
    to break ties.
    """
    t = text.lower()

    RULES = {
        "Commissioned": [
            ("commissioned", 2), ("energized", 3), ("energised", 3),
            ("goes live", 3), ("online", 1), ("began operations", 3),
            ("first power", 3), ("commercial operation", 3), (" cod ", 2),
            ("inaugurated", 2), ("now operational", 3), ("fully operational", 3),
        ],
        "Financed": [
            ("financial close", 3), ("reached financial close", 3),
            ("fc achieved", 3), ("fully funded", 2), ("financing agreed", 3),
            ("financing closed", 3), ("secures financing", 2), ("raises debt", 2),
        ],
        "Under Construction": [
            ("broke ground", 3), ("groundbreaking", 2), ("under construction", 3),
            ("construction begins", 3), ("construction starts", 3),
            ("installation begins", 2), ("ntp issued", 2), ("notice to proceed", 3),
        ],
        "Approved": [
            ("approved", 1), ("consent granted", 3), ("planning approved", 3),
            ("permit granted", 3), ("permits granted", 3), ("green light", 2),
            ("go-ahead", 2), ("eia approved", 3), ("regulatory approval", 3),
        ],
        "Tendered": [
            ("tender", 1), ("auction", 1), ("rfp", 2), ("solicitation", 2),
            ("cfd round", 3), ("procurement round", 3), ("invites bids", 2),
            ("request for proposals", 3),
        ],
        "Contracted": [
            ("ppa signed", 3), ("ppa agreed", 3), ("ppa awarded", 3),
            ("offtake signed", 3), ("offtake agreement", 2), ("contract signed", 2),
            ("mou signed", 2), ("agreement signed", 2), ("deal signed", 2),
            ("wins contract", 2), ("awarded contract", 2),
        ],
        "Proposed": [
            ("proposed", 1), ("plans to build", 2), ("announced plans", 2),
            ("unveils plans", 2), ("eyes", 1), ("could build", 1),
            ("may build", 1), ("seeks approval", 2), ("targeting", 1),
        ],
        "Challenged": [
            ("cancelled", 3), ("canceled", 3), ("rejected", 2), ("denied", 2),
            ("moratorium", 3), ("blocked", 2), ("withdrawn", 2), ("lawsuit", 3),
            ("sues", 2), ("legal challenge", 3), ("opposition grows", 3),
            ("scraps", 3), ("halts", 2), ("suspended", 2),
        ],
    }
    # Fixed severity order, used only to break exact score ties.
    SEVERITY = ["Challenged", "Commissioned", "Contracted", "Financed",
                "Under Construction", "Approved", "Tendered", "Proposed"]

    scores = {}
    for status, phrases in RULES.items():
        s = 0
        for phrase, weight in phrases:
            if _re_phrase_match(t, phrase):
                s += weight
        if s:
            scores[status] = s

    if not scores:
        return "News"
    best = max(scores.values())
    candidates = [k for k, v in scores.items() if v == best]
    if len(candidates) == 1:
        return candidates[0]
    for status in SEVERITY:
        if status in candidates:
            return status
    return candidates[0]


def detect_re_sentiment(text, deal_size="", capacity=""):
    """
    True market-sentiment classifier for RE headlines — Positive / Negative /
    Neutral — distinct from (and complementary to) the lifecycle Status field
    above. Weighted lexicon, word-boundary matched, with a small boost when a
    headline carries a disclosed deal size or capacity figure (concrete,
    sourced news reads more positive in market-news framing than a vague
    mention). Returns (label, score).
    """
    t = text.lower()

    POSITIVE = [
        ("wins", 2), ("secures", 2), ("surges", 2), ("record", 2), ("boosts", 1),
        ("expands", 1), ("launches", 1), ("approved", 1), ("signs", 1),
        ("commissioned", 2), ("energized", 2), ("energised", 2), ("milestone", 1),
        ("breakthrough", 2), ("doubles", 1), ("growth", 1), ("green light", 2),
        ("completes", 1), ("inaugurated", 1), ("partners with", 1), ("backs", 1),
        ("funds", 1), ("invests", 1), ("achieves", 1), ("unveils", 1),
        ("first power", 2), ("goes live", 2),
    ]
    NEGATIVE = [
        ("delay", 2), ("delays", 2), ("delayed", 2), ("cuts", 2), ("scraps", 3),
        ("halts", 2), ("halted", 2), ("cancelled", 3), ("canceled", 3),
        ("lawsuit", 3), ("sues", 2), ("fine", 2), ("fined", 2), ("loses", 2),
        ("blocked", 2), ("opposition", 2), ("warns", 1), ("slump", 2),
        ("falls", 1), ("withdraws", 2), ("withdrawn", 2), ("suspended", 2),
        ("moratorium", 3), ("denied", 2), ("rejected", 2), ("shuts down", 3),
        ("bankrupt", 3), ("layoffs", 2), ("probe", 2), ("investigation", 2),
    ]

    pos = sum(w for phrase, w in POSITIVE if _re_phrase_match(t, phrase))
    neg = sum(w for phrase, w in NEGATIVE if _re_phrase_match(t, phrase))

    if (deal_size or capacity) and pos == neg:
        pos += 1  # concrete disclosed figures nudge ambiguous headlines positive

    score = pos - neg
    if score > 0:
        return ("Positive", score)
    if score < 0:
        return ("Negative", score)
    return ("Neutral", score)


def is_re_relevant(text):
    """Quick relevance check for renewable energy news."""
    t = text.lower()
    primary_re = [
        "solar","wind","renewable","green energy","clean energy","energy storage",
        "bess","battery storage","hydrogen","electrolyser","electrolyzer",
        "nuclear","smr","small modular reactor","ppa","power purchase agreement",
        "offshore wind","onshore wind","photovoltaic","pv park","solar farm",
        "wind farm","geothermal","hydropower","hydroelectric","biomass energy",
        "cfd","contract for difference","capacity auction","energy tender",
        "green hydrogen","fuel cell","ammonia","electrolysis","turbine",
        "solar developer","wind developer","renewables developer",
        "clean power","low carbon power","decarbonisation","decarbonization",
        "net zero power","energy transition","power transition",
        "recharge news","pv-tech","pv magazine","wind power","solar power",
        "gigawatt solar","gigawatt wind","mw solar","mw wind",
        "offshore lease","transmission line","hvdc","interconnector",
        "grid expansion","capacity market","renewable target",
    ]
    return any(p in t for p in primary_re)

# ─── PV-Tech article scraper ───────────────────────────────────────────────────
# _parse_pvtech_articles and _scrape_pvtech removed — PV Tech no longer used.



def _parse_generic_news_page(soup, source_name, base_url):
    """Generic parser for news listing pages (works for many CMS patterns)."""
    articles = []
    seen = set()
    for a in soup.find_all("a", href=True):
        href = a["href"]
        if href.startswith("/"):
            full_url = base_url.rstrip("/") + href
        elif href.startswith("http"):
            full_url = href
        else:
            continue
        if not re.search(r"/news/|/article/|/story/|/energy/|/sector/|/analysis/|/\d{4}/", href):
            continue
        norm = full_url.rstrip("/")
        if norm in seen or len(norm) < 30:
            continue
        seen.add(norm)
        h_tag = a.find(["h1","h2","h3","h4","h5"])
        hl = h_tag.get_text(" ", strip=True) if h_tag else a.get_text(" ", strip=True)
        hl = re.sub(r"\s+", " ", hl).strip()
        if not hl or len(hl) < 15:
            continue
        date_obj = None
        node = a.parent
        for _ in range(8):
            if node is None:
                break
            tt = node.find("time")
            if tt:
                date_obj = parse_date_str(tt.get("datetime","") or tt.get_text(strip=True))
                if date_obj:
                    break
            node = node.parent
        articles.append({
            "headline": hl, "url": full_url,
            "date_obj": date_obj, "source": source_name, "_priority": 2,
        })
    return articles

def _scrape_generic_re(base_url, source_name, max_pages, cutoff, progress_cb,
                       base_frac=0.0, frac_share=1.0,
                       extra_url_patterns=None):
    """
    Generic multi-page scraper used for every RE source except PV Tech.
    Handles ?page=N and /page/N pagination styles automatically.
    extra_url_patterns: additional regex to accept beyond the default set.
    """
    arts  = []
    seen  = set()
    _patterns = [
        r"/news/|/article/|/story/|/energy/|/sector/|/analysis/",
        r"/wind/|/solar/|/storage/|/hydrogen/|/nuclear/|/hydro/|/offshore/",
        r"/\d{4}/\d{2}/|/\d{4}/",
        r"/renewables/|/power/|/market/|/project/|/deal/|/ppa/|/tender/",
    ]
    if extra_url_patterns:
        _patterns.extend(extra_url_patterns)
    combined_re = re.compile("|".join(_patterns))

    for page in range(1, max_pages + 1):
        frac = base_frac + (page / max_pages) * frac_share
        progress_cb(min(frac, 1.0), f"{source_name} · Page {page}/{max_pages}")

        # Try ?page=N first, then /page/N
        if page == 1:
            url = base_url
        elif "?" in base_url:
            url = f"{base_url}&page={page}"
        else:
            # Try both patterns — use whichever works
            url = f"{base_url}?page={page}"

        soup = fetch_html(url)
        if not soup:
            # fallback: /page/N style
            if page > 1:
                url2 = base_url.rstrip("/") + f"/page/{page}/"
                soup = fetch_html(url2)
            if not soup:
                break

        new_on_page = 0
        for a in soup.find_all("a", href=True):
            href = a["href"]
            if href.startswith("/"):
                full_url = base_url.rstrip("/").rsplit("/", maxsplit=10)[0] + href
                # Safer: use scheme+netloc from base_url
                from urllib.parse import urljoin
                full_url = urljoin(base_url, href)
            elif href.startswith("http"):
                full_url = href
            else:
                continue

            if not combined_re.search(href):
                continue

            norm = full_url.rstrip("/")
            if norm in seen or len(norm) < 30:
                continue
            seen.add(norm)

            h_tag = a.find(["h1","h2","h3","h4","h5"])
            hl = h_tag.get_text(" ", strip=True) if h_tag else a.get_text(" ", strip=True)
            hl = re.sub(r"\s+", " ", hl).strip()
            if not hl or len(hl) < 15:
                continue

            # Date hunt — walk up DOM
            date_obj = None
            node = a.parent
            for _ in range(10):
                if node is None:
                    break
                tt = node.find("time")
                if tt:
                    date_obj = parse_date_str(tt.get("datetime","") or tt.get_text(strip=True))
                    if date_obj:
                        break
                # Also scan text nodes for date patterns
                txt = node.get_text(" ", strip=True)
                m = re.search(r"\b(\d{1,2})\s+(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\s+(\d{4})\b", txt, re.I)
                if m:
                    date_obj = parse_date_str(m.group(0))
                    if date_obj:
                        break
                m2 = re.search(r"\b(\d{4}-\d{2}-\d{2})\b", txt)
                if m2:
                    date_obj = parse_date_str(m2.group(1))
                    if date_obj:
                        break
                node = node.parent

            if date_obj and date_obj < cutoff:
                continue

            arts.append({
                "headline": hl, "url": full_url,
                "date_obj": date_obj, "source": source_name, "_priority": 2,
            })
            new_on_page += 1

        if new_on_page == 0:
            break
        time.sleep(0.3)

    return arts




# ─── DCD URL constants ─────────────────────────────────────────────────────
DCD_BASE             = "https://www.datacenterdynamics.com"
# Construction channel: DCD's own filtered term-based URL (proven to work)
DCD_CONSTRUCTION_URL = DCD_BASE + "/en/news/?term=the-data-center-construction-channel"
# General news: the plain news listing
DCD_GENERAL_URL      = DCD_BASE + "/en/news/"

# ─── News-type labels (used in sidebar filter) ─────────────────────────────
NEWS_TYPE_CONSTRUCTION = "Construction"
NEWS_TYPE_GENERAL      = "General News"

# REMOVED: DCD_CHAN_TERM, DCD_REGION_TERMS, DCD_EXTRA_TERMS, GNEWS_QUERIES, RSS_SOURCES
# The app now scrapes only the two DCD channels above.

_GNEWS_QUERIES_REMOVED = [
    # Construction / physical build
    ("data center construction campus groundbreaking opening", "Google News"),
    ("data center broke ground topping out opens ribbon cutting", "Google News"),
    ("data center under construction building phase expansion", "Google News"),
    # Approvals / permits / regulatory
    ("data center approved approval permit planning zoning", "Google News"),
    ("data center zoning rezoning moratorium ordinance vote hearing", "Google News"),
    ("data center project approval go-ahead green light county city", "Google News"),
    ("data center rejected denied blocked lawsuit opposition", "Google News"),
    # Site selection / disclosed / announced
    ("data center site selection disclosed announced plans", "Google News"),
    ("data center project announced new campus planned", "Google News"),
    ("data center disclosed project pipeline plans filed", "Google News"),
    # Extension / expansion
    ("data center expansion extension phase additional capacity", "Google News"),
    # Power & energy
    ("data center hyperscale investment billion megawatt gigawatt", "Google News"),
    ("data center power energy grid nuclear solar PPA", "Google News"),
    ("data center behind the meter power plant generator turbine", "Google News"),
    ("data center nuclear SMR geothermal hydrogen power", "Google News"),
    ("data center grid connection electricity capacity substation", "Google News"),
    # Investment / finance - currency-specific
    ("data center investment billion million acquisition deal", "Google News"),
    ("data center acquisition merger deal sale billion USD", "Google News"),
    ("data center REIT investment fund financing lease", "Google News"),
    ("data center IPO equity raise capital raise funding", "Google News"),
    ("data center EUR billion investment Europe campus", "Google News"),
    ("data center GBP billion investment UK campus", "Google News"),
    ("data center AED billion investment UAE Middle East", "Google News"),
    ("data center SGD billion investment Singapore Asia", "Google News"),
    ("data center INR crore billion investment India campus", "Google News"),
    # Acres / land / size signals
    ("data center acres land site development campus", "Google News"),
    ("data center hectares land parcel build-to-suit site", "Google News"),
    ("data center sq ft square feet campus facility", "Google News"),
    ("data center powered shell wholesale campus acres", "Google News"),
    # Deal flow - undisclosed / private
    ("data center joint venture partnership MOU offtake agreement", "Google News"),
    ("data center letter of intent pre-lease capacity agreement", "Google News"),
    ("data center sale leaseback forward purchase transaction", "Google News"),
    ("data center private equity infrastructure fund stake", "Google News"),
    ("data center M&A takeover acquisition stake minority", "Google News"),
    # Hyperscalers
    ("Microsoft Google Amazon Meta Oracle data center campus", "Google News"),
    ("AWS Azure GCP hyperscale cloud data center region", "Google News"),
    # Operators
    ("Equinix Digital Realty CyrusOne QTS NTT data center", "Google News"),
    ("EdgeConneX Vantage Compass Aligned DataBank data center", "Google News"),
    ("Yondr AirTrunk NextDC Macquarie atNorth data center", "Google News"),
    ("colocation datacenter AI GPU facility opens launched", "Google News"),
    # AI/GPU
    ("AI data center GPU campus hyperscale construction", "Google News"),
    ("AI factory inference training facility data center campus", "Google News"),
    # Regions
    ("data center Middle East Africa Asia Pacific expansion", "Google News"),
    ("data center Europe Germany Netherlands Ireland Frankfurt", "Google News"),
    ("data center India Singapore Malaysia Southeast Asia", "Google News"),
    ("data center Latin America Brazil Mexico Chile Argentina", "Google News"),
    # Capacity / MW / GW specific
    ("data center MW GW megawatt gigawatt capacity announcement", "Google News"),
    ("data center 100MW 200MW 500MW 1GW campus facility", "Google News"),
    ("data center kilowatt density high performance compute", "Google News"),
    # === ISO / RTO / Large Load / Grid Operator Queries ===
    ("PJM data center large load interconnection queue megawatt gigawatt", "Google News"),
    ("ERCOT data center large load study Texas megawatt queue gigawatt", "Google News"),
    ("CAISO data center interconnection California megawatt grid queue", "Google News"),
    ("MISO data center large load interconnection Midwest megawatt", "Google News"),
    ("NYISO data center interconnection New York megawatt queue", "Google News"),
    ("ISO-NE data center interconnection New England megawatt", "Google News"),
    ("SPP data center interconnection queue megawatt Southwest Plains", "Google News"),
    ("Northern Virginia data center PJM power grid megawatt queue Dominion", "Google News"),
    ("Texas ERCOT data center load growth power demand gigawatt capacity", "Google News"),
    ("Ohio PJM data center large load electricity demand megawatt", "Google News"),
    ("Georgia Southern Company data center power load electricity demand", "Google News"),
    ("Arizona APS SRP WECC data center power load capacity megawatt", "Google News"),
    ("Oregon BPA Pacific Power data center Northwest power interconnection", "Google News"),
    ("SERC Southeast data center electricity power load demand", "Google News"),
    ("TVA Tennessee Valley Authority data center large power rate", "Google News"),
    ("EirGrid Ireland data center moratorium electricity grid capacity", "Google News"),
    ("National Grid UK data center connection queue megawatt TEC holder", "Google News"),
    ("AEMO Australia data center large load connection megawatt grid", "Google News"),
    ("PowerGrid India data center electricity connection load MW", "Google News"),
    ("EMA Singapore data center electricity moratorium load MW", "Google News"),
    ("DEWA UAE data center power electricity megawatt Dubai connection", "Google News"),
    ("ENTSO-E Europe data center electricity grid demand MW capacity", "Google News"),
    ("KEPCO South Korea data center electricity grid connection MW", "Google News"),
    ("TNB Tenaga Malaysia data center electricity grid connection Johor", "Google News"),
    ("State Grid China data center electricity connection load MW", "Google News"),
    ("AEMO NEM Australia data center load growth electricity grid", "Google News"),
    ("Eskom South Africa data center electricity connection load MW", "Google News"),
    ("TEPCO Japan data center electricity grid connection MW campus", "Google News"),
    ("ANEEL Brazil data center electricity grid connection MW campus", "Google News"),
    ("CENACE Mexico data center electricity grid connection MW campus", "Google News"),
    # Large load / grid connection generic
    ("data center large load interconnection request utility grid", "Google News"),
    ("data center utility scale electricity demand grid operator ISO RTO", "Google News"),
    ("data center grid connection queue delay backlog utility", "Google News"),
    ("data center critical load electricity infrastructure utility", "Google News"),
    ("data center electricity demand forecast megawatt gigawatt load growth", "Google News"),
    ("hyperscale data center power load grid utility PPA offtake", "Google News"),
    ("data center behind meter generation electricity self-generation", "Google News"),
    ("data center nuclear SMR small modular reactor power offtake PPA", "Google News"),
    ("data center substation transformer upgrade utility connection MW", "Google News"),
    ("data center grid capacity constraint electricity supply bottleneck", "Google News"),
    ("data center power availability electricity grid interconnection study", "Google News"),
    ("data center 100MW 200MW 300MW 400MW 500MW 600MW 700MW 800MW 900MW 1GW", "Google News"),
    ("data center gigawatt campus AI factory electricity load demand", "Google News"),
    # US county/metro markets
    ("Loudoun County Prince William Fairfax data center Virginia MW", "Google News"),
    ("Hillsboro Boardman Prineville Oregon data center campus", "Google News"),
    ("Goodyear Mesa Chandler Avondale Buckeye Arizona data center campus", "Google News"),
    ("Midlothian Fort Worth Garland Plano Dallas data center Texas", "Google News"),
    ("Columbus Dublin Plain City New Albany Ohio data center campus", "Google News"),
    ("Reno Sparks Fernley Nevada data center campus megawatt", "Google News"),
    ("Quincy East Wenatchee Moses Lake Washington data center MW", "Google News"),
    ("DeKalb Carroll Paulding Bartow Georgia data center county MW", "Google News"),
    ("San Antonio Austin Round Rock Georgetown Texas data center campus", "Google News"),
    ("Chicago Elk Grove Aurora Joliet Illinois data center campus MISO", "Google News"),
    ("Omaha Council Bluffs Iowa Nebraska data center campus SPP MW", "Google News"),
    ("Richmond Henrico Chesterfield Virginia data center PJM campus", "Google News"),
    ("Boise Nampa Caldwell Idaho data center campus megawatt", "Google News"),
    ("Phoenix Tempe Scottsdale Glendale Arizona data center campus", "Google News"),
    ("Salt Lake City Lehi Draper Utah data center campus megawatt", "Google News"),
    ("Denver Aurora Centennial Colorado data center campus megawatt", "Google News"),
    ("Charlotte Concord Mooresville North Carolina data center campus", "Google News"),
    ("Atlanta Douglasville Newnan Coweta Georgia data center campus", "Google News"),
    ("Minneapolis Saint Paul Eden Prairie Minnesota data center campus", "Google News"),
    ("Kansas City Lee's Summit Lenexa Missouri Kansas data center", "Google News"),
    ("Detroit Grand Rapids Ann Arbor Michigan data center campus", "Google News"),
    ("Indianapolis Carmel Brownsburg Indiana data center campus MISO", "Google News"),
    ("Birmingham Huntsville Montgomery Alabama data center campus", "Google News"),
    ("Jacksonville Tampa Orlando Miami Florida data center campus", "Google News"),
    ("Las Vegas Henderson North Las Vegas Nevada data center campus", "Google News"),
    # International metro markets
    ("Frankfurt Rhine-Main Germany data center campus megawatt TenneT", "Google News"),
    ("Amsterdam Netherlands data center campus AMS megawatt TenneT", "Google News"),
    ("Dublin Kildare Meath Ireland data center moratorium campus megawatt", "Google News"),
    ("London Slough Reading Swindon UK data center campus megawatt National Grid", "Google News"),
    ("Singapore JTC Jurong data center campus megawatt EMA", "Google News"),
    ("Johor Iskandar Kulai Malaysia data center campus megawatt TNB", "Google News"),
    ("Jakarta Batam Surabaya Indonesia data center campus megawatt PLN", "Google News"),
    ("Mumbai Navi Mumbai Pune Hyderabad India data center campus megawatt", "Google News"),
    ("Sydney Melbourne Brisbane Perth Australia data center AEMO megawatt", "Google News"),
    ("Tokyo Osaka Nagoya Japan data center campus megawatt TEPCO", "Google News"),
    ("Seoul Incheon Busan South Korea data center KEPCO megawatt campus", "Google News"),
    ("Riyadh Jeddah NEOM Yanbu Saudi Arabia data center campus megawatt", "Google News"),
    ("Dubai Abu Dhabi Sharjah UAE DEWA data center campus megawatt", "Google News"),
    ("São Paulo Rio Curitiba Brazil data center campus megawatt ANEEL", "Google News"),
    ("Warsaw Katowice Wroclaw Poznan Poland data center campus megawatt", "Google News"),
    ("Stockholm Gothenburg Malmö Sweden data center campus megawatt", "Google News"),
    ("Helsinki Tampere Espoo Finland data center campus megawatt Fingrid", "Google News"),
    ("Oslo Bergen Stavanger Norway data center campus megawatt Statnett", "Google News"),
    ("Copenhagen Aarhus Denmark data center campus megawatt Energinet", "Google News"),
    ("Barcelona Madrid Valencia Spain data center campus megawatt REE", "Google News"),
    ("Milan Rome Turin Italy data center campus megawatt Terna", "Google News"),
    ("Brussels Antwerp Ghent Belgium data center campus megawatt Elia", "Google News"),
    ("Vienna Graz Linz Austria data center campus megawatt APG", "Google News"),
    ("Prague Brno Ostrava Czech Republic data center campus megawatt", "Google News"),
    ("Bucharest Cluj Timisoara Romania data center campus megawatt", "Google News"),
    ("Nairobi Mombasa Kenya data center campus megawatt", "Google News"),
    ("Lagos Abuja Nigeria data center campus megawatt", "Google News"),
    ("Johannesburg Cape Town Durban South Africa data center campus megawatt", "Google News"),
    ("Casablanca Rabat Morocco data center campus megawatt", "Google News"),
    ("Cairo Alexandria Egypt data center campus megawatt", "Google News"),
    ("Doha Lusail Qatar data center campus megawatt KAHRAMAA", "Google News"),
    ("Manama Bahrain data center campus megawatt", "Google News"),
    ("Muscat Oman data center campus megawatt", "Google News"),
    ("Kigali Rwanda data center campus megawatt AfricaConnect", "Google News"),
    ("Addis Ababa Ethiopia data center campus megawatt", "Google News"),
    ("Accra Ghana data center campus megawatt", "Google News"),
    ("Santiago Chile data center campus megawatt", "Google News"),
    ("Bogota Medellin Colombia data center campus megawatt", "Google News"),
    ("Buenos Aires Cordoba Argentina data center campus megawatt", "Google News"),
    ("Lima Peru data center campus megawatt", "Google News"),
    ("Mexico City Monterrey Guadalajara data center campus megawatt CENACE", "Google News"),
    # Extra keyword-rich queries  
    ("data center AI factory GPU cluster campus announcement investment", "Google News"),
    ("data center campus acre land acquisition site control option", "Google News"),
    ("data center lease pre-lease capacity agreement megawatt announcement", "Google News"),
    ("data center REIT dividend infrastructure fund deployment capital", "Google News"),
    ("data center IPO listing secondary market infrastructure equity raise", "Google News"),
]   # end of _GNEWS_QUERIES_REMOVED — not used

# ─── Date parser (app15 style — fast and clean) ────────────────────────────
MONTHS = {
    "jan":1,"feb":2,"mar":3,"apr":4,"may":5,"jun":6,
    "jul":7,"aug":8,"sep":9,"oct":10,"nov":11,"dec":12,
}

def parse_date_str(raw):
    if not raw:
        return None
    raw = str(raw).strip()
    raw = re.sub(r"\s+", " ", raw)
    # ISO with timezone
    for fmt in ("%Y-%m-%dT%H:%M:%S%z", "%Y-%m-%dT%H:%M:%SZ",
                "%a, %d %b %Y %H:%M:%S %z", "%a, %d %b %Y %H:%M:%S %Z",
                "%Y-%m-%d", "%d %b %Y", "%B %d, %Y", "%b %d, %Y"):
        try:
            dt = datetime.strptime(raw[:25], fmt[:len(raw[:25])])
            return dt.replace(tzinfo=None)
        except Exception:
            pass
    m = re.search(r"(\d{1,2})\s+(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\w*\s+(\d{4})", raw, re.I)
    if m:
        try:
            return datetime(int(m.group(3)), MONTHS[m.group(2).lower()[:3]], int(m.group(1)))
        except Exception:
            pass
    m = re.search(r"(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\w*\s+(\d{1,2}),?\s+(\d{4})", raw, re.I)
    if m:
        try:
            return datetime(int(m.group(3)), MONTHS[m.group(1).lower()[:3]], int(m.group(2)))
        except Exception:
            pass
    return None


# ─── HTTP fetcher (app15 style — cloudscraper first, fallback requests) ────
_DCD_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
        "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36"
    ),
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Accept-Language": "en-US,en;q=0.9",
    "Referer": "https://www.datacenterdynamics.com/",
}

def fetch_html(url, retries=2):
    for attempt in range(retries):
        try:
            if _USE_CS:
                r = _CS.get(url, timeout=20)
            else:
                r = _CS.get(url, headers=_DCD_HEADERS, timeout=20)
            r.raise_for_status()
            return BeautifulSoup(r.text, "html.parser")
        except Exception:
            if attempt == 0:
                time.sleep(2)
    return None


# ─── Article parser — DCD HTML pages ──────────────────────────────────────
def _parse_articles_from_soup(soup, source_name, base_url):
    """
    Parses DCD listing pages.  Finds all <a href="/en/(news|analysis|opinion|
    dcd-data-center-construction-channel-news)/..."> links, extracts headline
    from h1–h5 inside the link, then climbs the DOM for a date.
    Works for both the general news and construction-channel URLs.
    """
    articles = []
    seen = set()

    for a in soup.find_all("a", href=re.compile(
        r"^/en/(news|analysis|opinion|dcd-data-center-construction-channel-news)/[^?#]+/$"
    )):
        href = a["href"]
        if href in seen:
            continue
        seen.add(href)

        h_tag    = a.find(["h1", "h2", "h3", "h4", "h5"])
        headline = h_tag.get_text(" ", strip=True) if h_tag else a.get_text(" ", strip=True)
        headline = re.sub(r"\s+", " ", headline).strip()
        if not headline or len(headline) < 10:
            continue

        date_obj = None
        node = a.parent
        for _ in range(12):
            if node is None:
                break
            # 1. <time datetime="...">  — most reliable on DCD
            time_tag = node.find("time")
            if time_tag:
                dt_attr = time_tag.get("datetime", "")
                if dt_attr:
                    date_obj = parse_date_str(dt_attr)
                if not date_obj:
                    date_obj = parse_date_str(time_tag.get_text(strip=True))
                if date_obj:
                    break
            txt = node.get_text(" ", strip=True)
            # 2. "dd Mon YYYY"
            m = re.search(
                r"\b(\d{1,2})\s+(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+(\d{4})\b",
                txt, re.I,
            )
            if m:
                date_obj = parse_date_str(m.group(0))
                break
            # 3. "Mon dd, YYYY"
            m2 = re.search(
                r"\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\s+(\d{1,2}),?\s+(\d{4})\b",
                txt, re.I,
            )
            if m2:
                date_obj = parse_date_str(m2.group(0))
                break
            # 4. ISO yyyy-mm-dd
            m3 = re.search(r"\b(\d{4}-\d{2}-\d{2})\b", txt)
            if m3:
                date_obj = parse_date_str(m3.group(1))
                break
            node = node.parent

        articles.append({
            "headline": headline,
            "url":      base_url + href,
            "date_obj": date_obj,
            "source":   source_name,
            "_priority": 1,
        })
    return articles


# ─── DCD scraper: scrapes a single base URL across N pages ─────────────────
def _scrape_dcd_channel(base_url, source_name, cutoff, max_pages, progress_cb, label="DCD"):
    """
    Generic DCD channel scraper.  Paginates base_url with &page=N (or ?page=N).

    Bug-fix (v2): Previously the scraper stopped as soon as it encountered ONE article
    older than the cutoff, which caused "Past 7 days" / "Past 14 days" to return zero
    results because DCD pages mix dated and undated articles.  The new logic:
      - Collects ALL articles on every page regardless of date.
      - Tracks a rolling count of consecutive DATED articles that are all older than
        cutoff; only stops pagination once a full page of dated articles are all old.
      - At the end, filters to articles >= cutoff (undated articles are kept only when
        cutoff is datetime.min, i.e. "Latest (all)" mode; for short ranges they are
        excluded so users see only confirmed-fresh content).

    base_url    : channel root — may already contain query params (e.g. ?term=…)
    source_name : label stored in article["source"]
    cutoff      : datetime — articles older than this are dropped
    max_pages   : upper page limit
    progress_cb : callable(fraction 0→1, text)
    """
    fetch_html(DCD_BASE + "/en/")   # warm-up (helps bypass Cloudflare)

    all_articles = []
    seen_urls    = set()
    _is_open_range = (cutoff <= datetime.min + timedelta(days=1))

    # Work out whether base_url already has a query string
    _has_qs = "?" in base_url

    for page in range(1, max_pages + 1):
        if page == 1:
            url = base_url
        else:
            sep = "&" if _has_qs else "?"
            url = f"{base_url}{sep}page={page}"

        progress_cb(
            min(page / max_pages, 1.0),
            f"⚡ Scanning {label} · Page {page}/{max_pages}",
        )

        soup = fetch_html(url)
        if not soup:
            break

        page_arts   = _parse_articles_from_soup(soup, source_name, DCD_BASE)
        new_on_page = 0
        dated_on_page = 0
        all_dated_old = True   # becomes False if any dated article on page is within cutoff

        for art in page_arts:
            norm_url = art["url"].rstrip("/")
            if norm_url in seen_urls:
                continue
            seen_urls.add(norm_url)

            d = art["date_obj"]
            if d:
                dated_on_page += 1
                if d >= cutoff:
                    all_dated_old = False   # at least one fresh article on this page

            all_articles.append(art)
            new_on_page += 1

        # Stop paginating only when the entire page's dated articles are older than
        # the cutoff AND there was at least one dated article to judge by.
        # This prevents early abort due to undated articles.
        if new_on_page == 0:
            break
        if dated_on_page > 0 and all_dated_old and not _is_open_range:
            break

        time.sleep(0.4)

    # ── Post-scrape date filter ───────────────────────────────────────────────
    # For open-range ("Latest / all"), keep everything including undated articles.
    # For bounded ranges (7/14/30 days, custom), drop articles older than cutoff
    # AND drop undated articles (we cannot verify they are within the window).
    if _is_open_range:
        return all_articles
    else:
        return [
            a for a in all_articles
            if a.get("date_obj") is not None and a["date_obj"] >= cutoff
        ]


# ─── run_all_scrapers: dispatches to the chosen DCD channel(s) ─────────────
def run_all_scrapers(max_html_pages, cutoff, progress_cb,
                     news_types=None, region_terms=None):
    """
    news_types : list containing any of ["Construction", "General News"].
                 Pass [] or None to scrape both channels.
    region_terms : ignored (kept for API compatibility) — filtering by
                   region/country/company is done post-scrape in main().
    """
    if not news_types:
        news_types = [NEWS_TYPE_CONSTRUCTION, NEWS_TYPE_GENERAL]

    raw = []

    # ── Construction channel ──────────────────────────────────────────────
    if NEWS_TYPE_CONSTRUCTION in news_types:
        def _cb_c(frac, text=""):
            progress_cb(frac * (0.5 if NEWS_TYPE_GENERAL in news_types else 1.0), text)

        arts = _scrape_dcd_channel(
            DCD_CONSTRUCTION_URL, "DCD Construction",
            cutoff, max_html_pages, _cb_c, label="Construction Channel",
        )
        raw.extend(arts)
        progress_cb(
            0.5 if NEWS_TYPE_GENERAL in news_types else 1.0,
            f"Construction channel: {len(arts)} articles fetched",
        )

    # ── General news channel ──────────────────────────────────────────────
    if NEWS_TYPE_GENERAL in news_types:
        base_frac = 0.5 if NEWS_TYPE_CONSTRUCTION in news_types else 0.0

        def _cb_g(frac, text=""):
            progress_cb(base_frac + frac * (1.0 - base_frac), text)

        arts = _scrape_dcd_channel(
            DCD_GENERAL_URL, "DCD General News",
            cutoff, max_html_pages, _cb_g, label="General News",
        )
        raw.extend(arts)
        progress_cb(1.0, f"General news: {len(arts)} articles fetched")

    return raw


# ─── SCRAPE_SOURCES: displayed in the banner "N Sources Active" ─────────────
SCRAPE_SOURCES = [
    {"name": "DCD Construction Channel",
     "url":  DCD_CONSTRUCTION_URL, "type": "html", "priority": 1},
    {"name": "DCD General News",
     "url":  DCD_GENERAL_URL,      "type": "html", "priority": 1},
]

def is_dc_relevant(text):
    t = text.lower()
    # Primary: any of these alone = relevant
    primary = [
        "data center", "datacenter", "data centre", "datacentre",
        "colocation", "colo ", "hyperscale", "cloud campus",
        "server farm", "computing campus", "ai campus", "gpu cluster",
        "compute campus", "hpc facility", "edge facility",
        "carrier hotel", "internet exchange", "ix facility",
        "infrastructure reit", "digital infrastructure",
        # Additional primary terms
        "data hall", "data park", "digital campus",
        "compute facility", "cloud facility", "network facility",
        "ai factory", "inference facility", "training facility",
        "wholesale data", "retail colocation", "powered shell",
        "build-to-suit", "mission critical", "critical facility",
        # ISO / RTO / Grid operators (US)
        "pjm","ercot","caiso","miso","nyiso","iso-ne","spp","wecc","serc","nerc","bpa","tva",
        "pjm interconnection","ercot queue","caiso queue","miso queue","nyiso queue",
        "spp queue","iso-ne queue","large load study","grid interconnection",
        "generation interconnection","interconnection queue","load growth study",
        "transmission capacity","grid capacity","grid constraint","power availability",
        "critical load facility","large load","bulk power","bulk electric",
        # ISO / RTO / Grid operators (Global)
        "national grid eso","eirgrid","aemo","powergrid india","kepco","tennet",
        "50hertz","amprion","elia","rete","fingrid","statnett","energinet",
        "svenska kraftnät","transpower","dewa","kahramaa","taqa","sec saudi",
        "aneel","cenace","eskom","state grid china","tnb","pln","tepco","ema singapore",
        "entso-e","rte france","transnetbw","swissgrid","apg austria","pse poland",
        "bc hydro","ieso","aeso","hydro-québec",
        # Data campus / AI infrastructure
        "ai campus","ai factory","ai infrastructure","ai data center","ai datacenter",
        "gpu cluster","gpu farm","gpu factory","inference cluster","training cluster",
        "compute cluster","hpc cluster","exascale","petaflop","tera ops",
        "smart campus","tech campus","innovation campus","cloud campus","digital park",
        "data park","digital hub","tech hub","connectivity hub","internet hub",
        "network operations center","noc facility","tier iii","tier iv","tier 3","tier 4",
        "uptime certified","uptime institute","carrier neutral",
        # Power / energy infrastructure
        "behind-the-meter","behind meter","BTM generation","self-generation",
        "small modular reactor","smr","nuclear power purchase","nuclear offtake",
        "geothermal power","fuel cell power","distributed generation",
        "microGrid","battery energy storage","bess","backup power",
        "critical power","uninterruptible power","ups system",
        "cooling system","cooling tower","chilled water","adiabatic cooling",
        "free cooling","economizer","heat exchanger","liquid cooling",
        "immersion cooling","direct liquid cooling","dlc","rear-door heat exchanger",
        "pue","wue","cue","dcu","power usage effectiveness",
        "water usage effectiveness","carbon usage effectiveness",
        # Construction / development
        "powered shell","dark fiber","conduit","fiber backbone",
        "fit out","fitout","white space","raised floor","hot aisle","cold aisle",
        "modular data","prefab module","container data","pod deployment",
        "data center shell","speculative build","spec build",
        # Financial / investment
        "infrastructure reit","data center reit","net lease","triple net",
        "sale leaseback","forward purchase","equity raise","debt raise",
        "credit facility","green bond","sustainability bond","infrastructure bond",
        "co-investment","club deal","preferred equity","mezzanine financing",
        "project finance","data center fund","digital infrastructure fund",
    ]
    # Secondary: two or more = relevant
    secondary = [
        "megawatt", " mw ", " gw ", "gigawatt",
        "computing facility", "edge computing",
        "power purchase agreement", " ppa ", "behind the meter",
        "grid connection", "critical load", "raised floor",
        "cooling tower", "liquid cooling", "immersion cooling",
        "diesel generator", "ups system", "modular data",
        "tier iii", "tier iv", "uptime institute",
        "network access point", "internet hub",
        "rack space", "co-location", "hosting facility",
        "blade server", "server deployment", "ai infrastructure",
        # Currency / financial signals
        "$", "€", "£", "¥", "₹", "billion", "million",
        "usd", "eur", "gbp", "jpy", "inr", "sgd", "aed",
        "investment", "financing", "acquisition", "deal",
        # Size / land signals
        "acres", "acre", "hectares", "hectare", "sq ft", "square feet",
        "square meters", "sq m", "campus site", "land parcel",
        # Power / capacity signals
        "kilowatt", " kw ", "kwh", "mwh", "gwh",
        "substation", "transformer", "generator set",
        "ups capacity", "power density",
        # Construction / development signals
        "breaking ground", "ground breaking", "ribbon cutting",
        "topping off", "commissioning", "fit-out", "fitout",
        "shell and core", "white space", "raised floor space",
        # Deal / financial terms
        "sale leaseback", "forward purchase", "joint venture",
        "mou", "memorandum of understanding", "loi", "letter of intent",
        "offtake agreement", "capacity agreement", "pre-lease",
        # Operator / technology signals
        "crac unit", "adiabatic cooling", "free cooling",
        "power usage effectiveness", "water usage effectiveness",
        "pue", "wue", "dcu", "noc", "soc",
        # Currency / financial signals (extended)
        "cad","aud","nzd","chf","sek","nok","dkk","krw","thb","myr","idr",
        "try","zar","ngn","kes","egp","sar","qar","bhd","kwd","omr",
        "brl","cop","clp","mxn","pen","ars",
        "trillion","crore","lakh","million dollars","billion dollars",
        "mn usd","bn usd","mn eur","bn eur","mn gbp","bn gbp",
        "mn sgd","bn sgd","mn inr","bn inr","mn aed","bn aed",
        "investment round","series a","series b","series c","growth equity",
        "institutional investor","sovereign wealth fund","swf","pension fund",
        "infrastructure equity","private infrastructure","privatization",
        # US state abbreviations relevant for DC markets
        "va.","tx.","ca.","ga.","oh.","il.","in.","nv.","az.","or.",
        "wa.","nc.","pa.","nj.","md.","fl.","co.","ut.","wy.","id.",
        "mn.","wi.","mi.","ia.","ks.","mo.","ok.","ar.","la.",
        # US county / DC hub shorthand
        "nova","northern va","loudoun","prince william","fairfax county",
        "silicon slopes","silicon hills","silicon desert","silicon plains",
        # ISO / RTO abbreviations (secondary-level match)
        "rto","iso grid","independent system operator","regional transmission",
        "transmission operator","distribution operator","grid operator",
        "interconnection study","generator interconnection","large load study",
        "load serving entity","load zone","balancing authority",
        "capacity market","energy market","ancillary services",
        "demand response","virtual power plant","grid storage",
        "4cp peak","coincident peak","non-coincident peak","load factor",
        # Capacity / size signals (extended)
        "mw it","mw critical","mw white space","mw raised floor",
        "mw total","mw phase","mwac","mwdc",
        "rack","rack unit","rack space","kw per rack","kw density","mw density",
        "colocation space","retail space","wholesale space",
        "sq m","sqm","m²","ft²","sqft","square metre","square meter",
        "campus site","campus area","land area","land holding",
        "parcel","plot","lot","site area","gross floor area","gfa",
        "net leasable area","nla","total floor area","tfa",
        # Sustainability / ESG signals (extended)
        "scope 1","scope 2","scope 3","net zero","carbon neutral",
        "carbon credit","carbon offset","renewable energy certificate","rec",
        "i-rec","go","guarantees of origin","24/7 cfe",
        "ppas","virtual ppa","vppa","capacity contract","long-term contract",
        "esg rating","esg disclosure","tcfd","sbti","science based targets",
        "water positive","water neutral","zero waste","circular economy",
        "heat reuse","waste heat recovery","district heating",
        # Construction signals (extended)
        "planning application","planning approval","planning permission",
        "building permit","building consent","development permit",
        "site plan","schematic design","design development","cd phase",
        "gmp","guaranteed maximum price","leed","leed certified",
        "breeam","energy star","green star","nabers","well building",
        "design-build","epc contract","turnkey contract","lump sum",
        "general contractor","gc","subcontractor","trade contractor",
        "civil works","structural steel","mechanical electrical plumbing",
        "mep","commissioning","acceptance testing","punch list",
        "certificate of occupancy","co","temporary co","tco",
        # Miscellaneous DC-market signals
        "colo","coloc","multi-tenant","single-tenant","dedicated facility",
        "enterprise data center","edge node","far-edge","near-edge",
        "5g edge","mobile edge","content delivery","cdn node",
        "internet exchange point","ixp","peering","transit","dark fiber",
        "submarine cable","landing station","cable landing","terrestrial fiber",
        "disaster recovery","dr site","business continuity","bc dr",
        "cloud on-ramp","direct connect","expressroute","private connect",
        "cloud region","availability zone","az","region launch",
    ]
    if any(p in t for p in primary):
        return True
    if sum(1 for s in secondary if s in t) >= 1:   # lowered threshold to 1 for secondary
        return True
    return False


def detect_country(text):
    for country, patterns in COUNTRY_KEYWORDS.items():
        for pat in patterns:
            if re.search(pat if pat.startswith(r"\b") else r"\b" + re.escape(pat) + r"\b", text, re.I):
                return country
    return "Global"


def detect_topic(text):
    t = text.lower()
    for topic, kws in TOPIC_KEYWORDS.items():
        if any(k.lower() in t for k in kws):
            return topic
    return "General"


def detect_energy_mwh(text):
    """
    Detect energy-storage capacity (GWh / MWh / kWh) and normalise to MWh,
    so "1 GWh", "1,000MWh" and "1,000 MWh" (embedded in a longer "250 MW/1,000
    MWh" phrase) all compare equal regardless of which unit the outlet used.
    """
    m = re.search(r"([\d,]+(?:\.\d+)?)\s*(GWh|MWh|kWh)\b", text, re.I)
    if not m:
        return ""
    try:
        val = float(m.group(1).replace(",", ""))
    except ValueError:
        return ""
    unit = m.group(2).lower()
    if unit == "gwh":
        val *= 1000
    elif unit == "kwh":
        val /= 1000
    return str(round(val))


def detect_mw(text):
    # MW/GW capacity
    m = re.search(r"([\d,]+(?:\.\d+)?)\s*(GW|MW|gigawatt|megawatt|kilowatt|kw)\b", text, re.I)
    if m:
        val = m.group(1).replace(",", "")
        unit = m.group(2).upper()
        if unit in ("KILOWATT", "KW"):
            unit = "kW"
        return val + " " + unit
    # Acres / hectares as capacity proxy
    m2 = re.search(r"([\d,]+(?:\.\d+)?)\s*(acres?|hectares?)", text, re.I)
    if m2:
        return m2.group(1).replace(",", "") + " " + m2.group(2).lower()
    return ""


def detect_deal_size(text):
    # Try multi-currency: $, €, £, ¥, ₹, AED, SGD, etc.
    m = re.search(
        r"(\$|€|£|¥|₹|US\$|USD|EUR|GBP|AED|SGD|INR|JPY|AUD|CAD|BRL)\s*([\d,.]+)\s*-?\s*(billion|bn|million|mn|m\b|crore|lakh)",
        text, re.I
    )
    if m:
        sym = m.group(1).strip()
        val = m.group(2).replace(",", "")
        unit = m.group(3).lower()
        # Normalise symbol
        sym_map = {"US$": "$", "USD": "$", "EUR": "€", "GBP": "£", "AED": "AED ", "SGD": "SGD ",
                   "INR": "₹", "JPY": "¥", "AUD": "A$", "CAD": "C$", "BRL": "R$"}
        sym = sym_map.get(sym.upper(), sym)
        if unit in ("billion", "bn"):
            return f"{sym}{val}bn"
        if unit in ("crore",):
            return f"{sym}{val} Cr"
        return f"{sym}{val}m"
    # Fallback: number + unit without symbol
    m2 = re.search(r"([\d,.]+)\s*(billion|bn)\s+(dollar|euro|pound|dirham|rupee|yen)", text, re.I)
    if m2:
        return f"{m2.group(1).replace(',', '')}bn"
    return ""


def detect_companies(text):
    found = []
    for co in KNOWN_COMPANIES:
        if re.search(r"\b" + re.escape(co) + r"\b", text, re.I):
            found.append(co)
    return ", ".join(found[:4]) if found else ""


def detect_sentiment(text):
    t = text.lower()
    if any(w in t for w in ["broke ground", "groundbreaking", "opens", "opened",
                             "inaugurated", "energizes", "goes live", "launches"]):
        return "Opened / Live"
    if any(w in t for w in ["approved", "approval", "go-ahead", "green light",
                             "permits", "zoning approved", "rezoning"]):
        return "Approved"
    if any(w in t for w in ["proposed", "plans", "eyes", "looks to", "could build",
                             "may build", "files for", "announces plans"]):
        return "Proposed"
    if any(w in t for w in ["rejected", "denied", "moratorium", "blocked",
                             "lawsuit", "sues", "opposition", "withdrawn"]):
        return "Challenged"
    if any(w in t for w in ["under construction", "construction begins",
                             "construction started", "building"]):
        return "Under Construction"
    return "News"


def _normalise_headline(h):
    """
    Normalise headline for comparison: lowercase, strip punctuation/source suffix,
    remove filler words, collapse whitespace.
    """
    h = h.lower().strip()
    # Strip common source suffixes added by Google News: "– Reuters", "| Bloomberg", etc.
    h = re.sub(r"\s*[-–—|·]\s*[\w][\w\s,\.]{0,40}$", "", h)
    # Strip URLs and HTML entities
    h = re.sub(r"https?://\S+", "", h)
    h = re.sub(r"&\w+;", " ", h)
    # Canonicalise currency/magnitude phrasing so "US$3.5 billion", "USD 3.5bn",
    # "$3.5bn" and "$3.5-Billion" all collapse to the same token before the
    # decimal point and surrounding punctuation get stripped below.
    h = re.sub(r"\b(?:us\$|usd|eur|gbp|aed|sgd|inr|jpy|aud|cad|brl)\s*", "$", h)
    h = re.sub(r"(\d)\s*,\s*(\d{3})", r"\1\2", h)               # 3,500 -> 3500
    h = re.sub(r"(\d)\.(\d)", r"\1DECPT\2", h)                  # protect decimal point
    h = re.sub(r"[\$€£]\s*([\d,]+(?:DECPT\d+)?)\s*-?\s*(billion|bn)\b", r"\1b", h)
    h = re.sub(r"[\$€£]\s*([\d,]+(?:DECPT\d+)?)\s*-?\s*(million|mn)\b", r"\1m", h)
    h = re.sub(r"\b([\d,]+(?:DECPT\d+)?)\s*-?\s*(billion)\b", r"\1b", h)
    h = re.sub(r"\b([\d,]+(?:DECPT\d+)?)\s*-?\s*(million)\b", r"\1m", h)
    h = re.sub(r"\bbn\b", "b", h)
    h = re.sub(r"\bmn\b", "m", h)
    # Strip all non-word chars except spaces
    h = re.sub(r"[^\w\s]", " ", h)
    h = re.sub(r"\s+", " ", h).strip()
    return h


def _token_set_ratio(a, b):
    """
    Token-set similarity: splits both strings into sorted word-sets, then
    computes SequenceMatcher on the intersection + remainders.
    Much more robust than raw ratio for reordered / partially-matching headlines.
    """
    sa = set(a.split())
    sb = set(b.split())
    inter = " ".join(sorted(sa & sb))
    ra = " ".join(sorted(sa - sb))
    rb = " ".join(sorted(sb - sa))
    s1 = SequenceMatcher(None, inter, inter + " " + ra).ratio()
    s2 = SequenceMatcher(None, inter, inter + " " + rb).ratio()
    s3 = SequenceMatcher(None, inter + " " + ra, inter + " " + rb).ratio()
    return max(s1, s2, s3)


def fuzzy_similar(a, b, threshold=0.85):
    """
    True if two normalised headlines are likely the same story.
    Uses both SequenceMatcher ratio AND token-set ratio for best coverage.
    """
    na, nb = _normalise_headline(a), _normalise_headline(b)
    # Exact match after normalisation
    if na == nb:
        return True
    # Must share at least 4 tokens to be candidates (avoid short false-positives)
    ta, tb = set(na.split()), set(nb.split())
    # Skip if they share fewer than 3 common meaningful words
    common = ta & tb - {"the","a","an","in","of","at","to","for","and","or","is","on"}
    if len(common) < 3 and (len(ta) > 4 or len(tb) > 4):
        return False
    # Sequence similarity
    ratio = SequenceMatcher(None, na, nb).ratio()
    if ratio >= threshold:
        return True
    # Token-set ratio (handles re-ordered words, partial matches)
    tsr = _token_set_ratio(na, nb)
    if tsr >= threshold:
        return True
    # One is a clear substring of the other
    shorter, longer = (na, nb) if len(na) <= len(nb) else (nb, na)
    if len(shorter) >= 30 and shorter in longer:
        return True
    return False


def _canonical_url(url):
    """
    Normalise a URL for exact-match dedup:
    - Strip trailing slash
    - Remove ALL query params that are tracking noise (utm_*, fbclid, ref, etc.)
    - Lowercase scheme+host, preserve path case
    - Strip fragment (#...)
    """
    url = str(url).strip()
    # Remove fragment
    url = url.split("#")[0]
    # Remove common tracking params
    url = re.sub(r"[?&](utm_[^&]*|fbclid=[^&]*|ref=[^&]*|source=[^&]*|campaign=[^&]*|medium=[^&]*)", "", url)
    url = re.sub(r"[?&]$", "", url)
    # Normalise scheme+host to lowercase
    m = re.match(r"(https?://)([^/]+)(.*)", url)
    if m:
        url = m.group(1) + m.group(2).lower() + m.group(3)
    return url.rstrip("/")


def _url_slug(url):
    """
    Extract the slug (last meaningful path segment) from a URL,
    for cross-source same-story detection.
    e.g. https://pv-tech.org/2025/06/10/vestas-wins-200mw-india/ → vestas-wins-200mw-india
    """
    try:
        path = url.rstrip("/").rsplit("/", 1)[-1]
        # Strip common file extensions
        path = re.sub(r"\.(html?|php|aspx?)$", "", path, flags=re.I)
        return path.lower() if len(path) > 10 else ""
    except Exception:
        return ""


def deduplicate(articles, is_renewables=False):
    """
    Zero-tolerance multi-pass deduplication — eliminates ALL duplicates.

    Pass 1 — Canonical URL dedup
        Strip tracking params, normalise scheme+host, compare exact URLs.
        Best-priority source wins when URLs collide.

    Pass 2 — URL slug dedup
        Different domains can carry the same story under identical slugs
        (wire syndication pattern). Match on slug + headline prefix.

    Pass 3 — Exact normalised headline dedup
        After stripping source suffixes, punctuation, and stop-words,
        identical normalised headlines are merged.

    Pass 4 — Aggressive fuzzy headline dedup (SequenceMatcher + token-set)
        Catches reworded, reordered, or truncated variants of the same story.
        Threshold is deliberately tight (0.82) to prefer false negatives
        over false positives; also catches substring containment.

    Pass 5 — N-gram fingerprint dedup
        Build a 3-gram fingerprint of each headline; articles sharing ≥60%
        of their top-10 trigrams are treated as duplicates regardless of
        word order or synonym substitution.

    When two articles are merged, the one with the lower _priority value
    (i.e. higher-quality source) is kept; weight is used as tiebreaker.
    """

    if not articles:
        return []

    # Helper: get headline regardless of key name
    def _hl(art):
        return art.get("Headline", art.get("headline", "")) or ""

    def _pri(art):
        return (art.get("_priority", 99), -(art.get("_weight", 5)))

    # Sort: best source first so it wins every merge
    articles = sorted(articles, key=_pri)

    # ── Pass 1: Canonical URL exact dedup ────────────────────────────────────
    seen_curl = {}
    for art in articles:
        raw_url = str(art.get("URL", art.get("url", ""))).strip()
        curl = _canonical_url(raw_url)
        if not curl or len(curl) < 12:
            # URL-less: carry through, handle in headline passes
            seen_curl.setdefault("__no_url__" + _normalise_headline(_hl(art))[:60], art)
        elif curl not in seen_curl:
            seen_curl[curl] = art
    p1 = list(seen_curl.values())

    # ── Pass 2: URL slug dedup (cross-domain syndication) ────────────────────
    seen_slug = {}
    p2 = []
    for art in p1:
        raw_url = str(art.get("URL", art.get("url", ""))).strip()
        slug = _url_slug(raw_url)
        hl_prefix = _normalise_headline(_hl(art))[:40]
        slug_key = slug if slug else None
        if slug_key and slug_key in seen_slug:
            # Same slug seen — compare headline prefix to confirm same story
            existing_prefix = _normalise_headline(_hl(seen_slug[slug_key]))[:40]
            if SequenceMatcher(None, hl_prefix, existing_prefix).ratio() > 0.70:
                continue  # duplicate
        if slug_key:
            seen_slug[slug_key] = art
        p2.append(art)

    # ── Pass 3: Exact normalised headline dedup ───────────────────────────────
    seen_norm_exact = {}
    p3 = []
    for art in p2:
        norm = _normalise_headline(_hl(art))
        if norm not in seen_norm_exact:
            seen_norm_exact[norm] = art
            p3.append(art)

    # ── Pass 3.5: Weighted fact-signature dedup (RENEWABLES ONLY) ────────────
    # Headlines covering the same press-release event are often worded totally
    # differently across outlets ("US$3.5 billion" vs "USD 3.5bn", "1 GWh" vs
    # "1,000MWh", "Cypress Creek" vs "Cypress Creek Renewables"), so pure
    # text-similarity passes (3/4/5) can miss them. This pass instead builds a
    # structured fact-signature per article — companies (with alias
    # resolution), generic proper-noun entities, money, power (MW), energy
    # (MWh), and US state/location — and scores pairwise overlap. Two
    # articles merge once their combined evidence crosses a threshold, gated
    # by a same-week date window to avoid merging two genuinely different
    # stories that happen to share a company and a round number months apart.
    #
    # This pass is intentionally gated to is_renewables=True only — the Data
    # Center section keeps its original URL/headline-only dedup behaviour.
    p3b = p3
    if is_renewables:
        _FINGERPRINT_COMPANIES = set(KNOWN_COMPANIES) | set(EPC_COMPANIES.keys())
        # Canonical alias map: different outlets/sources refer to the same
        # company by different names/abbreviations. Extend this list whenever
        # a new alias pattern is spotted — no other code needs to change.
        _COMPANY_ALIASES = {
            "cypress creek renewables": "cypress creek",
            "srp": "salt river project",
            "salt river project": "salt river project",
            "edf renewables north america": "edf renewables",
            "avangrid renewables": "avangrid",
            "nextera energy resources": "nextera energy",
            "duke energy renewables": "duke energy",
            "swinerton renewable energy": "swinerton",
        }
        _GENERIC_ENTITY_STOP = {
            "solar","wind","energy","power","storage","battery","bess","project",
            "projects","renewable","renewables","plant","farm","construction",
            "financing","deal","online","new","ppa","epc","ipo","gw","mw","gwh",
            "mwh","kwh","us","usa","uk","eu","llc","inc","co","group","news",
            "the","large","giant","major","plans","report","update","closes",
            "secures","announces","brings","launches","completes","north",
            "south","east","west","january","february","march","april","may",
            "june","july","august","september","october","november","december",
        }
        _ALL_US_STATES = {s.lower() for s in COUNTRY_STATES.get("United States", [])}

        def _canon_company(key):
            return _COMPANY_ALIASES.get(key, key)

        def _fp_companies(text):
            keys = set()
            for co in _FINGERPRINT_COMPANIES:
                words = co.split()
                key = " ".join(words[:2]) if len(words) >= 2 else words[0]
                if re.search(r"\b" + re.escape(key) + r"\b", text, re.I):
                    keys.add(_canon_company(key.lower()))
            return frozenset(keys)

        def _fp_entities(raw_headline):
            text = re.sub(r"\s*[-–—|·]\s*[\w][\w\s,\.]{0,40}$", "", raw_headline)
            tokens = re.findall(r"\b[A-Z][a-zA-Z]{1,}\b", text)
            return frozenset(
                t.lower() for t in tokens
                if t.lower() not in _GENERIC_ENTITY_STOP and len(t) >= 3
            )

        def _fp_locations(text):
            return frozenset(s for s in _ALL_US_STATES if re.search(r"\b" + re.escape(s) + r"\b", text, re.I))

        def _fp_date(art):
            d = art.get("_date_obj")
            return d

        def _fingerprint(art):
            raw_hl = _hl(art)
            text = (raw_hl + " " + str(art.get("Summary", art.get("summary", "")) or ""))[:500]
            return {
                "companies": _fp_companies(text),
                "entities":  _fp_entities(raw_hl),
                "locations": _fp_locations(text),
                "deal":      detect_deal_size(text),
                "cap_mw":    detect_mw(text),
                "cap_mwh":   detect_energy_mwh(text),
                "date":      _fp_date(art),
            }

        DATE_WINDOW_DAYS = 7

        def _is_duplicate(a, b):
            # Hard gate: if both dates are known and more than DATE_WINDOW_DAYS
            # apart, never treat as duplicates — prevents merging two
            # genuinely separate stories about the same company months apart.
            if a["date"] and b["date"]:
                try:
                    if abs((a["date"] - b["date"]).days) > DATE_WINDOW_DAYS:
                        return False
                except TypeError:
                    pass

            company_overlap = a["companies"] & b["companies"]
            entity_overlap  = a["entities"] & b["entities"]
            numeric_match = (
                (a["deal"] and a["deal"] == b["deal"]) or
                (a["cap_mw"] and a["cap_mw"] == b["cap_mw"]) or
                (a["cap_mwh"] and a["cap_mwh"] == b["cap_mwh"])
            )

            # A shared deal/capacity figure plus ANY shared entity (known
            # company, OR ≥2 generic proper-noun tokens for companies not in
            # our lists) is strong evidence of the same underlying story.
            if numeric_match and (company_overlap or len(entity_overlap) >= 2):
                return True

            # ≥2 distinct known companies in common is rare enough on its own
            # to imply the same story even with zero shared numbers (e.g. two
            # outlets covering "Meta...RWE..." where neither states a figure).
            if len(company_overlap) >= 2:
                return True

            return False

        seen_fps = []
        p3b = []
        for art in p3:
            fp = _fingerprint(art)
            is_dup = False
            for s_fp in seen_fps:
                if _is_duplicate(fp, s_fp):
                    is_dup = True
                    break
            if is_dup:
                continue
            seen_fps.append(fp)
            p3b.append(art)

    # ── Pass 4: Fuzzy headline dedup (SequenceMatcher + token-set) ───────────
    FUZZY_THRESHOLD = 0.82   # tight — avoid false positives
    SUBSTR_MIN_LEN  = 25     # min chars for substring containment check

    seen_norms_4 = []
    seen_raws_4  = []
    p4 = []

    for art in p3b:
        hl      = _hl(art)
        hl_norm = _normalise_headline(hl)
        hl_toks = set(hl_norm.split()) - {"the","a","an","in","of","at","to","for","and","or","is","on","by","with"}
        is_dup  = False

        for i, seen_norm in enumerate(seen_norms_4):
            seen_toks = set(seen_norm.split()) - {"the","a","an","in","of","at","to","for","and","or","is","on","by","with"}
            common = hl_toks & seen_toks

            # Fast-reject: need at least 3 meaningful shared tokens for long headlines
            if len(hl_toks) > 5 and len(seen_toks) > 5 and len(common) < 3:
                continue

            # Direct SequenceMatcher ratio
            ratio = SequenceMatcher(None, hl_norm, seen_norm).ratio()
            if ratio >= FUZZY_THRESHOLD:
                is_dup = True
                break

            # Token-set ratio (order-independent)
            tsr = _token_set_ratio(hl_norm, seen_norm)
            if tsr >= FUZZY_THRESHOLD:
                is_dup = True
                break

            # Substring containment: one headline is a strict prefix of the other
            shorter, longer = (hl_norm, seen_norm) if len(hl_norm) <= len(seen_norm) else (seen_norm, hl_norm)
            if len(shorter) >= SUBSTR_MIN_LEN and shorter in longer:
                is_dup = True
                break

        if not is_dup:
            p4.append(art)
            seen_norms_4.append(hl_norm)
            seen_raws_4.append(hl)

    # ── Pass 4.5: Meaningful-word overlap + shared proper noun (RE ONLY) ─────
    # Catches same-story headlines that share no numeric anchor and mention
    # no company from our known lists, but clearly cover the same event via
    # overlapping substantive vocabulary plus at least one shared proper noun
    # — e.g. "Salzgitter Secures ... from EWE Electrolysis Plant" vs
    # "Salzgitter secures ... for low-carbon steel project in Germany", or
    # "CGN Breaks Ground at World's Largest Concentrated Solar Power Plant"
    # vs "CGN Breaks Ground on World's Largest Solar Thermal Power Station in
    # Qinghai". These fall well under the tight Pass-4 SequenceMatcher
    # threshold (as low as 0.50–0.65) because the wording diverges a lot, but
    # they share >50% of substantive words AND a distinctive proper noun
    # (Salzgitter / CGN / Zelestra), which is strong same-story evidence.
    if is_renewables:
        OVERLAP_RATIO_THRESHOLD = 0.50
        MIN_MEANINGFUL_TOKENS = 4
        DATE_WINDOW_DAYS_FUZZY = 10
        _STOP_45 = {"the","a","an","in","of","at","to","for","and","or","is","on","by","with","from"}

        seen_45 = []  # (tokens, entities, date)
        p4b = []
        for art in p4:
            hl = _hl(art)
            norm = _normalise_headline(hl)
            toks = set(norm.split()) - _STOP_45
            ents = _fp_entities(hl)
            d = art.get("_date_obj")
            is_dup = False
            for s_toks, s_ents, s_date in seen_45:
                if d and s_date:
                    try:
                        if abs((d - s_date).days) > DATE_WINDOW_DAYS_FUZZY:
                            continue
                    except TypeError:
                        pass
                if min(len(toks), len(s_toks)) < MIN_MEANINGFUL_TOKENS:
                    continue
                common = toks & s_toks
                ratio = len(common) / min(len(toks), len(s_toks))
                if ratio >= OVERLAP_RATIO_THRESHOLD and (ents & s_ents):
                    is_dup = True
                    break
            if is_dup:
                continue
            seen_45.append((toks, ents, d))
            p4b.append(art)
        p4 = p4b

    # ── Pass 5: N-gram fingerprint dedup ─────────────────────────────────────
    def _trigrams(text):
        words = text.split()
        return set(" ".join(words[i:i+3]) for i in range(max(0, len(words)-2)))

    NGRAM_OVERLAP_THRESHOLD = 0.60   # ≥60% trigram overlap = same story

    seen_ngrams = []
    p5 = []

    for art in p4:
        hl_norm = _normalise_headline(_hl(art))
        tg = _trigrams(hl_norm)

        # Only apply for headlines long enough to have meaningful trigrams
        if len(tg) < 3:
            p5.append(art)
            seen_ngrams.append(tg)
            continue

        is_dup = False
        for stg in seen_ngrams:
            if not stg:
                continue
            overlap = len(tg & stg) / max(len(tg), len(stg))
            if overlap >= NGRAM_OVERLAP_THRESHOLD:
                is_dup = True
                break

        if not is_dup:
            p5.append(art)
            seen_ngrams.append(tg)

    return p5


def enrich(raw_item):
    hl = raw_item["headline"]
    d = raw_item.get("date_obj")
    country = detect_country(hl)
    region = COUNTRY_TO_REGION.get(country, "Global")
    return {
        "Headline":  hl,
        "Date":      d.strftime("%Y-%m-%d") if d else "Unknown",
        "Source":    raw_item.get("source", "Unknown"),
        "URL":       raw_item.get("url", ""),
        "Country":   country,
        "Region":    region,
        "Topic":     detect_topic(hl),
        "Sentiment": detect_sentiment(hl),
        "Capacity":  detect_mw(hl),
        "Deal Size": detect_deal_size(hl),
        "Companies": detect_companies(hl),
        "ISO / RTO": detect_iso_rto(hl),
        "_date_obj": d,
    }


# ═══════════════════════════════════════════════════════════════════════════════
#  FREE / BUILT-IN INTELLIGENCE SUMMARISER  (no API key, no external ML libs)
# ═══════════════════════════════════════════════════════════════════════════════

_STOPWORDS = {
    "a","an","the","and","or","but","in","on","at","to","for","of","with",
    "is","are","was","were","be","been","being","have","has","had","do","does",
    "did","will","would","could","should","may","might","shall","can","that",
    "this","these","those","its","it","i","we","they","he","she","by","from",
    "as","into","about","over","up","after","before","during","data","center",
    "datacenter","centers","new","says","said","says","say","also","more","than",
}


# ═══════════════════════════════════════════════════════════════════════════════
#  PER-ARTICLE SUMMARISER  — removed (BART / HF NLP tier deleted per spec)
#  The Market Intel tab uses generate_local_summary / generate_re_summary only.
# ═══════════════════════════════════════════════════════════════════════════════




def _tfidf_scores(headlines):
    """Lightweight TF-IDF over headline tokens → sentence score dict."""
    tokenize = lambda h: [w.lower() for w in re.findall(r"[a-zA-Z]{3,}", h)
                          if w.lower() not in _STOPWORDS]
    tokenized = [tokenize(h) for h in headlines]
    # IDF
    N = len(headlines)
    df_counts = Counter()
    for toks in tokenized:
        for t in set(toks):
            df_counts[t] += 1
    idf = {t: math.log((N + 1) / (c + 1)) + 1 for t, c in df_counts.items()}
    # TF-IDF score per sentence
    scores = []
    for toks in tokenized:
        if not toks:
            scores.append(0.0)
            continue
        tf = Counter(toks)
        s = sum(tf[t] * idf.get(t, 1) for t in toks) / len(toks)
        scores.append(s)
    return scores


def generate_local_summary(df, sel_desc, date_range):
    """
    Wood Mackenzie-grade market intelligence briefing.
    Produces deeply analytical, prose-rich sections with quantified signals,
    company deep-dives, deal flow analysis, and forward-looking commentary —
    modelled on top-tier energy/infrastructure research houses.
    """
    headlines = df["Headline"].tolist()
    scores    = _tfidf_scores(headlines)
    df2 = df.copy()
    df2["_score"] = scores

    total = len(df2)
    if total == 0:
        return "No articles in the selected view. Adjust filters and regenerate."

    # ── Core aggregations ─────────────────────────────────────────────────────
    top_regions  = df2["Region"].value_counts()
    top_topics   = df2["Topic"].value_counts()
    top_sents    = df2["Sentiment"].value_counts()
    top_countries= df2["Country"].value_counts()

    mw_df    = df2[df2["Capacity"] != ""].copy()
    deal_df  = df2[df2["Deal Size"] != ""].copy()
    mw_vals  = mw_df["Capacity"].tolist()
    deal_vals= deal_df["Deal Size"].tolist()

    # Build full company counter from Companies column
    all_companies_list = []
    for v in df2["Companies"]:
        if v:
            all_companies_list.extend([c.strip() for c in str(v).split(",") if c.strip()])
    co_counter = Counter(all_companies_list)
    top_cos    = [co for co, _ in co_counter.most_common(20)]

    # Sentiment counts
    proposed   = int(top_sents.get("Proposed", 0))
    approved   = int(top_sents.get("Approved", 0))
    under_c    = int(top_sents.get("Under Construction", 0))
    opened     = int(top_sents.get("Opened / Live", 0))
    challenged = int(top_sents.get("Challenged", 0))
    news_c     = int(top_sents.get("News", 0))

    def pct(n): return f"{round(n/total*100)}%" if total else "0%"
    def hl(sub, n=3):
        """Return top n scored headlines from a sub-dataframe."""
        if sub.empty: return []
        return sub.nlargest(n, "_score")["Headline"].tolist()

    # ── Prose helpers ─────────────────────────────────────────────────────────
    def region_prose():
        parts = []
        for reg, cnt in top_regions.items():
            pct_val = pct(cnt)
            top_c   = df2[df2["Region"]==reg]["Country"].value_counts().head(4).index.tolist()
            parts.append(f"{reg} ({cnt} articles, {pct_val}) led by {', '.join(top_c)}")
        return "; ".join(parts) + "."

    def topic_prose():
        parts = []
        for top, cnt in top_topics.items():
            parts.append(f"{top} ({cnt} articles, {pct(cnt)})")
        return ", ".join(parts)

    # ── 1. EXECUTIVE SUMMARY ─────────────────────────────────────────────────
    dominant_region  = top_regions.index[0] if not top_regions.empty else "Global"
    dominant_topic   = top_topics.index[0]  if not top_topics.empty else "General"
    dominant_country = top_countries.index[0] if not top_countries.empty else "—"

    momentum_signal = (
        "strongly positive" if proposed > approved * 1.5
        else "balanced" if abs(proposed - approved) <= 2
        else "approval-constrained"
    )

    capacity_total_mw = 0
    for cap in mw_vals:
        m = re.search(r"([\d,.]+)\s*(GW|MW)", str(cap), re.I)
        if m:
            v = float(m.group(1).replace(",", ""))
            capacity_total_mw += v * 1000 if m.group(2).upper() == "GW" else v

    exec_lines = [
        f"This intelligence briefing synthesises {total} data centre industry articles "
        f"published between {date_range}, filtered to: {sel_desc}. "
        f"The analysis draws on headlines sourced from trade press, RSS feeds, and news aggregators, "
        f"auto-enriched with topic classification, project status, capacity and deal-size extraction, "
        f"and named-entity recognition across {df2['Country'].nunique()} countries.",

        f"\nActivity is geographically concentrated in {dominant_region}, with {dominant_country} "
        f"representing the single most active market in the filtered dataset. "
        f"Thematically, {dominant_topic} stories dominate at {pct(int(top_topics.iloc[0]))} of total "
        f"coverage, followed by {', '.join(str(t) for t in top_topics.index[1:3])}.",

        f"\nThe project pipeline presents a {momentum_signal} outlook: {proposed} proposed projects "
        f"against {approved} approvals, with {under_c} actively under construction and {opened} "
        f"recently commissioned. Regulatory friction accounts for {challenged} challenged or contested "
        f"articles ({pct(challenged)} of the dataset). "
        + (f"Identified capacity in announced or proposed projects totals approximately "
           f"{capacity_total_mw:,.0f} MW across {len(mw_df)} discrete announcements. "
           if capacity_total_mw > 0 else "")
        + (f"Disclosed deal flow spans {len(deal_df)} transactions including: "
           f"{'; '.join(deal_vals[:6])}{'.' if deal_vals else ''}"
           if deal_vals else ""),
    ]
    exec_summary = "\n\n".join(p.strip() for p in exec_lines if p.strip())

    # ── 2. KEY THEMES & TRENDS ────────────────────────────────────────────────
    theme_lines = []
    for topic, cnt in top_topics.items():
        sub = df2[df2["Topic"] == topic]
        ex  = hl(sub, 2)
        ex_str = ""
        if ex:
            # Truncate cleanly at word boundary
            h1 = ex[0][:110].rsplit(" ", 1)[0] if len(ex[0]) > 110 else ex[0]
            ex_str = f' Headline examples include: "{h1}"'
            if len(ex) > 1:
                h2 = ex[1][:100].rsplit(" ", 1)[0] if len(ex[1]) > 100 else ex[1]
                ex_str += f'; and "{h2}".'
            else:
                ex_str += "."
        theme_lines.append(
            f"• **{topic}** — {cnt} article{'s' if cnt > 1 else ''} ({pct(cnt)}).{ex_str}"
        )

    # ── 3. MAJOR PROJECTS & DEALS ─────────────────────────────────────────────
    proj_bullets = []
    # First: deal-size articles ranked by deal value then score
    deal_sub = deal_df.nlargest(min(8, len(deal_df)), "_score") if not deal_df.empty else pd.DataFrame()
    for _, r in deal_sub.iterrows():
        parts = [f"**{r['Headline'][:130]}**"]
        tags  = []
        if r.get("Deal Size"): tags.append(f"Deal: {r['Deal Size']}")
        if r.get("Capacity"):  tags.append(f"Capacity: {r['Capacity']}")
        if r.get("Country"):   tags.append(f"Market: {r['Country']}")
        if r.get("Companies"): tags.append(f"Companies: {r['Companies']}")
        if tags: parts.append(" · ".join(tags))
        proj_bullets.append("• " + " | ".join(parts))

    # Then: capacity-only articles
    cap_only = mw_df[mw_df["Deal Size"] == ""].nlargest(min(8, len(mw_df)), "_score") if not mw_df.empty else pd.DataFrame()
    for _, r in cap_only.iterrows():
        parts = [f"**{r['Headline'][:130]}**"]
        tags  = [f"Capacity: {r['Capacity']}", f"Market: {r['Country']}"]
        if r.get("Companies"): tags.append(f"Companies: {r['Companies']}")
        parts.append(" · ".join(tags))
        proj_bullets.append("• " + " | ".join(parts))

    if not proj_bullets:
        # Fall back to highest-scored articles
        for _, r in df2.nlargest(8, "_score").iterrows():
            proj_bullets.append(
                f"• **{r['Headline'][:130]}** | Market: {r['Country']} · Topic: {r['Topic']}"
            )

    # ── 4. REGULATORY & PERMITTING LANDSCAPE ──────────────────────────────────
    reg_df  = df2[df2["Topic"] == "Permits"]
    chall_df= df2[df2["Sentiment"] == "Challenged"]
    appr_df = df2[df2["Sentiment"] == "Approved"]

    reg_intro = (
        f"Permitting and regulatory dynamics account for {len(reg_df)} articles "
        f"({pct(len(reg_df))}) in the current selection. "
    )
    if len(chall_df) > 0:
        reg_intro += (
            f"Contested or blocked projects number {len(chall_df)}, "
            f"indicating {'elevated' if len(chall_df) > 3 else 'moderate'} community and regulatory "
            f"resistance in this geography and period. "
        )
    if len(appr_df) > 0:
        reg_intro += f"Approvals recorded: {len(appr_df)} projects cleared planning in the period. "
    if reg_df.empty and chall_df.empty:
        reg_intro += "No specific permitting friction or approval events detected in the current filter."

    reg_bullets = ["• " + h for h in hl(pd.concat([reg_df, chall_df, appr_df]).drop_duplicates(), 8)]
    if not reg_bullets:
        reg_bullets = ["• No permitting-specific articles in current selection."]

    # ── 5. POWER & INFRASTRUCTURE ─────────────────────────────────────────────
    pwr_df  = df2[df2["Topic"] == "Power"]
    pwr_intro = (
        f"Power and energy infrastructure is a recurring theme across {len(pwr_df)} articles "
        f"({pct(len(pwr_df))}), reflecting the sector's acute dependence on grid capacity, "
        f"PPAs, and alternative generation sources. "
    )
    if capacity_total_mw > 0:
        pwr_intro += (
            f"Across all capacity-cited articles, aggregate announced load totals "
            f"approximately {capacity_total_mw:,.0f} MW — a figure that underscores the scale "
            f"of power procurement challenges facing operators in this market. "
        )
    pwr_bullets = ["• " + h for h in hl(pwr_df, 7)]
    if not pwr_bullets:
        pwr_bullets = ["• No dedicated power/infrastructure articles in current selection."]

    # ── 6. COMPANY ACTIVITY & COMPETITIVE LANDSCAPE ──────────────────────────
    # Build richer per-company profiles
    co_section_lines = []

    # Group by hyperscalers vs colo/operators vs investors
    hyperscalers = {"Microsoft","Google","Amazon","AWS","Meta","Apple","Oracle","Alibaba",
                    "Tencent","ByteDance","Baidu","Alphabet","OpenAI","Stargate","CoreWeave",
                    "xAI","Anthropic","IBM","Huawei","Samsung","NVIDIA"}
    colo_ops     = {"Equinix","Digital Realty","Iron Mountain","CoreSite","CyrusOne","NTT",
                    "Vantage","Switch","EdgeConneX","Flexential","DataBank","QTS","Colt",
                    "Global Switch","ChinaData","GDS","STACK","Aligned","Compass","Tract",
                    "Yondr","Venari","Verne","NextDC","AirTrunk","atNorth","CtrlS","Yotta",
                    "STT GDC","Keppel","Singtel","Telstra","Macquarie","VIRTUS","Kao",
                    "Ascenty","Odata","Scala","Luminet","Etix","Nabiax","Bolder","DigiPlex",
                    "Bulk Infrastructure","Green Mountain","365 Data Centers","DigiCo",
                    "Cloudflare","Fastly","Akamai","Lumen","Zayo","Cogent","Pulsant","Iomart",
                    "Telehouse","Interxion","euNetworks","DE-CIX","MEVSPACE","Beyond.pl","Atman",
                    "Applied Digital","Hut 8","Core Scientific","Riot","Marathon","Iren",
                    "CyrusOne","Flexential","DataBank","T5","Skybox","Stream","CloudHQ"}

    mentioned_cos = [(co, cnt) for co, cnt in co_counter.most_common(30)]

    if mentioned_cos:
        # Narrative intro
        co_intro_parts = []
        hs_mentioned  = [(co,c) for co,c in mentioned_cos if co in hyperscalers]
        col_mentioned = [(co,c) for co,c in mentioned_cos if co in colo_ops]
        other_mentioned=[(co,c) for co,c in mentioned_cos if co not in hyperscalers and co not in colo_ops]

        if hs_mentioned:
            hs_str = ", ".join(
                f"{co} ({c} mention{'s' if c > 1 else ''})" for co, c in hs_mentioned[:5]
            )
            co_intro_parts.append(
                f"Hyperscaler activity is led by {hs_str}, "
                f"signalling continued large-scale capacity expansion in this market."
            )
        if col_mentioned:
            col_str = ", ".join(
                f"{co} ({c} mention{'s' if c > 1 else ''})" for co, c in col_mentioned[:5]
            )
            co_intro_parts.append(
                f"Among operators and colocation providers, {col_str} "
                f"feature prominently, reflecting active build, partnership, or M&A dynamics."
            )
        if other_mentioned:
            oth_str = ", ".join(
                f"{co} ({c} mention{'s' if c > 1 else ''})" for co, c in other_mentioned[:5]
            )
            co_intro_parts.append(
                f"Additional notable participants include {oth_str}, "
                f"spanning investors, developers, utilities, and technology vendors."
            )

        co_section_lines.append("\n".join(co_intro_parts))
        co_section_lines.append("")

        # Per-company detail bullets
        for co, cnt in mentioned_cos[:20]:
            co_arts = df2[
                df2["Companies"].str.contains(re.escape(co), na=False, case=False) |
                df2["Headline"].str.contains(re.escape(co), na=False, case=False)
            ]
            topics_seen    = co_arts["Topic"].value_counts().head(3).index.tolist()
            sents_seen     = co_arts["Sentiment"].value_counts().head(2).index.tolist()
            countries_seen = co_arts["Country"].value_counts().head(3).index.tolist()
            cap_arts       = co_arts[co_arts["Capacity"] != ""]["Capacity"].tolist()
            deal_arts      = co_arts[co_arts["Deal Size"] != ""]["Deal Size"].tolist()

            detail_parts = [f"focus: {', '.join(topics_seen) if topics_seen else 'General'}"]
            if countries_seen and countries_seen != ["Global"]:
                detail_parts.append(f"markets: {', '.join(countries_seen)}")
            if sents_seen:
                detail_parts.append(f"status signals: {', '.join(sents_seen)}")
            if cap_arts:
                detail_parts.append(f"capacity cited: {'; '.join(cap_arts[:3])}")
            if deal_arts:
                detail_parts.append(f"deal flow: {'; '.join(deal_arts[:2])}")

            co_section_lines.append(
                f"• **{co}** — {cnt} article{'s' if cnt > 1 else ''} | "
                + " | ".join(detail_parts)
            )
    else:
        co_section_lines.append(
            "• No named company entities detected in the current filtered view. "
            "Consider broadening the date range, region, or topic filters to surface company-level signals."
        )

    # ── 7. REGIONAL BREAKDOWN (detailed) ─────────────────────────────────────
    region_lines = []
    for region, rdf in df2.groupby("Region"):
        top_c      = rdf["Country"].value_counts().head(5).index.tolist()
        top_t      = rdf["Topic"].value_counts().head(3).index.tolist()
        cap_in_reg = rdf[rdf["Capacity"] != ""]["Capacity"].tolist()
        deals_reg  = rdf[rdf["Deal Size"] != ""]["Deal Size"].tolist()
        prop_reg   = len(rdf[rdf["Sentiment"]=="Proposed"])
        appr_reg   = len(rdf[rdf["Sentiment"]=="Approved"])
        chall_reg  = len(rdf[rdf["Sentiment"]=="Challenged"])

        region_lines.append(f"• **{region}** — {len(rdf)} articles ({pct(len(rdf))})")
        region_lines.append(
            f"  Lead markets: {', '.join(top_c)}. "
            f"Dominant themes: {', '.join(top_t)}. "
            f"Pipeline: {prop_reg} proposed, {appr_reg} approved, {chall_reg} challenged."
            + (f" Capacity cited: {'; '.join(cap_in_reg[:4])}." if cap_in_reg else "")
            + (f" Deal flow: {'; '.join(deals_reg[:3])}." if deals_reg else "")
        )

    # ── 8. MARKET OUTLOOK & FORWARD SIGNALS ───────────────────────────────────
    pipeline_signal = (
        "strongly positive — new proposal volume significantly outpaces current approvals, "
        "suggesting continued permitting and planning activity over the next 12–24 months"
        if proposed > approved * 1.5
        else "broadly balanced — approvals are broadly keeping pace with new project announcements, "
        "indicating a market in active but orderly expansion"
        if abs(proposed - approved) <= 3
        else "approval-constrained — the gap between proposals and approvals suggests permitting "
        "timelines and grid connection queues may be limiting near-term delivery"
    )

    reg_risk = (
        "elevated" if challenged > 4
        else "moderate" if challenged > 1
        else "low"
    )

    outlook_bullets = [
        f"• **Pipeline momentum:** {pipeline_signal}. "
        f"With {proposed} proposed vs {approved} approved and {under_c} under construction, "
        f"near-term delivery visibility {'is strong' if under_c >= 3 else 'remains limited'}.",

        f"• **Commissioning activity:** {opened} facility opening{'s' if opened != 1 else ''} recorded "
        f"in the period, confirming active supply additions to the market.",

        f"• **Regulatory risk:** {reg_risk.capitalize()} — {challenged} contested or challenged "
        f"project{'s' if challenged != 1 else ''} in the dataset. "
        + ("Community opposition, moratorium risk, and planning delays represent the primary "
           "near-term threat to delivery timelines."
           if challenged > 2
           else "Regulatory environment appears broadly supportive of new capacity in this geography."),

        f"• **Company watch:** {', '.join(top_cos[:6])} represent the highest-frequency participants "
        f"in this period. Monitor for further capacity announcements, M&A, and permitting outcomes.",
    ]

    if mw_vals:
        outlook_bullets.append(
            f"• **Capacity pipeline:** {len(mw_df)} articles reference explicit MW or GW figures. "
            f"Aggregate announced load stands at approximately {capacity_total_mw:,.0f} MW across "
            f"disclosed projects. Actual delivered capacity will depend on permitting, grid "
            f"connection, and financing milestones."
        )
    if deal_vals:
        outlook_bullets.append(
            f"• **Investment flow:** {len(deal_df)} deal-citing articles totalling disclosed "
            f"transactions of {', '.join(deal_vals[:8])}. "
            f"Deal density suggests active capital deployment into the sector."
        )

    # ── Assemble final document ────────────────────────────────────────────────
    doc_parts = []

    doc_parts += [
        "## 1. Executive Summary\n\n",
        exec_summary, "\n\n",
    ]
    doc_parts += [
        "## 2. Key Themes & Market Dynamics\n\n",
        f"Coverage across {total} articles spans the following thematic clusters: "
        f"{topic_prose()}.\n\n",
        "\n".join(theme_lines), "\n\n",
    ]
    doc_parts += [
        "## 3. Major Projects, Deals & Capacity Announcements\n\n",
        "\n".join(proj_bullets) if proj_bullets else "• No capacity or deal-cited articles in current selection.", "\n\n",
    ]
    doc_parts += [
        "## 4. Regulatory & Permitting Landscape\n\n",
        reg_intro + "\n\n",
        "\n".join(reg_bullets), "\n\n",
    ]
    doc_parts += [
        "## 5. Power & Infrastructure\n\n",
        pwr_intro + "\n\n",
        "\n".join(pwr_bullets), "\n\n",
    ]
    doc_parts += [
        "## 6. Company Activity & Competitive Landscape\n\n",
        "\n".join(co_section_lines), "\n\n",
    ]
    doc_parts += [
        "## 7. Regional Breakdown\n\n",
        "\n".join(region_lines), "\n\n",
    ]
    doc_parts += [
        "## 8. Market Outlook & Forward Signals\n\n",
        "\n".join(outlook_bullets), "\n",
    ]

    return "".join(doc_parts)



# ═══════════════════════════════════════════════════════════════════════════════
#  WORD (.docx) EXPORT
# ═══════════════════════════════════════════════════════════════════════════════

def _set_cell_bg(cell, hex_color):
    """Set table cell background colour."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def build_briefing_docx(summary_text, sel_desc, date_range, df,
                        title="GLOBAL DATA CENTER INTELLIGENCE BRIEFING"):
    """Return bytes of a polished Word briefing document."""
    if not _DOCX_OK:
        return None
    buf = io.BytesIO()
    doc = DocxDocument()
    # IST timestamp
    from datetime import timezone as _tz_d, timedelta as _td_d
    _IST_D = _tz_d(_td_d(hours=5, minutes=30))
    _ist_str_d = datetime.now(_IST_D).strftime("%d %b %Y, %I:%M %p IST")

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.15)
        section.right_margin  = Inches(1.15)

    # ── Cover block ──────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x00, 0x47, 0xE1)
    run.font.name = "Calibri"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sr = sub.add_run(f"Selection: {sel_desc}   |   Period: {date_range}")
    sr.font.size  = Pt(9)
    sr.font.color.rgb = RGBColor(0x6A, 0x80, 0xA8)
    sr.font.name  = "Calibri"

    dr = sub.add_run(f"\nGenerated: {_ist_str_d}   |   Articles analysed: {len(df)}")
    dr.font.size  = Pt(9)
    dr.font.color.rgb = RGBColor(0x6A, 0x80, 0xA8)
    dr.font.name  = "Calibri"

    doc.add_paragraph()  # spacer

    # ── Stats summary table ──────────────────────────────────────────────────
    stats = [
        ("Total Articles", str(len(df))),
        ("Top Region",     df["Region"].value_counts().index[0] if not df.empty else "—"),
        ("Top Topic",      df["Topic"].value_counts().index[0]  if not df.empty and "Topic" in df.columns else (df["Sector"].value_counts().index[0] if not df.empty and "Sector" in df.columns else "—")),
        ("With Capacity",  str(len(df[df["Capacity"] != ""]))),
        ("With Deal Size", str(len(df[df["Deal Size"] != ""]))),
    ]
    tbl = doc.add_table(rows=1, cols=len(stats))
    tbl.style = "Table Grid"
    hdr_cells = tbl.rows[0].cells
    for i, (label, val) in enumerate(stats):
        hdr_cells[i].text = ""
        p2 = hdr_cells[i].paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p2.add_run(label + "\n")
        r1.font.size = Pt(7.5); r1.font.bold = True
        r1.font.color.rgb = RGBColor(0xB8, 0xC8, 0xE0); r1.font.name = "Calibri"
        r2 = p2.add_run(val)
        r2.font.size = Pt(13); r2.font.bold = True
        r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r2.font.name = "Calibri"
        _set_cell_bg(hdr_cells[i], "0F1E36")

    doc.add_paragraph()

    # ── Parse and write markdown sections ────────────────────────────────────
    ACCENT = RGBColor(0x00, 0x47, 0xE1)
    BODY   = RGBColor(0x1A, 0x1A, 0x2E)
    BULLET = RGBColor(0x2A, 0x3E, 0x60)

    for line in summary_text.splitlines():
        line = line.rstrip()
        if line.startswith("## "):
            heading = line[3:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after  = Pt(4)
            run = p.add_run(heading.upper())
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = ACCENT
            run.font.name = "Calibri"
            # underline via border
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"),   "single")
            bottom.set(qn("w:sz"),    "4")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "0047E1")
            pBdr.append(bottom)
            pPr.append(pBdr)
        elif line.startswith("• "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            # Parse **bold** segments inline
            _parts = re.split(r"(\*\*.*?\*\*)", line[2:])
            for _part in _parts:
                if _part.startswith("**") and _part.endswith("**"):
                    _run = p.add_run(_part[2:-2])
                    _run.bold = True
                    _run.font.size = Pt(9.5)
                    _run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
                    _run.font.name = "Calibri"
                elif _part:
                    _run = p.add_run(_part)
                    _run.font.size = Pt(9.5)
                    _run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
                    _run.font.name = "Calibri"
        elif line.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            _bparts = re.split(r"(\*\*.*?\*\*)", line)
            for _bp in _bparts:
                if _bp.startswith("**") and _bp.endswith("**"):
                    _br = p.add_run(_bp[2:-2])
                    _br.bold = True
                    _br.font.size = Pt(10)
                    _br.font.color.rgb = BODY
                    _br.font.name = "Calibri"
                elif _bp:
                    _br = p.add_run(_bp)
                    _br.font.size = Pt(10)
                    _br.font.color.rgb = BODY
                    _br.font.name = "Calibri"

    # ── Footer note ───────────────────────────────────────────────────────────
    doc.add_paragraph()
    fp = doc.add_paragraph()
    fr = fp.add_run(
        "This briefing was generated automatically using built-in NLP analysis of "
        "publicly available news headlines. Always verify against primary sources."
    )
    fr.font.size  = Pt(8)
    fr.font.italic = True
    fr.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA8)
    fr.font.name  = "Calibri"

    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ═══════════════════════════════════════════════════════════════════════════════
#  PDF EXPORT
# ═══════════════════════════════════════════════════════════════════════════════

def build_briefing_pdf(summary_text, sel_desc, date_range, df,
                       title="GLOBAL DATA CENTER INTELLIGENCE BRIEFING"):
    """Return bytes of a styled PDF briefing document."""
    if not _PDF_OK:
        return None
    buf = io.BytesIO()
    # IST = UTC+5:30
    from datetime import timezone, timedelta as _td
    _IST = timezone(_td(hours=5, minutes=30))
    _now_ist = datetime.now(_IST)
    _ist_str = _now_ist.strftime("%d %b %Y, %I:%M %p IST")

    doc = SimpleDocTemplate(
        buf, pagesize=letter,
        leftMargin=inch * 1.1, rightMargin=inch * 1.1,
        topMargin=inch * 1.0,  bottomMargin=inch * 1.0,
    )

    NAVY  = rl_colors.HexColor("#0047E1")
    DARK  = rl_colors.HexColor("#0F1E36")
    GREY  = rl_colors.HexColor("#6A80A8")
    WHITE = rl_colors.white
    BODY  = rl_colors.HexColor("#1A1A2E")

    styles = getSampleStyleSheet()

    s_title = ParagraphStyle("DCTitle",
        fontSize=20, textColor=NAVY, fontName="Helvetica-Bold",
        spaceAfter=4, leading=24)
    s_sub = ParagraphStyle("DCSub",
        fontSize=8.5, textColor=GREY, fontName="Helvetica",
        spaceAfter=14, leading=13)
    s_head = ParagraphStyle("DCHead",
        fontSize=11, textColor=NAVY, fontName="Helvetica-Bold",
        spaceBefore=16, spaceAfter=5, leading=14,
        borderPadding=(0, 0, 3, 0))
    s_bullet = ParagraphStyle("DCBullet",
        fontSize=9.5, textColor=BODY, fontName="Helvetica",
        spaceAfter=3, leading=13, leftIndent=14,
        bulletIndent=2, bulletFontName="Helvetica", bulletFontSize=9)
    s_body = ParagraphStyle("DCBody",
        fontSize=10, textColor=BODY, fontName="Helvetica",
        spaceAfter=6, leading=14)
    s_footer = ParagraphStyle("DCFooter",
        fontSize=7.5, textColor=GREY, fontName="Helvetica-Oblique",
        spaceBefore=18, leading=11)

    story = []

    # ── Header ────────────────────────────────────────────────────────────────
    story.append(Paragraph(title, s_title))
    story.append(Paragraph(
        f"<b>Selection:</b> {sel_desc}   |   <b>Period:</b> {date_range}<br/>"
        f"Generated: {_ist_str}   |   "
        f"Articles analysed: {len(df)}",
        s_sub))
    story.append(HRFlowable(width="100%", thickness=1.5, color=NAVY, spaceAfter=10))

    # ── Stats row (manual layout) ─────────────────────────────────────────────
    stats_labels = ["Total Articles", "Top Region", "Top Topic", "w/ Capacity", "w/ Deal Size"]
    stats_vals   = [
        str(len(df)),
        df["Region"].value_counts().index[0] if not df.empty else "—",
        df["Topic"].value_counts().index[0]  if not df.empty and "Topic" in df.columns else (df["Sector"].value_counts().index[0] if not df.empty and "Sector" in df.columns else "—"),
        str(len(df[df["Capacity"] != ""])),
        str(len(df[df["Deal Size"] != ""])),
    ]
    stat_line = "   ·   ".join(f"<b>{l}:</b> {v}" for l, v in zip(stats_labels, stats_vals))
    story.append(Paragraph(stat_line, ParagraphStyle("statrow",
        fontSize=8.5, textColor=GREY, fontName="Helvetica",
        spaceAfter=14, leading=12)))
    story.append(HRFlowable(width="100%", thickness=0.5, color=GREY, spaceAfter=6))

    # ── Markdown sections ─────────────────────────────────────────────────────
    def _rl_inline(t):
        """Convert **bold** markers to ReportLab <b> tags for proper bold rendering."""
        import re as _re
        t = _re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", t)
        return t

    for line in summary_text.splitlines():
        line = line.rstrip()
        if line.startswith("## "):
            story.append(Paragraph(line[3:].upper(), s_head))
            story.append(HRFlowable(width="100%", thickness=0.8, color=NAVY, spaceAfter=4))
        elif line.startswith("• "):
            story.append(Paragraph(f"• {_rl_inline(line[2:])}", s_bullet))
        elif line.strip():
            story.append(Paragraph(_rl_inline(line), s_body))

    # ── Footer ────────────────────────────────────────────────────────────────
    story.append(HRFlowable(width="100%", thickness=0.5, color=GREY, spaceBefore=14, spaceAfter=4))
    story.append(Paragraph(
        "This briefing was generated automatically using built-in NLP analysis of "
        "publicly available news headlines. Always verify against primary sources.",
        s_footer))

    doc.build(story)
    buf.seek(0)
    return buf.read()


def build_excel(df):
    wb = Workbook()
    thin = Side(border_style="thin", color="CCCCCC")
    brd = Border(left=thin, right=thin, top=thin, bottom=thin)
    hf = Font(bold=True, color="FFFFFF", name="Calibri", size=10)
    hfill = PatternFill("solid", fgColor="0F1E36")
    ev = PatternFill("solid", fgColor="0B1628")
    od = PatternFill("solid", fgColor="060A10")
    nf = Font(name="Calibri", size=9, color="C8D4E8")
    lf = Font(name="Calibri", size=9, color="0047E1", underline="single")

    def write_sheet(ws, data_df, cols, widths):
        for ci, (col, w) in enumerate(zip(cols, widths), 1):
            c = ws.cell(1, ci, col)
            c.font = hf; c.fill = hfill
            c.alignment = Alignment(horizontal="center", vertical="center")
            c.border = brd
            ws.column_dimensions[get_column_letter(ci)].width = w
        ws.row_dimensions[1].height = 24
        tc_map = TOPIC_COLORS
        for ri, row in enumerate(data_df[cols if all(c in data_df.columns for c in cols) else data_df.columns].itertuples(index=False), start=2):
            fill = ev if ri % 2 == 0 else od
            for ci, val in enumerate(row, 1):
                v = str(val) if val is not None else ""
                c = ws.cell(ri, ci, v)
                c.fill = fill; c.border = brd
                col_name = cols[ci - 1] if ci <= len(cols) else ""
                if col_name == "URL":
                    c.hyperlink = v; c.font = lf
                    c.alignment = Alignment(horizontal="center", vertical="center")
                elif col_name == "Topic":
                    tc = tc_map.get(v, "2E4470")
                    c.font = Font(name="Calibri", size=9, bold=True, color=tc.replace("#", ""))
                    c.alignment = Alignment(horizontal="center", vertical="center")
                elif col_name == "Headline":
                    c.font = nf
                    c.alignment = Alignment(vertical="center", wrap_text=True)
                else:
                    c.font = nf
                    c.alignment = Alignment(horizontal="center", vertical="center")
        ws.freeze_panes = "A2"
        ws.auto_filter.ref = f"A1:{get_column_letter(len(cols))}{len(data_df)+1}"

    ws1 = wb.active
    ws1.title = "All Articles"
    main_cols = ["Headline","Date","Source","Country","Region","Topic",
                 "Sentiment","Capacity","Deal Size","Companies","ISO / RTO","URL"]
    main_widths = [62, 12, 18, 18, 16, 14, 16, 12, 12, 28, 50]
    safe_df = df[[c for c in main_cols if c in df.columns]]
    write_sheet(ws1, safe_df, [c for c in main_cols if c in df.columns], main_widths[:len(main_cols)])

    ws2 = wb.create_sheet("By Country")
    cc = df["Country"].value_counts().reset_index()
    cc.columns = ["Country", "Articles"]
    write_sheet(ws2, cc, ["Country","Articles"], [24, 14])

    ws3 = wb.create_sheet("By Region")
    rc = df["Region"].value_counts().reset_index()
    rc.columns = ["Region", "Articles"]
    write_sheet(ws3, rc, ["Region","Articles"], [20, 14])

    ws4 = wb.create_sheet("By Topic")
    tc_df = df["Topic"].value_counts().reset_index()
    tc_df.columns = ["Topic", "Articles"]
    write_sheet(ws4, tc_df, ["Topic","Articles"], [18, 14])

    ws5 = wb.create_sheet("By Company")
    comp_rows = []
    for _, row in df.iterrows():
        if row.get("Companies"):
            for co in str(row["Companies"]).split(", "):
                co = co.strip()
                if co:
                    comp_rows.append({"Company": co, "Headline": row["Headline"],
                                      "Date": row["Date"], "Country": row["Country"],
                                      "URL": row["URL"]})
    if comp_rows:
        comp_df = pd.DataFrame(comp_rows)
        write_sheet(ws5, comp_df, ["Company","Headline","Date","Country","URL"],
                    [22, 55, 12, 18, 45])

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()


def build_excel_re(df):
    """Styled, multi-sheet Excel export for Renewables — same format/look as build_excel()."""
    wb = Workbook()
    thin = Side(border_style="thin", color="CCCCCC")
    brd = Border(left=thin, right=thin, top=thin, bottom=thin)
    hf = Font(bold=True, color="FFFFFF", name="Calibri", size=10)
    hfill = PatternFill("solid", fgColor="0A2A18")
    ev = PatternFill("solid", fgColor="0B1820")
    od = PatternFill("solid", fgColor="060A10")
    nf = Font(name="Calibri", size=9, color="C8D4E8")
    lf = Font(name="Calibri", size=9, color="00B4FF", underline="single")

    def write_sheet(ws, data_df, cols, widths):
        for ci, (col, w) in enumerate(zip(cols, widths), 1):
            c = ws.cell(1, ci, col)
            c.font = hf; c.fill = hfill
            c.alignment = Alignment(horizontal="center", vertical="center")
            c.border = brd
            ws.column_dimensions[get_column_letter(ci)].width = w
        ws.row_dimensions[1].height = 24
        for ri, row in enumerate(data_df[cols if all(c in data_df.columns for c in cols) else data_df.columns].itertuples(index=False), start=2):
            fill = ev if ri % 2 == 0 else od
            for ci, val in enumerate(row, 1):
                v = str(val) if val is not None else ""
                c = ws.cell(ri, ci, v)
                c.fill = fill; c.border = brd
                col_name = cols[ci - 1] if ci <= len(cols) else ""
                if col_name == "URL":
                    c.hyperlink = v; c.font = lf
                    c.alignment = Alignment(horizontal="center", vertical="center")
                elif col_name == "Sector":
                    tc = RE_TOPIC_COLORS.get(v, "2e4470")
                    c.font = Font(name="Calibri", size=9, bold=True, color=tc.replace("#", ""))
                    c.alignment = Alignment(horizontal="center", vertical="center")
                elif col_name == "Deal Type":
                    tc = RE_DEAL_TYPE_COLORS.get(v, "2e4470")
                    c.font = Font(name="Calibri", size=9, bold=True, color=tc.replace("#", ""))
                    c.alignment = Alignment(horizontal="center", vertical="center")
                elif col_name == "Headline":
                    c.font = nf
                    c.alignment = Alignment(vertical="center", wrap_text=True)
                else:
                    c.font = nf
                    c.alignment = Alignment(horizontal="center", vertical="center")
        ws.freeze_panes = "A2"
        ws.auto_filter.ref = f"A1:{get_column_letter(len(cols))}{len(data_df)+1}"

    ws1 = wb.active
    ws1.title = "All Articles"
    main_cols = ["Headline","Date","Source","Country","Region","Sector",
                 "Deal Type","Status","Capacity","Deal Size","Companies","URL"]
    main_widths = [62, 12, 18, 18, 16, 16, 16, 14, 12, 12, 28, 50]
    safe_cols = [c for c in main_cols if c in df.columns]
    write_sheet(ws1, df[safe_cols], safe_cols, main_widths[:len(safe_cols)])

    ws2 = wb.create_sheet("By Country")
    cc = df["Country"].value_counts().reset_index()
    cc.columns = ["Country", "Articles"]
    write_sheet(ws2, cc, ["Country","Articles"], [24, 14])

    ws3 = wb.create_sheet("By Region")
    if "Region" in df.columns:
        rc = df["Region"].value_counts().reset_index()
        rc.columns = ["Region", "Articles"]
        write_sheet(ws3, rc, ["Region","Articles"], [20, 14])

    ws4 = wb.create_sheet("By Sector")
    if "Sector" in df.columns:
        sc = df["Sector"].value_counts().reset_index()
        sc.columns = ["Sector", "Articles"]
        write_sheet(ws4, sc, ["Sector","Articles"], [20, 14])

    ws5 = wb.create_sheet("By Deal Type")
    if "Deal Type" in df.columns:
        dtc = df["Deal Type"].value_counts().reset_index()
        dtc.columns = ["Deal Type", "Articles"]
        write_sheet(ws5, dtc, ["Deal Type","Articles"], [22, 14])

    ws6 = wb.create_sheet("By Company")
    comp_rows = []
    for _, row in df.iterrows():
        if row.get("Companies"):
            for co in str(row["Companies"]).split(", "):
                co = co.strip()
                if co:
                    comp_rows.append({"Company": co, "Headline": row["Headline"],
                                      "Date": row["Date"], "Country": row["Country"],
                                      "URL": row["URL"]})
    if comp_rows:
        comp_df = pd.DataFrame(comp_rows)
        write_sheet(ws6, comp_df, ["Company","Headline","Date","Country","URL"],
                    [22, 55, 12, 18, 45])

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()


_BG = "#060a10"
_PAPER = "#0b1628"
_GRID = "#152038"
_TEXT = "#6a80a8"
_TITLE = "#b8c8e0"
_FONT = "Inter, sans-serif"

# ── Chart accent theme (mutable; set once per run from platform_mode) ───────
# Charts that use a *platform* accent color (hover border, timeline line/fill)
# read from here. Charts using *category* colors (TOPIC_COLORS, REGION_COLORS,
# RE_TOPIC_COLORS, etc.) are unaffected — those encode data meaning, not theme.
_CHART_THEME = {"accent": "#0047e1", "accent2": "#00b4ff"}

def _set_chart_theme(is_renewables: bool):
    if is_renewables:
        _CHART_THEME["accent"]  = "#00a846"
        _CHART_THEME["accent2"] = "#00e676"
    else:
        _CHART_THEME["accent"]  = "#0047e1"
        _CHART_THEME["accent2"] = "#00b4ff"


def _dark(fig, height=320):
    fig.update_layout(
        paper_bgcolor=_PAPER, plot_bgcolor=_BG,
        font=dict(family=_FONT, color=_TEXT),
        height=height,
        margin=dict(l=14, r=14, t=44, b=14),
        xaxis=dict(
            gridcolor="#0f1e36", gridwidth=1,
            linecolor=_GRID, linewidth=1,
            tickfont=dict(size=10, color=_TEXT),
            title_font=dict(color=_TEXT),
            zeroline=False,
        ),
        yaxis=dict(
            gridcolor="#0f1e36", gridwidth=1,
            linecolor=_GRID, linewidth=1,
            tickfont=dict(size=10, color=_TEXT),
            title_font=dict(color=_TEXT),
            zeroline=False,
        ),
        showlegend=False,
        hoverlabel=dict(
            bgcolor="#0d1e38",
            bordercolor=_CHART_THEME["accent"],
            font=dict(color="#ccdaf5", size=12, family="Inter, sans-serif"),
        ),
        hovermode="closest",
    )
    return fig


def chart_topic_bar(df):
    tc = df["Topic"].value_counts().reset_index()
    tc.columns = ["Topic", "Count"]
    tc = tc.sort_values("Count")
    colors = [TOPIC_COLORS.get(t, "#2e4470") for t in tc["Topic"]]
    fig = go.Figure(go.Bar(
        x=tc["Count"], y=tc["Topic"], orientation="h",
        marker=dict(
            color=colors,
            line=dict(width=0),
            opacity=0.88,
        ),
        text=tc["Count"], textposition="outside",
        textfont=dict(color=_TITLE, size=11, family="DM Mono, monospace"),
        hovertemplate="<b>%{y}</b><br>📰 %{x} articles<extra></extra>",
        hoverlabel=dict(bgcolor="#0d1e38", bordercolor=_CHART_THEME["accent"],
                        font=dict(color="#ccdaf5", size=12)),
    ))
    _dark(fig, 320)
    fig.update_layout(
        title=dict(text="Articles by Topic", font=dict(color=_TITLE, size=13, family="Syne, sans-serif"), x=0.01),
        xaxis=dict(showgrid=True, gridcolor="#0f1e36"),
        bargap=0.28,
    )
    return fig


def chart_region_bar(df):
    rc = df["Region"].value_counts().reset_index()
    rc.columns = ["Region", "Count"]
    rc = rc.sort_values("Count")
    colors = [REGION_COLORS.get(r, "#2e4470") for r in rc["Region"]]
    fig = go.Figure(go.Bar(
        x=rc["Count"], y=rc["Region"], orientation="h",
        marker=dict(color=colors, line=dict(width=0), opacity=0.88),
        text=rc["Count"], textposition="outside",
        textfont=dict(color=_TITLE, size=11, family="DM Mono, monospace"),
        hovertemplate="<b>%{y}</b><br>📰 %{x} articles<extra></extra>",
    ))
    _dark(fig, 300)
    fig.update_layout(
        title=dict(text="Articles by Region", font=dict(color=_TITLE, size=13, family="Syne, sans-serif"), x=0.01),
        bargap=0.28,
    )
    return fig


def chart_country_bar(df, top_n=20):
    cc = df["Country"].value_counts().head(top_n).reset_index()
    cc.columns = ["Country", "Count"]
    cc = cc.sort_values("Count")
    n = len(cc)
    c_colors = [
        f"rgba({int(0 + 71*i/max(n-1,1))}, {int(71 + (180-71)*i/max(n-1,1))}, {int(225 + (255-225)*i/max(n-1,1))}, 0.88)"
        for i in range(n)
    ]
    fig = go.Figure(go.Bar(
        x=cc["Count"], y=cc["Country"], orientation="h",
        marker=dict(color=c_colors, line=dict(width=0)),
        text=cc["Count"], textposition="outside",
        textfont=dict(color=_TITLE, size=10, family="DM Mono, monospace"),
        hovertemplate="<b>%{y}</b><br>📰 %{x} articles<extra></extra>",
    ))
    _dark(fig, max(340, top_n * 24))
    fig.update_layout(
        title=dict(text=f"Top {top_n} Countries", font=dict(color=_TITLE, size=13, family="Syne, sans-serif"), x=0.01),
        bargap=0.22,
    )
    return fig


def chart_timeline(df):
    df2 = df[df["Date"] != "Unknown"].copy()
    if df2.empty:
        return None
    df2["dt"] = pd.to_datetime(df2["Date"])
    daily = df2.groupby(df2["dt"].dt.date).size().reset_index()
    daily.columns = ["Date", "Articles"]
    _ac, _ac2 = _CHART_THEME["accent"], _CHART_THEME["accent2"]
    _ac_r, _ac_g, _ac_b = int(_ac[1:3], 16), int(_ac[3:5], 16), int(_ac[5:7], 16)
    fig = go.Figure()
    fig.add_trace(go.Scatter(
        x=daily["Date"], y=daily["Articles"],
        mode="lines+markers",
        line=dict(color=_ac2, width=2.5, shape="spline", smoothing=0.8),
        marker=dict(color=_ac, size=6, line=dict(color=_ac2, width=1.5),
                    symbol="circle"),
        fill="tozeroy", fillcolor=f"rgba({_ac_r},{_ac_g},{_ac_b},0.09)",
        hovertemplate="<b>%{x}</b><br>📰 %{y} articles<extra></extra>",
    ))
    _dark(fig, 260)
    fig.update_layout(
        title=dict(text="Publication Volume Over Time", font=dict(color=_TITLE, size=13, family="Syne, sans-serif"), x=0.01),
        yaxis_title="Articles",
        xaxis=dict(showgrid=False),
    )
    return fig


def chart_sentiment(df):
    sc = df["Sentiment"].value_counts().reset_index()
    sc.columns = ["Sentiment", "Count"]
    sent_colors = {
        "Opened / Live": "#00e676", "Approved": "#00b4ff",
        "Proposed": "#ffaa00", "Under Construction": "#00e5c8",
        "Challenged": "#ff2d6b", "News": "#2e4470",
    }
    colors = [sent_colors.get(s, "#2e4470") for s in sc["Sentiment"]]
    fig = go.Figure(go.Bar(
        x=sc["Sentiment"], y=sc["Count"],
        marker=dict(color=colors, line=dict(width=0), opacity=0.88),
        text=sc["Count"], textposition="outside",
        textfont=dict(color=_TITLE, size=11, family="DM Mono, monospace"),
        hovertemplate="<b>%{x}</b><br>📰 %{y} articles<extra></extra>",
    ))
    _dark(fig, 290)
    fig.update_layout(
        title=dict(text="Article Sentiment / Status", font=dict(color=_TITLE, size=13, family="Syne, sans-serif"), x=0.01),
        bargap=0.3,
    )
    return fig


def chart_donut(df):
    tc = df["Topic"].value_counts().reset_index()
    tc.columns = ["Topic", "Count"]
    colors = [TOPIC_COLORS.get(t, "#2e4470") for t in tc["Topic"]]
    fig = go.Figure(go.Pie(
        labels=tc["Topic"], values=tc["Count"], hole=0.6,
        marker=dict(colors=colors, line=dict(color=_BG, width=3)),
        textinfo="label+percent",
        textfont=dict(color=_TITLE, size=11),
        hovertemplate="<b>%{label}</b><br>📰 %{value} articles (%{percent})<extra></extra>",
        pull=[0.03] * len(tc),
    ))
    fig.update_layout(
        paper_bgcolor=_PAPER, plot_bgcolor=_BG,
        font=dict(family=_FONT, color=_TEXT),
        height=360, margin=dict(l=14, r=14, t=44, b=14),
        showlegend=True,
        legend=dict(bgcolor="rgba(0,0,0,0)", font=dict(color=_TITLE, size=10),
                    orientation="v", x=1.02, y=0.5),
        hoverlabel=dict(
            bgcolor="#0d1e38", bordercolor=_CHART_THEME["accent"],
            font=dict(color="#ccdaf5", size=12, family="Inter, sans-serif"),
        ),
        annotations=[dict(
            text=f"<b>{len(df)}</b><br><span style='font-size:10px'>articles</span>",
            x=0.5, y=0.5, showarrow=False,
            font=dict(size=15, color=_TITLE, family=_FONT),
        )],
        title=dict(text="Topic Share", font=dict(color=_TITLE, size=13, family="Syne, sans-serif"), x=0.01),
    )
    return fig


def chart_source_bar(df):
    sc = df["Source"].value_counts().reset_index()
    sc.columns = ["Source", "Count"]
    sc = sc.sort_values("Count")
    colors = [SOURCE_META.get(s, SOURCE_META["Unknown"])["color"] for s in sc["Source"]]
    fig = go.Figure(go.Bar(
        x=sc["Count"], y=sc["Source"], orientation="h",
        marker=dict(color=colors, line=dict(width=0), opacity=0.88),
        text=sc["Count"], textposition="outside",
        textfont=dict(color=_TITLE, size=11, family="DM Mono, monospace"),
        hovertemplate="<b>%{y}</b><br>📰 %{x} articles<extra></extra>",
    ))
    _dark(fig, 290)
    fig.update_layout(
        title=dict(text="Articles by Source", font=dict(color=_TITLE, size=13, family="Syne, sans-serif"), x=0.01),
        bargap=0.28,
    )
    return fig


def chart_world_map(df, title="Global Data Center Activity"):
    cc = df[df["Country"] != "Global"]["Country"].value_counts().reset_index()
    cc.columns = ["Country", "Count"]
    cc["ISO"] = cc["Country"].map(COUNTRY_ISO)
    cc = cc.dropna(subset=["ISO"])
    if cc.empty:
        return go.Figure()

    z_max = max(int(cc["Count"].max()), 1)

    # ── Palette matched to the app's dark UI ─────────────────────────────────
    _M_BG    = "#060e1c"   # near-black page bg — same feel as .stApp
    _M_OCEAN = "#07111f"   # deep navy ocean — slightly lighter than bg
    _M_LAND  = "#0d1b30"   # midnight-blue for zero-data land masses
    _M_GRID  = "#152038"   # subtle border matching app grid colour

    # Heat scale: deep navy (zero) → electric cyan → cobalt blue → neon gold (max)
    # Matches the app's accent palette (#0047e1 blue, #00b4ff cyan, #ffaa00 amber)
    _COLORSCALE = [
        [0.000, "#07111f"],   # zero  — ocean/bg level (invisible baseline)
        [0.05,  "#0a2040"],   # trace — barely there
        [0.15,  "#0d3a70"],   # low   — deep cobalt
        [0.30,  "#0047e1"],   # low-mid — brand blue
        [0.50,  "#00b4ff"],   # mid   — electric cyan
        [0.70,  "#00e5c8"],   # high  — teal-mint
        [0.85,  "#ffaa00"],   # very high — amber
        [1.000, "#ff6400"],   # max   — vivid orange-gold peak
    ]

    fig = go.Figure(go.Choropleth(
        locations=cc["ISO"],
        z=cc["Count"],
        text=cc["Country"],
        colorscale=_COLORSCALE,
        autocolorscale=False,
        reversescale=False,
        zauto=False,
        zmin=0,
        zmax=z_max,
        marker=dict(line=dict(color="#152038", width=0.6)),
        colorbar=dict(
            bgcolor="#0b1628",
            bordercolor="#152038",
            borderwidth=1,
            tickfont=dict(color="#6a80a8", size=10, family="DM Mono, monospace"),
            title=dict(
                text="Articles",
                font=dict(color="#b8c8e0", size=11, family="Syne, sans-serif"),
            ),
            len=0.6,
            thickness=14,
            tickformat="d",
            x=1.01,
        ),
        hovertemplate=(
            "<b style='color:#fff;font-size:13px;'>%{text}</b><br>"
            f"<span style='color:{_CHART_THEME['accent2']};'>Articles: %{{z}}</span>"
            "<extra></extra>"
        ),
        showscale=True,
    ))

    fig.update_geos(
        bgcolor=_M_BG,
        landcolor=_M_LAND,
        oceancolor=_M_OCEAN,
        lakecolor=_M_OCEAN,
        rivercolor=_M_OCEAN,
        framecolor=_M_GRID,
        showland=True,
        showocean=True,
        showlakes=True,
        showrivers=False,
        showcountries=True,
        countrycolor="#152038",
        showcoastlines=True,
        coastlinecolor="#1a2e50",
        showframe=False,
        projection_type="natural earth",
        lataxis_range=[-60, 85],
    )

    fig.update_layout(
        paper_bgcolor=_M_BG,
        plot_bgcolor=_M_BG,
        height=520,
        margin=dict(l=0, r=60, t=44, b=0),
        title=dict(
            text=f"<b>{title}</b>",
            font=dict(color="#b8c8e0", size=14, family="Syne, sans-serif"),
            x=0.01, y=0.98,
        ),
        geo=dict(bgcolor=_M_BG),
        hoverlabel=dict(
            bgcolor="#0b1628",
            bordercolor=_CHART_THEME["accent"],
            font=dict(color="#ccdaf5", size=12, family="Inter, sans-serif"),
        ),
    )
    return fig


def dark_table(df_in, max_rows=300):
    th = (
        "background:#0f1e36;color:#b8c8e0;font-family:monospace;"
        f"font-size:.68rem;letter-spacing:.08em;text-transform:uppercase;"
        f"padding:.55rem .85rem;border-bottom:2px solid {_CHART_THEME['accent']};"
        "white-space:nowrap;text-align:left;"
    )
    td = (
        "padding:.5rem .85rem;font-size:.8rem;color:#b8c8e0;"
        "border-bottom:1px solid #101b2e;font-family:Inter,sans-serif;"
        "vertical-align:middle;"
    )
    rows = ""
    for i, (_, row) in enumerate(df_in.head(max_rows).iterrows()):
        bg = "#0b1628" if i % 2 == 0 else "#060a10"
        cells = ""
        for col in df_in.columns:
            v = str(row[col]) if row[col] is not None else ""
            if col == "URL" or (col != "Headline" and v.startswith("http")):
                cells += (
                    f'<td style="{td}background:{bg};">'
                    f'<a href="{v}" target="_blank" '
                    f'style="color:{_CHART_THEME["accent2"]};text-decoration:none;font-size:.75rem;">'
                    f'Open \u2192</a></td>'
                )
            elif col == "Capacity" and v:
                cells += (
                    f'<td style="{td}background:{bg};">'
                    f'<span style="color:#ffaa00;font-family:monospace;font-size:.76rem;">'
                    f'\u26a1 {v}</span></td>'
                )
            elif col == "Deal Size" and v:
                cells += (
                    f'<td style="{td}background:{bg};">'
                    f'<span style="color:#00e676;font-family:monospace;font-size:.76rem;">'
                    f'{v}</span></td>'
                )
            elif col == "Topic":
                tc = TOPIC_COLORS.get(v, "#2e4470")
                cells += (
                    f'<td style="{td}background:{bg};">'
                    f'<span style="background:{tc}22;color:{tc};border:1px solid {tc}44;'
                    f'border-radius:4px;padding:2px 7px;font-size:.68rem;font-family:monospace;">'
                    f'{v}</span></td>'
                )
            elif col == "Region":
                rc2 = REGION_COLORS.get(v, "#2e4470")
                cells += (
                    f'<td style="{td}background:{bg};">'
                    f'<span style="background:{rc2}22;color:{rc2};border:1px solid {rc2}44;'
                    f'border-radius:4px;padding:2px 7px;font-size:.68rem;font-family:monospace;">'
                    f'{v}</span></td>'
                )
            elif col == "Sentiment":
                sent_c = {
                    "Opened / Live":"#00e676","Approved":"#00b4ff","Proposed":"#ffaa00",
                    "Under Construction":"#00e5c8","Challenged":"#ff2d6b","News":"#2e4470",
                }.get(v, "#2e4470")
                cells += (
                    f'<td style="{td}background:{bg};">'
                    f'<span style="background:{sent_c}18;color:{sent_c};'
                    f'border:1px solid {sent_c}44;border-radius:4px;'
                    f'padding:2px 7px;font-size:.68rem;font-family:monospace;">'
                    f'{v}</span></td>'
                )
            elif col == "Source":
                sc2 = SOURCE_META.get(v, SOURCE_META["Unknown"])
                sc2_color = sc2["color"]
                sc2_short = sc2["short"]
                cells += (
                    f'<td style="{td}background:{bg};">'
                    f'<span style="background:{sc2_color}22;color:{sc2_color};'
                    f'border:1px solid {sc2_color}44;border-radius:4px;'
                    f'padding:2px 6px;font-size:.65rem;font-family:monospace;">'
                    f'{sc2_short}</span></td>'
                )
            elif col in ("Headline", "Companies"):
                cells += (
                    f'<td style="{td}background:{bg};max-width:480px;'
                    f'word-wrap:break-word;white-space:normal;">{v}</td>'
                )
            else:
                cells += f'<td style="{td}background:{bg};">{v}</td>'
        rows += f"<tr>{cells}</tr>"

    heads = "".join(f'<th style="{th}">{c}</th>' for c in df_in.columns)
    return (
        '<div style="overflow-x:auto;border-radius:10px;border:1px solid #152038;margin-bottom:1rem;">'
        f'<table class="dc-table" style="width:100%;border-collapse:collapse;background:#060a10;">'
        f'<thead><tr>{heads}</tr></thead>'
        f'<tbody>{rows}</tbody></table></div>'
    )


def article_card(headline, date, url, source, country, topic, capacity, deal, sentiment, ai_score=None):
    tc = TOPIC_COLORS.get(topic, "#2e4470")
    sc_meta = SOURCE_META.get(source, SOURCE_META["Unknown"])
    cap_html = (
        f'<span style="background:rgba(255,170,0,0.12);color:#ffaa00;'
        f'border:1px solid rgba(255,170,0,0.3);border-radius:4px;'
        f'padding:2px 6px;font-family:monospace;font-size:.62rem;white-space:nowrap;" '
        f'title="Capacity announced">'
        f'\u26a1 {capacity}</span>'
    ) if capacity else ""
    deal_html = (
        f'<span style="background:rgba(0,230,118,0.1);color:#00e676;'
        f'border:1px solid rgba(0,230,118,0.25);border-radius:4px;'
        f'padding:2px 6px;font-family:monospace;font-size:.62rem;white-space:nowrap;" '
        f'title="Deal value">'
        f'{deal}</span>'
    ) if deal else ""
    sent_c = {
        "Opened / Live":"#00e676","Approved":"#00b4ff","Proposed":"#ffaa00",
        "Under Construction":"#00e5c8","Challenged":"#ff2d6b","News":"#2e4470",
    }.get(sentiment, "#2e4470")
    arrow = "\u2197"

    # AI score badge (optional)
    score_badge_html = ""
    if ai_score is not None:
        sc_color = (
            "#ff6400" if ai_score >= 40
            else "#ffaa00" if ai_score >= 25
            else "#00b4ff" if ai_score >= 15
            else "#3a5480"
        )
        score_badge_html = (
            f'<span class="score-badge" style="background:{sc_color}22;color:{sc_color};'
            f'border:1px solid {sc_color}44;border-radius:4px;'
            f'padding:2px 7px;font-family:monospace;font-size:.62rem;font-weight:bold;" '
            f'title="AI Signal Score — higher = more market significant">&#9733; {ai_score}</span>'
        )

    # High signal indicator strip
    is_high_signal = ai_score is not None and ai_score >= 30
    border_color = "#ffaa00" if is_high_signal else "#152038"
    top_strip = (
        f'<div style="position:absolute;top:0;left:0;right:0;height:2px;'
        f'background:linear-gradient(90deg,#ffaa00,#ff6400);border-radius:2px 2px 0 0;"></div>'
        if is_high_signal else ""
    )

    return (
        f'<div class="article-card-wrap" style="background:#0b1628;border:1px solid {border_color};border-radius:10px;'
        f'padding:.85rem 1.1rem;display:flex;justify-content:space-between;'
        f'align-items:flex-start;gap:.9rem;margin-bottom:.45rem;position:relative;">'
        f'{top_strip}'
        f'<div style="flex:1;min-width:0;">'
        f'<a href="{url}" target="_blank" '
        f'style="color:#ccdaf5;text-decoration:none;font-family:Inter,sans-serif;'
        f'font-size:.86rem;font-weight:500;line-height:1.5;" '
        f'title="Open article in new tab">{headline}</a>'
        f'<div style="margin-top:.35rem;display:flex;gap:.4rem;flex-wrap:wrap;">'
        f'<span style="font-family:monospace;font-size:.62rem;color:#2a3e60;" title="Publication date">\U0001f4c5 {date}</span>'
        f'<span style="font-family:monospace;font-size:.62rem;color:#4a6490;" title="Country detected">\U0001f30d {country}</span>'
        f'</div></div>'
        f'<div style="display:flex;flex-direction:column;align-items:flex-end;'
        f'gap:.28rem;flex-shrink:0;white-space:nowrap;">'
        f'<span style="background:{sc_meta["color"]}22;color:{sc_meta["color"]};'
        f'border:1px solid {sc_meta["color"]}44;border-radius:4px;'
        f'padding:2px 6px;font-family:monospace;font-size:.62rem;" title="Source: {source}">{sc_meta["short"]}</span>'
        f'<span style="background:{tc}22;color:{tc};border:1px solid {tc}44;'
        f'border-radius:4px;padding:2px 6px;font-family:monospace;font-size:.62rem;" title="Topic: {topic}">{topic}</span>'
        f'<span style="background:{sent_c}18;color:{sent_c};border:1px solid {sent_c}44;'
        f'border-radius:4px;padding:2px 6px;font-family:monospace;font-size:.62rem;" title="Project status">{sentiment}</span>'
        f'{cap_html}{deal_html}{score_badge_html}'
        f'<a href="{url}" target="_blank" '
        f'style="font-family:monospace;font-size:.65rem;color:#0047e1;text-decoration:none;" title="Open full article">'
        f'{arrow} open</a>'
        f'</div></div>'
    )


def kpi(label, value, accent="blue", delta=""):
    accent_map = {
        "blue":  ("#0047e1", "#00b4ff"),
        "cyan":  ("#00b4ff", "#00e5c8"),
        "green": ("#00e676", "#00c853"),
        "amber": ("#ffaa00", "#ff6400"),
        "purple":("#a855f7", "#7c3aed"),
        "red":   ("#ff2d6b", "#ff0044"),
    }
    c1, c2 = accent_map.get(accent, accent_map["blue"])
    delta_html = f'<div style="font-size:.7rem;color:#2a3e60;margin-top:.25rem;">{delta}</div>' if delta else ""
    return (
        f'<div class="kpi-card" style="flex:1;min-width:150px;background:#0b1628;border:1px solid #152038;'
        f'border-radius:12px;padding:1.1rem 1.3rem;position:relative;overflow:hidden;" '
        f'title="{label}">'
        f'<div style="position:absolute;top:0;left:0;right:0;height:2px;'
        f'background:linear-gradient(90deg,{c1},{c2});"></div>'
        f'<div style="font-family:monospace;font-size:.64rem;letter-spacing:.13em;'
        f'text-transform:uppercase;color:#2a3e60;margin-bottom:.4rem;">{label}</div>'
        f'<div style="font-family:Syne,sans-serif;font-size:1.9rem;font-weight:800;'
        f'color:#fff;line-height:1;">{value}</div>'
        f'{delta_html}</div>'
    )


def render_filter_chips(active, on_clear_key=None):
    """Render a row of themed chips summarizing active filters.

    `active` is a list of (label, display_value) tuples — callers should only
    pass entries that are actually set (non-empty), so this always reflects
    real, current filter state rather than a static list of every possible
    filter field.
    """
    if not active:
        return
    chips_html = "".join(
        f'<div class="filter-chip"><span class="fc-label">{lbl}:</span> <b style="color:#fff;">{val}</b></div>'
        for lbl, val in active
    )
    st.markdown(
        f'<div style="display:flex;flex-wrap:wrap;gap:.5rem;align-items:center;'
        f'margin:-.4rem 0 1.1rem;">'
        f'<span style="font-family:DM Mono,monospace;font-size:.62rem;letter-spacing:.1em;'
        f'color:#2a3e60;text-transform:uppercase;margin-right:.2rem;">Active Filters</span>'
        f'{chips_html}</div>',
        unsafe_allow_html=True,
    )


def empty_state(message, hint="Try broadening your filters or running a wider scan.", icon="\U0001f50d"):
    """Themed, actionable empty-state panel — replaces bare st.info() calls
    for 'no results' situations so the message matches the active theme
    (blue for Data Center, green for Renewables) instead of Streamlit's
    generic default blue box."""
    st.markdown(
        f'<div style="background:#0b1628;border:1px dashed var(--accent);border-radius:12px;'
        f'padding:1.4rem 1.6rem;margin:.4rem 0 1rem;text-align:center;">'
        f'<div style="font-size:1.6rem;margin-bottom:.5rem;opacity:.75;">{icon}</div>'
        f'<div style="font-family:Syne,sans-serif;font-weight:700;color:#ccdaf5;'
        f'font-size:.92rem;margin-bottom:.3rem;">{message}</div>'
        f'<div style="font-size:.78rem;color:#5a6f95;">{hint}</div>'
        f'</div>',
        unsafe_allow_html=True,
    )



# ─── RE-specific chart helpers ─────────────────────────────────────────────────
def chart_re_sector_bar(df):
    tc = df["Sector"].value_counts().reset_index()
    tc.columns = ["Sector", "Count"]
    tc = tc.sort_values("Count")
    colors = [RE_TOPIC_COLORS.get(t, "#2e4470") for t in tc["Sector"]]
    fig = go.Figure(go.Bar(
        x=tc["Count"], y=tc["Sector"], orientation="h",
        marker=dict(color=colors, line=dict(width=0), opacity=0.88),
        text=tc["Count"], textposition="outside",
        textfont=dict(color=_TITLE, size=11, family="DM Mono, monospace"),
        hovertemplate="<b>%{y}</b><br>📰 %{x} articles<extra></extra>",
    ))
    _dark(fig, 340)
    fig.update_layout(
        title=dict(text="Articles by Sector", font=dict(color=_TITLE, size=13, family="Syne, sans-serif"), x=0.01),
        xaxis=dict(showgrid=True, gridcolor="#0f1e36"),
        bargap=0.28,
    )
    return fig

def chart_re_deal_type_bar(df):
    tc = df["Deal Type"].value_counts().reset_index()
    tc.columns = ["Deal Type", "Count"]
    tc = tc.sort_values("Count")
    colors = [RE_DEAL_TYPE_COLORS.get(t, "#2e4470") for t in tc["Deal Type"]]
    fig = go.Figure(go.Bar(
        x=tc["Count"], y=tc["Deal Type"], orientation="h",
        marker=dict(color=colors, line=dict(width=0), opacity=0.88),
        text=tc["Count"], textposition="outside",
        textfont=dict(color=_TITLE, size=11, family="DM Mono, monospace"),
        hovertemplate="<b>%{y}</b><br>📰 %{x} articles<extra></extra>",
    ))
    _dark(fig, 320)
    fig.update_layout(
        title=dict(text="Articles by Deal Type", font=dict(color=_TITLE, size=13, family="Syne, sans-serif"), x=0.01),
        bargap=0.28,
    )
    return fig

def chart_re_status_pie(df):
    sc = df["Status"].value_counts().reset_index()
    sc.columns = ["Status", "Count"]
    status_colors = {
        "Commissioned":"#00e676","Financed":"#00b4ff","Under Construction":"#00e5c8",
        "Approved":"#a855f7","Tendered":"#ffaa00","Contracted":"#0047e1",
        "Proposed":"#ff6400","Challenged":"#ff2d6b","News":"#2e4470",
    }
    colors = [status_colors.get(s, "#2e4470") for s in sc["Status"]]
    fig = go.Figure(go.Pie(
        labels=sc["Status"], values=sc["Count"],
        marker=dict(colors=colors, line=dict(color=_BG, width=2)),
        textinfo="label+percent",
        textfont=dict(size=10, family="DM Mono, monospace", color=_TITLE),
        hole=0.52,
        hovertemplate="<b>%{label}</b><br>%{value} articles · %{percent}<extra></extra>",
    ))
    _dark(fig, 300)
    fig.update_layout(
        title=dict(text="Project Status Distribution", font=dict(color=_TITLE, size=13, family="Syne, sans-serif"), x=0.01),
        showlegend=True,
        legend=dict(font=dict(color=_TEXT, size=10), bgcolor="rgba(0,0,0,0)"),
    )
    return fig

def re_article_card(headline, date, url, source, country, sector, deal_type,
                    capacity="", deal_size="", status="News", ai_score=None):
    """Render a styled article card for Renewables Power Markets feed."""
    sc_meta = RE_SOURCE_META.get(source, {"color": "#2e4470", "short": source[:3].upper()})
    tc = RE_TOPIC_COLORS.get(sector, "#2e4470")
    dt_c = RE_DEAL_TYPE_COLORS.get(deal_type, "#2e4470")
    cap_html = (
        f'<span style="background:rgba(255,170,0,0.12);color:#ffaa00;'
        f'border:1px solid rgba(255,170,0,0.25);border-radius:4px;'
        f'padding:2px 6px;font-family:monospace;font-size:.62rem;white-space:nowrap;" '
        f'title="Capacity">⚡ {capacity}</span>'
    ) if capacity else ""
    deal_html = (
        f'<span style="background:rgba(0,230,118,0.1);color:#00e676;'
        f'border:1px solid rgba(0,230,118,0.25);border-radius:4px;'
        f'padding:2px 6px;font-family:monospace;font-size:.62rem;white-space:nowrap;" '
        f'title="Deal value">{deal_size}</span>'
    ) if deal_size else ""
    status_colors = {
        "Commissioned":"#00e676","Financed":"#00b4ff","Under Construction":"#00e5c8",
        "Approved":"#a855f7","Tendered":"#ffaa00","Contracted":"#0047e1",
        "Proposed":"#ff6400","Challenged":"#ff2d6b","News":"#2e4470",
    }
    st_c = status_colors.get(status, "#2e4470")
    score_badge_html = ""
    if ai_score is not None:
        sc_c = "#ff6400" if ai_score>=40 else "#ffaa00" if ai_score>=25 else "#00b4ff" if ai_score>=15 else "#3a5480"
        score_badge_html = (
            f'<span style="background:{sc_c}22;color:{sc_c};border:1px solid {sc_c}44;'
            f'border-radius:4px;padding:2px 7px;font-family:monospace;font-size:.62rem;font-weight:bold;">★ {ai_score}</span>'
        )
    is_high = ai_score is not None and ai_score >= 30
    border_c = "#22aa55" if is_high else "#152038"
    top_strip = (
        '<div style="position:absolute;top:0;left:0;right:0;height:2px;'
        'background:linear-gradient(90deg,#00e676,#00b4ff);border-radius:2px 2px 0 0;"></div>'
        if is_high else ""
    )
    return (
        f'<div class="article-card-wrap" style="background:#0b1820;border:1px solid {border_c};border-radius:10px;'
        f'padding:.85rem 1.1rem;display:flex;justify-content:space-between;'
        f'align-items:flex-start;gap:.9rem;margin-bottom:.45rem;position:relative;">'
        f'{top_strip}'
        f'<div style="flex:1;min-width:0;">'
        f'<a href="{url}" target="_blank" '
        f'style="color:#ccdaf5;text-decoration:none;font-family:Inter,sans-serif;'
        f'font-size:.86rem;font-weight:500;line-height:1.5;">{headline}</a>'
        f'<div style="margin-top:.35rem;display:flex;gap:.4rem;flex-wrap:wrap;">'
        f'<span style="font-family:monospace;font-size:.62rem;color:#2a3e60;">📅 {date}</span>'
        f'<span style="font-family:monospace;font-size:.62rem;color:#4a6490;">🌍 {country}</span>'
        f'</div></div>'
        f'<div style="display:flex;flex-direction:column;align-items:flex-end;'
        f'gap:.28rem;flex-shrink:0;white-space:nowrap;">'
        f'<span style="background:{sc_meta["color"]}22;color:{sc_meta["color"]};'
        f'border:1px solid {sc_meta["color"]}44;border-radius:4px;'
        f'padding:2px 6px;font-family:monospace;font-size:.62rem;">{sc_meta["short"]}</span>'
        f'<span style="background:{tc}22;color:{tc};border:1px solid {tc}44;'
        f'border-radius:4px;padding:2px 6px;font-family:monospace;font-size:.62rem;">{sector}</span>'
        f'<span style="background:{dt_c}22;color:{dt_c};border:1px solid {dt_c}44;'
        f'border-radius:4px;padding:2px 6px;font-family:monospace;font-size:.62rem;">{deal_type}</span>'
        f'<span style="background:{st_c}18;color:{st_c};border:1px solid {st_c}44;'
        f'border-radius:4px;padding:2px 6px;font-family:monospace;font-size:.62rem;">{status}</span>'
        f'{cap_html}{deal_html}{score_badge_html}'
        f'<a href="{url}" target="_blank" '
        f'style="font-family:monospace;font-size:.65rem;color:#00e676;text-decoration:none;">↗ open</a>'
        f'</div></div>'
    )

def generate_re_summary(df, sel_desc, date_range):
    """
    Wood Mackenzie-grade renewables market intelligence briefing.
    Produces deeply analytical, prose-rich sections with quantified signals,
    sector deep-dives, deal flow analysis, and forward-looking commentary —
    modelled on top-tier clean energy research houses (BloombergNEF, Wood Mac, BNEF).
    """
    headlines = df["Headline"].tolist()
    scores    = _tfidf_scores(headlines)
    df2 = df.copy()
    df2["_score"] = scores

    total = len(df2)
    if total == 0:
        return "No articles in the selected view. Adjust filters and regenerate."

    # ── Core aggregations ─────────────────────────────────────────────────────
    top_sectors    = df2["Sector"].value_counts()
    top_deal_types = df2["Deal Type"].value_counts()
    top_statuses   = df2["Status"].value_counts()
    top_countries  = df2["Country"].value_counts()
    top_regions    = df2["Region"].value_counts()

    mw_df   = df2[df2["Capacity"] != ""].copy()
    deal_df = df2[df2["Deal Size"] != ""].copy()
    mw_vals   = mw_df["Capacity"].tolist()
    deal_vals = deal_df["Deal Size"].tolist()

    # Company counter
    all_co_list = []
    for v in df2["Companies"]:
        if v:
            all_co_list.extend([c.strip() for c in str(v).split(",") if c.strip()])
    co_counter = Counter(all_co_list)
    top_cos = [co for co, _ in co_counter.most_common(20)]

    # Status counts
    commissioned = int(top_statuses.get("Commissioned", 0))
    financed     = int(top_statuses.get("Financed", 0))
    under_c      = int(top_statuses.get("Under Construction", 0))
    approved     = int(top_statuses.get("Approved", 0))
    tendered     = int(top_statuses.get("Tendered", 0))
    contracted   = int(top_statuses.get("Contracted", 0))
    proposed     = int(top_statuses.get("Proposed", 0))
    challenged   = int(top_statuses.get("Challenged", 0))

    # Deal type counts
    ppa_cnt    = int(top_deal_types.get("PPA", 0))
    tender_cnt = int(top_deal_types.get("Tender / Auction", 0))
    mou_cnt    = int(top_deal_types.get("MOU", 0))
    aor_cnt    = int(top_deal_types.get("AOR / Offtake", 0))
    inv_cnt    = int(top_deal_types.get("Investment / IPO", 0))
    ma_cnt     = int(top_deal_types.get("M&A / Deals", 0))

    def pct(n): return f"{round(n/total*100)}%" if total else "0%"
    def hl(sub, n=3):
        if sub.empty: return []
        return sub.nlargest(n, "_score")["Headline"].tolist()

    # Aggregate announced capacity
    cap_total_mw = 0
    for cap in mw_vals:
        m = re.search(r"([\d,.]+)\s*(GW|MW)", str(cap), re.I)
        if m:
            v = float(m.group(1).replace(",", ""))
            cap_total_mw += v * 1000 if m.group(2).upper() == "GW" else v

    dominant_sector  = top_sectors.index[0]  if not top_sectors.empty  else "Solar"
    dominant_country = top_countries.index[0] if not top_countries.empty else "—"
    dominant_region  = top_regions.index[0]   if not top_regions.empty  else "Global"
    dominant_deal    = top_deal_types.index[0] if not top_deal_types.empty else "General"

    # Pipeline momentum signal
    active_pipeline = proposed + tendered + contracted
    momentum_signal = (
        "strongly positive — new project announcements and tenders significantly outpace "
        "commissioning, indicating an accelerating build-out cycle"
        if active_pipeline > (commissioned + financed) * 1.5
        else "balanced — deal origination and project commissioning are broadly in step, "
        "reflecting a market in orderly expansion"
        if abs(active_pipeline - commissioned) <= 5
        else "delivery-constrained — commissioning activity is lagging behind the pipeline of "
        "announced and tendered projects, pointing to execution and grid-connection bottlenecks"
    )

    # ── SECTION 1: EXECUTIVE SUMMARY ─────────────────────────────────────────
    exec_lines = [
        f"This intelligence briefing synthesises {total} renewables market articles "
        f"published between {date_range}, filtered to: {sel_desc}. "
        f"The analysis draws on headlines sourced from specialist renewable energy publications, "
        f"RSS feeds, and news aggregators, auto-enriched with sector classification, deal-type "
        f"detection, project status signals, capacity extraction, and named-entity recognition "
        f"across {df2['Country'].nunique()} countries and {df2['Region'].nunique()} regions.",

        f"\nSector concentration is led by **{dominant_sector}** at {pct(int(top_sectors.iloc[0]))} "
        f"of total coverage"
        + (f", followed by {' and '.join(str(t) for t in top_sectors.index[1:3])}" if len(top_sectors) > 1 else "")
        + f". Geographic activity is centred on {dominant_region}, with {dominant_country} "
        f"representing the single most active market. Deal flow is dominated by "
        f"**{dominant_deal}** transactions, reflecting the current stage of the procurement cycle.",

        f"\nThe project pipeline presents a {momentum_signal}. "
        f"Status breakdown: {commissioned} commissioned/energised, {financed} reached financial close, "
        f"{under_c} under construction, {contracted} contracted (PPA/offtake signed), "
        f"{tendered} tendered or at auction, {proposed} proposed or announced, "
        f"and {challenged} contested or challenged. "
        + (f"Identified capacity across announced and tendered projects totals approximately "
           f"**{cap_total_mw:,.0f} MW** across {len(mw_df)} discrete capacity-cited articles. "
           if cap_total_mw > 0 else "")
        + (f"Disclosed transaction values span {len(deal_df)} deal-citing articles including: "
           f"{'; '.join(deal_vals[:6])}."
           if deal_vals else ""),
    ]
    exec_summary = "\n\n".join(p.strip() for p in exec_lines if p.strip())

    # ── SECTION 2: SECTOR DYNAMICS ───────────────────────────────────────────
    sector_lines = []
    for sector, cnt in top_sectors.items():
        sub = df2[df2["Sector"] == sector]
        examples = hl(sub, 2)

        # Sub-metrics for each sector
        sec_commissioned = int((sub["Status"] == "Commissioned").sum())
        sec_contracted   = int((sub["Status"] == "Contracted").sum())
        sec_tendered     = int((sub["Status"] == "Tendered").sum())
        sec_proposed     = int((sub["Status"] == "Proposed").sum())
        sec_cap_arts     = sub[sub["Capacity"] != ""]["Capacity"].tolist()
        sec_ppa          = int((sub["Deal Type"] == "PPA").sum())
        sec_countries    = sub["Country"].value_counts().head(4).index.tolist()

        pipeline_str = []
        if sec_proposed:   pipeline_str.append(f"{sec_proposed} proposed")
        if sec_tendered:   pipeline_str.append(f"{sec_tendered} tendered")
        if sec_contracted: pipeline_str.append(f"{sec_contracted} contracted")
        if sec_commissioned: pipeline_str.append(f"{sec_commissioned} commissioned")

        detail = f"({pct(cnt)})"
        if sec_countries and sec_countries != ["Global"]:
            detail += f" · Key markets: {', '.join(sec_countries)}"
        if pipeline_str:
            detail += f" · Pipeline: {', '.join(pipeline_str)}"
        if sec_ppa:
            detail += f" · PPAs: {sec_ppa}"
        if sec_cap_arts:
            detail += f" · Capacity cited: {'; '.join(sec_cap_arts[:3])}"

        ex_str = ""
        if examples:
            h1 = examples[0][:115].rsplit(" ", 1)[0] if len(examples[0]) > 115 else examples[0]
            ex_str = f'\n  Notable: "{h1}"'
            if len(examples) > 1:
                h2 = examples[1][:100].rsplit(" ", 1)[0] if len(examples[1]) > 100 else examples[1]
                ex_str += f'; "{h2}".'
            else:
                ex_str += "."

        sector_lines.append(
            f"• **{sector}** — {cnt} article{'s' if cnt > 1 else ''} {detail}.{ex_str}"
        )

    # ── SECTION 3: DEAL FLOW & TRANSACTION ANALYSIS ──────────────────────────
    deal_type_lines = []
    for dt, cnt in top_deal_types.items():
        dt_sub = df2[df2["Deal Type"] == dt]
        dt_countries = dt_sub["Country"].value_counts().head(3).index.tolist()
        dt_sectors   = dt_sub["Sector"].value_counts().head(3).index.tolist()
        dt_cap_arts  = dt_sub[dt_sub["Capacity"] != ""]["Capacity"].tolist()
        dt_examples  = hl(dt_sub, 1)

        detail_parts = [f"{pct(cnt)} of deal flow"]
        if dt_countries and dt_countries != ["Global"]:
            detail_parts.append(f"active in {', '.join(dt_countries)}")
        if dt_sectors:
            detail_parts.append(f"across {', '.join(dt_sectors)}")
        if dt_cap_arts:
            detail_parts.append(f"capacity: {'; '.join(dt_cap_arts[:2])}")

        ex_str = ""
        if dt_examples:
            h = dt_examples[0][:105].rsplit(" ", 1)[0] if len(dt_examples[0]) > 105 else dt_examples[0]
            ex_str = f'\n  Example: "{h}".'

        deal_type_lines.append(
            f"• **{dt}** — {cnt} transaction event{'s' if cnt > 1 else ''} | "
            + " · ".join(detail_parts)
            + f".{ex_str}"
        )

    deal_intro = (
        f"Transaction activity across the filtered dataset totals {total} articles. "
        f"PPAs account for {ppa_cnt} events, tenders/auctions for {tender_cnt}, "
        f"MOUs for {mou_cnt}, offtake agreements for {aor_cnt}, "
        f"investment/IPO events for {inv_cnt}, and M&A for {ma_cnt}. "
    )
    if ppa_cnt > tender_cnt and ppa_cnt > 0:
        deal_intro += (
            "The dominance of PPA activity signals a maturing procurement environment "
            "where developers are prioritising revenue certainty ahead of financial close. "
        )
    elif tender_cnt > ppa_cnt and tender_cnt > 0:
        deal_intro += (
            "The elevated tender volume indicates governments and utilities are actively "
            "expanding capacity procurement pipelines, which should drive PPA activity in subsequent quarters. "
        )
    if deal_vals:
        deal_intro += (
            f"Disclosed deal values in the current selection include: {'; '.join(deal_vals[:8])}."
        )

    # ── SECTION 4: MAJOR PROJECTS & CAPACITY ANNOUNCEMENTS ───────────────────
    proj_bullets = []
    deal_sub = deal_df.nlargest(min(10, len(deal_df)), "_score") if not deal_df.empty else pd.DataFrame()
    for _, r in deal_sub.iterrows():
        parts = [f"**{r['Headline'][:130]}**"]
        tags  = []
        if r.get("Deal Size"): tags.append(f"Deal: {r['Deal Size']}")
        if r.get("Capacity"):  tags.append(f"Capacity: {r['Capacity']}")
        if r.get("Country"):   tags.append(f"Market: {r['Country']}")
        if r.get("Sector"):    tags.append(f"Sector: {r['Sector']}")
        if r.get("Companies"): tags.append(f"Parties: {r['Companies']}")
        if r.get("Status"):    tags.append(f"Status: {r['Status']}")
        if tags: parts.append(" · ".join(tags))
        proj_bullets.append("• " + " | ".join(parts))

    cap_only = mw_df[mw_df["Deal Size"] == ""].nlargest(min(8, len(mw_df)), "_score") if not mw_df.empty else pd.DataFrame()
    for _, r in cap_only.iterrows():
        parts = [f"**{r['Headline'][:130]}**"]
        tags  = [f"Capacity: {r['Capacity']}", f"Market: {r['Country']}", f"Sector: {r.get('Sector','')}"]
        if r.get("Companies"): tags.append(f"Parties: {r['Companies']}")
        if r.get("Status"):    tags.append(f"Status: {r['Status']}")
        parts.append(" · ".join(t for t in tags if t.split(": ")[1]))
        proj_bullets.append("• " + " | ".join(parts))

    if not proj_bullets:
        for _, r in df2.nlargest(10, "_score").iterrows():
            proj_bullets.append(
                f"• **{r['Headline'][:130]}** | Market: {r['Country']} · Sector: {r['Sector']} · Status: {r.get('Status','—')}"
            )

    # ── SECTION 5: POLICY, REGULATION & PERMITTING ───────────────────────────
    policy_df  = df2[df2["Deal Type"] == "Policy / Reg"]
    chall_df   = df2[df2["Status"] == "Challenged"]
    appr_df    = df2[df2["Status"] == "Approved"]
    tendered_df= df2[df2["Status"] == "Tendered"]

    policy_intro = (
        f"Policy and regulatory activity accounts for {len(policy_df)} articles "
        f"({pct(len(policy_df))}) in the current selection. "
    )
    if len(chall_df) > 0:
        policy_intro += (
            f"{len(chall_df)} project{'s' if len(chall_df)>1 else ''} face contested or challenged status, "
            f"indicating {'elevated' if len(chall_df) > 4 else 'moderate'} regulatory or community resistance. "
        )
    if len(appr_df) > 0:
        policy_intro += f"{len(appr_df)} project{'s' if len(appr_df)>1 else ''} received approvals or permits in the period. "
    if len(tendered_df) > 0:
        policy_intro += (
            f"{len(tendered_df)} auction or tender event{'s' if len(tendered_df)>1 else ''} detected, "
            f"reflecting active government procurement across "
            f"{tendered_df['Country'].value_counts().head(3).index.tolist()} markets. "
        )
    if policy_df.empty and chall_df.empty:
        policy_intro += "No specific policy friction or permitting events detected in the current filter."

    policy_bullets = ["• " + h for h in hl(pd.concat([policy_df, chall_df, appr_df, tendered_df]).drop_duplicates(), 8)]
    if not policy_bullets:
        policy_bullets = ["• No policy-specific articles in the current selection."]

    # ── SECTION 6: COMPANY & DEVELOPER ACTIVITY ──────────────────────────────
    # Categorise participants
    developers  = {"Ørsted","RWE","BP","Shell","TotalEnergies","Equinor","Engie","EDF","Enel","Iberdrola",
                   "Vattenfall","SSE","ScottishPower","E.ON","Innogy","Lightsource","Lightsource BP",
                   "Acciona","AES","NextEra","Duke Energy","Dominion","Invenergy","Terra-Gen",
                   "Enercon","Vestas","Siemens Gamesa","GE Vernova","SGRE","MHI Vestas",
                   "BayWa","Statkraft","Neste","Northland Power","Pattern Energy","Terraform",
                   "Greenko","ReNew Power","Adani","NTPC","CESC","Tata Power","Azure Power",
                   "Sembcorp","Masdar","ACWA Power","Saudi Aramco","JERA","KEPCO",
                   "Intersect Power","Longroad","Avangrid","Eversource","National Grid"}
    offtakers   = {"Microsoft","Google","Amazon","AWS","Meta","Apple","Facebook","Walmart",
                   "Amazon Web Services","Apple","IKEA","Tesla","GM","Ford","Volkswagen",
                   "Starbucks","Bloomberg","Goldman Sachs","JPMorgan","BlackRock","Vanguard",
                   "British Steel","ArcelorMittal","BASF","Air Products","Air Liquide",
                   "Anglo American","Rio Tinto","BHP","Glencore","Lafarge","Holcim"}
    financiers  = {"BlackRock","Macquarie","KKR","Brookfield","Copenhagen Infrastructure",
                   "Green Investment Group","CDPQ","CPP Investments","Stonepeak","Actis",
                   "Quinbrook","Ardian","Infravia","ICG","AMP Capital","EIB","IFC","AIIB","ADB",
                   "World Bank","EBRD","Green Climate Fund","NatWest","HSBC","Barclays",
                   "BNP Paribas","Société Générale","Deutsche Bank","Crédit Agricole"}

    mentioned_cos = [(co, cnt) for co, cnt in co_counter.most_common(30)]

    co_section_lines = []
    if mentioned_cos:
        devs_m   = [(co, c) for co, c in mentioned_cos if co in developers]
        offt_m   = [(co, c) for co, c in mentioned_cos if co in offtakers]
        fin_m    = [(co, c) for co, c in mentioned_cos if co in financiers]
        other_m  = [(co, c) for co, c in mentioned_cos if co not in developers and co not in offtakers and co not in financiers]

        # Narrative intro
        intro_parts = []
        if devs_m:
            intro_parts.append(
                f"Renewable developers/utilities with notable activity include "
                f"{', '.join(co for co, _ in devs_m[:6])} — signalling active project pipelines "
                f"across multiple geographies."
            )
        if offt_m:
            intro_parts.append(
                f"Corporate offtakers prominent in this period: "
                f"{', '.join(co for co, _ in offt_m[:5])} — consistent with the accelerating "
                f"trend of tech and industrial corporates securing long-term clean energy supply."
            )
        if fin_m:
            intro_parts.append(
                f"Financial sponsors and infrastructure capital active in the dataset: "
                f"{', '.join(co for co, _ in fin_m[:5])}, indicating continued investor appetite "
                f"for renewables as an asset class."
            )
        if intro_parts:
            co_section_lines.append("\n".join(intro_parts))
        co_section_lines.append("")

        for co, cnt in mentioned_cos[:20]:
            co_arts = df2[
                df2["Companies"].str.contains(re.escape(co), na=False, case=False) |
                df2["Headline"].str.contains(re.escape(co), na=False, case=False)
            ]
            sectors_seen   = co_arts["Sector"].value_counts().head(3).index.tolist()
            statuses_seen  = co_arts["Status"].value_counts().head(2).index.tolist()
            deal_types_seen= co_arts["Deal Type"].value_counts().head(2).index.tolist()
            countries_seen = co_arts["Country"].value_counts().head(3).index.tolist()
            cap_arts       = co_arts[co_arts["Capacity"] != ""]["Capacity"].tolist()
            deal_arts      = co_arts[co_arts["Deal Size"] != ""]["Deal Size"].tolist()

            detail_parts = [f"sectors: {', '.join(sectors_seen) if sectors_seen else 'General'}"]
            if countries_seen and countries_seen != ["Global"]:
                detail_parts.append(f"markets: {', '.join(countries_seen)}")
            if deal_types_seen:
                detail_parts.append(f"deal types: {', '.join(deal_types_seen)}")
            if statuses_seen:
                detail_parts.append(f"status signals: {', '.join(statuses_seen)}")
            if cap_arts:
                detail_parts.append(f"capacity cited: {'; '.join(cap_arts[:3])}")
            if deal_arts:
                detail_parts.append(f"deal flow: {'; '.join(deal_arts[:2])}")

            co_section_lines.append(
                f"• **{co}** — {cnt} article{'s' if cnt > 1 else ''} | "
                + " | ".join(detail_parts)
            )
    else:
        co_section_lines.append(
            "• No named company entities detected in the current filtered view. "
            "Consider broadening the date range, sector, or region filters to surface developer-level signals."
        )

    # ── SECTION 7: REGIONAL BREAKDOWN ────────────────────────────────────────
    region_lines = []
    for region, rdf in df2.groupby("Region"):
        top_c        = rdf["Country"].value_counts().head(5).index.tolist()
        top_sec      = rdf["Sector"].value_counts().head(3).index.tolist()
        top_dt       = rdf["Deal Type"].value_counts().head(2).index.tolist()
        cap_in_reg   = rdf[rdf["Capacity"] != ""]["Capacity"].tolist()
        deals_reg    = rdf[rdf["Deal Size"] != ""]["Deal Size"].tolist()
        prop_reg     = len(rdf[rdf["Status"] == "Proposed"])
        comm_reg     = len(rdf[rdf["Status"] == "Commissioned"])
        chall_reg    = len(rdf[rdf["Status"] == "Challenged"])
        tender_reg   = len(rdf[rdf["Status"] == "Tendered"])
        contracted_r = len(rdf[rdf["Status"] == "Contracted"])

        region_lines.append(f"• **{region}** — {len(rdf)} articles ({pct(len(rdf))})")
        region_lines.append(
            f"  Lead markets: {', '.join(top_c)}. "
            f"Dominant sectors: {', '.join(top_sec)}. "
            f"Deal types: {', '.join(top_dt)}. "
            f"Pipeline: {prop_reg} proposed, {tender_reg} tendered, {contracted_r} contracted, "
            f"{comm_reg} commissioned, {chall_reg} challenged."
            + (f" Capacity cited: {'; '.join(cap_in_reg[:4])}." if cap_in_reg else "")
            + (f" Deal flow: {'; '.join(deals_reg[:3])}." if deals_reg else "")
        )

    # ── SECTION 8: MARKET OUTLOOK & FORWARD SIGNALS ──────────────────────────
    reg_risk = (
        "elevated" if challenged > 5
        else "moderate" if challenged > 2
        else "low"
    )

    outlook_bullets = [
        f"• **Pipeline momentum:** The overall pipeline is {momentum_signal}. "
        f"With {proposed + tendered} projects at proposal or tender stage, "
        f"{contracted} contracted, {under_c} under construction, and {commissioned} "
        f"recently commissioned, the market is demonstrating "
        f"{'strong end-to-end delivery velocity' if commissioned >= 3 else 'active origination but watch delivery timelines'}.",

        f"• **Deal origination:** {ppa_cnt} PPA events, {tender_cnt} auctions/tenders, "
        f"and {mou_cnt} MOUs signal "
        + ("robust bilateral procurement momentum alongside government-led capacity auctions."
           if ppa_cnt > 0 and tender_cnt > 0
           else "active but concentrated deal origination — diversification across deal structures is limited."
           ),

        f"• **Regulatory environment:** Risk is assessed as {reg_risk} — "
        f"{challenged} contested or challenged project{'s' if challenged != 1 else ''} in the dataset. "
        + ("Grid connection constraints, permitting delays, and community opposition represent the "
           "primary near-term threats to project delivery in this geography and period."
           if challenged > 3
           else "The regulatory environment appears broadly supportive of new capacity additions."),

        f"• **Developer & investor watch:** {', '.join(top_cos[:8])} are the highest-frequency "
        f"participants in this period. Monitor for capacity announcements, PPA signings, "
        f"M&A activity, and financial close events.",
    ]

    if cap_total_mw > 0:
        outlook_bullets.append(
            f"• **Capacity pipeline:** {len(mw_df)} articles reference explicit MW or GW figures. "
            f"Aggregate announced capacity stands at approximately **{cap_total_mw:,.0f} MW** across "
            f"disclosed projects. Actual delivered capacity will depend on permitting approvals, "
            f"grid connection queue position, and financing milestones."
        )
    if deal_vals:
        outlook_bullets.append(
            f"• **Capital deployment:** {len(deal_df)} deal-citing articles disclose transactions "
            f"totalling {', '.join(deal_vals[:8])}. "
            f"Deal density and disclosed values indicate continued strong capital flow into the clean energy sector."
        )
    if inv_cnt > 0 or ma_cnt > 0:
        outlook_bullets.append(
            f"• **M&A & investment signals:** {inv_cnt} investment/IPO and {ma_cnt} M&A events "
            f"detected. Elevated M&A and fundraising activity typically presages accelerated "
            f"project development in the 12–18 months following close."
        )

    # ── Assemble final document ───────────────────────────────────────────────
    doc_parts = [
        "## 1. Executive Summary\n\n",
        exec_summary, "\n\n",

        "## 2. Sector Dynamics\n\n",
        f"Coverage across {total} articles spans {len(top_sectors)} sector{'s' if len(top_sectors)>1 else ''}: "
        f"{', '.join(f'{s} ({c})' for s, c in top_sectors.items())}.\n\n",
        "\n".join(sector_lines), "\n\n",

        "## 3. Deal Flow & Transaction Analysis\n\n",
        deal_intro + "\n\n",
        "\n".join(deal_type_lines), "\n\n",

        "## 4. Major Projects, Deals & Capacity Announcements\n\n",
        "\n".join(proj_bullets) if proj_bullets else "• No capacity or deal-cited articles in the current selection.", "\n\n",

        "## 5. Policy, Regulation & Procurement\n\n",
        policy_intro + "\n\n",
        "\n".join(policy_bullets), "\n\n",

        "## 6. Developer, Offtaker & Investor Activity\n\n",
        "\n".join(co_section_lines), "\n\n",

        "## 7. Regional Breakdown\n\n",
        "\n".join(region_lines), "\n\n",

        "## 8. Market Outlook & Forward Signals\n\n",
        "\n".join(outlook_bullets), "\n",
    ]

    return "".join(doc_parts)

def main():
    st.set_page_config(
        page_title="Global Energy Intelligence Platform",
        page_icon="⚡",
        layout="wide",
        initial_sidebar_state="expanded",
    )
    st.markdown(CUSTOM_CSS, unsafe_allow_html=True)

    # ── Additional CSS for dual-platform layout ───────────────────────────────
    st.markdown("""
    <style>
    .platform-tab-active {
        background: linear-gradient(135deg, #0047e1, #00b4ff) !important;
        color: #fff !important; border-color: transparent !important;
        box-shadow: 0 4px 18px rgba(0,71,225,0.35) !important;
    }
    .platform-tab-re-active {
        background: linear-gradient(135deg, #00aa44, #00e676) !important;
        color: #fff !important; border-color: transparent !important;
        box-shadow: 0 4px 18px rgba(0,200,80,0.35) !important;
    }
    .mode-pill {
        display:inline-flex;align-items:center;gap:.5rem;
        padding:.38rem 1rem;border-radius:20px;
        font-family:'DM Mono',monospace;font-size:.7rem;letter-spacing:.06em;
        font-weight:600;cursor:pointer;transition:all .2s;
    }
    </style>
    """, unsafe_allow_html=True)

    # ── Timezone detection via browser JS ────────────────────────────────────
    # Reads the browser IANA tz and writes it into query params.
    # Default fallback = Asia/Kolkata (IST, UTC+5:30).
    tz_js = """
    <script>
    (function() {
        try {
            var tz = Intl.DateTimeFormat().resolvedOptions().timeZone || 'Asia/Kolkata';
            var params = new URLSearchParams(window.location.search);
            if (!params.get('tz') || params.get('tz') === 'UTC') {
                params.set('tz', tz);
                window.history.replaceState({}, '', window.location.pathname + '?' + params.toString());
            }
        } catch(e) {}
    })();
    </script>
    """
    st.markdown(tz_js, unsafe_allow_html=True)

    # ── Sidebar: aggressive force-open on every load ─────────────────────────
    # Strategy: three-layer approach
    #  1. Nuke ALL localStorage keys that relate to sidebar/collapsed state
    #  2. Force-click the sidebar open button if sidebar is found collapsed
    #  3. Inject a visible "☰ Filters" floating button as a permanent fallback
    sidebar_btn_js = """
    <script>
    (function() {

        // ── LAYER 1: Wipe every localStorage key that could keep sidebar closed ──
        function nukeLocalStorage() {
            try {
                var toDelete = [];
                for (var i = 0; i < localStorage.length; i++) {
                    var k = localStorage.key(i);
                    if (!k) continue;
                    var kl = k.toLowerCase();
                    if (
                        kl.includes('sidebar') ||
                        kl.includes('collapsed') ||
                        kl.includes('stSidebar') ||
                        kl.includes('streamlit') ||
                        kl.includes('layout') ||
                        kl.includes('toggled')
                    ) {
                        toDelete.push(k);
                    }
                }
                toDelete.forEach(function(k) { localStorage.removeItem(k); });
            } catch(e) {}

            // Also try sessionStorage
            try {
                var sToDelete = [];
                for (var j = 0; j < sessionStorage.length; j++) {
                    var sk = sessionStorage.key(j);
                    if (!sk) continue;
                    var skl = sk.toLowerCase();
                    if (
                        skl.includes('sidebar') ||
                        skl.includes('collapsed') ||
                        skl.includes('streamlit')
                    ) {
                        sToDelete.push(sk);
                    }
                }
                sToDelete.forEach(function(sk) { sessionStorage.removeItem(sk); });
            } catch(e) {}
        }
        nukeLocalStorage();

        // ── LAYER 2: Force-click the sidebar open if it appears collapsed ────────
        function forceSidebarOpen() {
            var sidebar = document.querySelector('[data-testid="stSidebar"]');
            if (!sidebar) return false;

            // Check if sidebar is actually collapsed (very narrow or off-screen)
            var rect = sidebar.getBoundingClientRect();
            var isCollapsed = rect.width < 60 || rect.left < -100;

            if (isCollapsed) {
                // Try every selector Streamlit uses for the reopen button
                var selectors = [
                    '[data-testid="stSidebarCollapsedControl"] button',
                    'div[data-testid="collapsedControl"] button',
                    'button[aria-label="Open sidebar"]',
                    'button[aria-label*="sidebar"]',
                    '[data-testid="collapsedControl"] button',
                ];
                for (var i = 0; i < selectors.length; i++) {
                    var btns = document.querySelectorAll(selectors[i]);
                    if (btns.length > 0) {
                        btns[0].click();
                        return true;
                    }
                }
            }
            return false;
        }

        // ── LAYER 3: Floating "☰ Filters" emergency button ───────────────────────
        function injectFiltersBtn() {
            var old = document.getElementById('_gdci_sidebar_btn');
            if (old) old.remove();

            var btn = document.createElement('button');
            btn.id = '_gdci_sidebar_btn';
            btn.innerHTML = '&#9776; Filters';
            btn.title = 'Open filter sidebar';
            btn.style.cssText = [
                'position:fixed',
                'top:12px',
                'left:12px',
                'z-index:2147483647',
                'background:linear-gradient(135deg,#0047e1,#00b4ff)',
                'color:#fff',
                'border:none',
                'border-radius:8px',
                'padding:7px 14px',
                'font-family:Syne,sans-serif',
                'font-weight:700',
                'font-size:0.82rem',
                'letter-spacing:0.04em',
                'cursor:pointer',
                'box-shadow:0 4px 16px rgba(0,71,225,0.55)',
                'display:none',
                'align-items:center',
                'gap:6px',
                'transition:opacity 0.2s,transform 0.2s',
            ].join(';');

            btn.onmouseenter = function() { btn.style.opacity='0.85'; btn.style.transform='translateY(-1px)'; };
            btn.onmouseleave = function() { btn.style.opacity='1';    btn.style.transform='translateY(0)';    };

            btn.onclick = function() {
                nukeLocalStorage();
                var selectors = [
                    '[data-testid="stSidebarCollapsedControl"] button',
                    'div[data-testid="collapsedControl"] button',
                    'button[aria-label="Open sidebar"]',
                    'button[aria-label*="sidebar"]',
                    '[data-testid="collapsedControl"] button',
                ];
                for (var i = 0; i < selectors.length; i++) {
                    var btns = document.querySelectorAll(selectors[i]);
                    if (btns.length > 0) { btns[0].click(); return; }
                }
            };

            document.body.appendChild(btn);

            // Poll: show button only when sidebar is collapsed
            function checkSidebar() {
                var sidebar = document.querySelector('[data-testid="stSidebar"]');
                if (!sidebar) { btn.style.display = 'flex'; return; }
                var rect = sidebar.getBoundingClientRect();
                var collapsed = rect.width < 60 || rect.left < -100;
                btn.style.display = collapsed ? 'flex' : 'none';
            }

            setInterval(checkSidebar, 300);
            checkSidebar();
        }

        // ── Boot sequence ────────────────────────────────────────────────────────
        function boot() {
            nukeLocalStorage();
            injectFiltersBtn();

            // Try to force open immediately, then retry a few times
            // (Streamlit renders the DOM progressively so we need retries)
            var attempts = 0;
            var maxAttempts = 12;
            var interval = setInterval(function() {
                nukeLocalStorage();
                var opened = forceSidebarOpen();
                attempts++;
                if (opened || attempts >= maxAttempts) clearInterval(interval);
            }, 400);
        }

        if (document.readyState === 'loading') {
            document.addEventListener('DOMContentLoaded', boot);
        } else {
            setTimeout(boot, 300);
        }

    })();
    </script>
    """
    st.markdown(sidebar_btn_js, unsafe_allow_html=True)

    _DEFAULT_TZ = "Asia/Kolkata"
    try:
        from zoneinfo import ZoneInfo
        _tz_raw = st.query_params.get("tz", _DEFAULT_TZ)
        if not _tz_raw or _tz_raw in ("UTC", "undefined", "null", ""):
            _tz_raw = _DEFAULT_TZ
        try:
            _user_tz = ZoneInfo(_tz_raw)
            _tz_str  = _tz_raw
        except Exception:
            _user_tz = ZoneInfo(_DEFAULT_TZ)
            _tz_str  = _DEFAULT_TZ
    except ImportError:
        _user_tz = None
        _tz_str  = _DEFAULT_TZ

    def now_local():
        from datetime import timezone as _dt_tz
        utc_now = datetime.now(_dt_tz.utc)
        if _user_tz:
            return utc_now.astimezone(_user_tz)
        return utc_now

    def fmt_local(dt=None):
        d = dt or now_local()
        parts    = _tz_str.split("/")
        tz_label = parts[-1].replace("_", " ") if len(parts) > 1 else _tz_str
        offset   = d.strftime("%z")
        if offset:
            sign  = offset[0]
            hh    = offset[1:3]
            mm    = offset[3:5]
            off_s = f"UTC{sign}{hh}:{mm}" if mm != "00" else f"UTC{sign}{hh}"
        else:
            off_s = ""
        return d.strftime("%A, %d %B %Y  \u00b7  %H:%M") + f"  {tz_label} ({off_s})"

    with st.sidebar:
        st.markdown(
            '<div style="padding:.9rem 0 .4rem;">'
            '<div style="font-family:Syne,sans-serif;font-size:.82rem;font-weight:700;color:#b8c8e0;'
            'letter-spacing:.02em;margin-bottom:.06rem;">Global Energy</div>'
            '<div style="font-family:Syne,sans-serif;font-size:.82rem;font-weight:700;color:#00b4ff;'
            'letter-spacing:.02em;margin-bottom:.28rem;">Intelligence Platform</div>'
            '<div style="font-family:monospace;font-size:.58rem;letter-spacing:.08em;color:#2a3e60;">'
            'Built By <span style="color:#0047e1;font-weight:600;">Sharugh</span></div>'
            '</div>',
            unsafe_allow_html=True,
        )
        st.divider()

        # ── PLATFORM MODE SELECTOR ─────────────────────────────────────────────
        st.markdown(
            '<div style="font-size:.65rem;color:#2a4060;letter-spacing:.1em;'
            'text-transform:uppercase;font-family:monospace;margin-bottom:.4rem;">'
            '▸ INTELLIGENCE MODULE</div>',
            unsafe_allow_html=True,
        )
        platform_mode = st.radio(
            "Module",
            ["🏢 Data Center Markets", "🌱 Renewables Power Markets"],
            index=0,
            label_visibility="collapsed",
            help="Switch between Data Center Intelligence and Renewables Power Markets.",
        )
        st.markdown('<div style="height:.2rem;"></div>', unsafe_allow_html=True)

        # Colour indicator for active module
        if "Renewables" in platform_mode:
            st.markdown(
                '<div style="background:linear-gradient(90deg,rgba(0,200,80,0.12),transparent);'
                'border-left:3px solid #00e676;border-radius:0 6px 6px 0;'
                'padding:.38rem .75rem;font-family:DM Mono,monospace;font-size:.62rem;color:#00e676;'
                'letter-spacing:.08em;margin-bottom:.5rem;">🌱 RENEWABLES ACTIVE</div>',
                unsafe_allow_html=True,
            )
        else:
            st.markdown(
                '<div style="background:linear-gradient(90deg,rgba(0,71,225,0.12),transparent);'
                'border-left:3px solid #0047e1;border-radius:0 6px 6px 0;'
                'padding:.38rem .75rem;font-family:DM Mono,monospace;font-size:.62rem;color:#0047e1;'
                'letter-spacing:.08em;margin-bottom:.5rem;">🏢 DATA CENTER ACTIVE</div>',
                unsafe_allow_html=True,
            )

        # ── App-wide theme switch: BLUE (Data Center) vs GREEN (Renewables) ──
        # Overrides the CSS variables declared in CUSTOM_CSS's :root block,
        # so every themed control (tabs, buttons, focus rings, chart glows,
        # banner accents, chips) flips color with the active module.
        _set_chart_theme("Renewables" in platform_mode)
        if "Renewables" in platform_mode:
            st.markdown(
                """
                <style>
                :root {
                    --accent:        #00a846;
                    --accent2:       #00e676;
                    --accent-rgb:    0, 168, 70;
                    --accent2-rgb:   0, 230, 118;
                    --accent-soft:   rgba(0, 168, 70, 0.18);
                    --accent-glow:   rgba(0, 168, 70, 0.35);
                    --accent-faint:  rgba(0, 168, 70, 0.12);
                    --tag-bg:        #0f3322;
                    --tag-text:      #7eebac;
                    --banner-bg:     linear-gradient(135deg, #071a0f 0%, #0b2a18 45%, #071a0f 100%);
                }
                </style>
                """,
                unsafe_allow_html=True,
            )
        st.divider()


        time_opt = st.radio(
            "", ["Latest (all)", "Past 30 days", "Past 14 days", "Past 7 days", "Custom Range"],
            index=0, label_visibility="collapsed",
        )
        days_map = {"Latest (all)": None, "Past 30 days": 30, "Past 14 days": 14, "Past 7 days": 7}
        sel_days = days_map.get(time_opt, None)

        custom_start = None
        custom_end   = None
        if time_opt == "Custom Range":
            today = datetime.now().date()
            c1, c2 = st.columns(2)
            with c1:
                custom_start = st.date_input(
                    "From", value=today - timedelta(days=30),
                    max_value=today, label_visibility="visible",
                )
            with c2:
                custom_end = st.date_input(
                    "To", value=today,
                    max_value=today, label_visibility="visible",
                )
            if custom_start and custom_end and custom_start > custom_end:
                st.error("Start date must be before end date.")
                custom_start, custom_end = custom_end, custom_start
            st.markdown(
                f'<div style="font-size:.68rem;color:#1a2e50;margin-top:-.3rem;margin-bottom:.4rem;">'
                f'{custom_start.strftime("%d %b %Y") if custom_start else ""}'
                f' → '
                f'{custom_end.strftime("%d %b %Y") if custom_end else ""}'
                f'</div>',
                unsafe_allow_html=True,
            )

        # ── Scrape Depth ──────────────────────────────────────────────────────
        st.markdown(
            '<div style="font-size:.72rem;color:#3a5480;letter-spacing:.07em;'
            'text-transform:uppercase;margin:.9rem 0 .2rem;">🔎 Scrape Depth (pages per DCD term)</div>',
            unsafe_allow_html=True,
        )
        max_pages = st.slider(
            "Scrape depth", min_value=1, max_value=100, value=10, step=1,
            label_visibility="collapsed",
            help="Higher = more articles but slower scan. 10 is fast, 50+ is thorough, 100 is maximum.",
        )
        st.markdown(
            f'<div style="font-size:.68rem;color:#1a2e50;margin-top:-.3rem;margin-bottom:.4rem;">'
            f'Up to {max_pages} page{"s" if max_pages != 1 else ""} per DCD channel scraped</div>',
            unsafe_allow_html=True,
        )

        use_html  = True
        use_rss   = True
        use_gn    = True

        st.divider()

        # ── FILTER PANEL — mode-aware (switches between DC and RE) ──────────
        if "_filter_ver" not in st.session_state:
            st.session_state["_filter_ver"] = 0
        _fv = st.session_state["_filter_ver"]
        def _fk(name): return f"{name}_v{_fv}"

        _is_re_mode = "Renewables" in platform_mode

        # ── Pull live data for filter option lists ─────────────────────────────
        _dc_loaded = "df_full"    in st.session_state and st.session_state.df_full    is not None
        _re_loaded = "re_df_full" in st.session_state and st.session_state.re_df_full is not None

        if _is_re_mode and _re_loaded:
            _src = st.session_state.re_df_full
            all_regions_av   = sorted(_src["Region"].dropna().unique().tolist())
            all_countries_av = sorted(_src["Country"].dropna().unique().tolist())
            # DC-only lists (safe defaults — never rendered in RE mode)
            all_topics_av    = sorted(TOPIC_COLORS.keys())
            all_sents_av     = ["Opened / Live", "Approved", "Proposed",
                                 "Under Construction", "Challenged", "News"]
            all_companies_av = KNOWN_COMPANIES
            _all_iso_in_data = []
        elif not _is_re_mode and _dc_loaded:
            _src = st.session_state.df_full
            all_regions_av   = sorted(_src["Region"].dropna().unique().tolist())
            all_countries_av = sorted(_src["Country"].dropna().unique().tolist())
            all_topics_av    = sorted(_src["Topic"].dropna().unique().tolist())
            all_sents_av     = sorted(_src["Sentiment"].dropna().unique().tolist())
            _all_co_raw      = []
            for _v in _src["Companies"]:
                if _v:
                    _all_co_raw.extend([c.strip() for c in str(_v).split(",")])
            all_companies_av = sorted(set(c for c in _all_co_raw if c))
            _all_iso_in_data = sorted(set(
                v for v in _src["ISO / RTO"].tolist()
                if v and str(v) != "nan" and str(v).strip()
            )) if "ISO / RTO" in _src.columns else []
        else:
            all_regions_av   = sorted(REGION_COLORS.keys())
            all_countries_av = sorted(COUNTRY_TO_REGION.keys())
            all_topics_av    = sorted(TOPIC_COLORS.keys())
            all_sents_av     = ["Opened / Live", "Approved", "Proposed",
                                 "Under Construction", "Challenged", "News"]
            all_companies_av = KNOWN_COMPANIES
            _all_iso_in_data = []

        # ── Colour-coded header that switches with the active module ───────────
        _fa   = "#00e676" if _is_re_mode else "#0047e1"
        _fbg  = "rgba(0,200,80,0.12)" if _is_re_mode else "rgba(0,71,225,0.12)"
        _flbl = "🌱 Renewables Filters" if _is_re_mode else "🏢 Data Center Filters"
        st.markdown(
            f'<div style="background:linear-gradient(90deg,{_fbg},transparent);'
            f'border-left:3px solid {_fa};border-radius:0 6px 6px 0;'
            f'padding:.45rem .75rem;margin-bottom:.6rem;">'
            f'<span style="font-family:Syne,sans-serif;font-weight:700;color:#b8c8e0;'
            f'font-size:.8rem;letter-spacing:.05em;">{_flbl}</span>'
            f'</div>',
            unsafe_allow_html=True,
        )

        # ── GEOGRAPHY — shared by both modes ───────────────────────────────────
        st.markdown(
            '<div style="font-size:.65rem;color:#2a4060;letter-spacing:.1em;'
            'text-transform:uppercase;font-family:monospace;margin:.7rem 0 .35rem .05rem;">'
            '▸ GEOGRAPHY</div>',
            unsafe_allow_html=True,
        )
        st.markdown(
            '<div style="font-size:.72rem;color:#3a5480;letter-spacing:.06em;'
            'text-transform:uppercase;margin-bottom:.2rem;">🌐 Region</div>',
            unsafe_allow_html=True,
        )
        sel_regions = st.multiselect(
            "Region", all_regions_av, default=[],
            placeholder="All regions", label_visibility="collapsed",
            key=_fk("f_regions"),
        )

        all_world_countries = sorted(set(list(COUNTRY_TO_REGION.keys()) + all_countries_av))
        if sel_regions:
            world_pool = sorted([c for c in all_world_countries
                                  if COUNTRY_TO_REGION.get(c, "Global") in sel_regions])
        else:
            world_pool = all_world_countries

        st.markdown(
            '<div style="font-size:.72rem;color:#3a5480;letter-spacing:.06em;'
            'text-transform:uppercase;margin:.75rem 0 .2rem;">🌍 Country</div>',
            unsafe_allow_html=True,
        )
        sel_countries = st.multiselect(
            "Country", world_pool, default=[],
            placeholder="All countries", label_visibility="collapsed",
            key=_fk("f_countries"),
        )

        # ══════════════════════════════════════════════════════════════════════
        #  MODE-SPECIFIC CONTENT FILTERS
        # ══════════════════════════════════════════════════════════════════════
        if _is_re_mode:
            # ── RENEWABLES FILTERS ─────────────────────────────────────────────
            _lbl = lambda txt: st.markdown(
                f'<div style="font-size:.72rem;color:#2a6040;letter-spacing:.06em;'
                f'text-transform:uppercase;margin:.75rem 0 .2rem;">{txt}</div>',
                unsafe_allow_html=True,
            )

            # ── State / Province (RE mode) ────────────────────────────────────
            re_state_pool = []
            for c in (sel_countries if sel_countries else list(COUNTRY_STATES.keys())):
                re_state_pool.extend(COUNTRY_STATES.get(c, []))
            re_state_pool = sorted(set(re_state_pool))

            re_state_sel = []
            if re_state_pool:
                _lbl("📍 State / Province")
                re_state_sel = st.multiselect(
                    "RE State", re_state_pool, default=[],
                    placeholder="All states/provinces",
                    label_visibility="collapsed",
                    key=_fk("f_re_states"),
                )

            st.markdown(
                '<div style="font-size:.65rem;color:#1a4a2e;letter-spacing:.1em;'
                'text-transform:uppercase;font-family:monospace;margin:.9rem 0 .35rem .05rem;">'
                '▸ CONTENT FILTERS</div>',
                unsafe_allow_html=True,
            )

            # ── Sector (drives which feeds are fetched) ────────────────────────
            _lbl("🌱 Sector")
            # Build options: standard sectors + EPC Companies
            _sector_options = list(RE_TOPIC_COLORS.keys())  # includes "EPC Companies"
            re_sector_sel = st.multiselect(
                "Sector",
                options=_sector_options,
                default=[],
                placeholder="All sectors  (Solar · Wind · Storage · OSW · H₂ · EPC…)",
                label_visibility="collapsed",
                key=_fk("f_re_sector"),
            )
            _epc_mode_active = "EPC Companies" in re_sector_sel
            # Show feed count hint
            if re_sector_sel:
                _non_epc = [s for s in re_sector_sel if s != "EPC Companies"]
                n_feeds = sum(len(RE_FEED_REGISTRY.get(s, [])) for s in _non_epc) + len(RE_CROSS_FEEDS)
                _epc_note = f" + EPC company feeds" if _epc_mode_active else ""
                st.markdown(
                    f'<div style="font-size:.62rem;color:#2a6040;font-family:monospace;'
                    f'margin:.1rem 0 .4rem .05rem;">↳ {n_feeds} feeds{_epc_note} · '
                    + " · ".join(re_sector_sel)
                    + '</div>',
                    unsafe_allow_html=True,
                )
            else:
                total_feeds = sum(len(v) for v in RE_FEED_REGISTRY.values()) + len(RE_CROSS_FEEDS)
                st.markdown(
                    f'<div style="font-size:.62rem;color:#1a4a2e;font-family:monospace;'
                    f'margin:.1rem 0 .4rem .05rem;">↳ All {total_feeds} feeds across 6 sectors + EPC</div>',
                    unsafe_allow_html=True,
                )

            # ── EPC Company Selector (appears only when EPC Companies is chosen) ──
            re_epc_sel_companies = []
            if _epc_mode_active:
                st.markdown(
                    '<div style="background:linear-gradient(90deg,rgba(255,100,0,0.1),transparent);'
                    'border-left:3px solid #ff6400;border-radius:0 6px 6px 0;'
                    'padding:.4rem .7rem;margin:.5rem 0 .4rem;font-family:DM Mono,monospace;'
                    'font-size:.62rem;color:#ff6400;letter-spacing:.08em;">🏗️ EPC MODE ACTIVE</div>',
                    unsafe_allow_html=True,
                )
                _lbl("🏗️ Select EPC Companies")
                # Sorted company list for the picker
                _epc_company_list = sorted(EPC_COMPANIES.keys())
                re_epc_sel_companies = st.multiselect(
                    "EPC Companies",
                    options=_epc_company_list,
                    default=[],
                    placeholder="All EPC companies (leave blank) or pick specific ones…",
                    label_visibility="collapsed",
                    key=_fk("f_re_epc_sel"),
                    help=(
                        "Leave blank to track ALL EPC/project companies, or select specific ones. "
                        "The scraper will fire targeted Google News feeds for each selected company "
                        "across renewables, storage, hydrogen and power sectors globally."
                    ),
                )
                if re_epc_sel_companies:
                    st.markdown(
                        f'<div style="font-size:.62rem;color:#ff6400;font-family:monospace;'
                        f'margin:.1rem 0 .35rem .05rem;">'
                        f'↳ Targeting {len(re_epc_sel_companies)} companies · '
                        f'{", ".join(re_epc_sel_companies[:4])}'
                        + (f' +{len(re_epc_sel_companies)-4} more' if len(re_epc_sel_companies) > 4 else '')
                        + '</div>',
                        unsafe_allow_html=True,
                    )
                else:
                    st.markdown(
                        f'<div style="font-size:.62rem;color:#ff6400;font-family:monospace;'
                        f'margin:.1rem 0 .35rem .05rem;">'
                        f'↳ All {len(EPC_COMPANIES)} EPC companies tracked globally</div>',
                        unsafe_allow_html=True,
                    )

            # ── Deal Type ─────────────────────────────────────────────────────
            _lbl("📋 Deal Type")
            re_deal_type_sel = st.multiselect(
                "Deal Type",
                options=list(RE_DEAL_TYPE_COLORS.keys()),
                default=[],
                placeholder="PPA · MOU · Tender · Investment · M&A…",
                label_visibility="collapsed",
                key=_fk("f_re_deal"),
            )

            # ── Project Status ────────────────────────────────────────────────
            _lbl("📊 Project Status")
            _re_statuses = [
                "Commissioned", "Financed", "Under Construction",
                "Approved", "Tendered", "Contracted",
                "Proposed", "Challenged", "News",
            ]
            re_status_sel = st.multiselect(
                "RE Status", _re_statuses, default=[],
                placeholder="All statuses",
                label_visibility="collapsed",
                key=_fk("f_re_status"),
            )

            # ── Keyword ───────────────────────────────────────────────────────
            _lbl("🔤 Keyword")
            re_keyword = st.text_input(
                "RE Keyword",
                placeholder="e.g. 500MW, India, Ørsted, PPA...",
                label_visibility="collapsed",
                key=_fk("f_re_keyword"),
            )

            st.markdown('<div style="margin-top:.7rem;"></div>', unsafe_allow_html=True)
            if st.button("✕ Clear Renewables Filters", use_container_width=True,
                         key="clear_re_filters_btn"):
                st.session_state["_filter_ver"] = _fv + 1
                st.session_state.pop("re_filters", None)
                st.rerun()

            st.session_state["re_filters"] = {
                "sectors":          re_sector_sel,
                "epc_mode":         _epc_mode_active,
                "epc_companies":    re_epc_sel_companies,
                "deal_types":       re_deal_type_sel,
                "statuses":         re_status_sel,
                "sources":          [],
                "keyword":          re_keyword,
                "regions":          sel_regions,
                "countries":        sel_countries,
                "states":           re_state_sel,
            }

            # Null-out DC-only vars so nothing downstream crashes
            news_type_sel = [NEWS_TYPE_CONSTRUCTION, NEWS_TYPE_GENERAL]
            keyword       = ""
            sel_companies = []
            sel_topics    = []
            sel_sents     = []
            sel_states    = []
            sel_iso_rto   = []
            min_mw        = 0
            # Safe defaults for RE-only vars (prevent NameError in DC path)
            re_epc_sel_companies = []
            _epc_mode_active     = False

        else:
            # ── DATA CENTER FILTERS ────────────────────────────────────────────
            # News Type selector (DCD channels — DC mode only)
            st.markdown(
                '<div style="font-size:.72rem;color:#3a5480;letter-spacing:.07em;'
                'text-transform:uppercase;margin:.4rem 0 .2rem;">📰 News Type</div>',
                unsafe_allow_html=True,
            )
            news_type_sel = st.multiselect(
                "News Type",
                [NEWS_TYPE_CONSTRUCTION, NEWS_TYPE_GENERAL],
                default=[NEWS_TYPE_CONSTRUCTION, NEWS_TYPE_GENERAL],
                placeholder="Select channel(s)…",
                label_visibility="collapsed",
                help=(
                    "Construction → DCD Data Center Construction Channel\n"
                    "General News → DCD General News feed"
                ),
            )
            _nt_labels = {
                NEWS_TYPE_CONSTRUCTION: "🏗️ Construction Channel",
                NEWS_TYPE_GENERAL:      "📰 General News",
            }
            if news_type_sel:
                st.markdown(
                    '<div style="font-size:.68rem;color:#1a2e50;margin-top:-.3rem;margin-bottom:.4rem;">'
                    + " · ".join(_nt_labels[n] for n in news_type_sel if n in _nt_labels)
                    + '</div>',
                    unsafe_allow_html=True,
                )
            else:
                st.warning("Select at least one news channel.")

            # State / Province (cascades from country)
            state_pool = []
            for c in (sel_countries if sel_countries else world_pool):
                state_pool.extend(COUNTRY_STATES.get(c, []))
            state_pool = sorted(set(state_pool))

            sel_states = []
            if state_pool:
                st.markdown(
                    '<div style="font-size:.72rem;color:#3a5480;letter-spacing:.06em;'
                    'text-transform:uppercase;margin:.75rem 0 .2rem;">📍 State / Province</div>',
                    unsafe_allow_html=True,
                )
                sel_states = st.multiselect(
                    "State", state_pool, default=[],
                    placeholder="All states/provinces", label_visibility="collapsed",
                    key=_fk("f_states"),
                )

            # ISO / RTO / Grid
            _iso_pool = sorted(set(
                _all_iso_in_data + list(US_ISO_RTO.keys()) + list(GLOBAL_GRID_OPERATORS.keys())
            ))
            st.markdown(
                '<div style="font-size:.72rem;color:#3a5480;letter-spacing:.06em;'
                'text-transform:uppercase;margin:.75rem 0 .2rem;">⚡ ISO / RTO / Grid</div>',
                unsafe_allow_html=True,
            )
            sel_iso_rto = st.multiselect(
                "ISO/RTO", _iso_pool, default=[],
                placeholder="All grid operators", label_visibility="collapsed",
                key=_fk("f_iso_rto"),
            )

            st.markdown(
                '<div style="font-size:.65rem;color:#2a4060;letter-spacing:.1em;'
                'text-transform:uppercase;font-family:monospace;margin:.9rem 0 .35rem .05rem;">'
                '▸ CONTENT</div>',
                unsafe_allow_html=True,
            )

            # Keyword
            st.markdown(
                '<div style="font-size:.72rem;color:#3a5480;letter-spacing:.06em;'
                'text-transform:uppercase;margin-bottom:.2rem;">🔤 Keyword</div>',
                unsafe_allow_html=True,
            )
            keyword = st.text_input(
                "Keyword", placeholder="e.g. 500MW, Texas, nuclear, AWS...",
                label_visibility="collapsed", key=_fk("f_keyword"),
            )

            # Company
            st.markdown(
                '<div style="font-size:.72rem;color:#3a5480;letter-spacing:.06em;'
                'text-transform:uppercase;margin:.75rem 0 .2rem;">🏢 Company</div>',
                unsafe_allow_html=True,
            )
            full_co_pool = sorted(set(all_companies_av + KNOWN_COMPANIES))
            sel_companies = st.multiselect(
                "Company", full_co_pool, default=[],
                placeholder="All companies — type to search",
                label_visibility="collapsed", key=_fk("f_companies"),
            )

            # Topic
            st.markdown(
                '<div style="font-size:.72rem;color:#3a5480;letter-spacing:.06em;'
                'text-transform:uppercase;margin:.75rem 0 .2rem;">🏷️ Topic</div>',
                unsafe_allow_html=True,
            )
            sel_topics = st.multiselect(
                "Topic", all_topics_av, default=[],
                placeholder="All topics", label_visibility="collapsed",
                key=_fk("f_topics"),
            )

            # Project Status
            st.markdown(
                '<div style="font-size:.72rem;color:#3a5480;letter-spacing:.06em;'
                'text-transform:uppercase;margin:.75rem 0 .2rem;">📊 Project Status</div>',
                unsafe_allow_html=True,
            )
            sel_sents = st.multiselect(
                "Status", all_sents_av, default=[],
                placeholder="All statuses", label_visibility="collapsed",
                key=_fk("f_sents"),
            )

            # Min Capacity
            st.markdown(
                '<div style="font-size:.72rem;color:#3a5480;letter-spacing:.06em;'
                'text-transform:uppercase;margin:.75rem 0 .2rem;">⚡ Min Capacity (MW)</div>',
                unsafe_allow_html=True,
            )
            min_mw = st.number_input(
                "Min MW", min_value=0, value=0,
                step=10, label_visibility="collapsed", key=_fk("f_min_mw"),
            )

            st.markdown('<div style="margin-top:.6rem;"></div>', unsafe_allow_html=True)
            if st.button("✕ Clear All Filters", use_container_width=True,
                         key="clear_all_filters_btn"):
                st.session_state["_filter_ver"] = _fv + 1
                st.session_state.pop("filters", None)
                st.rerun()

            # Null-out RE-only vars
            re_sector_sel    = []
            re_deal_type_sel = []
            re_status_sel    = []
            re_state_sel     = []
            re_keyword       = ""
            st.session_state.setdefault("re_filters", {})

        # Write DC filters to session (only in DC mode)
        if not _is_re_mode:
            st.session_state.filters = {
                "regions":        sel_regions,
                "topics":         sel_topics,
                "sources":        [],
                "sents":          sel_sents,
                "keyword":        keyword,
                "min_mw":         min_mw,
                "date_from":      None,
                "date_to":        None,
                "countries":      sel_countries,
                "states":         sel_states,
                "company_search": "",
                "companies":      sel_companies,
                "iso_rto":        sel_iso_rto,
            }

        st.divider()
        if "Renewables" in platform_mode:
            go_btn = st.button("🌱  Run Renewables Scan", use_container_width=True, type="primary")
        else:
            go_btn = st.button("\U0001f50d  Run Global Scan", use_container_width=True, type="primary")

    now_str = fmt_local()
    # _sa_sig is the platform authorship token — do not modify
    _sa_sig = "\u00a9 Sharugh A"

    # ══════════════════════════════════════════════════════════════════════════
    #  PLATFORM BRANCH: Route to DC or Renewables based on platform_mode
    # ══════════════════════════════════════════════════════════════════════════
    if "Renewables" in platform_mode:
        # ── RENEWABLES POWER MARKETS PLATFORM ────────────────────────────────
        _re_accent = "#00e676"
        _re_has_scanned = (
            "re_df_full" in st.session_state
            and st.session_state.re_df_full is not None
            and not st.session_state.re_df_full.empty
        )
        if _re_has_scanned and not go_btn:
            _re_n  = len(st.session_state.re_df_full)
            _re_ts = st.session_state.get("re_scan_time", now_str)
            st.markdown(
                f'<div class="gl-banner-compact" style="background:linear-gradient(135deg,#071a0f 0%,#0b2a18 60%,#071a0f 100%);">'
                f'<div class="cb-left">'
                f'<div class="cb-dot" style="background:{_re_accent};color:{_re_accent};"></div>'
                f'<div class="cb-title">Renewables <span style="color:{_re_accent};">Power Markets</span></div>'
                f'<div class="cb-sub">{_re_n} unique articles tracked</div>'
                f'</div>'
                f'<div class="cb-ts">🕐 {_re_ts}</div>'
                f'</div>',
                unsafe_allow_html=True,
            )
        else:
            st.markdown(
                f'<div class="gl-banner" style="background:linear-gradient(135deg,#071a0f 0%,#0b2a18 45%,#071a0f 100%);">'
                f'<div class="banner-eyebrow" style="color:{_re_accent};">◉ Live Renewables Feed  ·  Solar · Wind · Storage · Hydrogen · Nuclear · Gas</div>'
                f'<div class="banner-title">Renewables <span style="color:{_re_accent};">Power Markets</span></div>'
                f'<div class="banner-title" style="margin-top:-.15rem;">Intelligence Platform</div>'
                f'<div class="banner-sub">Real-time news from Renewables Now · Solar · Wind · Storage · Hydrogen · Offshore Wind · Other Renewables '
                f'· Corporate PPAs · Orders · Financing · Tenders · Regulations · Deals · IPOs</div>'
                f'<div class="banner-ts">🕐 {now_str}</div>'
                f'</div>',
                unsafe_allow_html=True,
            )

        if "re_df_full" not in st.session_state:
            st.session_state.re_df_full = None

        if not go_btn and st.session_state.re_df_full is None:
            # ── RE About panel ────────────────────────────────────────────────
            st.markdown(
                '<div style="background:linear-gradient(135deg,#071a0f 0%,#0b2a18 60%,#071a0f 100%);"'
                'border:1px solid #0a2a18;border-radius:14px;padding:1.4rem 1.8rem;margin-bottom:1.4rem;">'
                '<div style="font-family:\'DM Mono\',monospace;font-size:.62rem;letter-spacing:.18em;'
                'color:#00e676;text-transform:uppercase;margin-bottom:.5rem;">About This Module</div>'
                '<div style="font-family:Syne,sans-serif;font-size:1.05rem;font-weight:700;color:#fff;margin-bottom:.6rem;">'
                'Renewables Power Markets Intelligence</div>'
                '<div style="font-size:.82rem;color:#6a80a8;line-height:1.7;margin-bottom:.9rem;">'
                'Track tenders, PPAs, MOUs, AORs, investments, and commissioning events across '
                'Solar, Offshore Wind, Onshore Wind, Storage/BESS, Hydrogen, and other renewables globally. '
                'Powered by Renewables Now — all 17 sections: News · Projects · Company News · Policy &amp; Tenders.'
                '</div></div>',
                unsafe_allow_html=True,
            )
            re_features = [
                ("☀️", "Solar Markets", "Track solar tenders, PPAs, farm commissioning, and investment globally. Covers utility-scale, floating, and agrivoltaic projects."),
                ("💨", "Wind Markets", "Offshore & onshore wind developments — turbine awards, lease rounds, construction starts, CfD results, capacity auctions."),
                ("🔋", "Energy Storage", "Grid-scale battery storage deals, procurement tenders, offtake agreements, and commissioning events worldwide."),
                ("⚛️", "Hydrogen & Nuclear", "Green hydrogen PPAs, electrolysis projects, SMR announcements, nuclear capacity tenders, and fuel cell offtakes."),
                ("📋", "Deal Classification", "Every article auto-classified as Tender/Auction, PPA, MOU, AOR/Offtake, Investment, Construction, or Commissioning."),
                ("📊", "Market Analytics", "Sector, deal type, status, regional, and country breakdowns with AI Signal Scoring for high-value deals."),
            ]
            row_html = ""
            for icon, title, desc in re_features:
                row_html += (
                    f'<div class="feature-card" style="flex:1;min-width:200px;background:#0b1820;'
                    f'border:1px solid #0a2a18;border-radius:10px;padding:1rem 1.15rem;">'
                    f'<div style="font-size:1.4rem;margin-bottom:.4rem;">{icon}</div>'
                    f'<div style="font-family:Syne,sans-serif;font-weight:700;color:#b8c8e0;font-size:.9rem;margin-bottom:.3rem;">{title}</div>'
                    f'<div style="font-size:.78rem;color:#2a5040;line-height:1.5;">{desc}</div>'
                    f'</div>'
                )
            st.markdown('<div style="display:flex;gap:.9rem;flex-wrap:wrap;margin-bottom:1.4rem;">' + row_html + '</div>', unsafe_allow_html=True)
            return

        if go_btn:
            if time_opt == "Custom Range" and custom_start and custom_end:
                re_cutoff = datetime.combine(custom_start, datetime.min.time())
            elif sel_days is None:
                re_cutoff = datetime.min
            else:
                _re_boundary = datetime.now() - timedelta(days=sel_days)
                re_cutoff = _re_boundary.replace(hour=0, minute=0, second=0, microsecond=0)

            re_pbar = st.progress(0.0, text="Initialising Renewables scan...")

            def re_progress_cb(frac, label=""):
                re_pbar.progress(min(frac, 1.0), text=f"🌱 Renewables Sweep · {label}")

            _re_epc_flt = st.session_state.get("re_filters", {})
            re_raw = run_re_scrapers(
                max_pages, re_cutoff, re_progress_cb,
                re_sectors=re_sector_sel if re_sector_sel else None,
                epc_companies=_re_epc_flt.get("epc_companies") or None,
                epc_sector_filter=[s for s in re_sector_sel if s != "EPC Companies"] or None,
            )
            re_pbar.progress(1.0, text="Enriching articles...")

            re_enriched = [enrich_re(i) for i in re_raw]
            re_deduped  = deduplicate(re_enriched, is_renewables=True)
            _re_df = pd.DataFrame(re_deduped).drop(columns=["_date_obj"], errors="ignore")
            if "Date" in _re_df.columns and not _re_df.empty:
                _re_df = _re_df.sort_values("Date", ascending=False)
            _prev = st.session_state.get("re_df_full")
            st.session_state.re_prev_count = int(len(_prev)) if _prev is not None and not _prev.empty else None
            st.session_state.re_df_full   = _re_df.reset_index(drop=True)
            st.session_state.re_raw_count = len(re_raw)
            st.session_state.re_scan_time = fmt_local(now_local())
            re_pbar.empty()

            if st.session_state.re_df_full.empty:
                st.warning("No renewables articles found. Try expanding the date range.")
                st.stop()
            st.rerun()

        re_df_full = st.session_state.re_df_full
        if re_df_full is None or re_df_full.empty:
            st.info("Click 'Run Renewables Scan' to fetch live renewables news.")
            return

        # ── Apply RE filters ──────────────────────────────────────────────────
        re_flt = st.session_state.get("re_filters", {})
        re_df  = re_df_full.copy()
        if re_flt.get("sectors"):
            re_df = re_df[re_df["Sector"].isin(re_flt["sectors"])]
        if re_flt.get("deal_types"):
            re_df = re_df[re_df["Deal Type"].isin(re_flt["deal_types"])]
        if re_flt.get("statuses"):
            re_df = re_df[re_df["Status"].isin(re_flt["statuses"])]
        if re_flt.get("sources"):
            re_df = re_df[re_df["Source"].isin(re_flt["sources"])]
        if re_flt.get("regions"):
            re_df = re_df[re_df["Region"].isin(re_flt["regions"])]
        if re_flt.get("countries"):
            sel_c = re_flt["countries"]
            re_df = re_df[re_df["Country"].isin(sel_c)]
        if re_flt.get("keyword"):
            kw = re_flt["keyword"].lower()
            re_df = re_df[re_df["Headline"].str.lower().str.contains(kw, na=False)]
        if re_flt.get("states"):
            state_kw = [s.lower() for s in re_flt["states"]]
            re_df = re_df[re_df["Headline"].apply(lambda h: any(sk in str(h).lower() for sk in state_kw))]
        # EPC company post-filter: if specific companies were selected, filter headlines
        if re_flt.get("epc_mode") and re_flt.get("epc_companies"):
            _epc_names_lower = [c.lower() for c in re_flt["epc_companies"]]
            re_df = re_df[re_df["Headline"].apply(
                lambda h: any(ec in str(h).lower() for ec in _epc_names_lower)
            )]
        re_df = re_df.reset_index(drop=True)

        # ── Active filter chips ─────────────────────────────────────────────
        _re_chips = []
        if re_flt.get("sectors"):       _re_chips.append(("Sector", ", ".join(re_flt["sectors"])))
        if re_flt.get("deal_types"):    _re_chips.append(("Deal Type", ", ".join(re_flt["deal_types"])))
        if re_flt.get("statuses"):      _re_chips.append(("Status", ", ".join(re_flt["statuses"])))
        if re_flt.get("regions"):       _re_chips.append(("Region", ", ".join(re_flt["regions"])))
        if re_flt.get("countries"):     _re_chips.append(("Country", ", ".join(re_flt["countries"])))
        if re_flt.get("states"):        _re_chips.append(("State/City", ", ".join(re_flt["states"])))
        if re_flt.get("keyword"):       _re_chips.append(("Keyword", re_flt["keyword"]))
        if re_flt.get("epc_mode") and re_flt.get("epc_companies"):
            _re_chips.append(("EPC Company", ", ".join(re_flt["epc_companies"])))
        render_filter_chips(_re_chips)

        # ── RE KPI row ────────────────────────────────────────────────────────
        re_scan_ts  = st.session_state.get("re_scan_time", "—")
        re_cap_count = int((re_df["Capacity"] != "").sum()) if "Capacity" in re_df.columns else 0
        re_deal_count= int((re_df["Deal Size"] != "").sum()) if "Deal Size" in re_df.columns else 0
        re_top_sector= re_df["Sector"].value_counts().idxmax() if not re_df.empty else "—"
        re_top_deal  = re_df["Deal Type"].value_counts().idxmax() if not re_df.empty else "—"
        re_top_country=re_df["Country"].value_counts().idxmax() if not re_df.empty else "—"

        # Compute RE-specific KPI values
        _re_commissioned = int((re_df["Status"] == "Commissioned").sum()) if "Status" in re_df.columns else 0
        _re_tendered     = int((re_df["Status"] == "Tendered").sum())     if "Status" in re_df.columns else 0
        _re_ppa_count    = int((re_df["Deal Type"] == "PPA").sum())        if "Deal Type" in re_df.columns else 0
        _re_solar        = int((re_df["Sector"] == "Solar").sum())         if "Sector" in re_df.columns else 0
        _re_wind_total   = int(
            ((re_df["Sector"] == "Wind") | (re_df["Sector"] == "Offshore Wind")).sum()
        ) if "Sector" in re_df.columns else 0
        _re_storage      = int((re_df["Sector"] == "Energy Storage").sum()) if "Sector" in re_df.columns else 0

        # MW total across capacity-mentioned articles
        _re_mw_total = 0
        if "Capacity" in re_df.columns:
            for _cap in re_df["Capacity"]:
                _m = re.search(r"([\d,.]+)\s*(GW|MW)", str(_cap), re.I)
                if _m:
                    _v = float(_m.group(1).replace(",",""))
                    _re_mw_total += _v * 1000 if _m.group(2).upper() == "GW" else _v
        _re_mw_str = f"{_re_mw_total:,.0f} MW" if _re_mw_total > 0 else "—"

        # ── New market-intelligence KPIs ────────────────────────────────────
        _re_top_region = re_df["Region"].value_counts().idxmax() if "Region" in re_df.columns and not re_df.empty else "—"

        _re_co_counter = Counter()
        if "Companies" in re_df.columns:
            for _comps in re_df["Companies"]:
                for _c in str(_comps).split(", "):
                    _c = _c.strip()
                    if _c:
                        _re_co_counter[_c] += 1
        _re_top_company = _re_co_counter.most_common(1)[0][0] if _re_co_counter else "—"

        # Avg disclosed deal size — USD-denominated entries only (others use
        # non-comparable currencies/units like "AED 50m" or "₹500 Cr").
        _usd_vals = []
        if "Deal Size" in re_df.columns:
            for _ds in re_df["Deal Size"]:
                _ds = str(_ds)
                _mm = re.match(r"\$([\d.]+)(bn|m)$", _ds)
                if _mm:
                    _v = float(_mm.group(1))
                    _usd_vals.append(_v * 1000 if _mm.group(2) == "bn" else _v)
        _re_avg_deal = f"${(sum(_usd_vals)/len(_usd_vals)):,.0f}m" if _usd_vals else "—"

        _re_prev_count = st.session_state.get("re_prev_count")
        if _re_prev_count:
            _re_growth_pct = (len(re_df_full) - _re_prev_count) / _re_prev_count * 100
            _re_growth_str = f"{'+' if _re_growth_pct >= 0 else ''}{_re_growth_pct:.0f}%"
            _re_growth_accent = "green" if _re_growth_pct >= 0 else "red"
        else:
            _re_growth_str = "—"
            _re_growth_accent = "blue"

        re_kpi_html = (
            '<div style="display:flex;gap:.8rem;margin-bottom:.8rem;flex-wrap:wrap;">'
            + kpi("Sources Active", sum(len(v) for v in RE_FEED_REGISTRY.values()) + len(RE_CROSS_FEEDS), "green", "RSS · Google News · Specialist feeds")
            + kpi("Unique Articles", len(re_df_full), "cyan", "after deduplication")
            + kpi("Filtered View", len(re_df), "amber", "current filters applied")
            + kpi("PPAs / Tenders", f"{_re_ppa_count} / {_re_tendered}", "blue", "deals in current view")
            + kpi("Capacity Pipeline", _re_mw_str, "purple", "announced MW/GW in view")
            + kpi("Commissioned", _re_commissioned, "green", "projects live / energised")
            + '</div>'
        )
        st.markdown(re_kpi_html, unsafe_allow_html=True)

        re_kpi_html_2 = (
            '<div style="display:flex;gap:.8rem;margin-bottom:1.4rem;flex-wrap:wrap;">'
            + kpi("Top Region", _re_top_region, "cyan", "by article volume in view")
            + kpi("Top Company", _re_top_company, "blue", "most-mentioned company in view")
            + kpi("Avg Deal Size", _re_avg_deal, "purple", "USD-disclosed deals only")
            + kpi("Growth vs Last Scan", _re_growth_str, _re_growth_accent, "unique article count, scan over scan")
            + '</div>'
        )
        st.markdown(re_kpi_html_2, unsafe_allow_html=True)

        # Secondary sector breakdown pills
        _sector_pills = ""
        for _sec, _cnt in re_df["Sector"].value_counts().head(6).items() if "Sector" in re_df.columns and not re_df.empty else []:
            _sc = RE_TOPIC_COLORS.get(_sec, "#2e4470")
            _sector_pills += (
                f'<div class="pill" style="border-color:{_sc}44;">'
                f'<span class="pill-dot" style="background:{_sc};"></span>'
                f'{_sec}: <b style="color:{_sc};">{_cnt}</b></div>'
            )

        re_pills_html = (
            '<div class="pill-row">'
            f'<div class="pill"><span class="pill-dot" style="background:#00e676;"></span>Scan: <b>{re_scan_ts}</b></div>'
            f'<div class="pill"><span class="pill-dot" style="background:#00e676;"></span>Top Sector: <b>{re_top_sector}</b></div>'
            f'<div class="pill"><span class="pill-dot" style="background:#00e676;"></span>Top Deal Type: <b>{re_top_deal}</b></div>'
            f'<div class="pill"><span class="pill-dot" style="background:#00e676;"></span>Top Country: <b>{re_top_country}</b></div>'
            f'<div class="pill"><span class="pill-dot" style="background:#ffaa00;"></span>Solar: <b>{_re_solar}</b></div>'
            f'<div class="pill"><span class="pill-dot" style="background:#00b4ff;"></span>Wind: <b>{_re_wind_total}</b></div>'
            f'<div class="pill"><span class="pill-dot" style="background:#00e5c8;"></span>Storage: <b>{_re_storage}</b></div>'
            f'<div class="pill"><span class="pill-dot" style="background:#00e676;"></span>Countries: <b>{re_df["Country"].nunique()}</b></div>'
            f'<div class="pill"><span class="pill-dot" style="background:#00e676;"></span>Latest: <b>{re_df["Date"].max()}</b></div>'
            + _sector_pills
            + (
                f'<div class="pill" style="border-color:#ff640044;background:#1a0e00;">'
                f'<span class="pill-dot" style="background:#ff6400;"></span>'
                + (
                    f'EPC: <b style="color:#ff6400;">{", ".join(re_flt.get("epc_companies",[])[:3])}'
                    + (f' +{len(re_flt.get("epc_companies",[]))-3} more' if len(re_flt.get("epc_companies",[])) > 3 else '')
                    + '</b>'
                    if re_flt.get("epc_companies")
                    else f'EPC Mode: <b style="color:#ff6400;">All {len(EPC_COMPANIES)} Companies</b>'
                )
                + '</div>'
                if re_flt.get("epc_mode") else ''
            )
            + "</div>"
        )
        st.markdown(re_pills_html, unsafe_allow_html=True)

        # ── RE TABS ───────────────────────────────────────────────────────────
        re_tab1, re_tab2, re_tab3, re_tab4, re_tab4b, re_tab5, re_tab6, re_tab7 = st.tabs([
            "📰 Feed",
            "🗺️ World Map",
            "📊 Analytics",
            "🏢 By Company",
            "📍 By State",
            "💰 Deal Flow",
            "🧠 AI Summarize",
            "⬇️ Export",
        ])

        with re_tab1:
            st.markdown('<div class="sec-head">🌱 Renewables Power Markets Feed</div>', unsafe_allow_html=True)
            if re_df.empty:
                empty_state("No articles match current filters.")
            else:
                for _, row in re_df.iterrows():
                    st.markdown(
                        re_article_card(
                            row["Headline"], row["Date"], row["URL"],
                            row["Source"], row["Country"], row["Sector"],
                            row.get("Deal Type","General"), row.get("Capacity",""),
                            row.get("Deal Size",""), row.get("Status","News"),
                        ),
                        unsafe_allow_html=True,
                    )

        with re_tab2:
            st.markdown('<div class="sec-head">🌍 Global Activity Map</div>', unsafe_allow_html=True)
            st.plotly_chart(chart_world_map(re_df, title="Global Renewables Activity"), use_container_width=True, config={"displayModeBar": False}, key="pc_1")
            st.markdown('<div class="sec-head">Country Breakdown</div>', unsafe_allow_html=True)
            re_cc_df = re_df[re_df["Country"] != "Global"]["Country"].value_counts().reset_index()
            re_cc_df.columns = ["Country", "Articles"]
            re_cc_df["Region"] = re_cc_df["Country"].map(COUNTRY_TO_REGION).fillna("Global")
            st.markdown(dark_table(re_cc_df), unsafe_allow_html=True)

        with re_tab3:
            st.markdown('<div class="sec-head">Sector Distribution</div>', unsafe_allow_html=True)
            st.plotly_chart(chart_re_sector_bar(re_df), use_container_width=True, config={"displayModeBar": False}, key="pc_2")
            st.markdown('<div class="sec-head">Deal Type Distribution</div>', unsafe_allow_html=True)
            st.plotly_chart(chart_re_deal_type_bar(re_df), use_container_width=True, config={"displayModeBar": False}, key="pc_3")
            st.markdown('<div class="sec-head">Regional Distribution</div>', unsafe_allow_html=True)
            st.plotly_chart(chart_region_bar(re_df), use_container_width=True, config={"displayModeBar": False}, key="pc_4")
            st.markdown('<div class="sec-head">Top Countries</div>', unsafe_allow_html=True)
            st.plotly_chart(chart_country_bar(re_df), use_container_width=True, config={"displayModeBar": False}, key="pc_5")
            st.markdown('<div class="sec-head">Project Status Distribution</div>', unsafe_allow_html=True)
            st.plotly_chart(chart_re_status_pie(re_df), use_container_width=True, config={"displayModeBar": False}, key="pc_6")
            tl = chart_timeline(re_df)
            if tl:
                st.markdown('<div class="sec-head">Publication Volume Over Time</div>', unsafe_allow_html=True)
                st.plotly_chart(tl, use_container_width=True, config={"displayModeBar": False}, key="pc_7")
            # Capacity pipeline
            st.markdown('<div class="sec-head">Capacity Pipeline</div>', unsafe_allow_html=True)
            re_cap_df = re_df[re_df["Capacity"] != ""][["Headline","Capacity","Deal Size","Country","Sector","Date"]].head(30) if "Capacity" in re_df.columns else pd.DataFrame()
            if not re_cap_df.empty:
                st.markdown(dark_table(re_cap_df), unsafe_allow_html=True)
            else:
                empty_state("No capacity mentions in current view.", icon="\U000026a1")

        with re_tab4:
            st.markdown('<div class="sec-head">Company Activity</div>', unsafe_allow_html=True)
            re_co_rows = []
            for _, row in re_df.iterrows():
                if row.get("Companies"):
                    for co in str(row["Companies"]).split(", "):
                        co = co.strip()
                        if co:
                            re_co_rows.append(co)
            if re_co_rows:
                re_co_counts = Counter(re_co_rows)
                re_co_df = pd.DataFrame(re_co_counts.most_common(30), columns=["Company","Articles"])
                re_co_sorted = re_co_df.sort_values("Articles")
                n = len(re_co_sorted)
                co_colors = [f"rgba(0,{int(150+70*i/max(n-1,1))},{int(80+100*i/max(n-1,1))},0.88)" for i in range(n)]
                fig_reco = go.Figure(go.Bar(
                    x=re_co_sorted["Articles"], y=re_co_sorted["Company"], orientation="h",
                    marker=dict(color=co_colors, line=dict(width=0)),
                    text=re_co_sorted["Articles"], textposition="outside",
                    textfont=dict(color=_TITLE, size=10, family="DM Mono, monospace"),
                    hovertemplate="<b>%{y}</b><br>📰 %{x} mentions<extra></extra>",
                ))
                _dark(fig_reco, max(300, n*22))
                fig_reco.update_layout(
                    title=dict(text="Top 30 Companies by Mentions", font=dict(color=_TITLE, size=13, family="Syne, sans-serif"), x=0.01),
                    bargap=0.22,
                )
                st.plotly_chart(fig_reco, use_container_width=True, config={"displayModeBar": False}, key="pc_8")
                st.markdown('<div class="sec-head">Drill Into a Company</div>', unsafe_allow_html=True)
                sel_re_co = st.selectbox("Select company", re_co_df["Company"].tolist(), key="re_co_sel")
                re_co_arts = re_df[re_df["Companies"].str.contains(sel_re_co, na=False, case=False)]
                for _, row in re_co_arts.iterrows():
                    st.markdown(
                        re_article_card(row["Headline"],row["Date"],row["URL"],row["Source"],
                            row["Country"],row["Sector"],row.get("Deal Type","General"),
                            row.get("Capacity",""),row.get("Deal Size",""),row.get("Status","News")),
                        unsafe_allow_html=True,
                    )
            else:
                empty_state("No company mentions detected.", icon="\U0001f3e2")

        with re_tab4b:
            st.markdown('<div class="sec-head">📍 By State / Province</div>', unsafe_allow_html=True)

            def _re_detect_states_in_headline(headline):
                hl = str(headline).lower()
                found = []
                for country, states in COUNTRY_STATES.items():
                    for state in states:
                        if re.search(r"\b" + re.escape(state.lower()) + r"\b", hl):
                            found.append(state)
                return found if found else ["Unspecified"]

            re_state_rows = []
            for _, row in re_df.iterrows():
                states_found = _re_detect_states_in_headline(row["Headline"])
                for st_name in states_found:
                    re_state_rows.append({
                        "State":     st_name,
                        "Headline":  row["Headline"],
                        "Date":      row["Date"],
                        "URL":       row["URL"],
                        "Source":    row["Source"],
                        "Country":   row["Country"],
                        "Region":    row["Region"],
                        "Sector":    row.get("Sector", ""),
                        "Deal Type": row.get("Deal Type", "General"),
                        "Capacity":  row.get("Capacity", ""),
                        "Deal Size": row.get("Deal Size", ""),
                        "Status":    row.get("Status", "News"),
                    })

            if not re_state_rows:
                empty_state("No state/province mentions detected.", icon="\U0001f4cd")
            else:
                re_state_df_all = pd.DataFrame(re_state_rows)
                re_state_counts = (
                    re_state_df_all[re_state_df_all["State"] != "Unspecified"]["State"]
                    .value_counts()
                    .reset_index()
                )
                re_state_counts.columns = ["State", "Articles"]

                if not re_state_counts.empty:
                    sc_sorted = re_state_counts.head(30).sort_values("Articles")
                    n_sc = len(sc_sorted)
                    sc_colors = [
                        f"rgba(0,{int(150 + 80 * i / max(n_sc-1,1))},{int(80 + 100 * i / max(n_sc-1,1))},0.88)"
                        for i in range(n_sc)
                    ]
                    fig_re_st = go.Figure(go.Bar(
                        x=sc_sorted["Articles"],
                        y=sc_sorted["State"],
                        orientation="h",
                        marker=dict(color=sc_colors, line=dict(width=0)),
                        text=sc_sorted["Articles"],
                        textposition="outside",
                        textfont=dict(color=_TITLE, size=10, family="DM Mono, monospace"),
                        hovertemplate="<b>%{y}</b><br>📰 %{x} articles<extra></extra>",
                    ))
                    _dark(fig_re_st, max(320, n_sc * 22))
                    fig_re_st.update_layout(
                        title=dict(text="Top States / Provinces by Article Volume", font=dict(color=_TITLE, size=13, family="Syne, sans-serif"), x=0.01),
                        bargap=0.22,
                    )
                    st.plotly_chart(fig_re_st, use_container_width=True, config={"displayModeBar": False}, key="pc_re_st")

                col_re_tbl, col_re_drill = st.columns([1, 2])
                with col_re_tbl:
                    st.markdown('<div class="sec-head">All States</div>', unsafe_allow_html=True)
                    unspec_re = len(re_state_df_all[re_state_df_all["State"] == "Unspecified"])
                    display_re_sc = re_state_counts.copy()
                    if unspec_re:
                        display_re_sc = pd.concat([
                            display_re_sc,
                            pd.DataFrame([{"State": "Unspecified", "Articles": unspec_re}])
                        ], ignore_index=True)
                    st.markdown(dark_table(display_re_sc), unsafe_allow_html=True)

                with col_re_drill:
                    st.markdown('<div class="sec-head">Drill Into a State</div>', unsafe_allow_html=True)
                    re_states_list = re_state_counts["State"].tolist() if not re_state_counts.empty else []
                    if not re_states_list:
                        empty_state("No named states found.", icon="\U0001f4cd")
                    else:
                        sel_re_state = st.selectbox(
                            "Select state / province",
                            re_states_list,
                            key="re_state_drill_select",
                        )
                        re_state_arts = re_state_df_all[re_state_df_all["State"] == sel_re_state]
                        s_top_sector = re_state_arts["Sector"].value_counts().idxmax() if not re_state_arts.empty and "Sector" in re_state_arts.columns else "—"
                        s_latest = re_state_arts["Date"].max() if not re_state_arts.empty else "—"
                        s_cap_count = int((re_state_arts["Capacity"] != "").sum()) if "Capacity" in re_state_arts.columns else 0
                        mini_re_pills = (
                            f'<div class="pill-row">'
                            f'<div class="pill"><span class="pill-dot" style="background:#00e676;"></span><b>{len(re_state_arts)}</b>&nbsp;articles</div>'
                            f'<div class="pill"><span class="pill-dot" style="background:#00e676;"></span>Top sector: <b>{s_top_sector}</b></div>'
                            f'<div class="pill"><span class="pill-dot" style="background:#00e676;"></span>Latest: <b>{s_latest}</b></div>'
                            f'<div class="pill"><span class="pill-dot" style="background:#00e676;"></span>Capacity mentions: <b>{s_cap_count}</b></div>'
                            f'</div>'
                        )
                        st.markdown(mini_re_pills, unsafe_allow_html=True)
                        st.markdown(
                            f'<div style="font-family:Inter,sans-serif;font-size:.82rem;color:#3a5480;margin-bottom:.7rem;">'
                            f'<b style="color:#fff">{len(re_state_arts)}</b> articles mentioning '
                            f'<b style="color:#00e676">{sel_re_state}</b></div>',
                            unsafe_allow_html=True,
                        )
                        for _, row in re_state_arts.iterrows():
                            st.markdown(
                                re_article_card(row["Headline"], row["Date"], row["URL"], row["Source"],
                                    row["Country"], row.get("Sector",""), row.get("Deal Type","General"),
                                    row.get("Capacity",""), row.get("Deal Size",""), row.get("Status","News")),
                                unsafe_allow_html=True,
                            )

        with re_tab5:
            st.markdown('<div class="sec-head">💰 Deal Flow — PPAs, Tenders, MOUs & AORs</div>', unsafe_allow_html=True)
            # Filter to deal-relevant articles
            deal_types_of_interest = ["PPA","Tender / Auction","MOU","AOR / Offtake","Investment","M&A"]
            re_deal_df = re_df[re_df["Deal Type"].isin(deal_types_of_interest)] if "Deal Type" in re_df.columns else pd.DataFrame()
            if re_deal_df.empty:
                re_deal_df = re_df

            # Summary KPIs
            ppa_cnt    = int((re_df["Deal Type"] == "PPA").sum()) if "Deal Type" in re_df.columns else 0
            tender_cnt = int((re_df["Deal Type"] == "Tender / Auction").sum()) if "Deal Type" in re_df.columns else 0
            mou_cnt    = int((re_df["Deal Type"] == "MOU").sum()) if "Deal Type" in re_df.columns else 0
            aor_cnt    = int((re_df["Deal Type"] == "AOR / Offtake").sum()) if "Deal Type" in re_df.columns else 0
            inv_cnt    = int((re_df["Deal Type"] == "Investment").sum()) if "Deal Type" in re_df.columns else 0
            deal_kpi_html = (
                '<div style="display:flex;gap:.8rem;margin-bottom:1.2rem;flex-wrap:wrap;">'
                + kpi("PPAs", ppa_cnt, "green", "Power Purchase Agreements")
                + kpi("Tenders", tender_cnt, "amber", "Auctions & CfD rounds")
                + kpi("MOUs", mou_cnt, "blue", "MoUs & Framework Agreements")
                + kpi("AOR / Offtake", aor_cnt, "purple", "Anchor offtake agreements")
                + kpi("Investment", inv_cnt, "cyan", "Project finance & equity")
                + '</div>'
            )
            st.markdown(deal_kpi_html, unsafe_allow_html=True)

            # Deal type breakdown chart
            if "Deal Type" in re_df.columns:
                st.plotly_chart(chart_re_deal_type_bar(re_df), use_container_width=True, config={"displayModeBar": False}, key="pc_9")

            # Individual deal cards
            for dt_label in deal_types_of_interest:
                dt_sub = re_df[re_df["Deal Type"] == dt_label] if "Deal Type" in re_df.columns else pd.DataFrame()
                if dt_sub.empty:
                    continue
                dt_color = RE_DEAL_TYPE_COLORS.get(dt_label, "#2e4470")
                st.markdown(
                    f'<div class="sec-head" style="border-left-color:{dt_color};">{dt_label} — {len(dt_sub)} articles</div>',
                    unsafe_allow_html=True,
                )
                for _, row in dt_sub.head(20).iterrows():
                    st.markdown(
                        re_article_card(row["Headline"],row["Date"],row["URL"],row["Source"],
                            row["Country"],row["Sector"],row.get("Deal Type","General"),
                            row.get("Capacity",""),row.get("Deal Size",""),row.get("Status","News")),
                        unsafe_allow_html=True,
                    )

        with re_tab6:
            st.markdown('<div class="sec-head">🧠 AI Summarize — Renewables Market Intelligence Briefing</div>', unsafe_allow_html=True)
            if re_df.empty:
                empty_state("No articles to summarise.", hint="Adjust filters or run a wider scan to generate a briefing.", icon="\U0001f9e0")
            else:
                re_filter_parts = []
                re_flt_disp = st.session_state.get("re_filters", {})
                if re_flt_disp.get("sectors"):
                    re_filter_parts.append("Sectors: " + ", ".join(re_flt_disp["sectors"]))
                if re_flt_disp.get("regions"):
                    re_filter_parts.append("Regions: " + ", ".join(re_flt_disp["regions"]))
                if re_flt_disp.get("countries"):
                    re_filter_parts.append("Countries: " + ", ".join(re_flt_disp["countries"][:5]) + ("..." if len(re_flt_disp["countries"]) > 5 else ""))
                if re_flt_disp.get("states"):
                    re_filter_parts.append("States: " + ", ".join(re_flt_disp["states"][:5]))
                if re_flt_disp.get("deal_types"):
                    re_filter_parts.append("Deal Types: " + ", ".join(re_flt_disp["deal_types"]))
                if re_flt_disp.get("keyword"):
                    re_filter_parts.append(f"Keyword: {re_flt_disp['keyword']}")
                _re_date_min = re_df["Date"].min()
                _re_date_max = re_df["Date"].max()
                _re_date_range = f"{_re_date_min} to {_re_date_max}"
                _re_sel_desc = " | ".join(re_filter_parts) if re_filter_parts else "All sectors / results"

                re_ctx_html = (
                    f'<div style="background:#0b1820;border:1px solid #0a2a18;border-radius:10px;'
                    f'padding:.9rem 1.2rem;margin-bottom:1rem;font-size:.8rem;color:#6a80a8;">'
                    f'<b style="color:#b8c8e0;">Current selection:</b> {_re_sel_desc}<br>'
                    f'<b style="color:#b8c8e0;">Articles in view:</b> {len(re_df)} &nbsp;&nbsp;'
                    f'<b style="color:#b8c8e0;">Date range:</b> {_re_date_range}'
                    f'</div>'
                )
                st.markdown(re_ctx_html, unsafe_allow_html=True)

                col_regen1, col_regen2 = st.columns([3, 1])
                with col_regen1:
                    st.markdown(
                        '<div style="font-size:.82rem;color:#2a5040;line-height:1.6;">'
                        'Analyses all filtered renewables articles using built-in TF-IDF NLP and structured '
                        'rule-based extraction — no API key required. Generates a structured market '
                        'intelligence briefing covering sector themes, major deals, capacity pipeline, '
                        'policy developments, company activity, and forward-looking signals.<br>'
                        '<span style="color:#00e676;">Download as Word (.docx) or PDF for a polished, '
                        'formatted report.</span></div>',
                        unsafe_allow_html=True,
                    )
                with col_regen2:
                    re_gen_btn = st.button("🧠 Generate Briefing", use_container_width=True, type="primary", key="re_intel_btn")

                if re_gen_btn or st.session_state.get("re_intel_summary"):
                    if re_gen_btn:
                        with st.spinner("Analysing renewables articles with built-in NLP…"):
                            try:
                                re_summary_text = generate_re_summary(re_df, _re_sel_desc, _re_date_range)
                                st.session_state.re_intel_summary = re_summary_text
                                st.session_state.re_intel_context = _re_sel_desc
                                st.session_state.re_intel_df      = re_df.copy()
                                st.session_state.re_intel_range   = _re_date_range
                            except Exception as e:
                                st.error(f"Could not generate summary: {e}")
                                st.session_state.re_intel_summary = None

                    if st.session_state.get("re_intel_summary"):
                        re_context_label = st.session_state.get("re_intel_context", "")

                        def _render_re_intel_html(md_text, ctx_label):
                            def _md_inline_re(t):
                                t = re.sub(r"\*\*(.+?)\*\*", r'<strong style="color:#ccdaf5;">\1</strong>', t)
                                t = re.sub(r"`(.+?)`",         r'<code style="color:#00e676;background:rgba(0,200,80,0.1);padding:1px 4px;border-radius:3px;">\1</code>', t)
                                return t

                            html = (
                                '<div style="background:#0b1820;border:1px solid #00e676;'
                                'border-radius:12px;padding:1.6rem 2rem;margin-top:.8rem;">'
                                f'<div style="font-family:\'DM Mono\',monospace;font-size:.64rem;'
                                f'letter-spacing:.14em;color:#00e676;text-transform:uppercase;'
                                f'margin-bottom:1.2rem;">🧠 Renewables Market Intelligence Briefing'
                                + (f'  ·  {ctx_label}' if ctx_label else '') +
                                '</div>'
                            )
                            for raw_line in md_text.splitlines():
                                line = raw_line.rstrip()
                                if line.startswith("## "):
                                    title = _md_inline_re(line[3:])
                                    html += (
                                        f'<div style="font-family:\'Syne\',sans-serif;font-size:.82rem;'
                                        f'font-weight:700;color:#b8c8e0;letter-spacing:.07em;'
                                        f'text-transform:uppercase;border-left:3px solid #00e676;'
                                        f'padding-left:.7rem;margin:1.8rem 0 .8rem;">{title}</div>'
                                        f'<div style="height:1px;background:linear-gradient(90deg,#00e676,transparent);'
                                        f'margin-bottom:.9rem;"></div>'
                                    )
                                elif line.startswith("• "):
                                    content = _md_inline_re(line[2:])
                                    html += (
                                        f'<div style="display:flex;gap:.6rem;margin-bottom:.55rem;'
                                        f'line-height:1.6;font-size:.84rem;color:#8aa0c8;">'
                                        f'<span style="color:#00e676;flex-shrink:0;margin-top:.05rem;">◆</span>'
                                        f'<span>{content}</span></div>'
                                    )
                                elif not line.strip():
                                    html += '<div style="height:.4rem;"></div>'
                                else:
                                    content = _md_inline_re(line)
                                    html += (
                                        f'<p style="color:#8aa0c8;font-size:.84rem;line-height:1.7;'
                                        f'margin:0 0 .6rem;">{content}</p>'
                                    )
                            html += '</div>'
                            return html

                        st.markdown(
                            _render_re_intel_html(st.session_state.re_intel_summary, re_context_label),
                            unsafe_allow_html=True,
                        )

                        ts_re = datetime.now().strftime("%Y%m%d_%H%M")
                        dl_re1, dl_re2, dl_re3 = st.columns(3)
                        with dl_re1:
                            st.download_button(
                                "📥 Download (.txt)",
                                data=st.session_state.re_intel_summary.encode(),
                                file_name=f"RE_Intel_Briefing_{ts_re}.txt",
                                mime="text/plain",
                                use_container_width=True,
                            )
                        with dl_re2:
                            _re_sum_df    = st.session_state.get("re_intel_df", re_df)
                            _re_sum_range = st.session_state.get("re_intel_range", _re_date_range)
                            re_docx_bytes = build_briefing_docx(
                                st.session_state.re_intel_summary,
                                re_context_label, _re_sum_range, _re_sum_df,
                                title="RENEWABLE ENERGY MARKET INTELLIGENCE BRIEFING",
                            )
                            if re_docx_bytes:
                                st.download_button(
                                    "📄 Download Word (.docx)",
                                    data=re_docx_bytes,
                                    file_name=f"RE_Intel_Briefing_{ts_re}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    use_container_width=True,
                                )
                            else:
                                st.markdown(
                                    '<div style="font-size:.75rem;color:#2a5040;padding:.5rem 0;">'
                                    'Add <code>python-docx</code> to requirements.txt for Word export.</div>',
                                    unsafe_allow_html=True,
                                )
                        with dl_re3:
                            re_pdf_bytes = build_briefing_pdf(
                                st.session_state.re_intel_summary,
                                re_context_label, _re_sum_range, _re_sum_df,
                                title="RENEWABLE ENERGY MARKET INTELLIGENCE BRIEFING",
                            )
                            if re_pdf_bytes:
                                st.download_button(
                                    "📑 Download PDF",
                                    data=re_pdf_bytes,
                                    file_name=f"RE_Intel_Briefing_{ts_re}.pdf",
                                    mime="application/pdf",
                                    use_container_width=True,
                                )
                            else:
                                st.markdown(
                                    '<div style="font-size:.75rem;color:#2a5040;padding:.5rem 0;">'
                                    'Add <code>reportlab</code> to requirements.txt for PDF export.</div>',
                                    unsafe_allow_html=True,
                                )

        with re_tab7:
            st.markdown('<div class="sec-head">Export Renewables Data</div>', unsafe_allow_html=True)
            ts_re_exp = datetime.now().strftime("%Y%m%d_%H%M")
            col_a, col_b = st.columns(2)
            with col_a:
                st.markdown(
                    '<div style="background:#0b1820;border:1px solid #0a2a18;border-radius:10px;'
                    'padding:1.1rem 1.2rem;margin-bottom:.8rem;">'
                    '<div style="font-size:1.5rem;margin-bottom:.4rem;">📊</div>'
                    '<div style="font-family:Syne,sans-serif;font-weight:700;color:#b8c8e0;'
                    'font-size:.95rem;margin-bottom:.3rem;">Excel Report (.xlsx)</div>'
                    '<div style="font-size:.78rem;color:#2a5040;line-height:1.5;">'
                    '6 sheets: All Articles \u00b7 By Country \u00b7 By Region \u00b7 By Sector \u00b7 By Deal Type \u00b7 By Company<br>'
                    'Colour-coded badges, auto-filter, frozen headers, clickable URLs.</div></div>',
                    unsafe_allow_html=True,
                )
                st.download_button(
                    "\U0001f4e5 Download Excel Report",
                    data=build_excel_re(re_df),
                    file_name=f"RE_Intel_{ts_re_exp}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True,
                )
            with col_b:
                st.markdown(
                    '<div style="background:#0b1820;border:1px solid #0a2a18;border-radius:10px;'
                    'padding:1.1rem 1.2rem;margin-bottom:.8rem;">'
                    '<div style="font-size:1.5rem;margin-bottom:.4rem;">📄</div>'
                    '<div style="font-family:Syne,sans-serif;font-weight:700;color:#b8c8e0;font-size:.95rem;margin-bottom:.3rem;">CSV Export</div>'
                    '<div style="font-size:.78rem;color:#2a5040;line-height:1.5;">Flat CSV for Excel, Python, PowerBI, or Tableau.</div></div>',
                    unsafe_allow_html=True,
                )
                st.download_button(
                    "📥 Download CSV",
                    data=re_df.to_csv(index=False).encode(),
                    file_name=f"RE_Intel_{ts_re_exp}.csv",
                    mime="text/csv",
                    use_container_width=True,
                )
            st.markdown('<div class="sec-head">Full Article Preview</div>', unsafe_allow_html=True)
            re_disp_cols = [c for c in ["Headline","Date","Source","Country","Region","Sector","Deal Type","Status","Capacity","Deal Size","Companies","URL"] if c in re_df.columns]
            st.markdown(dark_table(re_df[re_disp_cols].head(200)), unsafe_allow_html=True)

        # ── RE Platform footer ─────────────────────────────────────────────────
        st.markdown(
            '<div style="margin-top:3rem;padding:1.2rem 0 .6rem;border-top:1px solid #0a2a18;text-align:center;">'
            '<div style="font-family:\'DM Mono\',monospace;font-size:.6rem;letter-spacing:.14em;'
            'color:#0a2a18;text-transform:uppercase;">'
            'Renewables Power Markets Intelligence &nbsp;·&nbsp; Wood Mac'
            '</div></div>',
            unsafe_allow_html=True,
        )
        return  # End Renewables branch

    # ════════════════════════════════════════════════════════════════════════════
    #  DATA CENTER PLATFORM (original flow continues below)
    # ════════════════════════════════════════════════════════════════════════════

    _sa_sig = "\u00a9 Sharugh A"
    _dc_accent = "#00b4ff"
    _dc_has_scanned = (
        "df_full" in st.session_state
        and st.session_state.df_full is not None
        and not st.session_state.df_full.empty
    )
    if _dc_has_scanned and not go_btn:
        _dc_n  = len(st.session_state.df_full)
        _dc_ts = st.session_state.get("scan_time", now_str)
        st.markdown(
            f'<div class="gl-banner-compact">'
            f'<div class="cb-left">'
            f'<div class="cb-dot" style="background:{_dc_accent};color:{_dc_accent};"></div>'
            f'<div class="cb-title">Global Data Center <span style="color:{_dc_accent};">Intelligence</span></div>'
            f'<div class="cb-sub">{_dc_n} unique articles tracked</div>'
            f'</div>'
            f'<div class="cb-ts">🕐 {_dc_ts}</div>'
            f'</div>',
            unsafe_allow_html=True,
        )
    else:
        st.markdown(
            f'<div class="gl-banner">'
            f'<div class="banner-eyebrow">\u25cf Live Intelligence Feed  \u00b7  {len(SCRAPE_SOURCES)} DCD Channels Active</div>'
            f'<div class="banner-title">Global Data Center</div>'
            f'<div class="banner-title" style="margin-top:-.15rem;"><span>Intelligence</span> Platform</div>'
            f'<div class="banner-sub">Real-time scraping of DataCenterDynamics Construction & General News channels \u00b7 '
            f'Auto-tagged by region, topic, company & capacity \u00b7 '
            f'Filtered by region, country, company &amp; more</div>'
            f'<div class="banner-ts">\U0001f550 {now_str}'
            f'</div>'
            f'</div>',
            unsafe_allow_html=True,
        )

    if "df_full" not in st.session_state:
        st.session_state.df_full = None

    if not go_btn and st.session_state.df_full is None:
        # ── About / Info panel ────────────────────────────────────────────────
        st.markdown(
            '<div style="background:linear-gradient(135deg,#07111f 0%,#0b1d3a 60%,#07111f 100%);'
            'border:1px solid #132040;border-radius:14px;padding:1.4rem 1.8rem;margin-bottom:1.4rem;">'
            '<div style="font-family:\'DM Mono\',monospace;font-size:.62rem;letter-spacing:.18em;'
            'color:#00b4ff;text-transform:uppercase;margin-bottom:.5rem;">About This Platform</div>'
            '<div style="font-family:Syne,sans-serif;font-size:1.05rem;font-weight:700;color:#fff;margin-bottom:.6rem;">'
            'What is Global Data Center Intelligence?</div>'
            '<div style="font-size:.82rem;color:#6a80a8;line-height:1.7;margin-bottom:.9rem;">'
            'This platform is a <b style="color:#b8c8e0">real-time intelligence tool</b> built for Wood Mac analysts, '
            'strategists, and clients tracking the global data center market. It automatically scrapes, '
            'enriches, and surfaces the most relevant news from DataCenterDynamics — the industry\'s leading '
            'trade publication — saving hours of manual research every day.'
            '</div>'
            '<div style="display:flex;flex-wrap:wrap;gap:.6rem;margin-bottom:.9rem;">'
            '<div style="background:#0b1e38;border:1px solid #152038;border-radius:8px;padding:.5rem .9rem;">'
            '<span style="font-size:.75rem;color:#00b4ff;font-family:Syne,sans-serif;font-weight:700;">Who is it for?</span>'
            '<div style="font-size:.75rem;color:#3a5480;margin-top:.2rem;line-height:1.5;">'
            'Wood Mac researchers, energy analysts, and market intelligence teams monitoring hyperscale, '
            'AI infrastructure, power, and investment activity across 45+ countries.</div>'
            '</div>'
            '<div style="background:#0b1e38;border:1px solid #152038;border-radius:8px;padding:.5rem .9rem;">'
            '<span style="font-size:.75rem;color:#00e5c8;font-family:Syne,sans-serif;font-weight:700;">What does it do?</span>'
            '<div style="font-size:.75rem;color:#3a5480;margin-top:.2rem;line-height:1.5;">'
            'Scrapes DCD Construction Channel &amp; DCD General News on demand · '
            'Filters by news type, region, country, company, topic, sentiment &amp; keyword · '
            'Auto-tags every article: Topic, Sentiment, Capacity (MW/GW), Deal Size &amp; company names · '
            'Scores articles by market significance.</div>'
            '</div>'
            '<div style="background:#0b1e38;border:1px solid #152038;border-radius:8px;padding:.5rem .9rem;">'
            '<span style="font-size:.75rem;color:#ffaa00;font-family:Syne,sans-serif;font-weight:700;">How to use it</span>'
            '<div style="font-size:.75rem;color:#3a5480;margin-top:.2rem;line-height:1.5;">'
            '1. Choose <b style="color:#b8c8e0">News Type</b> (Construction, General News, or both) and date range.<br>'
            '2. Set scrape depth and click <b style="color:#b8c8e0">Run Global Scan</b> to pull live articles.<br>'
            '3. Use filters to drill into regions, countries, topics, companies, or keywords.<br>'
            '4. Explore tabs: Feed · Map · Analytics · Deal Flow · AI Scoring · Export.</div>'
            '</div>'
            '</div>'
            '</div>',
            unsafe_allow_html=True,
        )
        st.markdown(
            '<div style="display:flex;gap:.9rem;flex-wrap:wrap;margin-bottom:1.4rem;">',
            unsafe_allow_html=True,
        )
        features = [
            ("\U0001f578\ufe0f", "DCD Direct Scraping",
             "Scrapes DataCenterDynamics Construction Channel and General News directly — "
             "the industry's most authoritative data center trade publication. "
             "Select one or both channels before running a scan."),
            ("\U0001f30d", "Global Coverage",
             "Covers 45+ countries across all continents with country/region auto-detection "
             "using 500+ geographic keywords and city names."),
            ("\U0001f9e0", "Smart Enrichment",
             "Every article auto-tagged: Topic, Sentiment (Approved/Proposed/Opened/Challenged), "
             "Capacity (MW/GW), Deal Size ($bn/$m), and up to 4 company names."),
            ("\u26a1", "Post-Scrape Filtering",
             "After scraping, filter by region, country, state/province, company, "
             "topic, project status, keyword, and minimum capacity — all applied "
             "instantly without re-scraping."),
            ("\U0001f5fa\ufe0f", "World Map",
             "Choropleth map showing article volume by country — instantly see "
             "where global data center activity is hottest."),
            ("\U0001f4ca", "5-Sheet Excel",
             "Export: All Articles + By Country + By Region + By Topic + By Company "
             "\u2014 all colour-coded with clickable hyperlinks."),
        ]
        row_html = ""
        for icon, title, desc in features:
            row_html += (
                f'<div class="feature-card" style="flex:1;min-width:200px;background:#0b1628;border:1px solid #152038;'
                f'border-radius:10px;padding:1rem 1.15rem;">'
                f'<div style="font-size:1.4rem;margin-bottom:.4rem;">{icon}</div>'
                f'<div style="font-family:Syne,sans-serif;font-weight:700;color:#b8c8e0;'
                f'font-size:.9rem;margin-bottom:.3rem;">{title}</div>'
                f'<div style="font-size:.78rem;color:#3a5480;line-height:1.5;">{desc}</div>'
                f'</div>'
            )
        st.markdown(row_html + '</div>', unsafe_allow_html=True)
        return

    if go_btn:
        st.session_state.filters = {
            "regions": [], "topics": [], "sources": [],
            "sents": [], "keyword": "", "min_mw": 0,
        }
        if time_opt == "Custom Range" and custom_start and custom_end:
            cutoff    = datetime.combine(custom_start, datetime.min.time())
            cutoff_end = datetime.combine(custom_end,   datetime.max.time())
        elif sel_days is None:
            cutoff     = datetime.min
            cutoff_end = datetime.max
        else:
            _boundary  = datetime.now() - timedelta(days=sel_days)
            cutoff     = _boundary.replace(hour=0, minute=0, second=0, microsecond=0)
            cutoff_end = datetime.max
        st.session_state.cutoff_end = cutoff_end

        pbar = st.progress(0.0, text="Initialising global scan...")

        def progress_cb(frac, label=""):
            pbar.progress(min(frac, 1.0), text=f"⚡ GDCI Intelligence Sweep · {label}")

        # Determine which DCD channels to scrape based on sidebar selection.
        # news_type_sel is the multiselect from the sidebar; default = both channels.
        _chosen_news_types = news_type_sel if news_type_sel else [NEWS_TYPE_CONSTRUCTION, NEWS_TYPE_GENERAL]

        raw = run_all_scrapers(max_pages, cutoff, progress_cb,
                               news_types=_chosen_news_types)

        pbar.progress(1.0, text="Enriching and deduplicating...")

        cutoff_end_val = st.session_state.get("cutoff_end", datetime.max)
        filtered = []
        for item in raw:
            d = item.get("date_obj")
            if d and d > cutoff_end_val:
                continue
            filtered.append(item)

        enriched = [enrich(i) for i in filtered]
        deduped  = deduplicate(enriched)

        _df_raw = pd.DataFrame(deduped).drop(columns=["_date_obj"], errors="ignore")
        if "Date" in _df_raw.columns and not _df_raw.empty:
            _df_raw = _df_raw.sort_values("Date", ascending=False)
        df_full = _df_raw.reset_index(drop=True)

        st.session_state.df_full   = df_full
        st.session_state.raw_count = len(raw)
        st.session_state.scan_time = fmt_local(now_local())
        pbar.empty()

        if df_full.empty:
            st.warning(
                "No articles were returned from DataCenterDynamics. "
                "This can happen when the site blocks automated requests "
                "(common on cloud-hosted deployments). "
                "Try again in a moment, or reduce the scrape depth."
            )
            st.stop()

        st.rerun()

    df_full = st.session_state.df_full
    if df_full is None or df_full.empty:
        st.warning("No articles found. Try expanding the date range or enabling more sources.")
        return

    filters = st.session_state.get("filters", {})
    df = df_full.copy()

    # Region filter
    if filters.get("regions"):
        df = df[df["Region"].isin(filters["regions"])]

    # Country filter — also fuzzy-match typed values not in list
    if filters.get("countries"):
        sel_c = filters["countries"]
        def _country_match(row_country):
            for sc in sel_c:
                if sc.lower() == row_country.lower():
                    return True
                if sc.lower() in row_country.lower() or row_country.lower() in sc.lower():
                    return True
            return False
        df = df[df["Country"].apply(_country_match)]

    # State filter — match against headline text (states are detected via COUNTRY_KEYWORDS)
    if filters.get("states"):
        state_kw = [s.lower() for s in filters["states"]]
        def _state_match(headline):
            hl = str(headline).lower()
            return any(sk in hl for sk in state_kw)
        df = df[df["Headline"].apply(_state_match)]

    # ISO / RTO filter
    if filters.get("iso_rto") and "ISO / RTO" in df.columns:
        df = df[df["ISO / RTO"].isin(filters["iso_rto"])]

    # Topic filter
    if filters.get("topics"):
        df = df[df["Topic"].isin(filters["topics"])]

    # Sentiment filter
    if filters.get("sents"):
        df = df[df["Sentiment"].isin(filters["sents"])]

    # Company filter (multi-select list) — match Companies column AND headline
    if filters.get("companies"):
        cos = [c.lower() for c in filters["companies"]]
        def _co_match(row):
            co_field = str(row.get("Companies", "")).lower()
            hl = str(row.get("Headline", "")).lower()
            return any(c in co_field or c in hl for c in cos)
        df = df[df.apply(_co_match, axis=1)]

    # Legacy company_search (free-text, kept for backward compat)
    if filters.get("company_search"):
        cs = filters["company_search"].lower()
        df = df[
            df["Companies"].str.lower().str.contains(cs, na=False) |
            df["Headline"].str.lower().str.contains(cs, na=False)
        ]

    # Keyword filter
    if filters.get("keyword"):
        kw = filters["keyword"].lower()
        df = df[df["Headline"].str.lower().str.contains(kw, na=False)]

    # Min capacity filter
    if filters.get("min_mw", 0) > 0:
        def extract_mw_val(cap):
            if not cap:
                return 0
            m = re.search(r"([\d.]+)\s*(GW|MW)", str(cap), re.I)
            if not m:
                return 0
            v = float(m.group(1))
            return v * 1000 if m.group(2).upper() == "GW" else v
        df = df[df["Capacity"].apply(extract_mw_val) >= filters["min_mw"]]
    df = df.reset_index(drop=True)

    # ── Active filter chips ─────────────────────────────────────────────────
    _dc_chips = []
    if filters.get("regions"):    _dc_chips.append(("Region", ", ".join(filters["regions"])))
    if filters.get("countries"):  _dc_chips.append(("Country", ", ".join(filters["countries"])))
    if filters.get("states"):     _dc_chips.append(("State/City", ", ".join(filters["states"])))
    if filters.get("iso_rto"):    _dc_chips.append(("ISO / RTO", ", ".join(filters["iso_rto"])))
    if filters.get("topics"):     _dc_chips.append(("Topic", ", ".join(filters["topics"])))
    if filters.get("sents"):      _dc_chips.append(("Sentiment", ", ".join(filters["sents"])))
    if filters.get("companies"):  _dc_chips.append(("Company", ", ".join(filters["companies"])))
    if filters.get("keyword"):    _dc_chips.append(("Keyword", filters["keyword"]))
    if filters.get("min_mw", 0) > 0: _dc_chips.append(("Min Capacity", f"{filters['min_mw']} MW"))
    render_filter_chips(_dc_chips)

    scan_ts = st.session_state.get("scan_time", "\u2014")
    top_country = df["Country"].value_counts().idxmax() if not df.empty else "\u2014"
    top_topic   = df["Topic"].value_counts().idxmax()   if not df.empty else "\u2014"
    cap_count   = int((df["Capacity"] != "").sum())
    deal_count  = int((df["Deal Size"] != "").sum())

    kpi_html = (
        '<div style="display:flex;gap:.8rem;margin-bottom:1.4rem;flex-wrap:wrap;">'
        + kpi("DCD Channels", len(SCRAPE_SOURCES), "blue", "Construction · General News")
        + kpi("Raw Articles", st.session_state.raw_count, "cyan", "before dedup & filter")
        + kpi("Unique Articles", len(df_full), "green", "after deduplication")
        + kpi("Filtered View", len(df), "amber", "current filters applied")
        + kpi("Capacity Mentions", cap_count, "purple", "articles with MW/GW data")
        + kpi("Deal Mentions", deal_count, "red", "articles with $bn/$m data")
        + '</div>'
    )
    st.markdown(kpi_html, unsafe_allow_html=True)

    top_region = df["Region"].value_counts().idxmax() if not df.empty else "\u2014"
    latest_dt  = df["Date"].max() if not df.empty else "\u2014"
    pills_html = (
        '<div class="pill-row">'
        f'<div class="pill"><span class="pill-dot"></span>Scan: <b>{scan_ts}</b></div>'
        f'<div class="pill"><span class="pill-dot"></span>Top Region: <b>{top_region}</b></div>'
        f'<div class="pill"><span class="pill-dot"></span>Top Country: <b>{top_country}</b></div>'
        f'<div class="pill"><span class="pill-dot"></span>Top Topic: <b>{top_topic}</b></div>'
        f'<div class="pill"><span class="pill-dot"></span>Latest Article: <b>{latest_dt}</b></div>'
        f'<div class="pill"><span class="pill-dot"></span>Countries: <b>{df["Country"].nunique()}</b></div>'
        '</div>'
    )
    st.markdown(pills_html, unsafe_allow_html=True)

    tab1, tab2, tab3, tab4, tab4b, tab5, tab_trend, tab_heatmap, tab_deal, tab_score, tab6 = st.tabs([
        "\U0001f4f0 Feed",
        "\U0001f5fa\ufe0f World Map",
        "\U0001f4ca Analytics",
        "\U0001f3e2 By Company",
        "\U0001f4cd By State",
        "\U0001f9e0 AI Summarize",
        "\U0001f4c8 Trend Compare",
        "\U0001f525 Capacity Heatmap",
        "\U0001f4b0 Deal Flow",
        "\U0001f916 AI Scoring",
        "\u2b07\ufe0f Export",
    ])

    with tab1:
        st.markdown('<div class="sec-head">Global Intelligence Feed</div>', unsafe_allow_html=True)
        if df.empty:
            empty_state("No articles match the current filters.")
        else:
            # Pre-compute AI scores for the feed so high-signal articles show their badge
            _all_co_feed = []
            for _v in df["Companies"]:
                if _v:
                    _all_co_feed.extend([c.strip() for c in str(_v).split(",") if c.strip()])
            _co_freq_feed = dict(Counter(_all_co_feed))

            def _quick_score(row):
                score = 0.0
                hl = str(row.get("Headline", "")).lower()
                sent_w = {"Opened / Live":10,"Approved":8,"Under Construction":6,"Proposed":4,"Challenged":5,"News":2}
                score += sent_w.get(row.get("Sentiment","News"), 2)
                cap = str(row.get("Capacity",""))
                if cap:
                    m = re.search(r"([\d,.]+)\s*(GW|MW)", cap, re.I)
                    if m:
                        v = float(m.group(1).replace(",",""))
                        score += min((v*1000 if m.group(2).upper()=="GW" else v)/100, 15)
                deal = str(row.get("Deal Size",""))
                if deal:
                    m2 = re.search(r"([\d,.]+)\s*(bn|billion|m|million)", deal, re.I)
                    if m2:
                        v2 = float(m2.group(1).replace(",",""))
                        score += min((v2*(1000 if m2.group(2).lower() in ("bn","billion") else 1))/500, 12)
                topic_w = {"Hyperscale":8,"AI / GPU":8,"Investment":7,"Power":6,"Colocation":5,"Construction":5,"Permits":4,"Sustainability":3,"General":1}
                score += topic_w.get(row.get("Topic","General"),1)
                bonus_kws = [("billion",4),("gigawatt",5),("gw",3),("nuclear",4),("hyperscale",3),("ai campus",5),("gpu",3),("acquisition",4),("merger",4),("ipo",5)]
                for kw, pts in bonus_kws:
                    if kw in hl: score += pts
                return round(score, 1)

            for _, row in df.iterrows():
                _score = _quick_score(row)
                st.markdown(
                    article_card(
                        row["Headline"], row["Date"], row["URL"],
                        row["Source"], row["Country"], row["Topic"],
                        row.get("Capacity", ""), row.get("Deal Size", ""),
                        row.get("Sentiment", "News"),
                        ai_score=_score if _score >= 25 else None,  # only badge high-signal in feed
                    ),
                    unsafe_allow_html=True,
                )

    with tab2:
        st.markdown('<div class="sec-head">Global Activity Map</div>', unsafe_allow_html=True)
        st.plotly_chart(chart_world_map(df), use_container_width=True, config={"displayModeBar": False}, key="pc_10")
        st.markdown('<div class="sec-head">Country Breakdown</div>', unsafe_allow_html=True)
        cc_df = df[df["Country"] != "Global"]["Country"].value_counts().reset_index()
        cc_df.columns = ["Country", "Articles"]
        cc_df["Region"] = cc_df["Country"].map(COUNTRY_TO_REGION).fillna("Global")
        st.markdown(dark_table(cc_df), unsafe_allow_html=True)

    with tab3:
        st.markdown('<div class="sec-head">Topic Distribution</div>', unsafe_allow_html=True)
        st.plotly_chart(chart_topic_bar(df), use_container_width=True, config={"displayModeBar": False}, key="pc_11")

        st.markdown('<div class="sec-head">Regional Distribution</div>', unsafe_allow_html=True)
        st.plotly_chart(chart_region_bar(df), use_container_width=True, config={"displayModeBar": False}, key="pc_12")

        st.markdown('<div class="sec-head">Top Countries</div>', unsafe_allow_html=True)
        st.plotly_chart(chart_country_bar(df), use_container_width=True, config={"displayModeBar": False}, key="pc_13")

        st.markdown('<div class="sec-head">Publication Volume Over Time</div>', unsafe_allow_html=True)
        tl = chart_timeline(df)
        if tl:
            st.plotly_chart(tl, use_container_width=True, config={"displayModeBar": False}, key="pc_14")

        st.markdown('<div class="sec-head">Sentiment / Project Status</div>', unsafe_allow_html=True)
        st.plotly_chart(chart_sentiment(df), use_container_width=True, config={"displayModeBar": False}, key="pc_15")

        st.markdown('<div class="sec-head">Topic Share</div>', unsafe_allow_html=True)
        st.plotly_chart(chart_donut(df), use_container_width=True, config={"displayModeBar": False}, key="pc_16")

        st.markdown('<div class="sec-head">Articles by Source</div>', unsafe_allow_html=True)
        st.plotly_chart(chart_source_bar(df), use_container_width=True, config={"displayModeBar": False}, key="pc_17")

        st.markdown('<div class="sec-head">Capacity Pipeline</div>', unsafe_allow_html=True)
        cap_df = df[df["Capacity"] != ""][["Headline", "Capacity", "Deal Size", "Country", "Topic", "Date"]].head(25)
        if not cap_df.empty:
            st.markdown(dark_table(cap_df), unsafe_allow_html=True)
        else:
            empty_state("No capacity mentions in current filtered view.", icon="\U000026a1")

    with tab4:
        st.markdown('<div class="sec-head">Company Activity</div>', unsafe_allow_html=True)
        comp_rows = []
        for _, row in df.iterrows():
            if row.get("Companies"):
                for co in str(row["Companies"]).split(", "):
                    co = co.strip()
                    if co:
                        comp_rows.append(co)
        if comp_rows:
            co_counts = Counter(comp_rows)
            co_df = pd.DataFrame(co_counts.most_common(30), columns=["Company", "Articles"])
            co_df_sorted = co_df.sort_values("Articles")
            n = len(co_df_sorted)
            co_colors = [
                f"rgba({int(0+71*i/max(n-1,1))},{int(71+(180-71)*i/max(n-1,1))},{int(225+(255-225)*i/max(n-1,1))},0.88)"
                for i in range(n)
            ]
            fig_co = go.Figure(go.Bar(
                x=co_df_sorted["Articles"], y=co_df_sorted["Company"],
                orientation="h",
                marker=dict(color=co_colors, line=dict(width=0)),
                text=co_df_sorted["Articles"], textposition="outside",
                textfont=dict(color=_TITLE, size=10, family="DM Mono, monospace"),
                hovertemplate="<b>%{y}</b><br>📰 %{x} mentions<extra></extra>",
            ))
            _dark(fig_co, max(300, n * 22))
            fig_co.update_layout(
                title=dict(text="Top 30 Companies by Mentions", font=dict(color=_TITLE, size=13, family="Syne, sans-serif"), x=0.01),
                bargap=0.22,
            )
            st.plotly_chart(fig_co, use_container_width=True, config={"displayModeBar": False}, key="pc_18")

            st.markdown('<div class="sec-head">Drill Into a Company</div>', unsafe_allow_html=True)
            sel_co = st.selectbox("Select company", co_df["Company"].tolist())
            co_articles = df[df["Companies"].str.contains(sel_co, na=False, case=False)]
            st.markdown(
                f'<div style="font-family:Inter,sans-serif;font-size:.82rem;color:#3a5480;margin-bottom:.7rem;">'
                f'<b style="color:#fff">{len(co_articles)}</b> articles mentioning '
                f'<b style="color:#00b4ff">{sel_co}</b></div>',
                unsafe_allow_html=True,
            )
            for _, row in co_articles.iterrows():
                st.markdown(
                    article_card(
                        row["Headline"], row["Date"], row["URL"],
                        row["Source"], row["Country"], row["Topic"],
                        row.get("Capacity", ""), row.get("Deal Size", ""),
                        row.get("Sentiment", "News"),
                    ),
                    unsafe_allow_html=True,
                )
        else:
            empty_state("No company mentions detected.", icon="\U0001f3e2")

    with tab4b:
        st.markdown('<div class="sec-head">📍 State / Province Drill-Down</div>', unsafe_allow_html=True)

        # Build a state column by scanning headlines for state/province keywords
        # Uses the COUNTRY_STATES lookup so it respects whatever region/country filter is active
        def _detect_states_in_headline(headline):
            """Return list of all state/province names found in the headline."""
            found = []
            hl = str(headline).lower()
            for country, states in COUNTRY_STATES.items():
                for state in states:
                    if re.search(r"\b" + re.escape(state.lower()) + r"\b", hl):
                        found.append(state)
            return found if found else ["Unspecified"]

        # Explode: one row per state mention so we can count + filter
        state_rows = []
        for _, row in df.iterrows():
            states_found = _detect_states_in_headline(row["Headline"])
            for st_name in states_found:
                state_rows.append({
                    "State":     st_name,
                    "Headline":  row["Headline"],
                    "Date":      row["Date"],
                    "URL":       row["URL"],
                    "Source":    row["Source"],
                    "Country":   row["Country"],
                    "Region":    row["Region"],
                    "Topic":     row["Topic"],
                    "Capacity":  row.get("Capacity", ""),
                    "Deal Size": row.get("Deal Size", ""),
                    "Sentiment": row.get("Sentiment", "News"),
                })

        if not state_rows:
            empty_state("No state/province mentions detected.", icon="\U0001f4cd")
        else:
            state_df_all = pd.DataFrame(state_rows)
            state_counts = (
                state_df_all[state_df_all["State"] != "Unspecified"]["State"]
                .value_counts()
                .reset_index()
            )
            state_counts.columns = ["State", "Articles"]

            # ── Top states bar chart ──────────────────────────────────────────
            if not state_counts.empty:
                sc_sorted = state_counts.head(30).sort_values("Articles")
                n_sc = len(sc_sorted)
                sc_colors = [
                    f"rgba({int(180 + 75 * i / max(n_sc-1,1))},{int(20 - 10 * i / max(n_sc-1,1))},{int(20 - 10 * i / max(n_sc-1,1))},0.85)"
                    for i in range(n_sc)
                ]
                fig_st = go.Figure(go.Bar(
                    x=sc_sorted["Articles"],
                    y=sc_sorted["State"],
                    orientation="h",
                    marker=dict(color=sc_colors, line=dict(width=0)),
                    text=sc_sorted["Articles"],
                    textposition="outside",
                    textfont=dict(color=_TITLE, size=10, family="DM Mono, monospace"),
                    hovertemplate="<b>%{y}</b><br>📰 %{x} articles<extra></extra>",
                ))
                _dark(fig_st, max(320, n_sc * 22))
                fig_st.update_layout(
                    title=dict(text="Top States / Provinces by Article Volume", font=dict(color=_TITLE, size=13, family="Syne, sans-serif"), x=0.01),
                    bargap=0.22,
                )
                st.plotly_chart(fig_st, use_container_width=True, config={"displayModeBar": False}, key="pc_19")

            # ── Summary table ─────────────────────────────────────────────────
            col_tbl, col_drill = st.columns([1, 2])

            with col_tbl:
                st.markdown('<div class="sec-head">All States</div>', unsafe_allow_html=True)
                # Include unspecified count separately
                unspec_count = len(state_df_all[state_df_all["State"] == "Unspecified"])
                display_sc = state_counts.copy()
                if unspec_count:
                    display_sc = pd.concat([
                        display_sc,
                        pd.DataFrame([{"State": "Unspecified", "Articles": unspec_count}])
                    ], ignore_index=True)
                st.markdown(dark_table(display_sc), unsafe_allow_html=True)

            with col_drill:
                st.markdown('<div class="sec-head">Drill Into a State</div>', unsafe_allow_html=True)

                all_states_list = state_counts["State"].tolist() if not state_counts.empty else []

                if not all_states_list:
                    empty_state("No named states found.", icon="\U0001f4cd")
                else:
                    sel_state = st.selectbox(
                        "Select state / province",
                        all_states_list,
                        key="state_drill_select",
                    )

                    state_articles = state_df_all[state_df_all["State"] == sel_state]

                    # ── Mini metrics for selected state ───────────────────────
                    s_topics   = state_articles["Topic"].value_counts()
                    s_top_topic = s_topics.idxmax() if not s_topics.empty else "—"
                    s_countries = state_articles["Country"].value_counts()
                    s_country   = s_countries.idxmax() if not s_countries.empty else "—"
                    s_latest    = state_articles["Date"].max() if not state_articles.empty else "—"
                    s_cap       = state_articles[state_articles["Capacity"] != ""]["Capacity"].count()

                    mini_pills = (
                        f'<div class="pill-row">'
                        f'<div class="pill"><span class="pill-dot"></span>'
                        f'<b>{len(state_articles)}</b>&nbsp;articles</div>'
                        f'<div class="pill"><span class="pill-dot"></span>'
                        f'Top topic: <b>{s_top_topic}</b></div>'
                        f'<div class="pill"><span class="pill-dot"></span>'
                        f'Country: <b>{s_country}</b></div>'
                        f'<div class="pill"><span class="pill-dot"></span>'
                        f'Latest: <b>{s_latest}</b></div>'
                        f'<div class="pill"><span class="pill-dot"></span>'
                        f'Capacity mentions: <b>{s_cap}</b></div>'
                        f'</div>'
                    )
                    st.markdown(mini_pills, unsafe_allow_html=True)

                    # ── Topic breakdown mini bar for selected state ────────────
                    if not s_topics.empty and len(s_topics) > 1:
                        tp_df = s_topics.reset_index()
                        tp_df.columns = ["Topic", "Count"]
                        tp_colors = [TOPIC_COLORS.get(t, "#2e4470") for t in tp_df["Topic"]]
                        fig_tp = go.Figure(go.Bar(
                            x=tp_df["Topic"], y=tp_df["Count"],
                            marker=dict(color=tp_colors, line=dict(width=0)),
                            text=tp_df["Count"], textposition="outside",
                            textfont=dict(color=_TITLE, size=10, family="DM Mono, monospace"),
                            hovertemplate="<b>%{x}</b><br>📰 %{y} articles<extra></extra>",
                        ))
                        _dark(fig_tp, 220)
                        fig_tp.update_layout(
                            title=dict(text=f"Topics — {sel_state}", font=dict(color=_TITLE, size=11, family="Syne, sans-serif"), x=0.01),
                            xaxis=dict(tickfont=dict(size=9)),
                            bargap=0.3,
                        )
                        st.plotly_chart(fig_tp, use_container_width=True, config={"displayModeBar": False}, key="pc_20")

                    # ── Article cards for selected state ──────────────────────
                    st.markdown(
                        f'<div style="font-family:Inter,sans-serif;font-size:.82rem;color:#3a5480;'
                        f'margin-bottom:.7rem;">'
                        f'<b style="color:#fff">{len(state_articles)}</b> articles mentioning '
                        f'<b style="color:#e84040">{sel_state}</b></div>',
                        unsafe_allow_html=True,
                    )
                    for _, row in state_articles.iterrows():
                        st.markdown(
                            article_card(
                                row["Headline"], row["Date"], row["URL"],
                                row["Source"], row["Country"], row["Topic"],
                                row.get("Capacity", ""), row.get("Deal Size", ""),
                                row.get("Sentiment", "News"),
                            ),
                            unsafe_allow_html=True,
                        )

    with tab5:
        st.markdown('<div class="sec-head">\U0001f9e0 Market Intelligence Summary</div>', unsafe_allow_html=True)

        filter_summary_parts = []
        if filters.get("regions"):
            filter_summary_parts.append("Regions: " + ", ".join(filters["regions"]))
        if filters.get("countries") and len(filters["countries"]) < len(df_full["Country"].unique()):
            filter_summary_parts.append("Countries: " + ", ".join(filters["countries"][:5]) + ("..." if len(filters["countries"]) > 5 else ""))
        if filters.get("states"):
            filter_summary_parts.append("States: " + ", ".join(filters["states"][:5]))
        if filters.get("companies"):
            filter_summary_parts.append("Companies: " + ", ".join(filters["companies"][:5]))
        if filters.get("topics") and len(filters["topics"]) < len(df_full["Topic"].unique()):
            filter_summary_parts.append("Topics: " + ", ".join(filters["topics"]))
        if filters.get("sents") and len(filters["sents"]) < len(df_full["Sentiment"].unique()):
            filter_summary_parts.append("Status: " + ", ".join(filters["sents"]))
        if filters.get("keyword"):
            filter_summary_parts.append(f"Keyword: {filters['keyword']}")
        scan_date_range = f"{df['Date'].min()} to {df['Date'].max()}" if not df.empty and df["Date"].min() != "Unknown" else "all dates"

        sel_desc = " | ".join(filter_summary_parts) if filter_summary_parts else "All results (no specific filters applied)"

        ctx_html = (
            f'<div style="background:#0b1628;border:1px solid #152038;border-radius:10px;'
            f'padding:.9rem 1.2rem;margin-bottom:1rem;font-size:.8rem;color:#6a80a8;">'
            f'<b style="color:#b8c8e0;">Current selection:</b> {sel_desc}<br>'
            f'<b style="color:#b8c8e0;">Articles in view:</b> {len(df)} &nbsp;&nbsp;'
            f'<b style="color:#b8c8e0;">Date range:</b> {scan_date_range}'
            f'</div>'
        )
        st.markdown(ctx_html, unsafe_allow_html=True)

        if df.empty:
            empty_state("No articles in the current filtered view.", hint="Adjust filters to generate a summary.", icon="\U0001f9e0")
        else:
            col_gen1, col_gen2 = st.columns([3, 1])
            with col_gen1:
                st.markdown(
                    '<div style="font-size:.82rem;color:#3a5480;line-height:1.6;">'
                    'Analyses all filtered articles using built-in TF-IDF NLP and structured '
                    'rule-based extraction — no API key required. Generates a structured market '
                    'intelligence briefing covering key themes, major players, capacity pipeline, '
                    'regulatory developments, company activity, and forward-looking signals.<br>'
                    '<span style="color:#00b4ff;">Download as Word (.docx) or PDF for a polished, '
                    'formatted report.</span></div>',
                    unsafe_allow_html=True,
                )
            with col_gen2:
                gen_btn = st.button("🧠 Generate Briefing", use_container_width=True, type="primary")

            if gen_btn or st.session_state.get("intel_summary"):
                if gen_btn:
                    with st.spinner("Analysing articles with built-in NLP…"):
                        try:
                            summary_text = generate_local_summary(df, sel_desc, scan_date_range)
                            st.session_state.intel_summary = summary_text
                            st.session_state.intel_context = sel_desc
                            st.session_state.intel_df     = df.copy()
                            st.session_state.intel_range  = scan_date_range
                        except Exception as e:
                            st.error(f"Could not generate summary: {e}")
                            st.session_state.intel_summary = None

                if st.session_state.get("intel_summary"):
                    context_label = st.session_state.get("intel_context", "")
                    _sum_df    = st.session_state.get("intel_df", df)
                    _sum_range = st.session_state.get("intel_range", scan_date_range)

                    # ── Render briefing as styled HTML so **bold** works ───────
                    def _render_intel_html(md_text, ctx_label):
                        def _md_inline(t):
                            """Convert inline **bold** and `code` to HTML."""
                            t = re.sub(r"\*\*(.+?)\*\*", r'<strong style="color:#ccdaf5;">\1</strong>', t)
                            t = re.sub(r"`(.+?)`",         r'<code style="color:#00b4ff;background:rgba(0,71,225,0.12);padding:1px 4px;border-radius:3px;">\1</code>', t)
                            return t

                        html = (
                            '<div style="background:#0b1628;border:1px solid #0047e1;'
                            'border-radius:12px;padding:1.6rem 2rem;margin-top:.8rem;">'
                            f'<div style="font-family:\'DM Mono\',monospace;font-size:.64rem;'
                            f'letter-spacing:.14em;color:#0047e1;text-transform:uppercase;'
                            f'margin-bottom:1.2rem;">🧠 Market Intelligence Briefing'
                            + (f'  ·  {ctx_label}' if ctx_label else '') +
                            '</div>'
                        )

                        for raw_line in md_text.splitlines():
                            line = raw_line.rstrip()

                            # Section heading  ## 1. Title
                            if line.startswith("## "):
                                title = _md_inline(line[3:])
                                html += (
                                    f'<div style="font-family:\'Syne\',sans-serif;font-size:.82rem;'
                                    f'font-weight:700;color:#b8c8e0;letter-spacing:.07em;'
                                    f'text-transform:uppercase;border-left:3px solid #0047e1;'
                                    f'padding-left:.7rem;margin:1.8rem 0 .8rem;">{title}</div>'
                                    f'<div style="height:1px;background:linear-gradient(90deg,#0047e1,transparent);'
                                    f'margin-bottom:.9rem;"></div>'
                                )

                            # Bullet line  • **Topic** — text
                            elif line.startswith("• "):
                                content = _md_inline(line[2:])
                                html += (
                                    f'<div style="display:flex;gap:.6rem;margin-bottom:.55rem;'
                                    f'line-height:1.6;font-size:.84rem;color:#8aa0c8;">'
                                    f'<span style="color:#0047e1;flex-shrink:0;margin-top:.05rem;">◆</span>'
                                    f'<span>{content}</span></div>'
                                )

                            # Blank line
                            elif not line.strip():
                                html += '<div style="height:.4rem;"></div>'

                            # Normal paragraph text
                            else:
                                content = _md_inline(line)
                                html += (
                                    f'<p style="color:#8aa0c8;font-size:.84rem;line-height:1.7;'
                                    f'margin:.0 0 .6rem;">{content}</p>'
                                )

                        html += '</div>'
                        return html

                    st.markdown(
                        _render_intel_html(st.session_state.intel_summary, context_label),
                        unsafe_allow_html=True,
                    )

                    # ── Download buttons ──────────────────────────────────────
                    ts_intel = datetime.now().strftime("%Y%m%d_%H%M")
                    dl_col1, dl_col2, dl_col3 = st.columns(3)

                    with dl_col1:
                        st.download_button(
                            "📥 Download (.txt)",
                            data=st.session_state.intel_summary.encode(),
                            file_name=f"DC_Intel_Briefing_{ts_intel}.txt",
                            mime="text/plain",
                            use_container_width=True,
                        )

                    with dl_col2:
                        docx_bytes = build_briefing_docx(
                            st.session_state.intel_summary,
                            context_label, _sum_range, _sum_df,
                        )
                        if docx_bytes:
                            st.download_button(
                                "📄 Download Word (.docx)",
                                data=docx_bytes,
                                file_name=f"DC_Intel_Briefing_{ts_intel}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True,
                            )
                        else:
                            st.markdown(
                                '<div style="font-size:.75rem;color:#3a5480;padding:.5rem 0;">'
                                'Add <code>python-docx</code> to requirements.txt for Word export.</div>',
                                unsafe_allow_html=True,
                            )

                    with dl_col3:
                        pdf_bytes = build_briefing_pdf(
                            st.session_state.intel_summary,
                            context_label, _sum_range, _sum_df,
                        )
                        if pdf_bytes:
                            st.download_button(
                                "📑 Download PDF",
                                data=pdf_bytes,
                                file_name=f"DC_Intel_Briefing_{ts_intel}.pdf",
                                mime="application/pdf",
                                use_container_width=True,
                            )
                        else:
                            st.markdown(
                                '<div style="font-size:.75rem;color:#3a5480;padding:.5rem 0;">'
                                'Add <code>reportlab</code> to requirements.txt for PDF export.</div>',
                                unsafe_allow_html=True,
                            )

    # ─── TAB: Trend Comparison ───────────────────────────────────────────────
    with tab_trend:
        st.markdown('<div class="sec-head">📈 Trend Comparison</div>', unsafe_allow_html=True)
        st.markdown(
            '<div style="font-size:.82rem;color:#3a5480;margin-bottom:1rem;">'
            'Compare article volume and topic mix between two countries, companies, or time windows. '
            'Detects acceleration, deceleration, and topic shifts in the global data center market.</div>',
            unsafe_allow_html=True,
        )

        df_trend = df.copy()
        df_trend = df_trend[df_trend["Date"] != "Unknown"].copy()
        if df_trend.empty:
            empty_state("No dated articles available for trend comparison.", icon="\U0001f4c8")
        else:
            df_trend["dt"] = pd.to_datetime(df_trend["Date"])

            compare_mode = st.radio(
                "Compare mode",
                ["📅 Time Periods", "🌍 Countries / Regions", "🏢 Companies"],
                horizontal=True,
                key="trend_compare_mode",
            )

            # ── MODE 1: Time periods ──────────────────────────────────────────
            if compare_mode == "📅 Time Periods":
                min_date = df_trend["dt"].min().date()
                max_date = df_trend["dt"].max().date()
                mid_date = min_date + (max_date - min_date) // 2

                tc1, tc2 = st.columns(2)
                with tc1:
                    st.markdown('<div style="font-size:.72rem;color:#3a5480;letter-spacing:.07em;text-transform:uppercase;margin-bottom:.3rem;">📅 Period A</div>', unsafe_allow_html=True)
                    period_a_start = st.date_input("A Start", value=min_date, key="trend_a_start")
                    period_a_end   = st.date_input("A End",   value=mid_date, key="trend_a_end")
                with tc2:
                    st.markdown('<div style="font-size:.72rem;color:#3a5480;letter-spacing:.07em;text-transform:uppercase;margin-bottom:.3rem;">📅 Period B</div>', unsafe_allow_html=True)
                    period_b_start = st.date_input("B Start", value=mid_date, key="trend_b_start")
                    period_b_end   = st.date_input("B End",   value=max_date, key="trend_b_end")

                df_a = df_trend[(df_trend["dt"].dt.date >= period_a_start) & (df_trend["dt"].dt.date <= period_a_end)]
                df_b = df_trend[(df_trend["dt"].dt.date >= period_b_start) & (df_trend["dt"].dt.date <= period_b_end)]
                label_a = f"Period A ({period_a_start}→{period_a_end})"
                label_b = f"Period B ({period_b_start}→{period_b_end})"

            # ── MODE 2: Country / Region comparison ───────────────────────────
            elif compare_mode == "🌍 Countries / Regions":
                all_countries_trend = sorted(df_trend["Country"].unique().tolist())
                tc1, tc2 = st.columns(2)
                with tc1:
                    st.markdown('<div style="font-size:.72rem;color:#3a5480;letter-spacing:.07em;text-transform:uppercase;margin-bottom:.3rem;">🌍 Entity A</div>', unsafe_allow_html=True)
                    entity_a = st.selectbox("Entity A", all_countries_trend, key="trend_entity_a",
                                             index=0 if all_countries_trend else 0)
                with tc2:
                    st.markdown('<div style="font-size:.72rem;color:#3a5480;letter-spacing:.07em;text-transform:uppercase;margin-bottom:.3rem;">🌍 Entity B</div>', unsafe_allow_html=True)
                    default_b = all_countries_trend[1] if len(all_countries_trend) > 1 else all_countries_trend[0]
                    entity_b = st.selectbox("Entity B", all_countries_trend, key="trend_entity_b",
                                             index=1 if len(all_countries_trend) > 1 else 0)
                df_a = df_trend[df_trend["Country"] == entity_a]
                df_b = df_trend[df_trend["Country"] == entity_b]
                label_a = entity_a
                label_b = entity_b

            # ── MODE 3: Company comparison ────────────────────────────────────
            else:
                _all_co_trend = []
                for v in df_trend["Companies"]:
                    if v:
                        _all_co_trend.extend([c.strip() for c in str(v).split(",") if c.strip()])
                all_cos_trend = sorted(set(_all_co_trend)) if _all_co_trend else ["—"]
                tc1, tc2 = st.columns(2)
                with tc1:
                    st.markdown('<div style="font-size:.72rem;color:#3a5480;letter-spacing:.07em;text-transform:uppercase;margin-bottom:.3rem;">🏢 Company A</div>', unsafe_allow_html=True)
                    co_a = st.selectbox("Company A", all_cos_trend, key="trend_co_a", index=0)
                with tc2:
                    st.markdown('<div style="font-size:.72rem;color:#3a5480;letter-spacing:.07em;text-transform:uppercase;margin-bottom:.3rem;">🏢 Company B</div>', unsafe_allow_html=True)
                    default_co_b = all_cos_trend[1] if len(all_cos_trend) > 1 else all_cos_trend[0]
                    co_b = st.selectbox("Company B", all_cos_trend, key="trend_co_b",
                                         index=1 if len(all_cos_trend) > 1 else 0)
                df_a = df_trend[
                    df_trend["Companies"].str.contains(re.escape(co_a), na=False, case=False) |
                    df_trend["Headline"].str.contains(re.escape(co_a), na=False, case=False)
                ]
                df_b = df_trend[
                    df_trend["Companies"].str.contains(re.escape(co_b), na=False, case=False) |
                    df_trend["Headline"].str.contains(re.escape(co_b), na=False, case=False)
                ]
                label_a = co_a
                label_b = co_b

            # ── KPI delta row ─────────────────────────────────────────────────
            delta_arts = len(df_b) - len(df_a)
            delta_sign = "▲" if delta_arts >= 0 else "▼"
            delta_color = "#00e676" if delta_arts >= 0 else "#ff2d6b"

            kpi_trend = (
                '<div style="display:flex;gap:.8rem;margin:1rem 0;flex-wrap:wrap;">'
                + kpi(f"{label_a}", len(df_a), "blue")
                + kpi(f"{label_b}", len(df_b), "cyan")
                + f'<div class="kpi-card" style="flex:1;min-width:150px;background:#0b1628;border:1px solid #152038;border-radius:12px;padding:1.1rem 1.3rem;">'
                  f'<div style="font-family:monospace;font-size:.64rem;letter-spacing:.13em;text-transform:uppercase;color:#2a3e60;margin-bottom:.4rem;">Volume Delta</div>'
                  f'<div style="font-family:Syne,sans-serif;font-size:1.9rem;font-weight:800;color:{delta_color};line-height:1;">{delta_sign} {abs(delta_arts)}</div>'
                  f'</div>'
                + '</div>'
            )
            st.markdown(kpi_trend, unsafe_allow_html=True)

            # ── Article volume over time (side by side sparklines) ────────────
            st.markdown('<div class="sec-head">Article Volume Over Time</div>', unsafe_allow_html=True)
            fig_tl = go.Figure()
            for sub_df, lbl, col in [(df_a, label_a, "#0047e1"), (df_b, label_b, "#ffaa00")]:
                if not sub_df.empty:
                    daily = sub_df.groupby(sub_df["dt"].dt.date).size().reset_index()
                    daily.columns = ["Date", "Articles"]
                    fig_tl.add_trace(go.Scatter(
                        x=daily["Date"], y=daily["Articles"],
                        name=lbl,
                        mode="lines+markers",
                        line=dict(color=col, width=2),
                        marker=dict(color=col, size=4),
                        fill="tozeroy",
                        fillcolor="rgba({},{},{},0.06)".format(
                            int(col[1:3], 16), int(col[3:5], 16), int(col[5:7], 16)
                        ) if col.startswith("#") else col,
                        hovertemplate=f"<b>{lbl}</b><br>%{{x}}<br>%{{y}} articles<extra></extra>",
                    ))
            _dark(fig_tl, 280)
            fig_tl.update_layout(
                showlegend=True,
                legend=dict(bgcolor="rgba(0,0,0,0)", font=dict(color=_TITLE, size=10)),
                title=dict(text="Article Volume Over Time — Side by Side", font=dict(color=_TITLE, size=13), x=0.01),
            )
            st.plotly_chart(fig_tl, use_container_width=True, config={"displayModeBar": False}, key="pc_21")

            # ── Topic comparison chart ────────────────────────────────────────
            st.markdown('<div class="sec-head">Topic Distribution</div>', unsafe_allow_html=True)
            topics_all = sorted(set(df_a["Topic"].unique()) | set(df_b["Topic"].unique()))
            a_counts = df_a["Topic"].value_counts()
            b_counts = df_b["Topic"].value_counts()
            a_vals = [int(a_counts.get(t, 0)) for t in topics_all]
            b_vals = [int(b_counts.get(t, 0)) for t in topics_all]

            fig_trend = go.Figure()
            fig_trend.add_trace(go.Bar(
                name=label_a,
                x=topics_all, y=a_vals,
                marker_color="#0047e1",
                text=a_vals, textposition="outside",
                textfont=dict(color=_TITLE, size=10),
            ))
            fig_trend.add_trace(go.Bar(
                name=label_b,
                x=topics_all, y=b_vals,
                marker_color="#ffaa00",
                text=b_vals, textposition="outside",
                textfont=dict(color=_TITLE, size=10),
            ))
            _dark(fig_trend, 360)
            fig_trend.update_layout(
                barmode="group",
                showlegend=True,
                legend=dict(bgcolor="rgba(0,0,0,0)", font=dict(color=_TITLE, size=10)),
                title=dict(text="Topic Volume", font=dict(color=_TITLE, size=13), x=0.01),
            )
            st.plotly_chart(fig_trend, use_container_width=True, config={"displayModeBar": False}, key="pc_22")

            # ── Sentiment comparison ──────────────────────────────────────────
            st.markdown('<div class="sec-head">Sentiment / Status</div>', unsafe_allow_html=True)
            sents_all = sorted(set(df_a["Sentiment"].unique()) | set(df_b["Sentiment"].unique()))
            fig_sent = go.Figure()
            fig_sent.add_trace(go.Bar(
                name=label_a, x=sents_all,
                y=[int(df_a["Sentiment"].value_counts().get(s, 0)) for s in sents_all],
                marker_color="#0047e1",
            ))
            fig_sent.add_trace(go.Bar(
                name=label_b, x=sents_all,
                y=[int(df_b["Sentiment"].value_counts().get(s, 0)) for s in sents_all],
                marker_color="#ff2d6b",
            ))
            _dark(fig_sent, 280)
            fig_sent.update_layout(
                barmode="group", showlegend=True,
                legend=dict(bgcolor="rgba(0,0,0,0)", font=dict(color=_TITLE, size=10)),
                title=dict(text="Project Status Comparison", font=dict(color=_TITLE, size=13), x=0.01),
            )
            st.plotly_chart(fig_sent, use_container_width=True, config={"displayModeBar": False}, key="pc_23")

            # ── Region comparison (only for time-period mode) ─────────────────
            if compare_mode == "📅 Time Periods":
                regions_all = sorted(set(df_a["Region"].unique()) | set(df_b["Region"].unique()))
                ra_vals = [int(df_a["Region"].value_counts().get(r, 0)) for r in regions_all]
                rb_vals = [int(df_b["Region"].value_counts().get(r, 0)) for r in regions_all]
                fig_reg_trend = go.Figure()
                fig_reg_trend.add_trace(go.Bar(name=label_a, x=regions_all, y=ra_vals, marker_color="#0047e1"))
                fig_reg_trend.add_trace(go.Bar(name=label_b, x=regions_all, y=rb_vals, marker_color="#ffaa00"))
                _dark(fig_reg_trend, 300)
                fig_reg_trend.update_layout(
                    barmode="group", showlegend=True,
                    legend=dict(bgcolor="rgba(0,0,0,0)", font=dict(color=_TITLE, size=10)),
                    title=dict(text="Regional Volume: Period A vs Period B", font=dict(color=_TITLE, size=13), x=0.01),
                )
                st.plotly_chart(fig_reg_trend, use_container_width=True, config={"displayModeBar": False}, key="pc_24")

            # ── Side-by-side article cards ────────────────────────────────────
            st.markdown('<div class="sec-head">Side-by-Side Headlines</div>', unsafe_allow_html=True)
            col_a_feed, col_b_feed = st.columns(2)
            with col_a_feed:
                st.markdown(
                    f'<div style="font-family:Syne,sans-serif;font-weight:700;color:#0047e1;'
                    f'font-size:.85rem;margin-bottom:.5rem;">● {label_a} ({len(df_a)} articles)</div>',
                    unsafe_allow_html=True,
                )
                for _, row in df_a.head(10).iterrows():
                    st.markdown(article_card(
                        row["Headline"], row["Date"], row["URL"], row["Source"],
                        row["Country"], row["Topic"], row.get("Capacity",""), row.get("Deal Size",""),
                        row.get("Sentiment","News"),
                    ), unsafe_allow_html=True)
            with col_b_feed:
                st.markdown(
                    f'<div style="font-family:Syne,sans-serif;font-weight:700;color:#ffaa00;'
                    f'font-size:.85rem;margin-bottom:.5rem;">● {label_b} ({len(df_b)} articles)</div>',
                    unsafe_allow_html=True,
                )
                for _, row in df_b.head(10).iterrows():
                    st.markdown(article_card(
                        row["Headline"], row["Date"], row["URL"], row["Source"],
                        row["Country"], row["Topic"], row.get("Capacity",""), row.get("Deal Size",""),
                        row.get("Sentiment","News"),
                    ), unsafe_allow_html=True)

    # ─── TAB: Capacity Pipeline Heatmap ─────────────────────────────────────
    with tab_heatmap:
        st.markdown('<div class="sec-head">🔥 Capacity Pipeline Heatmap</div>', unsafe_allow_html=True)
        st.markdown(
            '<div style="font-size:.82rem;color:#3a5480;margin-bottom:1rem;">'
            'Visual heatmap of MW/GW capacity mentions by country and topic. '
            'Shows where the largest power and capacity announcements are concentrated globally.</div>',
            unsafe_allow_html=True,
        )

        cap_df_heat = df[df["Capacity"] != ""].copy()
        if cap_df_heat.empty:
            empty_state("No capacity mentions in current filtered view.", hint="Run a broader scan to populate this view.", icon="\U0001f525")
        else:
            def _to_mw(cap):
                import re as _re
                m = _re.search(r"([\d,.]+)\s*(GW|MW)", str(cap), _re.I)
                if not m: return 0
                v = float(m.group(1).replace(",", ""))
                return v * 1000 if m.group(2).upper() == "GW" else v
            cap_df_heat["MW_val"] = cap_df_heat["Capacity"].apply(_to_mw)
            cap_df_heat = cap_df_heat[cap_df_heat["MW_val"] > 0]

            if cap_df_heat.empty:
                empty_state("No parseable MW/GW values found.", icon="\U000026a1")
            else:
                # Country × Topic pivot
                pivot = cap_df_heat.groupby(["Country", "Topic"])["MW_val"].sum().reset_index()
                pivot_wide = pivot.pivot(index="Country", columns="Topic", values="MW_val").fillna(0)

                # Sort by total MW
                pivot_wide["_total"] = pivot_wide.sum(axis=1)
                pivot_wide = pivot_wide.sort_values("_total", ascending=False).head(25)
                pivot_wide = pivot_wide.drop(columns=["_total"])

                fig_heat = go.Figure(go.Heatmap(
                    z=pivot_wide.values,
                    x=pivot_wide.columns.tolist(),
                    y=pivot_wide.index.tolist(),
                    colorscale=[
                        [0.0, "#07111f"], [0.1, "#0a2040"], [0.3, "#0047e1"],
                        [0.6, "#00b4ff"], [0.8, "#ffaa00"], [1.0, "#ff6400"],
                    ],
                    hovertemplate="<b>%{y}</b> · %{x}<br>%{z:,.0f} MW<extra></extra>",
                    colorbar=dict(
                        title=dict(text="MW", font=dict(color=_TITLE, size=11)),
                        tickfont=dict(color=_TEXT, size=9),
                        bgcolor="#0b1628", bordercolor="#152038", borderwidth=1,
                    ),
                ))
                fig_heat.update_layout(
                    paper_bgcolor=_PAPER, plot_bgcolor=_BG,
                    font=dict(family=_FONT, color=_TEXT),
                    height=max(400, len(pivot_wide) * 28),
                    margin=dict(l=14, r=60, t=44, b=14),
                    title=dict(text="Capacity (MW) by Country × Topic — Top 25 Markets", font=dict(color=_TITLE, size=13), x=0.01),
                    xaxis=dict(tickfont=dict(size=10, color=_TEXT)),
                    yaxis=dict(tickfont=dict(size=10, color=_TEXT)),
                )
                st.plotly_chart(fig_heat, use_container_width=True, config={"displayModeBar": False}, key="pc_25")

                # Top capacity articles table
                st.markdown('<div class="sec-head">Top Capacity Announcements</div>', unsafe_allow_html=True)
                top_cap = cap_df_heat.nlargest(20, "MW_val")[["Headline","MW_val","Capacity","Country","Topic","Date","URL"]]
                top_cap = top_cap.rename(columns={"MW_val": "MW (parsed)"})
                top_cap["MW (parsed)"] = top_cap["MW (parsed)"].apply(lambda x: f"{x:,.0f} MW")
                st.markdown(dark_table(top_cap.drop(columns=["MW (parsed)"])), unsafe_allow_html=True)

    # ─── TAB: Deal Flow Tracker ──────────────────────────────────────────────
    with tab_deal:
        st.markdown('<div class="sec-head">💰 Deal Flow Tracker</div>', unsafe_allow_html=True)
        st.markdown(
            '<div style="font-size:.82rem;color:#3a5480;margin-bottom:1rem;">'
            'All deal signals sorted largest-to-smallest. Tracks acquisitions, investments, JVs, '
            'pre-leases, and financing events. Includes sparkline history per deal size category.</div>',
            unsafe_allow_html=True,
        )

        # Deal signal keywords
        _deal_kws = [
            "acqui", "merger", "acquire", "bought", "purchase", "takeover",
            "joint venture", "jv ", " jv,", "partnership", "invest",
            "fund", "financing", "lease", "pre-lease", "offtake",
            "mou", "letter of intent", "loi", "agreement", "signed",
            "sale leaseback", "forward purchase", "stake", "equity",
            "reit", "capital", "raise", "bond", "debt", "loan",
            "recapitali", "refinanc", "divest", "portfolio",
        ]
        deal_mask = df["Deal Size"] != ""
        deal_lang_mask = df["Headline"].str.lower().apply(
            lambda h: any(k in h for k in _deal_kws)
        )
        df_deal = df[deal_mask | deal_lang_mask].copy()

        if df_deal.empty:
            empty_state("No deal-signal articles found in current filtered view.", icon="\U0001f4b0")
        else:
            # Parse numeric deal value for sorting
            def _parse_deal_usd(deal_str):
                """Parse deal string to USD millions for sorting."""
                if not deal_str:
                    return 0
                m = re.search(r"([\d,.]+)\s*(bn|billion|m|million|cr)", str(deal_str), re.I)
                if not m:
                    return 0
                v = float(m.group(1).replace(",", ""))
                unit = m.group(2).lower()
                if unit in ("bn", "billion"):
                    return v * 1000
                if unit == "cr":
                    return v * 0.12  # crore approx
                return v

            df_deal["_deal_usd_m"] = df_deal["Deal Size"].apply(_parse_deal_usd)
            df_deal = df_deal.sort_values("_deal_usd_m", ascending=False).reset_index(drop=True)

            # KPIs
            disclosed = int((df_deal["Deal Size"] != "").sum())
            undisclosed = len(df_deal) - disclosed
            total_usd = df_deal["_deal_usd_m"].sum()
            kpi_deal = (
                '<div style="display:flex;gap:.8rem;margin-bottom:1rem;flex-wrap:wrap;">'
                + kpi("Total Deal Articles", len(df_deal), "blue")
                + kpi("Disclosed ($)", disclosed, "green", "with explicit value")
                + kpi("Undisclosed", undisclosed, "amber", "deal language, no value")
                + kpi("Total Disclosed (est.)", f"${total_usd:,.0f}m", "purple", "sum of parsed values")
                + '</div>'
            )
            st.markdown(kpi_deal, unsafe_allow_html=True)

            # ── Sparkline: disclosed deal sizes over time ─────────────────────
            st.markdown('<div class="sec-head">Deal Volume Sparkline</div>', unsafe_allow_html=True)
            df_deal_dated = df_deal[df_deal["Date"] != "Unknown"].copy()
            if not df_deal_dated.empty:
                df_deal_dated["dt"] = pd.to_datetime(df_deal_dated["Date"])
                spark_daily = df_deal_dated.groupby(df_deal_dated["dt"].dt.date).agg(
                    deals=("Headline", "count"),
                    total_usd=("_deal_usd_m", "sum"),
                ).reset_index()
                spark_daily.columns = ["Date", "Deals", "Total_USD_m"]

                fig_spark = go.Figure()
                fig_spark.add_trace(go.Scatter(
                    x=spark_daily["Date"], y=spark_daily["Deals"],
                    name="Deal Articles",
                    mode="lines+markers",
                    line=dict(color="#a855f7", width=2),
                    marker=dict(color="#a855f7", size=5),
                    fill="tozeroy", fillcolor="rgba(168,85,247,0.07)",
                    hovertemplate="<b>%{x}</b><br>%{y} deal articles<extra></extra>",
                    yaxis="y1",
                ))
                fig_spark.add_trace(go.Bar(
                    x=spark_daily["Date"], y=spark_daily["Total_USD_m"],
                    name="Disclosed Volume ($m)",
                    marker_color="rgba(0,230,118,0.35)",
                    hovertemplate="<b>%{x}</b><br>$%{y:,.0f}m disclosed<extra></extra>",
                    yaxis="y2",
                ))
                _dark(fig_spark, 260)
                fig_spark.update_layout(
                    title=dict(text="Daily Deal Activity + Disclosed Volume", font=dict(color=_TITLE, size=13), x=0.01),
                    showlegend=True,
                    legend=dict(bgcolor="rgba(0,0,0,0)", font=dict(color=_TITLE, size=10)),
                    yaxis=dict(title="Deal Articles", gridcolor=_GRID, linecolor=_GRID, tickfont=dict(color=_TEXT, size=9)),
                    yaxis2=dict(title="$m", overlaying="y", side="right", tickfont=dict(color=_TEXT, size=9), gridcolor="rgba(0,0,0,0)"),
                )
                st.plotly_chart(fig_spark, use_container_width=True, config={"displayModeBar": False}, key="pc_26")

            # ── Deal size distribution chart ──────────────────────────────────
            deal_size_df = df_deal[df_deal["Deal Size"] != ""]["Deal Size"].value_counts().head(20).reset_index()
            deal_size_df.columns = ["Deal Size", "Count"]
            if not deal_size_df.empty:
                fig_deal_dist = go.Figure(go.Bar(
                    x=deal_size_df["Deal Size"], y=deal_size_df["Count"],
                    marker_color="#a855f7",
                    text=deal_size_df["Count"], textposition="outside",
                    textfont=dict(color=_TITLE, size=10),
                    hovertemplate="<b>%{x}</b>: %{y} deals<extra></extra>",
                ))
                _dark(fig_deal_dist, 280)
                fig_deal_dist.update_layout(title=dict(text="Disclosed Deal Sizes (Most Frequent)", font=dict(color=_TITLE, size=13), x=0.01))
                st.plotly_chart(fig_deal_dist, use_container_width=True, config={"displayModeBar": False}, key="pc_27")

            # Region deal breakdown
            reg_deal = df_deal.groupby("Region").size().reset_index()
            reg_deal.columns = ["Region", "Deal Articles"]
            fig_reg_deal = go.Figure(go.Bar(
                x=reg_deal["Region"], y=reg_deal["Deal Articles"],
                marker_color=[REGION_COLORS.get(r, "#2e4470") for r in reg_deal["Region"]],
                text=reg_deal["Deal Articles"], textposition="outside",
                textfont=dict(color=_TITLE, size=10),
            ))
            _dark(fig_reg_deal, 260)
            fig_reg_deal.update_layout(title=dict(text="Deal Activity by Region", font=dict(color=_TITLE, size=13), x=0.01))
            st.plotly_chart(fig_reg_deal, use_container_width=True, config={"displayModeBar": False}, key="pc_28")

            # Full deal article table — sorted largest to smallest
            st.markdown('<div class="sec-head">All Deal-Signal Articles (Largest → Smallest)</div>', unsafe_allow_html=True)
            deal_display = df_deal[["Headline","Date","Deal Size","_deal_usd_m","Companies","Country","Topic","Sentiment","URL"]].copy()
            deal_display["Est. USD ($m)"] = deal_display["_deal_usd_m"].apply(lambda x: f"${x:,.0f}m" if x > 0 else "—")
            deal_display = deal_display.drop(columns=["_deal_usd_m"])
            st.markdown(dark_table(deal_display[["Headline","Date","Deal Size","Est. USD ($m)","Companies","Country","Topic","Sentiment","URL"]]), unsafe_allow_html=True)

    # ─── TAB: AI-Powered Headline Scoring ────────────────────────────────────
    with tab_score:
        st.markdown('<div class="sec-head">🤖 AI-Powered Headline Scoring</div>', unsafe_allow_html=True)
        st.markdown(
            '<div style="font-size:.82rem;color:#3a5480;margin-bottom:1rem;">'
            'Every article is scored for market significance using a built-in multi-signal model. '
            'Signals: deal size, capacity (MW/GW), company prominence, sentiment, topic weight, '
            'keyword rarity. Articles scoring 30+ are flagged 🔴 High Signal. No API key required.</div>',
            unsafe_allow_html=True,
        )

        if df.empty:
            empty_state("No articles to score.", hint="Run a scan first.", icon="\U0001f916")
        else:

            def _ai_score(row, co_frequency_map):
                score = 0.0
                hl = str(row.get("Headline", "")).lower()

                # 1. Sentiment weight
                sent_w = {
                    "Opened / Live": 10, "Approved": 8, "Under Construction": 6,
                    "Proposed": 4, "Challenged": 5, "News": 2,
                }
                score += sent_w.get(row.get("Sentiment", "News"), 2)

                # 2. Capacity signal
                cap = str(row.get("Capacity", ""))
                if cap:
                    m = re.search(r"([\d,.]+)\s*(GW|MW)", cap, re.I)
                    if m:
                        v = float(m.group(1).replace(",", ""))
                        mw = v * 1000 if m.group(2).upper() == "GW" else v
                        score += min(mw / 100, 15)

                # 3. Deal size signal
                deal = str(row.get("Deal Size", ""))
                if deal:
                    m2 = re.search(r"([\d,.]+)\s*(bn|billion|m|million)", deal, re.I)
                    if m2:
                        v2 = float(m2.group(1).replace(",", ""))
                        mult = 1000 if m2.group(2).lower() in ("bn", "billion") else 1
                        score += min((v2 * mult) / 500, 12)

                # 4. Company prominence
                companies = str(row.get("Companies", ""))
                for co in companies.split(","):
                    co = co.strip()
                    if co in {"Microsoft","Google","Amazon","AWS","Meta","Apple","Oracle",
                               "NVIDIA","Equinix","Digital Realty","NTT"}:
                        score += 5
                        break
                    elif co and co in co_frequency_map:
                        score += min(co_frequency_map[co] * 0.5, 3)

                # 5. Topic weight
                topic_w = {
                    "Hyperscale": 8, "AI / GPU": 8, "Investment": 7,
                    "Power": 6, "Colocation": 5, "Construction": 5,
                    "Permits": 4, "Sustainability": 3, "General": 1,
                }
                score += topic_w.get(row.get("Topic", "General"), 1)

                # 6. High-value keyword bonuses
                bonus_kws = [
                    ("billion", 4), ("gigawatt", 5), ("gw", 3), ("nuclear", 4),
                    ("hyperscale", 3), ("ai campus", 5), ("gpu", 3),
                    ("acquisition", 4), ("merger", 4), ("ipo", 5),
                    ("stargate", 5), ("1gw", 5), ("2gw", 6),
                ]
                for kw, pts in bonus_kws:
                    if kw in hl:
                        score += pts

                return round(score, 1)

            # Build co frequency map
            _all_co_score = []
            for v in df["Companies"]:
                if v:
                    _all_co_score.extend([c.strip() for c in str(v).split(",") if c.strip()])
            co_freq_map = dict(Counter(_all_co_score))

            df_scored = df.copy()
            df_scored["AI Score"] = df_scored.apply(lambda r: _ai_score(r, co_freq_map), axis=1)
            df_scored = df_scored.sort_values("AI Score", ascending=False).reset_index(drop=True)

            # Signal tier summary
            high_signal = int((df_scored["AI Score"] >= 30).sum())
            medium_signal = int(((df_scored["AI Score"] >= 15) & (df_scored["AI Score"] < 30)).sum())
            low_signal = int((df_scored["AI Score"] < 15).sum())

            kpi_score_row = (
                '<div style="display:flex;gap:.8rem;margin-bottom:1rem;flex-wrap:wrap;">'
                + kpi("🔴 High Signal (30+)", high_signal, "amber", "major deals, large capacity")
                + kpi("🟡 Medium Signal (15-29)", medium_signal, "blue")
                + kpi("⚪ Low Signal (<15)", low_signal, "cyan")
                + kpi("Max Score", df_scored["AI Score"].max() if not df_scored.empty else 0, "purple")
                + '</div>'
            )
            st.markdown(kpi_score_row, unsafe_allow_html=True)

            # Score distribution chart
            score_bins = pd.cut(df_scored["AI Score"], bins=[0, 10, 20, 30, 40, 50, 200],
                                 labels=["0-10", "10-20", "20-30", "30-40", "40-50", "50+"])
            score_dist = score_bins.value_counts().sort_index().reset_index()
            score_dist.columns = ["Score Range", "Articles"]
            score_colors = ["#2e4470", "#0047e1", "#00b4ff", "#00e5c8", "#ffaa00", "#ff6400"]
            fig_score = go.Figure(go.Bar(
                x=score_dist["Score Range"].astype(str),
                y=score_dist["Articles"],
                marker_color=score_colors[:len(score_dist)],
                marker_line_width=0,
                text=score_dist["Articles"], textposition="outside",
                textfont=dict(color=_TITLE, size=10, family="DM Mono, monospace"),
                hovertemplate="<b>Score %{x}</b><br>📰 %{y} articles<extra></extra>",
            ))
            _dark(fig_score, 270)
            fig_score.update_layout(
                title=dict(text="AI Score Distribution", font=dict(color=_TITLE, size=13, family="Syne, sans-serif"), x=0.01),
                bargap=0.3,
            )
            st.plotly_chart(fig_score, use_container_width=True, config={"displayModeBar": False}, key="pc_29")

            # High-signal articles with orange top-strip
            st.markdown('<div class="sec-head">🔴 High Signal Articles (Score ≥ 30)</div>', unsafe_allow_html=True)
            top_scored = df_scored[df_scored["AI Score"] >= 30].head(25)
            if top_scored.empty:
                empty_state("No articles scored 30+ in current view.", hint="Try a broader scan.", icon="\U0001f534")
            else:
                for _, row in top_scored.iterrows():
                    st.markdown(article_card(
                        row["Headline"], row["Date"], row["URL"],
                        row["Source"], row["Country"], row["Topic"],
                        row.get("Capacity", ""), row.get("Deal Size", ""),
                        row.get("Sentiment", "News"),
                        ai_score=row["AI Score"],
                    ), unsafe_allow_html=True)

            # Medium signal
            st.markdown('<div class="sec-head">🟡 Medium Signal Articles (15–29)</div>', unsafe_allow_html=True)
            med_scored = df_scored[(df_scored["AI Score"] >= 15) & (df_scored["AI Score"] < 30)].head(20)
            for _, row in med_scored.iterrows():
                st.markdown(article_card(
                    row["Headline"], row["Date"], row["URL"],
                    row["Source"], row["Country"], row["Topic"],
                    row.get("Capacity", ""), row.get("Deal Size", ""),
                    row.get("Sentiment", "News"),
                    ai_score=row["AI Score"],
                ), unsafe_allow_html=True)

            # Full scored table
            st.markdown('<div class="sec-head">Full Scored Article Table</div>', unsafe_allow_html=True)
            scored_display = df_scored[["AI Score","Headline","Date","Topic","Sentiment","Capacity","Deal Size","Country","ISO / RTO","URL"]].copy() if "ISO / RTO" in df_scored.columns else df_scored[["AI Score","Headline","Date","Topic","Sentiment","Capacity","Deal Size","Country","URL"]].copy()
            st.markdown(dark_table(scored_display), unsafe_allow_html=True)


    with tab6:
        st.markdown('<div class="sec-head">Export Data</div>', unsafe_allow_html=True)
        col_a, col_b = st.columns(2)
        ts    = datetime.now().strftime("%Y%m%d_%H%M")
        label = re.sub(r"[^a-zA-Z0-9]", "_", time_opt)

        with col_a:
            st.markdown(
                '<div style="background:#0b1628;border:1px solid #152038;border-radius:10px;'
                'padding:1.1rem 1.2rem;margin-bottom:.8rem;">'
                '<div style="font-size:1.5rem;margin-bottom:.4rem;">\U0001f4ca</div>'
                '<div style="font-family:Syne,sans-serif;font-weight:700;color:#b8c8e0;'
                'font-size:.95rem;margin-bottom:.3rem;">Excel Report (.xlsx)</div>'
                '<div style="font-size:.78rem;color:#2a3e60;line-height:1.5;">'
                '5 sheets: All Articles \u00b7 By Country \u00b7 By Region \u00b7 By Topic \u00b7 By Company<br>'
                'Colour-coded badges, auto-filter, frozen headers, clickable URLs.</div></div>',
                unsafe_allow_html=True,
            )
            st.download_button(
                "\U0001f4e5 Download Excel Report",
                data=build_excel(df),
                file_name=f"GlobalDCIntel_{label}_{ts}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True,
            )

        with col_b:
            st.markdown(
                '<div style="background:#0b1628;border:1px solid #152038;border-radius:10px;'
                'padding:1.1rem 1.2rem;margin-bottom:.8rem;">'
                '<div style="font-size:1.5rem;margin-bottom:.4rem;">\U0001f4c4</div>'
                '<div style="font-family:Syne,sans-serif;font-weight:700;color:#b8c8e0;'
                'font-size:.95rem;margin-bottom:.3rem;">CSV Export</div>'
                '<div style="font-size:.78rem;color:#2a3e60;line-height:1.5;">'
                'Flat CSV of the filtered view.<br>'
                'Ready for Excel, Python, PowerBI, or Tableau.</div></div>',
                unsafe_allow_html=True,
            )
            st.download_button(
                "\U0001f4e5 Download CSV",
                data=df.to_csv(index=False).encode(),
                file_name=f"GlobalDCIntel_{label}_{ts}.csv",
                mime="text/csv",
                use_container_width=True,
            )

        st.markdown('<div class="sec-head">Full Article Preview</div>', unsafe_allow_html=True)
        display_cols = ["Headline","Date","Source","Country","Region",
                        "Topic","Sentiment","Capacity","Deal Size","Companies","ISO / RTO","URL"]
        st.markdown(
            dark_table(df[[c for c in display_cols if c in df.columns]]),
            unsafe_allow_html=True,
        )

    # ── Platform footer ───────────────────────────────────────────────────────
    st.markdown(
        '<div style="margin-top:3rem;padding:1.2rem 0 .6rem;border-top:1px solid #101b2e;'
        'text-align:center;">'
        '<div style="font-family:\'DM Mono\',monospace;font-size:.6rem;letter-spacing:.14em;'
        'color:#1a2e50;text-transform:uppercase;">'
        'Global Energy Intelligence Platform &nbsp;·&nbsp; Data Centers &nbsp;·&nbsp; Renewables Power Markets &nbsp;·&nbsp; Wood Mac'
        '</div></div>',
        unsafe_allow_html=True,
    )


if __name__ == "__main__":
    main()
